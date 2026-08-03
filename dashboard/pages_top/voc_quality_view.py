from __future__ import annotations

import ast
import base64
import json
import re
import time
from copy import deepcopy
from concurrent.futures import ThreadPoolExecutor
from datetime import date, datetime, timedelta
from functools import partial
from html import escape
from io import BytesIO
from pathlib import Path

import altair as alt
import pandas as pd
import streamlit as st

from components.integration_status import load_integration_status, render_integration_status
from components.quality_report_template import build_voc_quality_report_html
from core.paths import JIRA_REGISTERED_ISSUES_FILE
from core.storage import load_json_file, save_json_file
from services.jira_client import (
    JiraConfigurationError,
    JiraIssueCreateError,
    create_jira_issue,
    jira_environment_snapshot,
)
from services.voc_background_job_service import (
    background_job_snapshot,
    discard_background_job,
    start_background_job,
    update_background_job,
)
from services.voc_quality_state_model import (
    VOC_STATUS_DISPLAY_LABELS,
    validity_human_review_readiness,
    voc_case_next_action,
    voc_run_next_action,
    voc_status_label,
)
from services.voc_quality_service import (
    REPORT_CATEGORIES,
    QUALITY_RUBRIC_SPECS,
    a2a_trace_snapshot,
    agent_status_snapshot,
    audit_summary,
    batch_preflight,
    build_voc_acceptance_snapshot,
    build_voc_quality_report,
    check_anthropic_agent_credential,
    check_gemini_agent_credential,
    check_openai_agent_credential,
    compare_voc_runs,
    compare_voc_improvement_answers,
    create_voc_defect,
    delete_voc_run_history,
    download_voc_run_evidence,
    execute_batch_run,
    get_batch_run_progress,
    judge_independence_preview,
    judge_provider_options,
    list_reports,
    list_voc_defects,
    list_voc_run_history,
    latest_voc_full_run_id,
    load_guide,
    load_improvement_validity_rubric,
    list_improvement_validity_candidates,
    load_independent_judge_rubric,
    load_quality_test_catalog,
    load_quality_rubric,
    load_unified_quality_cases,
    load_voc_case_history_detail,
    load_voc_defect,
    load_voc_run_history_detail,
    load_system_rubric,
    pipeline_trace_events,
    read_report,
    reevaluate_voc_run_case,
    evaluate_voc_improvement_validity,
    generate_voc_acceptance_evidence,
    generate_voc_quality_report,
    review_voc_improvement_validity,
    run_agent_action,
    run_diagnostics,
    run_test_case,
    run_voc_analysis,
    runtime_health,
    save_quality_rubric,
    save_quality_test_catalog,
    save_voc_rubric_reevaluation_plan,
    save_voc_validity_supplement,
    test_case_summary,
    test_agent_rpc,
    transition_voc_defect,
    request_batch_stop,
    start_batch_run,
    validate_quality_rubric,
    validate_quality_test_catalog,
    validity_provider_options,
)


RUBRIC_STAGE_OPTIONS = (
    "내부 파이프라인 품질",
    "독립 LLM 평가",
    "개선안 타당성 평가",
)
RUBRIC_STAGE_TYPES = {
    RUBRIC_STAGE_OPTIONS[0]: "internal_pipeline",
    RUBRIC_STAGE_OPTIONS[1]: "independent_judge",
    RUBRIC_STAGE_OPTIONS[2]: "improvement_validity",
}
RUBRIC_STAGE_DISPLAY_ALIASES = {
    "내부 Pipeline 품질": RUBRIC_STAGE_OPTIONS[0],
    "독립 LLM Judge": RUBRIC_STAGE_OPTIONS[1],
    "개선안 타당성": RUBRIC_STAGE_OPTIONS[2],
}
RUBRIC_STAGE_TYPES.update({
    old_label: RUBRIC_STAGE_TYPES[new_label]
    for old_label, new_label in RUBRIC_STAGE_DISPLAY_ALIASES.items()
})

MANUAL_JUDGE_PROVIDERS = (
    {
        "provider": "openai",
        "label": "OpenAI",
        "model": "gpt-5.2",
        "number": 1,
    },
    {
        "provider": "anthropic",
        "label": "Anthropic",
        "model": "claude-haiku-4-5",
        "number": 2,
    },
    {
        "provider": "gemini",
        "label": "Gemini",
        "model": "gemini-3.5-flash-lite",
        "number": 3,
    },
)

MANUAL_PROVIDER_LABELS = {
    "openai": "OpenAI",
    "anthropic": "Anthropic",
    "gemini": "Google Gemini",
    "google": "Google Gemini",
}


def _manual_provider_label(provider: str | None) -> str:
    provider_key = str(provider or "").strip().lower()
    return MANUAL_PROVIDER_LABELS.get(provider_key, provider_key.upper() if provider_key else "-")

MANUAL_PREPARATION_STEPS = (
    "Agent 실행 상태 점검",
    "Run 폴더 생성",
    "Rubric과 Test Case 스냅샷 저장",
    "증적 파일 준비",
    "별도 Python 프로세스 시작",
)
MANUAL_EVENT_CARD_HEIGHT = 154


AGENT_PIPELINE = (
    ("Interpreter", "질문 의도 해석", "6101"),
    ("Retriever", "관련 VOC 검색", "6102"),
    ("Summarizer", "요약 후보 생성", "6103"),
    ("Evaluator", "최적 후보 평가", "6104"),
    ("Critic", "요약·정책 검토", "6105"),
    ("Improver", "개선안 생성", "6106"),
)
AGENT_DISPLAY_NAMES_BY_KEY = {name.lower(): name for name, _role, _port in AGENT_PIPELINE}

VOC_RUN_STATUS_COLORS = {
    "PASS": "#155A96",
    "REVIEW_REQUIRED": "#2F75B5",
    "FAIL": "#5599D2",
    "ERROR": "#7EAED4",
    "NOT_RUN": "#A9CAE7",
}
VOC_HISTORY_COLORS = {
    "통과율": "#155A96",
    "검토율": "#5599D2",
    "실패·오류율": "#A9CAE7",
}
VOC_OVERVIEW_PANEL_HEIGHT = 390

VOC_STATUS_LABELS = {
    **VOC_STATUS_DISPLAY_LABELS,
    "PASS": "통과",
    "FAIL": "실패",
    "ERROR": "오류",
    "REVIEW_REQUIRED": "검토 필요",
    "NOT_RUN": "미실행",
    "RUNNING": "진행 중",
    "STARTED": "시작",
    "COMPLETED": "완료",
    "ENDED": "종료",
    "INTERRUPTED": "중단됨",
    "SUCCESS": "성공",
    "DRAFT": "초안",
    "PENDING": "대기",
    "CONFIRMED": "확인됨",
    "OPEN": "접수",
    "ANALYZED": "분석 완료",
    "FIXED": "조치 완료",
    "RETESTED": "재시험 완료",
    "CLOSED": "종결",
    "RESOLVED": "해결",
    "IMPLEMENTED": "실행 구현 완료",
    "DEFINED": "정의됨 · 후속 구현",
    "MANUAL": "수동 수행",
    "BATCH": "일괄 수행",
    "RETEST": "재시험",
    "VOC": "VOC",
    "FAULT": "장애 시험",
    "AI_PASS": "AI 평가 통과",
    "AI_REVIEWED": "AI 평가 완료",
    "QA_REVIEWED": "QA 검토 완료",
    "REVISION_REQUIRED": "보완 필요",
    "REJECTED": "반려",
    "APPROVE": "승인",
    "APPROVED": "승인 완료",
    "FORMAL_APPROVED": "정식 승인",
    "FORMAL_QUALITY_APPROVED": "정식 품질 승인",
    "NOT_APPROVED": "미승인",
    "BUSINESS_APPROVED": "업무 승인 완료",
    "BUSINESS_REVIEW_REQUIRED": "업무 검토 필요",
    "HUMAN_REVIEW_REQUIRED": "사람 검토 필요",
    "REMAINING_CASE_REVIEW_REQUIRED": "잔여 Case 검토 필요",
    "PARTIALLY_APPROVED": "일부 승인",
    "READY_FOR_UAT": "UAT 준비 완료",
    "HOLD": "보류",
    "EVIDENCE_DRAFT": "증적 초안",
    "NOT_CONFIGURED": "미설정",
    "CONFIGURED": "설정됨",
    "NOT_AVAILABLE": "확인 불가",
    "NOT_EVALUATED": "평가 전",
    "UNKNOWN": "미확인",
    "STOPPED": "중지",
    "STARTING/FAILED": "시작 실패",
    "CRITICAL": "치명",
    "HIGH": "높음",
    "MEDIUM": "중간",
    "LOW": "낮음",
    "INTERFACE_BRANCH": "연계·분기",
    "API_RATE_LIMIT": "API 제한",
    "AGENT_FAILURE": "Agent 장애",
    "DATA": "데이터",
    "PERFORMANCE": "성능",
    "OTHER": "기타",
    "VALIDITY_EVALUATION_REQUIRED": "개선안 타당성 평가 필요",
    "REWORK_REQUIRED": "보완·재시험 필요",
    "QA_REVIEW": "QA 검토 가능",
    "BUSINESS_APPROVAL": "업무 승인 가능",
    "NO_ACTION": "추가 조치 없음",
    "미판정": "미판정",
}
VOC_EXECUTION_TYPE_LABELS = {
    "voc_pipeline": "VOC 파이프라인",
    "fault_proxy": "장애 대리 실행",
    "isolated_fault": "격리 장애",
    "agent_role_quality": "Agent 역할 품질",
    "quality_gate": "품질 게이트",
    "defined_only": "정의만 있음",
}
JUDGE_INDEPENDENCE_GRADE_LABELS = {
    "A": "독립성 좋음 · Provider 분리",
    "B": "독립성 보통 · 모델 분리",
    "C": "독립성 낮음 · 편향 검토 필요",
}
VALIDITY_RUN_TYPE_FILTERS = ("전체", "수동 수행", "일괄 수행", "재시험")
VALIDITY_RUN_TYPE_FILTER_VALUES = {
    "수동 수행": "MANUAL",
    "일괄 수행": "BATCH",
    "재시험": "RETEST",
}
VALIDITY_STATUS_FILTER_BY_ACTION = {
    "RUN_VALIDITY": "전체",
    "REWORK_AND_RETEST": "평가 완료",
    "QA_REVIEW": "QA 검토 가능",
    "BUSINESS_APPROVAL": "업무 승인 가능",
    "FORMAL_APPROVED": "정식 승인",
    "REPORT_READY": "정식 승인",
}
VALIDITY_CRITERIA_LABELS = {
    "complaint_to_root_cause": "불만↔근본 원인",
    "root_cause_to_action": "원인↔개선 조치",
    "expected_customer_impact": "고객 영향 개선",
    "voc_id_reference": "VOC ID 근거",
    "trace_and_agent_reference": "실행 Trace·Agent 근거",
    "no_unsupported_evidence": "미확인 근거 배제",
    "process_feasibility": "업무 프로세스 실행성",
    "technical_feasibility": "기술 실행성",
    "resource_and_dependency_awareness": "자원·의존성 인식",
    "responsible_owner": "담당자 명확성",
    "target_schedule": "목표 일정",
    "measurable_kpi": "측정 KPI",
    "priority": "우선순위",
    "customer_and_operational_risk": "고객·운영 리스크",
    "privacy_and_security": "개인정보·보안",
    "compliance_and_escalation": "규제·에스컬레이션",
}
VALIDITY_HOLD_RULE_LABELS = {
    "missing_voc_or_trace_evidence": "VOC·실행 Trace 근거 부족",
    "unsafe_or_noncompliant_action": "안전·규제 부적합 조치",
    "unresolved_high_or_critical_defect": "미종결 High/Critical 결함",
    "judge_error_or_not_run": "독립 LLM 평가 미통과 또는 미수행",
    "safety_regression_against_baseline": "기존 대비 안전성 회귀",
}


def _voc_status_label(value, default: str = "-") -> str:
    text = voc_status_label(value, default=default)
    return VOC_STATUS_LABELS.get(text, VOC_STATUS_LABELS.get(str(text).upper(), text))


def _judge_independence_grade_label(value, default: str = "-") -> str:
    text = str(value or "").strip().upper()
    if not text:
        return default
    return JUDGE_INDEPENDENCE_GRADE_LABELS.get(text, text)


def _voc_status_counts_for_display(counts: dict | None) -> dict:
    return {_voc_status_label(key): value for key, value in (counts or {}).items()}


def _voc_display_term(value: object) -> str:
    """Normalize legacy product terms only at UI display boundaries."""
    text = "" if value is None else str(value)
    replacements = {
        "VOC·Trace": "VOC·실행 Trace",
        "Agent Pipeline": "Agent 파이프라인",
        "A2A Pipeline": "Agent 파이프라인",
        "독립 Judge": "독립 LLM 평가",
        "독립 LLM Judge": "독립 LLM 평가",
        "Judge Provider": "독립 LLM 평가 Provider",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text


VOC_UI_DESIGN_RULES = (
    "화면에는 지금 판단해야 하는 정보만 먼저 보여준다.",
    "상세 로그·원본 Trace·산식은 접힘/팝업/상세 화면으로 보낸다.",
    "Run ID보다 Case, 질문, 상태, 다음 액션을 먼저 보이게 한다.",
    "상태는 한글 라벨, 색상 배지, Material Symbols 아이콘을 함께 사용한다.",
    "같은 성격의 카드·리스트·팝업은 같은 높이와 같은 정보 순서를 유지한다.",
)
VOC_UI_BADGE_TONES = {
    "green": "green",
    "blue": "blue",
    "orange": "orange",
    "red": "red",
    "gray": "gray",
}


def _voc_ui_badge(label: object, tone: str = "gray") -> str:
    label_text = _voc_display_term(label or "확인")
    tone_name = VOC_UI_BADGE_TONES.get(str(tone or "").lower(), "gray")
    return f":{tone_name}-badge[{label_text}]"


def _voc_bool_tone(value: bool, *, pending_tone: str = "orange") -> str:
    return "green" if bool(value) else pending_tone


def _render_voc_section_heading(
    title: str,
    caption: str = "",
    *,
    icon: str = "info",
    badges: list[tuple[str, str]] | tuple[tuple[str, str], ...] = (),
    right_caption: str = "",
) -> None:
    left, right = st.columns([2.4, 1.15], gap="small", vertical_alignment="center")
    with left:
        st.markdown(f"#### :material/{icon}: {title}")
        if caption:
            st.caption(_voc_display_term(caption))
    with right:
        if badges:
            st.markdown(" ".join(_voc_ui_badge(label, tone) for label, tone in badges), text_alignment="right")
        if right_caption:
            st.caption(_voc_display_term(right_caption), text_alignment="right")


def _render_voc_summary_card_styles() -> None:
    st.markdown(
        """
        <style>
        .vqd-action-card-grid{display:grid;grid-template-columns:repeat(var(--vqd-action-cols,3),minmax(0,1fr));gap:8px;margin:0 0 8px;font-family:'Malgun Gothic','Apple SD Gothic Neo',sans-serif}
        .vqd-action-card{border:1px solid #c8d9ee;border-left:4px solid #7b8797;border-radius:8px;background:linear-gradient(145deg,#fff,#f8fbff);box-shadow:0 3px 10px rgba(22,78,128,.05);padding:9px 10px;box-sizing:border-box;display:grid;grid-template-rows:22px 1fr 17px;gap:2px;overflow:hidden;min-width:0}
        .vqd-action-card-head{display:flex;align-items:center;justify-content:space-between;gap:8px;min-width:0}
        .vqd-action-label{display:flex;align-items:center;gap:5px;min-width:0;color:#40536d;font-size:11px;font-weight:800;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
        .vqd-action-icon{display:flex;width:16px;min-width:16px;color:#155a96;flex:0 0 auto}.vqd-action-icon svg{width:16px;height:16px}
        .vqd-action-badge{flex:0 0 auto;display:inline-flex;align-items:center;justify-content:center;height:20px;padding:0 7px;border-radius:999px;background:#eef2f7;color:#64748b;border:1px solid #d8e2ee;font-size:9px;font-weight:850;white-space:nowrap}
        .vqd-action-card strong{align-self:center;color:#073b72;font-size:21px;line-height:1.12;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
        .vqd-action-card small{color:#728095;font-size:9px;line-height:1.2;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
        .vqd-action-card.blue{border-left-color:#155a96}.vqd-action-card.blue strong,.vqd-action-card.blue .vqd-action-icon{color:#155a96}.vqd-action-card.blue .vqd-action-badge{background:#eaf3fb;color:#155a96;border-color:#b9d2ec}
        .vqd-action-card.green{border-left-color:#299049}.vqd-action-card.green strong,.vqd-action-card.green .vqd-action-icon{color:#299049}.vqd-action-card.green .vqd-action-badge{background:#eaf7ef;color:#176b35;border-color:#a9d7b8}
        .vqd-action-card.orange{border-left-color:#b36a08}.vqd-action-card.orange strong,.vqd-action-card.orange .vqd-action-icon{color:#b36a08}.vqd-action-card.orange .vqd-action-badge{background:#fff7e6;color:#92550a;border-color:#e8c47b}
        .vqd-action-card.red{border-left-color:#d83f36}.vqd-action-card.red strong,.vqd-action-card.red .vqd-action-icon{color:#d83f36}.vqd-action-card.red .vqd-action-badge{background:#fff0ee;color:#b42318;border-color:#efaaa4}
        .vqd-action-card.gray{border-left-color:#9aa5b1}.vqd-action-card.gray strong,.vqd-action-card.gray .vqd-action-icon{color:#718096}.vqd-action-card.gray .vqd-action-badge{background:#f1f3f5;color:#7b8797;border-color:#d8dee5}
        @media(max-width:1100px){.vqd-action-card-grid{grid-template-columns:repeat(2,minmax(0,1fr))}}
        @media(max-width:720px){.vqd-action-card-grid{grid-template-columns:repeat(2,1fr)}}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _render_voc_summary_cards(
    cards: list[dict],
    *,
    columns: int | None = None,
    height: int = 118,
    header_badge: bool = True,
) -> None:
    if not cards:
        return
    column_count = max(1, min(columns or len(cards), len(cards), 5))
    if header_badge:
        _render_voc_summary_card_styles()
        for start in range(0, len(cards), column_count):
            row_cards = cards[start:start + column_count]
            card_html = "".join(
                _voc_summary_card_html(card, height=height)
                for card in row_cards
            )
            st.markdown(
                (
                    f"<div class='vqd-action-card-grid' "
                    f"style='--vqd-action-cols:{column_count};--vqd-action-height:{int(height)}px'>"
                    f"{card_html}</div>"
                ),
                unsafe_allow_html=True,
            )
        return
    for start in range(0, len(cards), column_count):
        row_cards = cards[start:start + column_count]
        for column, card in zip(st.columns(len(row_cards), gap="small"), row_cards, strict=False):
            tone = str(card.get("tone") or "gray")
            badge = card.get("badge")
            with column.container(border=True, height=height):
                st.caption(f":material/{card.get('icon', 'info')}: {_voc_display_term(card.get('label', '-'))}")
                st.markdown(f"#### {_voc_display_term(card.get('value', '-'))}")
                detail = card.get("detail")
                if detail:
                    st.caption(_voc_display_term(detail))
                if badge:
                    st.markdown(_voc_ui_badge(badge, tone))


def _voc_summary_card_svg_icon(name: object) -> str:
    icon_name = str(name or "").strip()
    if icon_name.startswith(":material/") and icon_name.endswith(":"):
        icon_name = icon_name.removeprefix(":material/").removesuffix(":")
    paths = {
        "account_tree": "<path d='M6 5h6v5H6zM14 14h6v5h-6zM3 14h6v5H3z'/><path d='M9 10v2h8v2M6 10v4'/>",
        "approval": "<path d='M5 4h14v16H5z'/><path d='m8 13 2.5 2.5L16 9'/><path d='M8 7h8'/>",
        "article": "<path d='M6 3h9l4 4v14H6z'/><path d='M15 3v5h5M9 12h7m-7 4h7'/>",
        "assignment_ind": "<path d='M6 3h12v18H6z'/><circle cx='12' cy='10' r='2.5'/><path d='M8.5 17c.7-2 2-3 3.5-3s2.8 1 3.5 3'/>",
        "block": "<circle cx='12' cy='12' r='9'/><path d='M5.6 5.6 18.4 18.4'/>",
        "check_circle": "<circle cx='12' cy='12' r='9'/><path d='m8 12 3 3 6-7'/>",
        "checklist": "<path d='m4 7 1.5 1.5L8 5.5M11 7h9M4 13l1.5 1.5L8 11.5M11 13h9M4 19l1.5 1.5L8 17.5M11 19h9'/>",
        "conversion_path": "<path d='M4 6h7a4 4 0 0 1 4 4v8'/><path d='m12 15 3 3 3-3M4 18h4'/>",
        "dashboard": "<rect x='3' y='3' width='8' height='8' rx='1'/><rect x='13' y='3' width='8' height='5' rx='1'/><rect x='13' y='10' width='8' height='11' rx='1'/><rect x='3' y='13' width='8' height='8' rx='1'/>",
        "draft": "<path d='M6 3h9l4 4v14H6z'/><path d='M15 3v5h5M9 15h6'/>",
        "fact_check": "<path d='M5 4h14v16H5z'/><path d='m8 10 2 2 5-5m-7 9h8'/>",
        "edit_note": "<path d='M4 19h8'/><path d='M14 4l6 6-8 8H6v-6z'/><path d='m15 5 4 4'/>",
        "error": "<circle cx='12' cy='12' r='9'/><path d='M12 7v6m0 4h.01'/>",
        "format_list_numbered": "<path d='M10 6h10M10 12h10M10 18h10M4 6h1v3M4 12h2l-2 3h2M4 18h2'/>",
        "help": "<circle cx='12' cy='12' r='9'/><path d='M9.5 9a2.6 2.6 0 1 1 4.2 2c-.9.6-1.7 1.2-1.7 2.5M12 17h.01'/>",
        "history": "<path d='M4 12a8 8 0 1 0 2.3-5.7L4 8.6'/><path d='M4 4v4h4M12 8v5l3 2'/>",
        "hub": "<circle cx='12' cy='12' r='2.5'/><circle cx='5' cy='6' r='2'/><circle cx='19' cy='6' r='2'/><circle cx='19' cy='18' r='2'/><path d='M7 7l3.2 3.2M14 10.2 17.2 7M14 13.8 17.2 17'/>",
        "inventory_2": "<path d='M4 7 12 3l8 4-8 4z'/><path d='M4 7v10l8 4 8-4V7M12 11v10'/>",
        "low_priority": "<path d='M4 7h9M4 12h6M4 17h3M14 7h6m-3-3 3 3-3 3M12 17h8m-3-3 3 3-3 3'/>",
        "menu_book": "<path d='M4 5.5A3.5 3.5 0 0 1 7.5 2H12v18H7.5A3.5 3.5 0 0 0 4 23zM20 5.5A3.5 3.5 0 0 0 16.5 2H12v18h4.5A3.5 3.5 0 0 1 20 23z'/>",
        "pending_actions": "<path d='M6 3h12v18H6z'/><path d='M9 7h6M9 11h3'/><circle cx='16' cy='16' r='4'/><path d='M16 14v2l1.5 1'/>",
        "play_circle": "<circle cx='12' cy='12' r='9'/><path d='m10 8 6 4-6 4z'/>",
        "playlist_play": "<path d='M4 6h10M4 11h10M4 16h6'/><path d='m14 15 6 4-6 4z'/>",
        "psychology": "<path d='M9 18H7a4 4 0 0 1 0-8 5 5 0 0 1 10 1 3.5 3.5 0 0 1-1 6.8V21h-5v-3'/><path d='M10 10h.01M14 10h.01M11 14h3'/>",
        "published_with_changes": "<path d='M12 3a9 9 0 1 1-8.2 5.3'/><path d='M3 4v4h4M8 12l3 3 6-7'/>",
        "query_stats": "<path d='M4 19V5m0 14h16'/><path d='M7 15l3-4 3 2 4-7'/><circle cx='17' cy='6' r='2'/>",
        "rate_review": "<path d='M4 5h16v11H8l-4 4z'/><path d='m8 10 2 2 5-5m-7 7h8'/>",
        "replay": "<path d='M4 7v5h5'/><path d='M5.5 12a7 7 0 1 0 1.8-4.7L4 10'/>",
        "rule": "<path d='M5 4h14v16H5z'/><path d='m8 9 2 2 4-4M8 16h8'/>",
        "schedule": "<circle cx='12' cy='12' r='9'/><path d='M12 7v5l3 2'/>",
        "score": "<path d='M4 19h16'/><path d='M6 16V9m6 7V5m6 11v-4'/>",
        "shield": "<path d='M12 3 4 6v6c0 5 3.4 8.3 8 10 4.6-1.7 8-5 8-10V6z'/><path d='m8 12 3 3 5-6'/>",
        "smart_toy": "<rect x='5' y='8' width='14' height='10' rx='3'/><path d='M12 4v4M8 13h.01M16 13h.01M9 18v2h6v-2'/>",
        "summarize": "<path d='M6 3h9l4 4v14H6z'/><path d='M15 3v5h5M9 12h7M9 16h5'/>",
        "sync_alt": "<path d='M4 7h13m-4-4 4 4-4 4M20 17H7m4-4-4 4 4 4'/>",
        "task_alt": "<circle cx='12' cy='12' r='9'/><path d='m8 12 3 3 6-7'/>",
        "touch_app": "<path d='M8 11V6a2 2 0 1 1 4 0v6'/><path d='M12 12v-2a2 2 0 1 1 4 0v3'/><path d='M16 13v-1a2 2 0 1 1 4 0v4a5 5 0 0 1-5 5h-3a5 5 0 0 1-4.4-2.6L5 14a1.8 1.8 0 0 1 3-2z'/>",
        "tune": "<path d='M4 7h10M18 7h2M4 17h2M10 17h10M8 14v6M16 4v6'/>",
        "verified": "<path d='M12 3 4 6v6c0 5 3.4 8.3 8 10 4.6-1.7 8-5 8-10V6z'/><path d='m8.5 12 2.5 2.5 5-6'/>",
        "bug_report": "<path d='M8 7h8v10a4 4 0 0 1-8 0z'/><path d='M9 7 7 4m8 3 2-3M4 13h4m8 0h4M5 19l3-2m8 0 3 2'/>",
        "warning": "<path d='M12 3 2.8 20h18.4L12 3Z'/><path d='M12 9v5m0 3h.01'/>",
    }
    path = paths.get(icon_name, "<circle cx='12' cy='12' r='8'/><path d='M8 12h8'/>")
    return (
        "<svg viewBox='0 0 24 24' aria-hidden='true' fill='none' stroke='currentColor' "
        "stroke-width='1.8' stroke-linecap='round' stroke-linejoin='round'>"
        + path
        + "</svg>"
    )


def _voc_summary_card_html(card: dict, *, height: int) -> str:
    tone = str(card.get("tone") or "gray").lower()
    if tone not in {"blue", "green", "orange", "red", "gray"}:
        tone = "gray"
    icon = _voc_summary_card_svg_icon(card.get("icon"))
    label = escape(_voc_display_term(card.get("label", "-")))
    value = escape(_voc_display_term(card.get("value", "-")))
    detail = escape(_voc_display_term(card.get("detail", "")))
    badge = escape(_voc_display_term(card.get("badge", "")))
    badge_html = f"<span class='vqd-action-badge'>{badge}</span>" if badge else ""
    return (
        f"<article class='vqd-action-card {tone}' style='height:{int(height)}px'>"
        "<div class='vqd-action-card-head'>"
        f"<span class='vqd-action-label'><span class='vqd-action-icon'>{icon}</span>{label}</span>"
        f"{badge_html}"
        "</div>"
        f"<strong>{value}</strong>"
        f"<small>{detail}</small>"
        "</article>"
    )


def _release_scope_basis_cards(release_scope: dict, release_decision: str = "") -> list[dict]:
    if not isinstance(release_scope, dict) or not release_scope:
        return []
    voc_counts = release_scope.get("voc_counts", {}) if isinstance(release_scope.get("voc_counts"), dict) else {}
    fault_counts = release_scope.get("fault_counts", {}) if isinstance(release_scope.get("fault_counts"), dict) else {}
    pending_counts = release_scope.get("pending_counts", {}) if isinstance(release_scope.get("pending_counts"), dict) else {}
    voc_total = int(release_scope.get("voc_count") or 0)
    fault_total = int(release_scope.get("fault_count") or 0)
    pending_total = int(release_scope.get("pending_count") or 0)
    voc_pass = int(voc_counts.get("PASS") or 0)
    fault_confirmed = int(fault_counts.get("PASS") or 0) + int(fault_counts.get("REVIEW_REQUIRED") or 0)
    pending_approved = int(pending_counts.get("NOT_RUN") or 0)
    linked_retest_count = int(release_scope.get("linked_retest_count") or 0)
    decision_label = _voc_status_label(release_decision or release_scope.get("release_decision") or "NOT_APPROVED")
    is_formal = str(release_decision or release_scope.get("release_decision") or "").upper() == "FORMAL_APPROVED"
    return [
        {
            "icon": "task_alt",
            "label": "VOC 개선 Case",
            "value": f"{voc_pass}/{voc_total} PASS",
            "detail": "답변 품질·독립 LLM·업무 승인 완료",
            "tone": _voc_bool_tone(voc_total > 0 and voc_pass >= voc_total),
            "badge": "통과",
        },
        {
            "icon": "shield",
            "label": "장애 검증 Case",
            "value": f"{fault_confirmed}/{fault_total} 실행 확인",
            "detail": "보호 동작 확인 대상은 REVIEW_REQUIRED로 관리",
            "tone": _voc_bool_tone(fault_total == 0 or fault_confirmed >= fault_total),
            "badge": "실행 확인",
        },
        {
            "icon": "pending_actions",
            "label": "후속 구현 Case",
            "value": f"{pending_approved}/{pending_total} 승인",
            "detail": "이번 회차 NOT_RUN이 정상 상태",
            "tone": _voc_bool_tone(pending_total == 0 or pending_approved >= pending_total),
            "badge": "후속 승인",
        },
        {
            "icon": "sync_alt",
            "label": "보완 RETEST",
            "value": f"{linked_retest_count}건 연결",
            "detail": "부족 Case의 승인 증적을 원본 Run에 반영",
            "tone": "green" if linked_retest_count else "gray",
            "badge": "연결",
        },
        {
            "icon": "verified",
            "label": "최종 판정",
            "value": decision_label,
            "detail": "품질 보고서와 최종 인수 게이트 연결",
            "tone": "green" if is_formal else "orange",
            "badge": "최종 상태",
        },
    ]


def _render_release_scope_basis(
    release_scope: dict,
    *,
    release_decision: str = "",
    title: str = "35건 최종 인수 기준",
    caption: str = "최종 판정은 실행 결과, 승인 증적, 후속 구현 범위를 분리해서 판단합니다.",
) -> None:
    cards = _release_scope_basis_cards(release_scope, release_decision)
    if not cards:
        return
    badges = [
        (_voc_status_label(release_decision), "green" if str(release_decision).upper() == "FORMAL_APPROVED" else "orange")
    ] if release_decision else []
    with st.container(border=True):
        _render_voc_section_heading(
            title,
            caption,
            icon="rule",
            badges=badges,
            right_caption=f"전체 {int(release_scope.get('selected_count') or release_scope.get('catalog_total_cases') or 0)}건",
        )
        _render_voc_summary_cards(cards, columns=5, height=132)
        st.caption(
            "시연 멘트 기준: VOC 개선 Case는 PASS, 장애 검증 Case는 실행 확인, "
            "후속 구현 Case는 승인된 NOT_RUN, 보완 RETEST는 원본 Run 반영으로 설명합니다."
        )

VOC_PAGE_META = {
    "Dashboard": {
        "icon": "dashboard",
        "title": "VOC 품질 Dashboard",
        "description": "실행 환경부터 품질 판정, 독립 LLM 평가, 결함과 Agent 파이프라인 연결 상태를 한눈에 확인합니다.",
        "group": "현황",
        "flow": ("기간 설정", "품질 비교", "이슈 확인"),
    },
    "수동 TC 수행": {
        "icon": "fact_check",
        "title": "수동 TC 수행",
        "description": "Test Case를 선택해 Agent 파이프라인과 독립 LLM 평가 근거를 단계별로 확인합니다.",
        "group": "실행",
        "flow": (),
    },
    "일괄 TC 수행": {
        "icon": "playlist_play",
        "title": "일괄 TC 수행",
        "description": "다수 Test Case를 백그라운드로 실행하고 단계·예상시간·결과를 추적합니다.",
        "group": "실행",
        "flow": (),
    },
    "수행 이력": {
        "icon": "history",
        "title": "수행 이력",
        "description": "Run별 품질 판정과 Case 증적을 비교하고 재평가·다운로드까지 연결합니다.",
        "group": "추적",
        "flow": (),
    },
    "개선안 타당성 검증": {
        "icon": "verified",
        "title": "개선안 타당성 검증",
        "description": "VOC 개선안의 원인 연결성, 실행 가능성, 책임·일정·위험을 독립적으로 검증합니다.",
        "group": "평가",
        "flow": (),
    },
    "Agent 관리": {
        "icon": "smart_toy",
        "title": "Agent 관리",
        "description": "6개 Agent의 실행 상태와 포트·PID를 확인하고 안전하게 제어합니다.",
        "group": "운영",
        "flow": ("상태 확인", "제어 승인", "실행 결과"),
    },
    "VOC 분석": {
        "icon": "query_stats",
        "title": "VOC 분석",
        "description": "자연어 VOC를 Agent 파이프라인에 전달해 요약, 정책 개선안과 실행 Trace를 생성합니다.",
        "group": "분석",
        "flow": ("질문 입력", "Agent 분석", "개선안 확인"),
    },
    "테스트케이스": {
        "icon": "checklist",
        "title": "VOC 테스트케이스",
        "description": "의도·키워드·필수 요소·금지 요소 기반의 품질 Test Case를 조회합니다.",
        "group": "기준",
        "flow": (),
    },
    "품질 평가 기준": {
        "icon": "tune",
        "title": "품질 평가 기준 수립",
        "description": "내부 파이프라인·독립 LLM 평가·개선안 타당성 평가의 배점과 판정 구간을 시각적으로 관리합니다.",
        "group": "기준",
        "flow": (),
    },
    "장애·결함 관리": {
        "icon": "bug_report",
        "title": "장애·결함 관리",
        "description": "품질 결함과 격리 장애시험을 등록하고 상태·심각도·증적을 추적합니다.",
        "group": "결함",
        "flow": ("유형 선택", "결함 처리", "상태 추적"),
    },
    "A2A Trace": {
        "icon": "hub",
        "title": "Agent 파이프라인 Trace",
        "description": "Agent 간 실제 호출 경로와 성공·실패, 처리시간 및 전달 정보를 확인합니다.",
        "group": "추적",
        "flow": ("실행 Trace 집계", "연결 진단", "Report 확인"),
    },
    "품질 보고서": {
        "icon": "article",
        "title": "품질 보고서",
        "description": "Run 증적을 정량 분석하고 승인 판단에 필요한 보고서를 생성합니다.",
        "group": "보고",
        "flow": ("Run 선택", "수치 대조", "증적 생성"),
    },
    "사용자 가이드": {
        "icon": "menu_book",
        "title": "사용자 가이드",
        "description": "VOC 품질진단의 실행·이식·운영 절차를 목적별로 확인합니다.",
        "group": "안내",
        "flow": ("가이드 선택", "절차 확인", "실행 적용"),
    },
    "최종 인수·시연": {
        "icon": "approval",
        "title": "최종 인수·시연",
        "description": "전체 Run의 품질 게이트와 잔여 위험을 대조해 최종 UAT 준비 상태를 판단합니다.",
        "group": "승인",
        "flow": ("Run 연결", "Gate 검증", "인수 증적"),
    },
}


def _render_voc_design_system() -> None:
    st.html(
        """
        <style>
        .st-key-voc_page_hero {
            background:linear-gradient(118deg,#f4f9ff 0%,#ffffff 62%,#edf5fd 100%);
            border:1px solid #bfd4e9!important;border-left:5px solid #155a96!important;
            border-radius:14px!important;padding:16px 20px!important;margin:0 0 14px;
            box-shadow:0 7px 20px rgba(21,90,150,.08);
        }
        .st-key-voc_page_hero h2{color:#0b4478!important;font-size:25px!important;letter-spacing:-.6px;margin:0!important}
        .st-key-voc_page_hero [data-testid="stCaptionContainer"]{color:#526a83!important;line-height:1.5}
        .st-key-voc_page_flow{background:rgba(255,255,255,.72);border:1px solid #d3e2f0;border-radius:10px;padding:9px 12px!important}
        .st-key-voc_page_content>div[data-testid="stVerticalBlock"]{gap:.72rem}
        .st-key-voc_page_content div[data-testid="stVerticalBlockBorderWrapper"]{
            border-color:#c8d9ee!important;border-radius:11px!important;
            background:linear-gradient(145deg,#ffffff 0%,#fbfdff 100%);
            box-shadow:0 3px 12px rgba(21,90,150,.045);
        }
        .st-key-voc_page_content [data-testid="stMetric"]{
            min-height:82px;padding:11px 13px;border:1px solid #d1dfed;border-radius:10px;
            background:linear-gradient(145deg,#fff,#f5f9fd);box-shadow:0 2px 8px rgba(21,90,150,.04)
        }
        .st-key-voc_page_content [data-testid="stMetricLabel"]{color:#4a6078;font-weight:700}
        .st-key-voc_page_content [data-testid="stMetricValue"]{color:#0b4f91;letter-spacing:-.5px}
        .st-key-voc_page_content div[data-testid="stForm"]{
            border:1px solid #d2e0ee;border-radius:11px;padding:12px 14px;background:#f8fbfe;
        }
        .st-key-voc_page_content [data-testid="stWidgetLabel"] p{font-weight:650;color:#334f6c}
        .st-key-voc_page_content [data-testid="stDataFrame"]{border:1px solid #cfdeec;border-radius:10px;overflow:hidden}
        .st-key-voc_page_content [data-testid="stAlert"]{border-radius:9px;border-left-width:4px}
        .st-key-voc_page_content [data-testid="stExpander"]{border-color:#cfdeec!important;border-radius:10px!important;background:#fbfdff}
        .st-key-voc_page_content [data-testid="stProgress"] [role="progressbar"]{height:10px;border-radius:8px}
        .st-key-voc_page_content button{transition:transform .12s ease,box-shadow .12s ease}
        .st-key-voc_page_content button:hover{transform:translateY(-1px);box-shadow:0 3px 9px rgba(21,90,150,.11)}
        .st-key-voc_page_content h3{color:#153f6d!important;letter-spacing:-.3px;margin-top:.3rem!important}
        .st-key-voc_page_content h4{color:#24557f!important;letter-spacing:-.2px}
        @media(max-width:800px){.st-key-voc_page_hero{padding:14px!important}.st-key-voc_page_hero h2{font-size:21px!important}}
        </style>
        """
    )


def _render_voc_page_header(sub_menu: str) -> None:
    meta = VOC_PAGE_META[sub_menu]
    with st.container(border=True, key="voc_page_hero"):
        if sub_menu == "Dashboard":
            st.session_state["voc_dashboard_header_rendered"] = True
            header_col, control_col = st.columns([1.45, 1.1], vertical_alignment="bottom")
            with header_col:
                st.markdown(f"## :material/{meta['icon']}: {meta['title']}")
                st.caption(meta["description"])
            with control_col:
                today = date.today()
                selected_range = st.session_state.get(
                    "voc_dashboard_filter_range",
                    (today - timedelta(days=6), today),
                )
                with st.form("voc_dashboard_filters", border=False):
                    filter_columns = st.columns([2.2, 0.9, 0.95], vertical_alignment="bottom")
                    with filter_columns[0]:
                        selected_range = st.date_input(
                            "기간",
                            value=selected_range,
                            max_value=today,
                            key="voc_dashboard_filter_range",
                        )
                    with filter_columns[1]:
                        submitted = st.form_submit_button(
                            "조회",
                            icon=":material/search:",
                            type="primary",
                            width="stretch",
                        )
                    with filter_columns[2]:
                        refresh_requested = st.form_submit_button(
                            "새로고침",
                            icon=":material/refresh:",
                            width="stretch",
                        )
                st.session_state["voc_dashboard_filter_submitted"] = bool(submitted)
                st.session_state["voc_dashboard_filter_refresh_requested"] = bool(refresh_requested)
        else:
            if sub_menu == "Agent 관리":
                with st.container(
                    horizontal=True,
                    horizontal_alignment="distribute",
                    vertical_alignment="center",
                    gap="small",
                ):
                    st.markdown(f"## :material/{meta['icon']}: {meta['title']}")
                    if st.button(
                        "상태 새로고침",
                        type="primary",
                        width="content",
                        icon=":material/refresh:",
                        key="agent_header_refresh",
                    ):
                        _clear_agent_control_messages()
                        _load_agent_management_snapshot.clear()
                        _load_goal_monitor_snapshot.clear()
                        st.rerun()
                st.caption(
                    f"{meta['description']} 전체 시작은 Interpreter 등 6개 Agent 프로세스만 기동하며 "
                    "Test Case나 VOC 품질진단을 실행하지 않습니다. 전체 또는 개별 제어는 관리 스크립트가 "
                    "생성한 PID만 대상으로 하며, 외부 프로세스가 점유한 포트는 종료하지 않습니다."
                )
            else:
                st.markdown(f"## :material/{meta['icon']}: {meta['title']}")
                st.caption(meta["description"])
                flow = tuple(meta.get("flow", ()))
                if flow:
                    st.markdown(
                        " ".join(
                            f":blue-badge[{item}]"
                            for item in flow
                        )
                    )


def _new_manual_preparation_progress() -> dict:
    return {
        "status": "RUNNING",
        "current_step": 1,
        "steps": [
            {
                "number": index,
                "label": label,
                "status": "active" if index == 1 else "waiting",
            }
            for index, label in enumerate(MANUAL_PREPARATION_STEPS, start=1)
        ],
    }


def _update_manual_preparation(job_id: str, step_number: int, status: str) -> None:
    if not job_id:
        return
    job = background_job_snapshot(job_id) or {}
    preparation = deepcopy(
        job.get("progress", {}).get("preparation")
        or _new_manual_preparation_progress()
    )
    for step in preparation["steps"]:
        if step["number"] == step_number:
            step["status"] = status
        elif step["number"] < step_number and step["status"] in {"waiting", "active"}:
            step["status"] = "success"
    preparation["current_step"] = step_number
    if status == "failure":
        preparation["status"] = "ERROR"
    elif all(step["status"] == "success" for step in preparation["steps"]):
        preparation["status"] = "COMPLETED"
    else:
        preparation["status"] = "RUNNING"
    update_background_job(job_id, progress={"preparation": preparation})


def _execute_goal_testcase(job_id: str, case_id: str | None = None) -> dict:
    """Agent 사전 점검과 TC 실행을 모두 백그라운드에서 수행합니다."""
    if case_id is None:
        case_id, job_id = job_id, ""
    _update_manual_preparation(job_id, 1, "active")
    try:
        agent_snapshot = agent_status_snapshot()
        _update_manual_preparation(job_id, 1, "success")
        timeout_seconds = 180 if agent_snapshot.get("all_running") else 20
        judge_config = {
            "enabled": False,
            "provider": "anthropic",
            "model": "claude-haiku-4-5",
        }
        if job_id:
            testcase_result = run_test_case(
                case_id,
                timeout_seconds,
                judge_config,
                progress_callback=lambda step, status: _update_manual_preparation(
                    job_id, step, status
                ),
            )
        else:
            testcase_result = run_test_case(case_id, timeout_seconds, judge_config)
        return {"testcase_result": testcase_result, "agent_snapshot": agent_snapshot}
    except Exception:
        job = background_job_snapshot(job_id) or {}
        preparation = job.get("progress", {}).get("preparation", {})
        current_step = int(preparation.get("current_step") or 1)
        _update_manual_preparation(job_id, current_step, "failure")
        raise


def _execute_goal_judge(
    _job_id: str,
    run_id: str,
    case_id: str,
    judge_config: dict,
) -> dict:
    return reevaluate_voc_run_case(run_id, case_id, judge_config)


def _execute_history_judge_reevaluation(
    job_id: str,
    run_id: str,
    case_id: str,
    judge_config: dict,
) -> dict:
    update_background_job(
        job_id,
        progress={
            "percent": 12,
            "stage": "재평가 준비",
            "detail": "저장된 Run·Case 증적과 Agent 파이프라인 결과를 확인합니다.",
        },
    )
    update_background_job(
        job_id,
        progress={
            "percent": 34,
            "stage": "평가 입력 구성",
            "detail": "기존 개선안은 변경하지 않고 독립 LLM 평가 입력만 다시 구성합니다.",
        },
    )
    update_background_job(
        job_id,
        progress={
            "percent": 52,
            "stage": "독립 LLM 요청",
            "detail": "선택한 Provider와 모델로 동일 결과를 재평가하고 있습니다.",
        },
    )
    reevaluated = reevaluate_voc_run_case(run_id, case_id, judge_config)
    update_background_job(
        job_id,
        progress={
            "percent": 92,
            "stage": "결과 저장",
            "detail": "재평가 결과와 이전 평가 이력을 Run 증적에 저장했습니다.",
        },
    )
    return reevaluated


@st.cache_resource
def _batch_executor():
    return ThreadPoolExecutor(max_workers=1, thread_name_prefix="voc-batch")


def _judge_config_controls(key_prefix: str, *, fault_only: bool = False) -> dict:
    options = judge_provider_options()
    enabled = st.toggle(
        "독립 LLM 평가",
        key=f"{key_prefix}_judge_enabled",
        help="Agent 파이프라인 성공 후 별도 LLM이 100점 Rubric으로 최종 결과를 평가합니다.",
    )
    default = next((item for item in options if item["provider"] == "anthropic"), options[0])
    if not enabled:
        return {"enabled": False, "provider": default["provider"], "model": default["default_model"]}

    provider = st.selectbox(
        "독립 LLM 평가 Provider",
        [item["provider"] for item in options],
        format_func=lambda value: next(item["label"] for item in options if item["provider"] == value),
        key=f"{key_prefix}_judge_provider",
    )
    selected = next(item for item in options if item["provider"] == provider)
    model = st.text_input(
        "독립 LLM 평가 모델",
        value=selected["default_model"],
        key=f"{key_prefix}_judge_model_{provider}",
    )
    independence = judge_independence_preview(provider, model)
    if selected["credential_configured"]:
        st.caption(
            f"자격 증명 설정됨 · {_judge_independence_grade_label(independence['grade'])} · {independence['reason']}"
        )
    else:
        st.error(f"{selected['label']} API 자격 증명이 설정되지 않았습니다.")
    if fault_only:
        st.info("격리 장애 Case는 개선안이 없으므로 독립 LLM 평가가 미실행으로 기록됩니다.")
    return {
        "enabled": enabled,
        "provider": provider,
        "model": model,
        "timeout_seconds": 90,
        "max_retries": 2,
    }


def _judge_config_summary(judge_config: dict) -> dict:
    options = {item["provider"]: item for item in judge_provider_options()}
    provider = str(judge_config.get("provider") or "anthropic")
    option = options.get(provider, {})
    provider_label = str(option.get("label") or provider)
    model = str(judge_config.get("model") or option.get("default_model") or "-")
    enabled = bool(judge_config.get("enabled"))
    return {
        "enabled": enabled,
        "provider": provider,
        "provider_label": provider_label,
        "model": model,
        "label": f"{provider_label} · {model}" if enabled else "독립 LLM 평가 미실행",
    }


def _html_status_chip_label(icon_name: str) -> str:
    return {
        "gavel": "평가",
        "block": "미실행",
        "hub": "연결",
        "fact_check": "평가",
        "hourglass_top": "진행",
        "progress_activity": "진행",
        "check_circle": "완료",
        "done_all": "완료",
        "warning": "확인",
        "error": "오류",
        "pending": "대기",
        "pending_actions": "대기",
        "play_circle": "실행",
        "lock": "잠금",
    }.get(str(icon_name or "").strip(), "확인")


def _render_batch_judge_selection_badge(judge_config: dict) -> None:
    summary = _judge_config_summary(judge_config)
    tone = "on" if summary["enabled"] else "off"
    icon = "gavel" if summary["enabled"] else "block"
    icon_label = _html_status_chip_label(icon)
    st.markdown(
        f"""
        <div class="vqa-batch-judge-summary {tone}">
            <span>{escape(icon_label)}</span>
            <div>
                <small>선택된 독립 LLM 평가</small>
                <strong>{escape(summary["label"])}</strong>
            </div>
        </div>
        <style>
        .vqa-batch-judge-summary{{
            height:46px;margin:7px 0 8px;padding:7px 10px;border-radius:10px;
            display:flex;align-items:center;gap:8px;box-sizing:border-box;
            border:1px solid #c8d9ee;background:linear-gradient(135deg,#f8fbff,#ffffff);
            font-family:'Segoe UI','Malgun Gothic',sans-serif;overflow:hidden;
        }}
        .vqa-batch-judge-summary span{{
            flex:0 0 38px;height:24px;border-radius:999px;display:flex;align-items:center;justify-content:center;
            font-size:11px;font-weight:900;color:#155a96;background:#e7f1fb;border:1px solid #bfd7ef;
        }}
        .vqa-batch-judge-summary.off span{{color:#6f7c8c;background:#eef1f4;border-color:#d3d9e0}}
        .vqa-batch-judge-summary small{{
            display:block;font-size:9px;line-height:1.1;color:#708096;font-weight:800;
        }}
        .vqa-batch-judge-summary strong{{
            display:block;margin-top:2px;font-size:12px;line-height:1.18;color:#173f68;
            white-space:nowrap;overflow:hidden;text-overflow:ellipsis;max-width:100%;
        }}
        .vqa-batch-judge-summary.off strong{{color:#617083}}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _select_manual_judge(state_key: str, provider: str):
    st.session_state[state_key] = provider


@st.cache_data(max_entries=2, show_spinner=False)
def _manual_judge_logo_data_uri(provider: str) -> str:
    logo_path = Path(__file__).resolve().parents[1] / "assets" / "providers" / f"{provider}.svg"
    encoded = base64.b64encode(logo_path.read_bytes()).decode("ascii")
    return f"data:image/svg+xml;base64,{encoded}"


def _manual_pipeline_generation_info(pipeline_snapshot: dict | None) -> dict:
    snapshot = pipeline_snapshot if isinstance(pipeline_snapshot, dict) else {}
    raw_provider = str(snapshot.get("provider") or "").strip().lower()
    raw_model = str(snapshot.get("model") or "").strip()
    recorded_provider = raw_provider and raw_provider != "internal_pipeline"
    recorded_model = bool(raw_model and raw_model != "-")
    if recorded_provider or recorded_model:
        return {
            "provider": _manual_provider_label(raw_provider) if recorded_provider else str(snapshot.get("provider_label") or "내부 Agent 파이프라인"),
            "model": raw_model if recorded_model else "모델 확인 필요",
            "source": "실제 수행 Trace 기준",
        }
    try:
        fallback = judge_independence_preview("anthropic", "claude-haiku-4-5")
    except Exception:
        fallback = {}
    return {
        "provider": _manual_provider_label(fallback.get("generator_provider")) or "내부 Agent 파이프라인",
        "model": str(fallback.get("generator_model") or "모델 확인 필요"),
        "source": "현재 파이프라인 설정 기준",
    }


def _manual_judge_config_controls(
    key_prefix: str,
    *,
    fault_only: bool = False,
    pipeline_snapshot: dict | None = None,
) -> dict:
    options = {item["provider"]: item for item in judge_provider_options()}
    state_key = f"{key_prefix}_judge_provider"
    st.session_state.setdefault(state_key, "anthropic")
    selected_provider = st.session_state[state_key]
    if selected_provider not in {item["provider"] for item in MANUAL_JUDGE_PROVIDERS}:
        selected_provider = "anthropic"
        st.session_state[state_key] = selected_provider
    selected_option = options.get(selected_provider, {})
    selected = next(item for item in MANUAL_JUDGE_PROVIDERS if item["provider"] == selected_provider)
    selected_model = str(selected_option.get("default_model") or selected["model"])
    selected_independence = judge_independence_preview(selected_provider, selected_model)

    card_classes = []
    for item in MANUAL_JUDGE_PROVIDERS:
        state = "selected" if item["provider"] == selected_provider else "inactive"
        logo_uri = _manual_judge_logo_data_uri(item["provider"])
        card_classes.append(
            f".st-key-{key_prefix}_judge_select_{item['provider']} button{{"
            "height:164px;border-radius:14px;display:flex;flex-direction:column;"
            "justify-content:flex-end;align-items:flex-start;padding:18px 20px;"
            "font-size:14px;font-weight:700;white-space:pre-line;transition:all .18s ease;"
            f"background-image:url('{logo_uri}');background-repeat:no-repeat;"
            "background-position:18px 18px;background-size:auto 48px;"
            + (
                "background-color:#f4f9ff;border:2px solid #1d65a6;"
                "color:#174f85;box-shadow:0 6px 18px rgba(29,101,166,.13)"
                if state == "selected"
                else
                "background-color:#eef1f4;border:1px solid #ccd3da;"
                "color:#697684;box-shadow:none;filter:grayscale(1);opacity:.72"
            )
            + "}"
        )
    st.html(
        "<style>"
        + "".join(card_classes)
        + "</style>"
    )

    if pipeline_snapshot:
        generation = _manual_pipeline_generation_info(pipeline_snapshot)
        pipeline_provider = escape(generation["provider"])
        pipeline_model = escape(generation["model"])
        pipeline_source = escape(generation["source"])
        selected_label = escape(_manual_provider_label(selected_provider))
        selected_model_text = escape(selected_model)
        selected_independence_label = escape(_judge_independence_grade_label(selected_independence.get("grade")))
        st.html(
            f"""
            <div class="vqa-cross-judge-guide">
                <div class="vqa-cross-judge-card source">
                    <span>연결</span>
                    <div>
                        <small>Agent 파이프라인 생성 LLM</small>
                        <strong>{pipeline_provider}</strong>
                        <p>{pipeline_model} · {pipeline_source}</p>
                    </div>
                </div>
                <div class="vqa-cross-judge-arrow">독립 검증</div>
                <div class="vqa-cross-judge-card judge">
                    <span>평가</span>
                    <div>
                        <small>선택 평가 LLM · {selected_independence_label}</small>
                        <strong>{selected_label}</strong>
                        <p>{selected_model_text}</p>
                    </div>
                </div>
            </div>
            <style>
            .vqa-cross-judge-guide{{
                display:grid;grid-template-columns:1fr auto 1fr;gap:10px;align-items:stretch;
                margin:8px 0 12px;
                font-family:'Segoe UI','Malgun Gothic',sans-serif;
            }}
            .vqa-cross-judge-card{{
                display:flex;gap:11px;align-items:center;padding:12px 14px;border-radius:14px;
                border:1px solid #c9d8e8;background:#f8fbff;color:#244b72;min-height:76px;
            }}
            .vqa-cross-judge-card.judge{{background:#ffffff;border-color:#b9d2ed}}
            .vqa-cross-judge-card span{{
                flex:0 0 38px;height:28px;border-radius:999px;display:flex;align-items:center;justify-content:center;
                font-size:11px;font-weight:900;color:#1f6fb2;background:#e7f1fb;border:1px solid #bfd7ef;
            }}
            .vqa-cross-judge-card small{{display:block;font-size:10px;color:#6a7b8f;margin-bottom:2px}}
            .vqa-cross-judge-card strong{{display:block;font-size:15px;color:#174f85}}
            .vqa-cross-judge-card p{{margin:2px 0 0;font-size:11px;color:#65788c}}
            .vqa-cross-judge-card p{{white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}
            .vqa-cross-judge-arrow{{
                align-self:center;border-radius:999px;background:#eaf3fc;color:#1f6fb2;
                font-size:11px;font-weight:800;padding:6px 10px;white-space:nowrap;
            }}
            </style>
            """,
        )

    st.markdown("#### 독립 LLM Provider 선택")
    columns = st.columns(len(MANUAL_JUDGE_PROVIDERS), gap="medium")
    for column, item in zip(columns, MANUAL_JUDGE_PROVIDERS):
        provider = item["provider"]
        selected = provider == selected_provider
        option = options.get(provider, {})
        display_model = str(option.get("default_model") or item["model"])
        independence = judge_independence_preview(provider, display_model)
        with column:
            st.button(
                f"{_manual_provider_label(provider)}\n\n{display_model}\n\n"
                + ("✓ 현재 선택" if selected else "카드를 클릭하여 선택"),
                key=f"{key_prefix}_judge_select_{provider}",
                width="stretch",
                on_click=_select_manual_judge,
                args=(state_key, provider),
                help=f"{item['label']} {display_model}을 독립 LLM 평가 Provider로 사용합니다.",
            )
            credential_text = "API 자격 증명 설정됨" if option.get("credential_configured") else "API 자격 증명 미설정"
            st.caption(f"{credential_text} · {_judge_independence_grade_label(independence['grade'])}")

    credential_configured = bool(selected_option.get("credential_configured"))
    if fault_only:
        st.info(
            "격리 장애 Case는 개선안이 없어 선택한 Provider를 호출하지 않고 독립 LLM 평가 결과를 미실행으로 기록합니다.",
            icon=":material/info:",
        )
    elif not credential_configured:
        st.error(
            f"{selected['label']} API 자격 증명이 설정되지 않아 실행할 수 없습니다.",
            icon=":material/key_off:",
        )
    return {
        "enabled": not fault_only,
        "provider": selected_provider,
        "model": selected_model,
        "timeout_seconds": 90,
        "max_retries": 2,
        "credential_configured": credential_configured or fault_only,
    }


def _validity_config_controls(key_prefix: str, *, compact: bool = False) -> dict:
    options = validity_provider_options()
    provider_options = [item["provider"] for item in options]
    if compact:
        provider_col, model_col, state_col = st.columns([1.05, 1.55, 0.9], gap="small", vertical_alignment="bottom")
        with provider_col:
            provider = st.selectbox(
                "Provider",
                provider_options,
                format_func=lambda value: next(item["label"] for item in options if item["provider"] == value),
                key=f"{key_prefix}_validity_provider",
            )
    else:
        provider = st.selectbox(
            "개선안 타당성 평가 Provider",
            provider_options,
            format_func=lambda value: next(item["label"] for item in options if item["provider"] == value),
            key=f"{key_prefix}_validity_provider",
        )
    selected = next(item for item in options if item["provider"] == provider)
    if compact:
        with model_col:
            model = st.text_input(
                "평가 모델",
                value=selected["default_model"],
                key=f"{key_prefix}_validity_model_{provider}",
            )
        with state_col:
            if selected["credential_configured"]:
                st.markdown(":green-badge[키 설정됨]", text_alignment="right")
            else:
                st.markdown(":red-badge[키 확인 필요]", text_alignment="right")
                st.caption(f"{selected['label']} API 키 필요", text_alignment="right")
    else:
        model = st.text_input(
            "개선안 타당성 평가 모델",
            value=selected["default_model"],
            key=f"{key_prefix}_validity_model_{provider}",
        )
        if selected["credential_configured"]:
            st.caption("자격 증명 설정됨 · 독립 LLM 평가 결과와 별도로 개선안 실행 타당성을 평가합니다.")
        else:
            st.error(f"{selected['label']} API 자격 증명이 설정되지 않았습니다.")
    return {
        "provider": provider,
        "model": model,
        "credential_configured": selected["credential_configured"],
    }


def _keyword_text(values) -> str:
    safe = [escape(str(value)) for value in (values or [])[:6] if str(value).strip()]
    return ", ".join(safe) if safe else "-"


def _parse_json_mapping(value) -> dict:
    if isinstance(value, dict):
        return value
    if not isinstance(value, str) or not value.strip():
        return {}
    try:
        parsed = json.loads(value)
    except (TypeError, ValueError, json.JSONDecodeError):
        return {}
    return parsed if isinstance(parsed, dict) else {}


def _parse_pipeline_trace_summary(value) -> dict:
    values = {}
    flags = []
    for item in str(value or "").split(";"):
        token = item.strip()
        if not token:
            continue
        if "=" in token:
            key, raw_value = token.split("=", 1)
            values[key.strip()] = raw_value.strip()
        else:
            flags.append(token)
    return {"values": values, "flags": flags}


def _manual_pipeline_llm_snapshot(test_execution: dict | None) -> dict:
    """Agent 파이프라인에서 개선안 생성에 사용한 LLM 정보를 화면용으로 추출합니다."""
    test_execution = test_execution if isinstance(test_execution, dict) else {}
    execution = test_execution.get("execution", {}) if isinstance(test_execution.get("execution"), dict) else {}
    result = execution.get("result", {}) if isinstance(execution.get("result"), dict) else {}
    trace_summary = _parse_pipeline_trace_summary(result.get("trace"))
    trace_values = trace_summary.get("values", {})

    provider = (
        trace_values.get("policy_provider")
        or result.get("policy_provider")
        or execution.get("policy_provider")
        or test_execution.get("policy_provider")
        or "internal_pipeline"
    )
    model = (
        trace_values.get("policy_model")
        or result.get("policy_model")
        or execution.get("policy_model")
        or test_execution.get("policy_model")
        or "-"
    )
    if str(provider).lower() == "internal_pipeline":
        provider_label = "내부 Agent 파이프라인"
    else:
        provider_label = _manual_provider_label(provider)
    return {
        "provider": str(provider or "-"),
        "provider_label": provider_label,
        "model": str(model or "-"),
    }


def _manual_judge_comparison_rows(judge_result: dict | None) -> list[dict]:
    """현재 독립 LLM 평가 결과와 이전 Provider 재평가 이력을 비교 표 데이터로 변환합니다."""
    if not isinstance(judge_result, dict) or not judge_result:
        return []
    candidates = []
    for previous in judge_result.get("evaluation_history", []) or []:
        if isinstance(previous, dict) and previous.get("decision") != "NOT_RUN":
            candidates.append(previous)
    if judge_result.get("decision") != "NOT_RUN":
        candidates.append(judge_result)

    def evaluated_at_text(value) -> str:
        text = str(value or "").strip()
        if not text:
            return "-"
        try:
            return datetime.fromisoformat(text.replace("Z", "+00:00")).strftime("%Y-%m-%d %H:%M:%S")
        except (TypeError, ValueError):
            normalized = text.replace("T", " ")
            return normalized[:19] if len(normalized) >= 19 else normalized

    def duration_text(value) -> str:
        try:
            seconds = float(value)
        except (TypeError, ValueError):
            return "-"
        return f"{seconds:g}초"

    rows = []
    seen = set()
    for index, item in enumerate(candidates, start=1):
        provider = str(item.get("provider") or "-").lower()
        model = str(item.get("model") or "-")
        evaluated_at = str(item.get("evaluated_at") or item.get("recorded_at") or "-")
        raw_score = item.get("total_score")
        try:
            score = float(raw_score) if raw_score is not None and str(raw_score).strip() != "" else None
        except (TypeError, ValueError):
            score = None
        if score is not None and score.is_integer():
            score = int(score)
        duration = duration_text(item.get("duration_seconds"))
        identity = (provider, model, evaluated_at, score, duration)
        if identity in seen:
            continue
        seen.add(identity)
        rows.append(
            {
                "순서": index,
                "평가 Provider": _manual_provider_label(provider),
                "모델": model,
                "판정": _voc_status_label(item.get("decision", "NOT_RUN")),
                "점수": score,
                "독립성": _judge_independence_grade_label(item.get("independence_grade", "-")),
                "수행 시간": duration,
                "평가 시각": evaluated_at_text(evaluated_at),
            }
        )
    return rows


def _pipeline_trace_event_rows(trace: dict) -> list[dict]:
    operation_labels = {
        "ParseQuestion": "질문 의도 해석",
        "IntentDataTransfer": "분석 요청·결과 전달",
        "Retrieve": "관련 VOC 검색",
        "VOCDataTransfer": "검색 결과 전달",
        "Evaluate": "요약 후보 평가",
        "ReviewSummary": "요약 품질 검토",
        "ReviewPolicy": "개선안 품질 검토",
        "Refine": "요약 보완",
        "RefinePolicy": "개선안 보완",
        "Improve": "개선안 생성",
        "RunPolicyPipeline": "개선안 생성·검토",
    }
    status_labels = {"success": "성공", "failure": "실패", "started": "시작"}
    completed_events = [
        event
        for event in (trace.get("events", []) if isinstance(trace, dict) else [])
        if isinstance(event, dict) and event.get("status") in {"success", "failure"}
    ]
    rows = []
    for index, event in enumerate(completed_events, start=1):
        error_text = str(event.get("error") or "").strip()
        error = error_text.splitlines()[0] if error_text else ""
        clues = event.get("output_keywords") or event.get("keywords") or []
        clue_text = ", ".join(str(item) for item in clues[:6])
        if event.get("item_count") is not None:
            clue_text = f"VOC {event['item_count']}건 전달" + (f" · {clue_text}" if clue_text else "")
        if error:
            clue_text = error[:160]
        rows.append({
            "순서": index,
            "Agent 연결": f"{event.get('source', '-')} → {event.get('target', '-')}",
            "처리 내용": operation_labels.get(event.get("operation"), event.get("operation", "-")),
            "결과": status_labels.get(event.get("status"), event.get("status", "-")),
            "처리시간(ms)": round(float(event.get("duration_ms") or 0), 2),
            "판단 단서": clue_text or "-",
        })
    return rows


TRACE_FLOW_EXPLANATIONS = {
    "ParseQuestion": (
        "start",
        "질문 해석",
        "Agent 1이 사용자 질문을 의도·검색 조건·수행 작업으로 구조화합니다.",
    ),
    "IntentDataTransfer": (
        "handoff",
        "분석 이관",
        "Agent 1이 해석한 의도를 조정 역할의 Agent 3에 전달해 전체 파이프라인을 시작합니다.",
    ),
    "Retrieve": (
        "lookup",
        "근거 조회",
        "Agent 3이 답변과 요약의 근거가 되는 VOC 원문을 확보하기 위해 Agent 2를 호출합니다.",
    ),
    "VOCDataTransfer": (
        "return",
        "검색 결과 반환",
        "Agent 2가 검색한 VOC 데이터와 건수를 이후 단계를 조정하는 Agent 3에 반환합니다.",
    ),
    "Evaluate": (
        "selection",
        "후보 평가",
        "Agent 3이 생성한 후보 중 가장 적합한 결과를 선택하기 위해 Agent 4의 평가를 요청합니다.",
    ),
    "ReviewSummary": (
        "review",
        "요약 검토",
        "선택된 요약의 누락·왜곡·보완 필요 여부를 확인하기 위해 Agent 5가 검토합니다.",
    ),
    "ImprovePolicy": (
        "generation",
        "개선안 생성",
        "검토된 요약을 실행 가능한 정책 개선안으로 바꾸기 위해 Agent 6을 호출합니다.",
    ),
    "ReviewPolicy": (
        "feedback",
        "개선안 재검토",
        "Agent 6이 만든 개선안을 확정하기 전에 품질과 실행 가능성을 확인하기 위해 Agent 5로 되돌아갑니다.",
    ),
    "RefinePolicy": (
        "rework",
        "수정 요청 반영",
        "Agent 5가 개선안에 수정이 필요하다고 판단해 전달한 보완 의견을 Agent 6이 반영합니다.",
    ),
    "Improve": (
        "generation",
        "개선안 생성",
        "검토 결과를 실행 가능한 개선안으로 전환하기 위해 개선 Agent를 호출합니다.",
    ),
    "Refine": (
        "rework",
        "요약 보완",
        "Critic이 요청한 수정 사항을 반영하기 위해 요약 생성 단계를 다시 수행합니다.",
    ),
    "RunPolicyPipeline": (
        "handoff",
        "정책 파이프라인 실행",
        "검토된 요약을 기준으로 정책 개선과 재검토 흐름을 실행합니다.",
    ),
}


def _trace_agent_number(agent_name: str) -> int | None:
    return next(
        (
            index
            for index, (name, _, _) in enumerate(AGENT_PIPELINE, start=1)
            if name == agent_name
        ),
        None,
    )


def _trace_flow_explanation(
    event: dict,
    previous_event: dict | None = None,
) -> dict:
    operation = str(event.get("operation") or "작업")
    status = str(event.get("status") or "")
    current_number = _trace_agent_number(str(event.get("target") or ""))
    previous_number = _trace_agent_number(
        str((previous_event or {}).get("target") or "")
    )
    if previous_number and current_number:
        transition = f"Agent {previous_number} → Agent {current_number}"
    elif current_number:
        transition = f"시작 → Agent {current_number}"
    else:
        transition = "Agent 흐름"

    if status == "failure":
        error_text = str(event.get("error") or "").strip()
        short_error = error_text.splitlines()[0][:140] if error_text else ""
        return {
            "kind": "failure",
            "label": "호출 실패",
            "transition": transition,
            "reason": short_error
            or f"{operation} 처리 중 오류가 발생해 다음 단계로 진행하지 못했습니다.",
            "inferred": False,
        }

    configured = TRACE_FLOW_EXPLANATIONS.get(operation)
    if configured:
        kind, label, reason = configured
        return {
            "kind": kind,
            "label": label,
            "transition": transition,
            "reason": reason,
            "inferred": False,
        }

    if previous_number and current_number and current_number < previous_number:
        label = "이전 단계 재호출"
        direction_reason = f"{operation} 처리를 위해 앞 단계 Agent를 다시 호출했습니다."
    elif (
        previous_number
        and current_number
        and current_number > previous_number + 1
    ):
        label = "분기 호출"
        direction_reason = (
            f"{operation} 처리에 중간 단계가 필요하지 않아 해당 Agent로 바로 분기했습니다."
        )
    elif previous_number == current_number and current_number:
        label = "동일 단계 재처리"
        direction_reason = f"{operation} 결과를 보완하기 위해 같은 Agent가 다시 처리했습니다."
    else:
        label = "다음 단계 호출"
        direction_reason = f"{operation} 처리를 위해 다음 Agent를 호출했습니다."
    return {
        "kind": "inferred",
        "label": label,
        "transition": transition,
        "reason": (
            f"{direction_reason} 실행 Trace에 상세 분기 사유가 없어 작업 유형과 이동 방향을 기준으로 표시했습니다."
        ),
        "inferred": True,
    }


def _trace_reason(event: dict) -> str:
    """기존 호출부 호환용으로 작업 자체의 설명만 반환합니다."""
    return _trace_flow_explanation(event)["reason"]


def _trace_event_display_statuses(events: list[dict], *, running: bool) -> list[str]:
    """원본 started 이벤트가 후속 단계 진행 뒤에도 '진행'으로 남지 않게 표시 상태를 보정합니다."""
    display_statuses: list[str] = []
    for index, event in enumerate(events):
        status = str(event.get("status") or "unknown")
        if status != "started":
            display_statuses.append(status)
            continue

        signature = (event.get("source"), event.get("target"), event.get("operation"))
        later_events = events[index + 1 :]
        has_matching_terminal = any(
            (later.get("source"), later.get("target"), later.get("operation")) == signature
            and later.get("status") in {"success", "failure"}
            for later in later_events
        )
        # 같은 작업의 종료 로그가 있거나 다음 단계가 시작됐다면 이 started 기록은 이미 지난 단계입니다.
        display_statuses.append(
            "completed" if has_matching_terminal or bool(later_events) else ("started" if running else "ended")
        )
    return display_statuses


def _trace_display_events(events: list[dict]) -> list[dict]:
    """started/success 한 쌍을 하나의 단계 카드로 합쳐 상태 중복을 제거합니다."""
    display_events: list[dict] = []
    pending: dict[tuple, list[int]] = {}
    for event in events:
        item = dict(event)
        signature = (item.get("source"), item.get("target"), item.get("operation"))
        status = item.get("status")
        if status == "started":
            pending.setdefault(signature, []).append(len(display_events))
            display_events.append(item)
            continue
        if status in {"success", "failure"} and pending.get(signature):
            started_index = pending[signature].pop(0)
            if not pending[signature]:
                pending.pop(signature)
            started_event = display_events[started_index]
            item["started_at"] = started_event.get("timestamp", "")
            if not item.get("input_keywords"):
                item["input_keywords"] = started_event.get("input_keywords", [])
            display_events[started_index] = item
            continue
        display_events.append(item)
    return display_events


def _manual_pipeline_timestamp_text(value: str | None) -> str:
    text = str(value or "").strip()
    if not text:
        return "-"
    try:
        parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
        if parsed.tzinfo is not None:
            parsed = parsed.astimezone()
        return parsed.strftime("%Y-%m-%d %H:%M:%S")
    except (TypeError, ValueError):
        normalized = text.replace("T", " ")
        return normalized[:19] if len(normalized) >= 19 else normalized


def _manual_pipeline_compact_text(value: str | None, limit: int = 78) -> str:
    text = " ".join(str(value or "").split())
    if not text:
        return "-"
    return text if len(text) <= limit else f"{text[: max(0, limit - 1)]}…"


JIRA_ACTION_POPOVER_LABEL = "Jira 등록"
JIRA_CONTEXT_ISSUE_TYPES = ("작업", "버그", "스토리", "Task", "Bug")
JIRA_CONTEXT_PRIORITIES = ("미지정", "Highest", "High", "Medium", "Low", "Lowest")


def _jira_context_key(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_\-]+", "_", str(value or "context")).strip("_")[:120] or "context"


def _jira_context_labels(*values: object) -> list[str]:
    labels = ["voc-quality"]
    for value in values:
        text = str(value or "").strip()
        if not text or text == "-":
            continue
        clean = re.sub(r"[^A-Za-z0-9가-힣_\-]+", "-", text).strip("-")
        if clean:
            labels.append(clean[:80])
    return list(dict.fromkeys(labels))


def _jira_context_line_has_positive_signal(line: str) -> bool:
    match = re.search(r":\s*(\d+)\s*건", str(line or ""))
    return int(match.group(1)) > 0 if match else True


def _jira_context_diagnostic_model(
    *,
    status_label: str = "",
    extra_detail: str = "",
) -> dict:
    status_text = str(status_label or "").strip()
    detail_lines = [
        re.sub(r"^\s*[-•]\s*", "", str(line or "").strip())
        for line in str(extra_detail or "").splitlines()
        if str(line or "").strip()
    ]
    status_lines = []
    if status_text:
        status_lines.append(f"현재 상태/다음 액션: {status_text}")
    for line in detail_lines:
        if ":" in line:
            status_lines.append(line)
    status_lines = list(dict.fromkeys(status_lines))

    signal_source = "\n".join([status_text, *detail_lines]).lower()
    positive_signal_lines = [
        line
        for line in status_lines
        if _jira_context_line_has_positive_signal(line)
    ]
    positive_signal_text = "\n".join(positive_signal_lines).lower()

    error_markers = ("오류", "error", "exception", "traceback", "deadline_exceeded")
    fail_markers = ("실패", "fail", "failed")
    review_markers = (
        "검토 필요",
        "보완",
        "수정 필요",
        "반려",
        "revision",
        "review_required",
        "revision_required",
        "rejected",
    )
    not_run_markers = ("미실행", "not_run", "대기")

    issue_lines = [
        line
        for line in positive_signal_lines
        if any(marker in line.lower() for marker in [*error_markers, *fail_markers, *review_markers, *not_run_markers])
    ]

    if any(marker in positive_signal_text for marker in error_markers):
        category = "오류/버그 확인 필요"
        issue_type = "버그"
        priority = "High"
        labels = ["voc-error", "voc-bug"]
    elif any(marker in positive_signal_text for marker in fail_markers):
        category = "Fail 원인 확인 필요"
        issue_type = "버그"
        priority = "High"
        labels = ["voc-fail", "voc-bug"]
    elif any(marker in signal_source for marker in review_markers):
        category = "검토/보완 필요"
        issue_type = "작업"
        priority = "Medium"
        labels = ["voc-review"]
    elif any(marker in signal_source for marker in not_run_markers):
        category = "미실행/대기 확인 필요"
        issue_type = "작업"
        priority = "Medium"
        labels = ["voc-not-run"]
    else:
        category = "상태 확인"
        issue_type = "작업"
        priority = "미지정"
        labels = ["voc-check"]

    return {
        "category": category,
        "issue_type": issue_type,
        "priority": priority,
        "labels": labels,
        "status_lines": status_lines,
        "issue_lines": list(dict.fromkeys(issue_lines)),
    }


def _jira_context_summary(
    *,
    area_label: str,
    target_label: str,
    run_id: str = "",
    case_id: str = "",
    status_label: str = "",
    question: str = "",
    extra_detail: str = "",
) -> str:
    parts = ["VOC"]
    diagnostic = _jira_context_diagnostic_model(status_label=status_label, extra_detail=extra_detail)
    if diagnostic.get("category") and diagnostic.get("category") != "상태 확인":
        parts.append(str(diagnostic["category"]))
    if case_id:
        parts.append(case_id)
    elif run_id:
        parts.append(run_id)
    parts.append(target_label or area_label or "선택 항목")
    if status_label:
        parts.append(status_label)
    if question:
        parts.append(_manual_pipeline_compact_text(question, 70))
    return " · ".join(str(part) for part in parts if str(part or "").strip())[:255]


def _jira_context_description(
    *,
    area_label: str,
    target_label: str,
    run_id: str = "",
    case_id: str = "",
    status_label: str = "",
    question: str = "",
    extra_detail: str = "",
) -> str:
    diagnostic = _jira_context_diagnostic_model(
        status_label=status_label,
        extra_detail=extra_detail,
    )
    lines = [
        "VOC 품질진단 화면에서 선택한 표 Row 기준으로 등록한 Jira 이슈입니다.",
        "",
        f"- 화면: {area_label or '-'}",
        f"- 대상: {target_label or '-'}",
    ]
    if run_id:
        lines.append(f"- Run ID: {run_id}")
    if case_id:
        lines.append(f"- Case ID: {case_id}")
    if status_label:
        lines.append(f"- 현재 상태/다음 액션: {status_label}")
    lines.extend(["", "등록 진단 요약:", f"- 구분: {diagnostic['category']}"])
    status_lines = diagnostic.get("status_lines") or []
    if status_lines:
        lines.append("- Run/Case 상태 정보:")
        lines.extend(f"  - {line}" for line in status_lines[:10])
    issue_lines = diagnostic.get("issue_lines") or []
    if issue_lines:
        lines.append("- 오류/버그/Fail 확인 포인트:")
        lines.extend(f"  - {line}" for line in issue_lines[:8])
    lines.extend(
        [
            "- 조치 확인 기준: 증적 파일, Trace, 평가 결과, 재현 조건을 함께 확인합니다.",
        ]
    )
    if question:
        lines.extend(["", "질문/요약:", str(question).strip()])
    if extra_detail:
        lines.extend(["", "상세:", str(extra_detail).strip()])
    lines.extend(
        [
            "",
            "확인 요청:",
            "- 선택한 Run/Case의 증적과 평가 결과를 확인합니다.",
            "- 필요한 경우 보완, 재평가, 결함 처리 또는 승인 후속 조치를 진행합니다.",
        ]
    )
    return "\n".join(lines)


def _jira_context_run_detail_extra(run_id: str) -> str:
    try:
        detail = load_voc_run_history_detail(run_id)
    except Exception:
        return "Run 상태: 상세 정보를 불러오지 못했습니다."
    manifest = detail.get("manifest", {}) if isinstance(detail.get("manifest"), dict) else {}
    summary = detail.get("summary", {}) if isinstance(detail.get("summary"), dict) else {}
    counts = summary.get("counts") if isinstance(summary.get("counts"), dict) else {}
    judge_counts = summary.get("judge_counts") if isinstance(summary.get("judge_counts"), dict) else {}
    total = summary.get("total") or summary.get("selected_count") or sum(
        int(counts.get(key, 0) or 0)
        for key in ("PASS", "REVIEW_REQUIRED", "FAIL", "ERROR", "NOT_RUN")
    )
    run_context = {
        "run_id": run_id,
        "status": manifest.get("status") or summary.get("status"),
        "run_type": manifest.get("run_type") or summary.get("run_type"),
        "selected_count": total,
        "counts": counts,
        "judge_counts": judge_counts,
        "validity_state": summary.get("validity_state"),
        "deployment_decision": summary.get("deployment_decision"),
    }
    next_action = voc_run_next_action(run_context)
    return "\n".join(
        [
            f"유형: {_voc_status_label(run_context.get('run_type'))}",
            f"Run 상태: {_voc_status_label(run_context.get('status'))}",
            f"다음 조치: {next_action.get('label', '-')}",
            f"대상: {total or 0}건",
            f"통과: {int(counts.get('PASS', 0) or 0)}건",
            f"검토 필요: {int(counts.get('REVIEW_REQUIRED', 0) or 0)}건",
            f"실패: {int(counts.get('FAIL', 0) or 0)}건",
            f"오류: {int(counts.get('ERROR', 0) or 0)}건",
            f"독립 LLM 통과: {int(judge_counts.get('PASS', 0) or 0)}건",
            f"독립 LLM 검토 필요: {int(judge_counts.get('REVIEW_REQUIRED', 0) or 0)}건",
            f"독립 LLM 실패: {int(judge_counts.get('FAIL', 0) or 0)}건",
            f"독립 LLM 오류: {int(judge_counts.get('ERROR', 0) or 0)}건",
            f"개선안 타당성: {_voc_status_label(summary.get('validity_state'))}",
            f"배포 판정: {_voc_status_label(summary.get('deployment_decision'))}",
        ]
    )


def _ensure_contextual_jira_history() -> list[dict]:
    if "jira_registered_issues" not in st.session_state:
        st.session_state.jira_registered_issues = load_json_file(JIRA_REGISTERED_ISSUES_FILE, [])
    return st.session_state.jira_registered_issues


def _record_contextual_jira_issue(history_item: dict) -> None:
    history = _ensure_contextual_jira_history()
    history.insert(0, history_item)
    save_json_file(JIRA_REGISTERED_ISSUES_FILE, history)


def _render_contextual_jira_action_menu(
    *,
    area_label: str,
    target_label: str,
    run_id: str = "",
    case_id: str = "",
    status_label: str = "",
    question: str = "",
    extra_detail: str = "",
    key: str,
) -> None:
    """Register a Jira issue from a selected table row context."""
    safe_key = _jira_context_key(key)
    snapshot = jira_environment_snapshot()
    default_summary = _jira_context_summary(
        area_label=area_label,
        target_label=target_label,
        run_id=run_id,
        case_id=case_id,
        status_label=status_label,
        question=question,
        extra_detail=extra_detail,
    )
    default_description = _jira_context_description(
        area_label=area_label,
        target_label=target_label,
        run_id=run_id,
        case_id=case_id,
        status_label=status_label,
        question=question,
        extra_detail=extra_detail,
    )
    result_key = f"voc_context_jira_result_{safe_key}"
    diagnostic = _jira_context_diagnostic_model(
        status_label=status_label,
        extra_detail=extra_detail,
    )
    default_issue_type = (
        diagnostic.get("issue_type")
        if diagnostic.get("issue_type") in JIRA_CONTEXT_ISSUE_TYPES
        else JIRA_CONTEXT_ISSUE_TYPES[0]
    )
    default_priority = (
        diagnostic.get("priority")
        if diagnostic.get("priority") in JIRA_CONTEXT_PRIORITIES
        else JIRA_CONTEXT_PRIORITIES[0]
    )
    action_container = st.popover(
        JIRA_ACTION_POPOVER_LABEL,
        icon=":material/bug_report:",
        type="secondary",
        width="content",
        key=f"voc_context_action_{safe_key}",
        help="선택한 표 Row를 Jira 이슈로 등록합니다.",
    )
    with action_container:
        st.html(
            """
            <div class="voc-jira-action-head">
                <span class="voc-jira-mini-logo">J</span>
                <div class="voc-jira-action-copy">
                    <strong>Jira 이슈 등록</strong>
                    <small>선택한 표 Row를 Jira 작업으로 연결합니다</small>
                </div>
                <span class="voc-jira-action-badge">Jira</span>
            </div>
            <style>
            .voc-jira-action-head{
                display:flex;align-items:center;gap:9px;margin:8px 0 8px;
                padding:9px 10px;border:1px solid #c7d8f7;border-radius:11px;
                background:linear-gradient(135deg,#f4f8ff,#ffffff);
                font-family:'Segoe UI','Malgun Gothic',sans-serif;
            }
            .voc-jira-mini-logo{
                display:grid;place-items:center;flex:0 0 24px;width:24px;height:24px;
                border-radius:7px;background:#0c66e4;color:#ffffff;
                font-size:13px;font-weight:900;line-height:1;
                box-shadow:0 3px 8px rgba(12,102,228,.22);
            }
            .voc-jira-action-copy{min-width:0;flex:1}
            .voc-jira-action-copy strong{
                display:block;color:#172b4d;font-size:12px;font-weight:900;line-height:1.15;
            }
            .voc-jira-action-copy small{
                display:block;color:#5e6c84;font-size:10px;line-height:1.2;margin-top:2px;
                white-space:nowrap;overflow:hidden;text-overflow:ellipsis;
            }
            .voc-jira-action-badge{
                flex:0 0 auto;padding:3px 7px;border-radius:999px;
                background:#deebff;color:#0747a6;border:1px solid #b3d4ff;
                font-size:9px;font-weight:900;
            }
            </style>
            """
        )
        if not snapshot["ready"]:
            missing = ", ".join(snapshot.get("missing") or [])
            st.warning(
                f"Jira API 설정이 필요합니다. 누락: {missing or '-'}",
                icon=":material/warning:",
            )
            return

        latest_result = st.session_state.get(result_key)
        if isinstance(latest_result, dict) and latest_result.get("issue_key"):
            st.success(
                f"최근 등록: {latest_result['issue_key']}",
                icon=":material/check_circle:",
            )
            if latest_result.get("issue_url"):
                st.link_button("Jira에서 열기", latest_result["issue_url"], icon=":material/open_in_new:")

        with st.form(f"voc_context_jira_form_{safe_key}", border=False):
            summary = st.text_input(
                "요약",
                value=default_summary,
                max_chars=255,
                key=f"voc_context_jira_summary_{safe_key}",
            )
            issue_type = st.segmented_control(
                "유형",
                JIRA_CONTEXT_ISSUE_TYPES,
                default=default_issue_type,
                key=f"voc_context_jira_type_{safe_key}",
                width="stretch",
            ) or default_issue_type
            priority = st.segmented_control(
                "우선순위",
                JIRA_CONTEXT_PRIORITIES,
                default=default_priority,
                key=f"voc_context_jira_priority_{safe_key}",
                width="stretch",
            ) or default_priority
            description = st.text_area(
                "설명",
                value=default_description,
                height=170,
                key=f"voc_context_jira_description_{safe_key}",
            )
            submitted = st.form_submit_button(
                "Jira 이슈 등록",
                icon=":material/add_task:",
                type="primary",
                width="stretch",
            )

        if submitted:
            jira_priority = "" if priority == "미지정" else priority
            try:
                created = create_jira_issue(
                    summary=summary,
                    description=description,
                    issue_type=issue_type,
                    priority=jira_priority,
                    labels=_jira_context_labels(
                        area_label,
                        target_label,
                        run_id,
                        case_id,
                        *diagnostic.get("labels", []),
                    ),
                )
            except (JiraConfigurationError, JiraIssueCreateError) as exc:
                st.error(str(exc), icon=":material/error:")
                return
            history_item = {
                "created_at": created.get("created_at", datetime.now().isoformat(timespec="seconds")),
                "issue_key": created.get("key", "-"),
                "issue_url": created.get("url", ""),
                "summary": summary,
                "issue_type": issue_type,
                "priority": jira_priority or "-",
                "diagnostic_category": diagnostic.get("category", ""),
                "case_id": case_id,
                "run_id": run_id,
                "source": area_label,
            }
            _record_contextual_jira_issue(history_item)
            st.session_state[result_key] = history_item
            st.toast(f"Jira 이슈 {created.get('key', '-')}를 등록했습니다.", icon=":material/check_circle:")
            st.success(f"Jira 이슈 {created.get('key', '-')}를 등록했습니다.", icon=":material/check_circle:")
            if created.get("url"):
                st.link_button("Jira에서 열기", created["url"], icon=":material/open_in_new:")


def _manual_pipeline_case_context(snapshot: dict, *, running: bool) -> dict:
    result = st.session_state.get("goal_testcase_result") or {}
    result_case = result.get("case", {}) if isinstance(result.get("case"), dict) else {}
    running_case_id = str(st.session_state.get("goal_testcase_running_case_id") or "").strip()
    selected_case_id = str(st.session_state.get("goal_testcase_selected_case_id") or "").strip()
    result_case_id = str(result_case.get("case_id") or "").strip()
    case_id = (
        running_case_id
        if running and running_case_id
        else selected_case_id
        or result_case_id
        or "-"
    )

    question = (
        result_case.get("question")
        or result_case.get("name")
        or (st.session_state.get("goal_testcase_running_question") if running else "")
        or ""
    )
    if not question:
        selected_case = _selected_goal_testcase() or {}
        if str(selected_case.get("case_id") or "") == case_id:
            question = selected_case.get("question") or selected_case.get("name") or ""
    if not question and case_id != "-":
        catalog_case = next(
            (
                case
                for case in load_unified_quality_cases().get("cases", [])
                if str(case.get("case_id") or "") == case_id
            ),
            {},
        )
        question = catalog_case.get("question") or catalog_case.get("name") or ""
    question = question or "-"
    source_events = snapshot.get("events", []) if isinstance(snapshot.get("events"), list) else []
    started_at = st.session_state.get("goal_testcase_started_at", "")
    completed_at = st.session_state.get("goal_testcase_completed_at", "")
    last_event_at = next(
        (
            str(event.get("timestamp") or "")
            for event in reversed(source_events)
            if isinstance(event, dict) and event.get("timestamp")
        ),
        "",
    )
    has_recent = bool(result or source_events or started_at or completed_at)
    timestamp_source = (
        started_at
        if running
        else completed_at
        or last_event_at
        or started_at
    )
    state_label = "현재 실행 중" if running else ("최근 수행" if has_recent else "선택 Case")
    timestamp_label = "시작 일시" if running else "최근 수행 일시"
    return {
        "case_id": case_id,
        "question": str(question or "-"),
        "question_short": _manual_pipeline_compact_text(str(question or "-"), 88),
        "state_label": state_label,
        "state_class": "running" if running else ("recent" if has_recent else "selected"),
        "timestamp_label": timestamp_label,
        "timestamp": _manual_pipeline_timestamp_text(timestamp_source),
        "has_recent": has_recent,
    }


def _pipeline_run_summary(snapshot: dict, *, running: bool) -> dict:
    events = _trace_display_events(snapshot.get("events", []))
    result = st.session_state.get("goal_testcase_result") or {}
    selected_case_id = st.session_state.get("goal_testcase_selected_case_id", "")
    started_at = st.session_state.get("goal_testcase_started_at", "")
    completed_at = st.session_state.get("goal_testcase_completed_at", "")

    execution = result.get("execution", {}) if isinstance(result, dict) else {}
    result_payload = execution.get("result", {}) if isinstance(execution, dict) else {}
    execution_ok = bool(execution.get("ok"))
    if result.get("mode") == "voc":
        execution_ok = execution_ok and bool(result_payload.get("ok"))

    failures = sum(event.get("status") == "failure" for event in events)
    successes = sum(event.get("status") == "success" for event in events)
    active_event = next(
        (event for event in reversed(events) if event.get("status") == "started"),
        {},
    )
    active_agent_name = str(active_event.get("target") or "")
    active_agent_number = next(
        (
            index
            for index, (name, _, _) in enumerate(AGENT_PIPELINE, start=1)
            if name == active_agent_name
        ),
        None,
    )
    if running:
        state = "preparing" if not events else "running"
        if not events:
            label = "테스트 수행 준비"
        elif active_agent_number:
            label = f"Agent {active_agent_number} · {active_agent_name} 수행 중"
        else:
            label = "파이프라인 결과 저장 중"
    elif result:
        state = "completed" if execution_ok else "failed"
        label = "수행 완료" if execution_ok else "수행 실패"
    elif failures:
        state, label = "failed", "최근 수행 실패"
    elif events and events[-1].get("operation") == "IntentDataTransfer" and events[-1].get("status") == "success":
        state, label = "completed", "최근 수행 완료"
    elif events:
        state, label = "recorded", "최근 수행 기록"
    else:
        state, label = "waiting", "수행 대기"

    duration_seconds = 0.0
    try:
        start = datetime.fromisoformat(started_at) if started_at else None
        end = datetime.fromisoformat(completed_at) if completed_at else datetime.now().astimezone()
        if start:
            duration_seconds = max(0.0, (end - start).total_seconds())
    except (TypeError, ValueError):
        pass
    if not duration_seconds and events:
        try:
            first = datetime.fromisoformat(str(events[0].get("started_at") or events[0].get("timestamp")))
            last = datetime.fromisoformat(str(events[-1].get("timestamp")))
            duration_seconds = max(0.0, (last - first).total_seconds())
        except (TypeError, ValueError):
            pass

    return {
        "state": state,
        "label": label,
        "case_id": selected_case_id or result.get("case", {}).get("case_id") or st.session_state.get("goal_testcase_running_case_id", "-"),
        "steps": len(events),
        "successes": successes,
        "failures": failures,
        "duration_seconds": duration_seconds,
        "active_agent_number": active_agent_number,
        "active_agent_name": active_agent_name,
    }


def _render_agent_pipeline(snapshot: dict, running: bool):
    events = snapshot.get("events", [])
    agent_snapshot = st.session_state.get("goal_testcase_agent_snapshot")
    if not isinstance(agent_snapshot, dict):
        agent_snapshot = {"agents": []}
    agent_statuses = {agent["name"]: agent for agent in agent_snapshot.get("agents", [])}
    operation_states = {}
    for index, event in enumerate(events):
        key = (event.get("source"), event.get("target"), event.get("operation"))
        operation_states[key] = (index, event)
    active_events = [item for item in operation_states.values() if item[1].get("status") == "started"]
    current_agent = max(active_events, default=(-1, {}), key=lambda item: item[0])[1].get("target")

    cards = []
    for agent_number, (name, role, port) in enumerate(AGENT_PIPELINE, start=1):
        agent_status = agent_statuses.get(name, {})
        is_enabled = agent_status.get("healthy", True)
        related = [event for event in events if event.get("target") == name]
        failures = [event for event in related if event.get("status") == "failure"]
        successes = [event for event in related if event.get("status") == "success"]
        if not is_enabled:
            state, label = "disabled", "비활성"
        elif failures and (not successes or failures[-1].get("timestamp", "") >= successes[-1].get("timestamp", "")):
            state, label = "error", "오류"
        elif running and current_agent == name:
            state, label = "active", "작업 중"
        elif successes:
            state, label = "done", "완료"
        else:
            state, label = "waiting", "대기"
        last_in = next((event for event in reversed(related) if event.get("input_keywords")), {})
        last_out = next((event for event in reversed(related) if event.get("output_keywords")), {})
        duration = last_out.get("duration_ms") or (successes[-1].get("duration_ms") if successes else 0)
        status_detail = escape(str(agent_status.get("status", "점검 중" if running else "미확인")))
        cards.append(f"""
          <div class="agent-card {state}">
            <div class="agent-head"><span class="agent-icon">{agent_number}</span><span><b>{name}</b><small>{role} · {port}</small></span><em>{label}</em></div>
            <div class="io in"><b>IN</b><span>{_keyword_text(last_in.get('input_keywords'))}</span></div>
            <div class="io out"><b>OUT</b><span>{_keyword_text(last_out.get('output_keywords'))}</span></div>
            <div class="agent-meta">상태 {status_detail} · 처리시간 {float(duration or 0):,.0f} ms</div>
          </div>""")

    context = _manual_pipeline_case_context(snapshot, running=running)
    status_text = "현재 실행 중 Agent 파이프라인" if running else "최근 수행 Agent 파이프라인"
    case_meta = escape(f"{context['case_id']} · {context['question_short']}")
    time_meta = escape(f"{context['timestamp_label']} {context['timestamp']}")
    st.html(f"""
    <style>
      .pipeline-wrap{{border:1px solid #d8e3f0;border-radius:16px;padding:16px;background:linear-gradient(180deg,#f8fbff,#fff);margin:4px 0 18px}}
      .pipeline-title{{display:flex;justify-content:space-between;align-items:center;margin-bottom:12px;color:#17355f;font:600 13px 'Segoe UI','Malgun Gothic',sans-serif}}
      .pipeline-title span:last-child{{display:flex;gap:8px;align-items:center;color:#6d7d91;font-size:10px;font-weight:600;min-width:0}}
      .pipeline-title span:last-child b{{max-width:520px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;color:#17355f;font-size:11px}}
      .pipeline-title span:last-child em{{font-style:normal;color:#8593a5;font-size:9px}}
      .pipeline-grid{{display:grid;grid-template-columns:repeat(6,minmax(170px,1fr));gap:24px;overflow-x:auto;padding:4px}}
      .agent-card{{position:relative;min-width:170px;border:2px solid #d6deea;border-radius:13px;background:#fff;padding:12px;box-shadow:0 4px 12px rgba(26,56,96,.07)}}
      .agent-card:not(:last-child):after{{content:'→';position:absolute;right:-22px;top:47%;color:#4470aa;font-size:22px;font-weight:800}}
      .agent-card.active{{border-color:#1b6fd1;box-shadow:0 0 0 4px rgba(27,111,209,.12),0 8px 20px rgba(27,111,209,.18);animation:pulse 1.4s infinite}}
      .agent-card.done{{border-color:#36a269;background:#f7fff9}} .agent-card.error{{border-color:#dc4c4c;background:#fff8f8}}
      .agent-card.disabled{{border-color:#b8bec7;background:#eceff2;box-shadow:none;filter:grayscale(1);opacity:.72}}
      .agent-card.disabled:not(:last-child):after{{color:#aeb4bd}}
      .agent-head{{display:flex;gap:8px;align-items:center;font:13px 'Segoe UI','Malgun Gothic',sans-serif;color:#172b48}}
      .agent-head small{{display:block;color:#75849a;font-size:10px;margin-top:2px}} .agent-head em{{margin-left:auto;font-style:normal;font-size:10px;padding:3px 7px;border-radius:10px;background:#edf2f8}}
      .active .agent-head em{{background:#1b6fd1;color:#fff}} .done .agent-head em{{background:#daf4e3;color:#16713d}} .error .agent-head em{{background:#fde0e0;color:#a62525}}
      .disabled .agent-head{{color:#59616c}} .disabled .agent-head em{{background:#7a828d;color:#fff}}
      .agent-icon{{display:grid;place-items:center;width:30px;height:30px;border-radius:9px;background:#173f75;color:#fff;font-weight:800;font-size:10px}}
      .io{{display:grid;grid-template-columns:30px 1fr;gap:5px;margin-top:10px;border-radius:8px;padding:7px;font:10px 'Segoe UI','Malgun Gothic',sans-serif;line-height:1.35}}
      .io b{{color:#fff;text-align:center;border-radius:5px;padding:2px}} .io span{{color:#465871;overflow-wrap:anywhere}}
      .io.in{{background:#eef5ff}} .io.in b{{background:#2f6eb5}} .io.out{{background:#eefaf2}} .io.out b{{background:#32935b}}
      .disabled .agent-icon,.disabled .io b{{background:#858c95}} .disabled .io{{background:#dde1e5}} .disabled .io span{{color:#737a84}}
      .agent-meta{{text-align:right;color:#8794a8;font-size:9px;margin-top:7px}}
      @keyframes pulse{{50%{{transform:translateY(-2px)}}}} @media(max-width:900px){{.pipeline-grid{{grid-template-columns:repeat(6,180px)}}}}
    </style>
    <div class="pipeline-wrap"><div class="pipeline-title"><span>{status_text}</span><span><b>{case_meta}</b><em>{time_meta}</em></span></div><div class="pipeline-grid">{''.join(cards)}</div></div>
    """)


def _render_agent_pipeline_v2(
    snapshot: dict,
    running: bool,
    preparation: dict | None = None,
):
    source_events = (
        snapshot.get("events", [])
        if isinstance(snapshot.get("events", []), list)
        else []
    )
    raw_event_count = len(source_events)
    raw_events = [
        event
        if isinstance(event, dict)
        else {
            "operation": "형식 미확인 원본 로그",
            "status": "unknown",
        }
        for event in source_events
    ]
    events = _trace_display_events(raw_events)
    agent_numbers = {name: index for index, (name, _, _) in enumerate(AGENT_PIPELINE, start=1)}
    display_statuses = _trace_event_display_statuses(events, running=running)
    run_summary = _pipeline_run_summary(snapshot, running=running)
    context = _manual_pipeline_case_context(snapshot, running=running)
    case_context = escape(f"{context['case_id']} · {context['question_short']}")
    event_title = "실시간 실행 이벤트(Agent 호출)" if running else "최근 수행 이벤트(Agent 호출)"
    raw_event_title = "실시간 실행 이벤트(원본 로그)" if running else "최근 수행 이벤트(원본 로그)"
    event_meta = f"Agent 호출 {len(events)}건"
    raw_event_meta = f"원본 로그 {raw_event_count}건 · 전체 표시"
    if context["timestamp"] != "-":
        event_meta += f" · {context['timestamp_label']} {context['timestamp']}"
        raw_event_meta += f" · {context['timestamp_label']} {context['timestamp']}"
    event_meta = escape(event_meta)
    raw_event_meta = escape(raw_event_meta)

    status_labels = {
        "started": "진행",
        "completed": "완료",
        "ended": "종료",
        "success": "성공",
        "failure": "실패",
    }
    raw_status_labels = {
        "started": "시작",
        "success": "성공",
        "failure": "실패",
    }

    def build_trace_history(
        event_items: list[dict],
        statuses: list[str],
        *,
        raw: bool = False,
    ) -> str:
        trace_rows = []
        for event_index, (event, display_status) in enumerate(
            zip(event_items, statuses),
            start=1,
        ):
            source = escape(str(event.get("source") or "-"))
            target = escape(str(event.get("target") or "-"))
            target_number = agent_numbers.get(event.get("target"), "-")
            operation = escape(str(event.get("operation") or "작업"))
            status = str(display_status or "unknown")
            status_class = status if status in status_labels else "unknown"
            status_label = (
                raw_status_labels.get(status, escape(status))
                if raw
                else status_labels.get(status, "기록")
            )
            timestamp = str(event.get("timestamp") or "-")
            time_text = escape(timestamp[11:19] if len(timestamp) >= 19 else timestamp)
            duration = float(event.get("duration_ms") or 0)
            duration_text = f"{duration:,.0f} ms" if duration else "-"
            previous_event = event_items[event_index - 2] if event_index > 1 else None
            flow = _trace_flow_explanation(event, previous_event)
            flow_kind = escape(str(flow["kind"]))
            flow_label = escape(str(flow["label"]))
            transition = escape(str(flow["transition"]))
            reason = escape(str(flow["reason"]))
            inferred_badge = " · 추정" if flow.get("inferred") else ""
            reason_text = (
                f"<b>{transition} · {flow_label}{inferred_badge}</b>"
                f"<span>{reason}</span>"
            )
            reason_class = f"flow2-reason flow2-{flow_kind}"
            trace_rows.append(f"""
          <div class="flow2-trace-row {status_class}">
            <span class="flow2-seq">#{event_index:02d}</span>
            <span class="flow2-agent-no">Agent {target_number}</span>
            <span class="flow2-time">{time_text}</span>
            <span class="flow2-route"><b>{source}</b><i>→</i><b>{target}</b><small>{operation}</small></span>
            <em>{status_label}</em>
            <span class="flow2-duration">{duration_text}</span>
            <span class="{reason_class}">{reason_text}</span>
          </div>""")
        return "".join(trace_rows)

    trace_history = build_trace_history(events, display_statuses)
    raw_trace_history = build_trace_history(
        raw_events,
        [str(event.get("status") or "unknown") for event in raw_events],
        raw=True,
    )
    if not raw_trace_history:
        raw_trace_history = (
            '<div class="flow2-trace-empty">아직 기록된 원본 로그가 없습니다.</div>'
        )
    has_run_context = bool(
        events
        or st.session_state.get("goal_testcase_started_at")
        or st.session_state.get("goal_testcase_result")
    )
    if not preparation:
        preparation = _new_manual_preparation_progress()
        if events or not running:
            preparation["status"] = "COMPLETED"
            for step in preparation["steps"]:
                step["status"] = "success"
        else:
            preparation["steps"][0]["status"] = "active"
    preparation_status = preparation.get("status", "RUNNING")
    preparation_state = {
        "COMPLETED": "completed",
        "ERROR": "failed",
    }.get(preparation_status, "active")
    preparation_label = {
        "COMPLETED": "성공",
        "ERROR": "실패",
    }.get(preparation_status, "진행 중")
    preparation_icon = {
        "COMPLETED": "✓",
        "ERROR": "!",
    }.get(preparation_status, "●")
    step_status_labels = {
        "waiting": "대기",
        "active": "진행",
        "success": "완료",
        "failure": "실패",
    }
    preparation_steps = list(preparation.get("steps", []))
    completed_preparation_steps = sum(
        step.get("status") == "success"
        for step in preparation_steps
    )
    preparation_progress_text = (
        "준비 완료"
        if preparation_status == "COMPLETED"
        else f"{completed_preparation_steps}/5 완료 · 순서대로 처리 중"
    )
    current_preparation_step = next(
        (
            str(step.get("label") or "-")
            for step in preparation_steps
            if step.get("status") in {"active", "failure"}
        ),
        "Agent 파이프라인 호출 준비 완료"
        if preparation_status == "COMPLETED"
        else "다음 준비 단계 확인 중",
    )
    preparation_summary_card = (
        f"""
      <article class="flow2-preparation {preparation_state} flow2-preparation-summary">
        <div class="flow2-preparation-head">
          <span class="flow2-preparation-icon">{preparation_icon}</span>
          <span><strong>테스트 수행 준비</strong><small>Agent 이벤트 생성 전 선행 작업</small></span>
          <em>{preparation_label}</em>
        </div>
        <div class="flow2-preparation-progress">{preparation_progress_text}</div>
        <div class="flow2-preparation-current">
          <b>{'완료 결과' if preparation_status == 'COMPLETED' else '현재 단계'}</b>
          <span>{escape(current_preparation_step)}</span>
        </div>
      </article>"""
        if has_run_context
        else ""
    )

    def build_preparation_step_card(group: list[dict]) -> str:
        statuses = [str(step.get("status") or "waiting") for step in group]
        if "failure" in statuses:
            group_state, group_label, group_icon = "failed", "실패", "!"
        elif "active" in statuses:
            group_state, group_label, group_icon = "active", "진행", "●"
        elif statuses and all(status == "success" for status in statuses):
            group_state, group_label, group_icon = "completed", "완료", "✓"
        else:
            group_state, group_label, group_icon = "waiting", "대기", "○"
        first_number = int(group[0].get("number") or 1)
        last_number = int(group[-1].get("number") or first_number)
        group_rows = "".join(
            (
                f"<li class='{escape(str(step.get('status', 'waiting')))}'>"
                f"<span>{int(step.get('number') or index)}</span>"
                f"<b>{escape(str(step.get('label') or '-'))}</b>"
                f"<em>{step_status_labels.get(step.get('status'), '대기')}</em></li>"
            )
            for index, step in enumerate(group, start=first_number)
        )
        completed_in_group = sum(status == "success" for status in statuses)
        return f"""
      <article class="flow2-preparation {group_state} flow2-preparation-step">
        <div class="flow2-preparation-head">
          <span class="flow2-preparation-icon">{group_icon}</span>
          <span><strong>준비 단계 {first_number}–{last_number}</strong><small>{completed_in_group}/{len(group)}단계 완료</small></span>
          <em>{group_label}</em>
        </div>
        <ol>{group_rows}</ol>
      </article>"""

    preparation_step_cards = (
        "".join(
            build_preparation_step_card(group)
            for group in (preparation_steps[:3], preparation_steps[3:])
            if group
        )
        if has_run_context
        else ""
    )
    preparation_cards = preparation_summary_card + preparation_step_cards
    duration_text = (
        f"{run_summary['duration_seconds']:.1f}초"
        if run_summary["duration_seconds"]
        else "측정 중"
    )
    summary_card = f"""
      <article class="flow2-run-summary {run_summary['state']}">
        <span class="flow2-summary-icon">{'✓' if run_summary['state'] == 'completed' else '!' if run_summary['state'] == 'failed' else '●'}</span>
        <strong>{escape(run_summary['label'])}</strong>
        <small>Case {escape(str(run_summary['case_id']))}</small>
        <dl>
          <div><dt>총 스텝</dt><dd>{run_summary['steps']}단계</dd></div>
          <div><dt>성공</dt><dd>{run_summary['successes']}건</dd></div>
          <div><dt>실패</dt><dd>{run_summary['failures']}건</dd></div>
          <div><dt>수행시간</dt><dd>{duration_text}</dd></div>
        </dl>
      </article>""" if has_run_context else ""

    st.html(f"""
    <style>
      .flow2-wrap{{border:1px solid #d8e3f0;border-radius:16px;padding:16px;background:#fff;margin:4px 0 18px}}
      .flow2-trace{{font-family:'Segoe UI','Malgun Gothic',sans-serif}}
      .flow2-trace + .flow2-trace{{margin-top:14px;padding-top:14px;border-top:1px solid #e3e9f1}}
      .flow2-trace-head{{display:flex;justify-content:space-between;align-items:center;margin-bottom:8px;color:#334d70;font-size:11px;font-weight:700}}
      .flow2-trace-head span:last-child{{color:#7a889b;font-size:9px;font-weight:500}}
      .flow2-head-main{{display:flex;align-items:center;gap:7px;min-width:0}}
      .flow2-head-main strong{{font-size:11px;color:#334d70;white-space:nowrap}}
      .flow2-head-main b{{max-width:620px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;color:#17355f;font-size:11px}}
      .flow2-context-badge{{display:inline-grid;place-items:center;border-radius:999px;padding:3px 8px;background:#eef3f7;color:#66798c;font-size:9px;font-style:normal;white-space:nowrap}}
      .flow2-context-badge.running{{background:#dfeeff;color:#175b9d}} .flow2-context-badge.recent{{background:#eaf8ef;color:#247147}}
      .flow2-head-meta{{white-space:nowrap}}
      .flow2-raw-summary{{display:grid;grid-template-columns:1fr auto auto;gap:10px;align-items:center;color:#334d70;font-size:11px;font-weight:700;cursor:pointer;list-style:none}}
      .flow2-raw-summary::-webkit-details-marker{{display:none}}
      .flow2-raw-summary span:nth-child(2){{color:#7a889b;font-size:9px;font-weight:500}}
      .flow2-raw-summary:after{{content:'펼치기 ＋';min-width:58px;padding:4px 8px;border:1px solid #c9d8e8;border-radius:8px;background:#f3f7fb;color:#42688f;font-size:9px;text-align:center}}
      .flow2-raw-trace[open] .flow2-raw-summary:after{{content:'접기 −'}}
      .flow2-raw-content{{margin-top:8px}}
      .flow2-raw-guide{{margin:-2px 0 7px;color:#7a889b;font-size:8px}}
      .flow2-legend{{display:flex;flex-wrap:wrap;gap:5px;margin:0 0 7px}} .flow2-legend span{{padding:3px 7px;border-radius:10px;background:#eef5fd;color:#355b87;font-size:8px}}
      .flow2-legend .return{{background:#edf9f2;color:#17643b}} .flow2-legend .feedback{{background:#fff8e8;color:#8a540c}} .flow2-legend .failure{{background:#fff0f0;color:#982d2d}} .flow2-legend .inferred{{background:#f7f2fb;color:#654186}}
      .flow2-trace-list{{overflow-x:auto;padding:10px;border:1px solid #e1e7ef;border-radius:10px;background:#f9fbfd;direction:rtl}}
      .flow2-trace-track{{display:flex;align-items:flex-start;gap:28px;width:max-content;min-width:100%;direction:ltr}}
      .flow2-trace-row{{position:relative;display:grid;grid-template-columns:38px 56px 1fr 42px;grid-template-areas:'seq agent time status' 'route route route duration' 'reason reason reason reason';gap:8px;align-items:center;width:280px;min-width:280px;height:{MANUAL_EVENT_CARD_HEIGHT}px;box-sizing:border-box;padding:10px;border:1px solid #dfe6ef;border-radius:10px;background:#fff;box-shadow:0 3px 9px rgba(30,59,96,.06);color:#52637a;font-size:9px}}
      .flow2-trace-row:not(:last-child):after{{content:'→';position:absolute;right:-22px;top:46%;color:#4b75a9;font-size:19px;font-weight:800}}
      .flow2-trace-row:hover{{border-color:#9db7d5;background:#f7fbff}}
      .flow2-seq{{grid-area:seq;color:#95a1b1;font-family:Consolas,monospace}} .flow2-time{{grid-area:time;color:#6f7f93;font-family:Consolas,monospace}}
      .flow2-agent-no{{display:inline-grid;place-items:center;border-radius:8px;background:#173f75;color:#fff;padding:3px 4px;font-weight:800}}
      .flow2-agent-no{{grid-area:agent}} .flow2-route{{grid-area:route;display:grid;grid-template-columns:auto 14px auto 1fr;align-items:center;gap:3px;color:#263d5d}}
      .flow2-route i{{font-style:normal;color:#4b75a9;text-align:center}} .flow2-route small{{color:#7f8da0;margin-left:5px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}}
      .flow2-trace-row em{{grid-area:status;font-style:normal;text-align:center;border-radius:8px;padding:3px 4px;background:#e8eef5;color:#52637a}}
      .flow2-trace-row.started em{{background:#dfeeff;color:#175b9d}} .flow2-trace-row.completed em,.flow2-trace-row.success em{{background:#ddf4e6;color:#247147}} .flow2-trace-row.ended em{{background:#eef1f4;color:#687687}} .flow2-trace-row.failure em{{background:#fde3e3;color:#a52d2d}}
      .flow2-trace-row.failure{{background:#fff8f8}} .flow2-duration{{grid-area:duration;text-align:right;color:#6e7c8e;font-family:Consolas,monospace}}
      .flow2-reason{{grid-area:reason;display:grid;gap:3px;min-height:38px;padding:6px 7px;border-left:3px solid #4c7fba;border-radius:6px;background:#eef5fd;color:#355b87;line-height:1.35}}
      .flow2-reason b{{font-size:9px}} .flow2-reason span{{font-size:8px;color:#60748d}}
      .flow2-return{{border-left-color:#2f9660;background:#edf9f2;color:#17643b}} .flow2-feedback,.flow2-rework{{border-left-color:#b7791f;background:#fff8e8;color:#8a540c}}
      .flow2-failure{{border-left-color:#c84646;background:#fff0f0;color:#982d2d}} .flow2-inferred{{border-left-color:#8b67b1;background:#f7f2fb;color:#654186}}
      .flow2-failure span{{color:#a84a4a}} .flow2-feedback span,.flow2-rework span{{color:#8d6a32}} .flow2-inferred span{{color:#766383}}
      .flow2-trace-empty{{padding:18px;text-align:center;color:#8996a7;font-size:10px}}
      .flow2-preparation{{position:relative;width:280px;min-width:280px;height:{MANUAL_EVENT_CARD_HEIGHT}px;box-sizing:border-box;padding:10px;border:1px solid #9db7d5;border-radius:10px;background:#f4f8fd;box-shadow:0 3px 9px rgba(30,59,96,.06);color:#1c4f82}}
      .flow2-preparation:after{{content:'→';position:absolute;right:-22px;top:46%;color:#4b75a9;font-size:19px;font-weight:800}}
      .flow2-preparation-head{{display:grid;grid-template-columns:32px 1fr auto;gap:8px;align-items:center}}
      .flow2-preparation-icon{{display:grid;place-items:center;width:30px;height:30px;border-radius:9px;background:#2f75b5;color:#fff;font-size:13px;font-weight:800}}
      .flow2-preparation-head strong{{display:block;font-size:12px}} .flow2-preparation-head small{{display:block;margin-top:2px;color:#71869d;font-size:8px}}
      .flow2-preparation-head em{{font-style:normal;border-radius:8px;padding:3px 6px;background:#dcecff;color:#175b9d;font-size:9px;white-space:nowrap}}
      .flow2-preparation-progress{{margin-top:7px;padding:5px 7px;border-radius:7px;background:rgba(255,255,255,.72);font-size:9px;font-weight:700}}
      .flow2-preparation-current{{display:grid;gap:3px;margin-top:8px;padding:8px;border-left:3px solid #4c7fba;border-radius:6px;background:#eef5fd}}
      .flow2-preparation-current b{{font-size:9px}} .flow2-preparation-current span{{color:#60748d;font-size:8px}}
      .flow2-preparation ol{{display:grid;gap:3px;margin:8px 0 0;padding:0;list-style:none}}
      .flow2-preparation li{{display:grid;grid-template-columns:19px 1fr 30px;gap:6px;align-items:center;min-height:19px;padding:2px 6px;border-radius:6px;background:rgba(255,255,255,.72);font-size:9px;color:#98a3b1}}
      .flow2-preparation li span{{display:grid;place-items:center;width:16px;height:16px;border-radius:50%;background:#dceaff;color:#285f99;font:700 8px Consolas,monospace}}
      .flow2-preparation li b{{font-weight:600}} .flow2-preparation li em{{font-style:normal;text-align:right;font-size:8px}}
      .flow2-preparation li.active{{color:#175b9d;background:#e8f3ff}} .flow2-preparation li.active span{{background:#2f75b5;color:#fff}}
      .flow2-preparation li.success{{color:#247147;background:#edf9f2}} .flow2-preparation li.success span{{background:#2f9660;color:#fff}}
      .flow2-preparation li.failure{{color:#a52d2d;background:#fff0f0}} .flow2-preparation li.failure span{{background:#c84646;color:#fff}}
      .flow2-preparation-step.active{{animation:flow2-preparation-pulse 1.5s ease-in-out infinite}}
      .flow2-preparation.completed{{border-color:#2f9660;background:#edf9f2;color:#17643b;box-shadow:0 4px 12px rgba(47,150,96,.12)}}
      .flow2-preparation.completed .flow2-preparation-icon{{background:#2f9660}} .flow2-preparation.completed .flow2-preparation-head em{{background:#d9f1e3;color:#247147}}
      .flow2-preparation.completed li span{{background:#d9f1e3;color:#247147}}
      .flow2-preparation.failed{{border-color:#c84646;background:#fff0f0;color:#982d2d;animation:none}} .flow2-preparation.failed .flow2-preparation-icon{{background:#c84646}} .flow2-preparation.failed .flow2-preparation-head em{{background:#fde3e3;color:#a52d2d}}
      .flow2-preparation.waiting{{border-color:#d8e0e9;background:#f7f8fa;color:#69798d;box-shadow:none}}
      .flow2-preparation.waiting .flow2-preparation-icon{{background:#9aa8b7}} .flow2-preparation.waiting .flow2-preparation-head em{{background:#e9edf1;color:#69798d}}
      .flow2-run-summary{{position:relative;display:grid;grid-template-columns:34px 1fr;grid-template-areas:'icon title' 'icon case' 'stats stats';gap:4px 9px;width:280px;min-width:280px;height:{MANUAL_EVENT_CARD_HEIGHT}px;box-sizing:border-box;padding:12px;border:2px solid #7da8d1;border-radius:12px;background:#edf6ff;color:#174f85;box-shadow:0 6px 16px rgba(23,79,133,.15)}}
      .flow2-run-summary .flow2-summary-icon{{grid-area:icon;display:grid;place-items:center;width:32px;height:32px;border-radius:50%;background:#2f75b5;color:#fff;font-size:16px;font-weight:800}}
      .flow2-run-summary strong{{grid-area:title;font-size:13px}} .flow2-run-summary small{{grid-area:case;color:#58728f;font-size:9px}}
      .flow2-run-summary dl{{grid-area:stats;display:grid;grid-template-columns:repeat(4,1fr);gap:5px;margin:10px 0 0}}
      .flow2-run-summary dl div{{padding:6px 4px;border-radius:7px;background:rgba(255,255,255,.66);text-align:center}} .flow2-run-summary dt{{font-size:8px;color:#71859c}} .flow2-run-summary dd{{margin:2px 0 0;font-size:10px;font-weight:700}}
      .flow2-run-summary.completed{{border-color:#2f9660;background:#eaf8ef;color:#17643b}} .flow2-run-summary.completed .flow2-summary-icon{{background:#2f9660}}
      .flow2-run-summary.failed{{border-color:#c84646;background:#fff0f0;color:#982d2d}} .flow2-run-summary.failed .flow2-summary-icon{{background:#c84646}}
      .flow2-run-summary.waiting,.flow2-run-summary.recorded{{border-color:#aeb9c5;background:#f1f3f5;color:#5c6978;box-shadow:none}} .flow2-run-summary.waiting .flow2-summary-icon,.flow2-run-summary.recorded .flow2-summary-icon{{background:#7f8a96}}
      @keyframes flow2-preparation-pulse{{50%{{transform:translateY(-2px);box-shadow:0 0 0 4px rgba(47,117,181,.10),0 7px 16px rgba(35,91,148,.18)}}}}
    </style>
    <div class="flow2-wrap">
      <div class="flow2-trace">
        <div class="flow2-trace-head"><span class="flow2-head-main"><em class="flow2-context-badge {context['state_class']}">{escape(context['state_label'])}</em><strong>{event_title}</strong><b>{case_context}</b></span><span class="flow2-head-meta">{event_meta}</span></div>
        <div class="flow2-legend"><span>분석·조회·평가</span><span class="return">결과 반환</span><span class="feedback">재검토·보완</span><span class="failure">호출 실패</span><span class="inferred">실행 Trace 사유 미기록·추정</span></div>
        <div class="flow2-trace-list"><div class="flow2-trace-track">{preparation_cards}{trace_history}{summary_card}</div></div>
      </div>
      <details class="flow2-trace flow2-raw-trace">
        <summary class="flow2-raw-summary"><span class="flow2-head-main"><em class="flow2-context-badge {context['state_class']}">{escape(context['state_label'])}</em><strong>{raw_event_title}</strong><b>{case_context}</b></span><span>{raw_event_meta}</span></summary>
        <div class="flow2-raw-content">
          <div class="flow2-raw-guide">실행 Trace에 저장된 started·success·failure 로그를 병합하거나 상태를 보정하지 않고 그대로 표시합니다.</div>
          <div class="flow2-trace-list"><div class="flow2-trace-track" data-event-count="{raw_event_count}">{raw_trace_history}</div></div>
        </div>
      </details>
    </div>
    """)


def _render_agent_pipeline_comparison(
    snapshot: dict,
    running: bool,
    preparation: dict | None = None,
):
    _render_agent_pipeline_v2(snapshot, running, preparation)


@st.fragment(run_every="2s")
def _live_testcase_pipeline():
    job_id = st.session_state.get("goal_testcase_job_id")
    started_at = st.session_state.get("goal_testcase_started_at", "")
    if not job_id:
        selected_case = _selected_goal_testcase()
        if selected_case:
            _sync_goal_testcase_recent_artifacts(selected_case)
        _render_agent_pipeline_comparison(
            pipeline_trace_events(
                st.session_state.get("goal_testcase_started_at", ""),
                st.session_state.get("goal_testcase_trace_id", ""),
            ),
            running=False,
            preparation=st.session_state.get("goal_testcase_preparation"),
        )
        return
    job = background_job_snapshot(job_id)
    if not job:
        st.session_state.pop("goal_testcase_job_id", None)
        st.session_state.goal_testcase_result = {
            "case": {"case_id": st.session_state.get("goal_testcase_running_case_id", "-")},
            "mode": "voc",
            "execution": {"result": {"ok": False, "error": "백그라운드 작업 상태를 찾을 수 없습니다."}},
        }
        st.rerun(scope="app")
        return
    preparation = (
        job.get("progress", {}).get("preparation")
        or st.session_state.get("goal_testcase_preparation")
        or _new_manual_preparation_progress()
    )
    st.session_state.goal_testcase_preparation = preparation
    if job.get("done"):
        if job.get("status") == "COMPLETED":
            completed = job.get("result") or {}
            st.session_state.goal_testcase_result = completed["testcase_result"]
            st.session_state.goal_testcase_agent_snapshot = completed["agent_snapshot"]
        else:
            st.session_state.goal_testcase_result = {
                "case": {"case_id": st.session_state.get("goal_testcase_running_case_id", "-")},
                "mode": "voc",
                "execution": {"result": {"ok": False, "error": job.get("error", "백그라운드 실행 실패")}},
            }
        st.session_state.goal_testcase_preparation = preparation
        st.session_state.goal_testcase_completed_at = datetime.now().astimezone().isoformat()
        st.session_state.goal_testcase_focus_result = True
        st.session_state.pop("goal_testcase_job_id", None)
        discard_background_job(job_id)
        _load_goal_monitor_snapshot.clear()
        _load_validity_candidates.clear()
        st.rerun(scope="app")
        return
    trace_id = st.session_state.get("goal_testcase_trace_id", "")
    trace = pipeline_trace_events(started_at, trace_id)
    if trace.get("trace_id") and not trace_id:
        st.session_state.goal_testcase_trace_id = trace["trace_id"]
    _render_agent_pipeline_comparison(
        trace,
        running=job.get("status") == "RUNNING",
        preparation=preparation,
    )


@st.fragment(run_every="2s")
def _live_manual_judge():
    job_id = st.session_state.get("goal_judge_job_id")
    if not job_id:
        return
    job = background_job_snapshot(job_id)
    if not job:
        st.session_state.pop("goal_judge_job_id", None)
        st.session_state.pop("goal_judge_focus_running_once", None)
        st.session_state.goal_judge_error = "백그라운드 독립 LLM 평가 작업 상태를 찾을 수 없습니다."
        st.rerun()
    if job.get("status") == "RUNNING":
        return
    if job.get("status") == "COMPLETED":
        reevaluated = job.get("result") or {}
        testcase_result = deepcopy(st.session_state.get("goal_testcase_result") or {})
        testcase_result["judge_result"] = reevaluated.get("judge_result", {})
        testcase_result["evidence_status"] = next(
            (
                item.get("status", testcase_result.get("evidence_status", "-"))
                for item in reevaluated.get("summary", {}).get("case_results", [])
                if item.get("case_id") == reevaluated.get("case_id")
            ),
            testcase_result.get("evidence_status", "-"),
        )
        st.session_state.goal_testcase_result = testcase_result
        st.session_state.pop("goal_judge_error", None)
    else:
        st.session_state.goal_judge_error = job.get("error", "독립 LLM 평가 실패")
    st.session_state.pop("goal_judge_job_id", None)
    st.session_state.pop("goal_judge_running_case_id", None)
    st.session_state.pop("goal_judge_focus_running_once", None)
    discard_background_job(job_id)
    _load_voc_history_rows.clear()
    _load_validity_candidates.clear()
    st.session_state.goal_judge_result_focus_once = True
    st.rerun()


@st.cache_data(ttl=5, max_entries=1, show_spinner=False)
def _load_goal_monitor_snapshot():
    return agent_status_snapshot(), a2a_trace_snapshot()


@st.cache_data(ttl=5, max_entries=1, show_spinner=False)
def _load_agent_management_snapshot():
    return agent_status_snapshot()


def _title(title, description):
    st.markdown(f"## {title}")
    st.caption(description)


def _show_command_result(*, show_success: bool = True):
    result = st.session_state.get("voc_command_result")
    if not result:
        return
    if result.get("ok"):
        if not show_success:
            st.session_state.pop("voc_command_result", None)
            return
        st.success(f"실행 성공 · {result.get('duration_seconds', 0)}초")
    else:
        st.error(f"실행 실패 · 종료 코드 {result.get('return_code')}")
    st.code(result.get("output", "출력 없음"), language="text")


def _run_and_store(callback, *args):
    with st.spinner("VOC 품질진단 작업을 수행하고 있습니다..."):
        st.session_state.voc_command_result = callback(*args)


def _agent_control_progress_message(action: str, agent_name: str | None = None) -> str:
    action_labels = {
        "start": "시작",
        "restart": "재기동",
        "stop": "중지",
    }
    if action not in action_labels:
        raise ValueError(f"허용되지 않은 Agent 제어 작업: {action}")
    target = f"{agent_name} Agent" if agent_name else "Interpreter 등 6개 Agent"
    return f"{target} 프로세스를 {action_labels[action]}하고 있습니다..."


def _agent_control_action_label(action: str) -> str:
    return {
        "start": "시작",
        "restart": "재시작",
        "stop": "중지",
    }.get(action, action)


def _agent_control_confirmation_key() -> str:
    nonce = int(st.session_state.get("agent_control_confirm_nonce", 0) or 0)
    return f"agent_control_confirmed_{nonce}"


def _reset_agent_control_confirmation() -> None:
    st.session_state["agent_control_confirm_nonce"] = (
        int(st.session_state.get("agent_control_confirm_nonce", 0) or 0) + 1
    )


AGENT_CONTROL_JOB_KEY = "agent_control_job_id"
AGENT_CREDENTIAL_RESULT_KEYS = (
    "agent_openai_credential_result",
    "agent_anthropic_credential_result",
    "agent_gemini_credential_result",
)


def _has_agent_credential_feedback() -> bool:
    return any(st.session_state.get(key) for key in AGENT_CREDENTIAL_RESULT_KEYS)


def _clear_agent_control_messages() -> None:
    for key in (
        "agent_control_feedback",
        "agent_control_latest_snapshot",
        "agent_control_log",
        "voc_command_result",
        "agent_openai_credential_result",
        "agent_anthropic_credential_result",
        "agent_gemini_credential_result",
        AGENT_CONTROL_JOB_KEY,
    ):
        st.session_state.pop(key, None)


def _clear_agent_quick_test_results(agent_key: str | None = None) -> None:
    if agent_key:
        st.session_state.pop(f"agent_quick_test_result_{agent_key}", None)
        return
    for key in list(st.session_state.keys()):
        if str(key).startswith("agent_quick_test_result_"):
            st.session_state.pop(key, None)


def _agent_control_target_label(agent_name: str | None, display_name: str | None = None) -> str:
    if display_name:
        return f"{display_name} Agent"
    if agent_name:
        return f"{agent_name} Agent"
    return "전체 Agent"


def _agent_control_display_name(agent_key: str | None) -> str | None:
    if not agent_key:
        return None
    return AGENT_DISPLAY_NAMES_BY_KEY.get(str(agent_key).lower())


def _agent_management_all_running(snapshot: dict) -> bool:
    agents = snapshot.get("agents", [])
    total = int(snapshot.get("total") or len(agents) or 0)
    running = int(snapshot.get("running") or 0)
    if total <= 0 or running != total:
        return False
    return all(
        str(agent.get("status") or "") == "RUNNING" and bool(agent.get("healthy", True))
        for agent in agents
    )


def _agent_control_output_lines(output: str, *, max_lines: int = 80) -> list[str]:
    lines = [line.strip() for line in str(output or "").splitlines() if line.strip()]
    if len(lines) <= max_lines:
        return lines
    hidden_count = len(lines) - max_lines
    return [*lines[:max_lines], f"... {hidden_count}줄 생략"]


def _agent_control_log_title(action: str, target: str) -> str:
    return f"{target} {_agent_control_action_label(action)} 처리 로그"


def _store_agent_control_log(
    *,
    action: str,
    target: str,
    result: dict,
    snapshot: dict,
    ok: bool,
) -> None:
    command_ok = bool(result.get("ok"))
    action_label = _agent_control_action_label(action)
    lines = [
        f"요청 접수 · {target} {action_label}",
        f"명령 실행 · {'정상 종료' if command_ok else '오류 종료'}",
    ]
    lines.extend(_agent_control_output_lines(result.get("output", "")))
    lines.append(
        f"상태 확인 · 실행 중 {snapshot.get('running', 0)} / {snapshot.get('total', 0)}"
    )
    lines.append("최종 판정 · 완료" if ok else "최종 판정 · 확인 필요")
    st.session_state["agent_control_log"] = {
        "ok": ok,
        "command_ok": command_ok,
        "action": action,
        "target": target,
        "title": _agent_control_log_title(action, target),
        "duration_seconds": result.get("duration_seconds", "-"),
        "checked_at": snapshot.get("checked_at", "-"),
        "lines": lines,
    }


def _find_agent_status(snapshot: dict, agent_key: str | None) -> dict | None:
    if not agent_key:
        return None
    for agent in snapshot.get("agents", []):
        if agent.get("key") == agent_key:
            return agent
    return None


def _agent_control_reached_desired_state(
    action: str,
    snapshot: dict,
    agent_key: str | None = None,
) -> bool:
    agents = snapshot.get("agents", [])
    total = int(snapshot.get("total") or len(agents) or 0)
    running = int(snapshot.get("running") or 0)
    if agent_key:
        agent = _find_agent_status(snapshot, agent_key)
        status = str((agent or {}).get("status") or "")
        if action == "stop":
            return status == "STOPPED"
        return status == "RUNNING"
    if action == "stop":
        return total > 0 and running == 0
    return total > 0 and running == total


def _parse_agent_control_datetime(value: object) -> datetime | None:
    text = str(value or "").strip()
    if not text or text == "-":
        return None
    try:
        return datetime.fromisoformat(text.replace("Z", "+00:00"))
    except ValueError:
        return None


def _agent_control_job_context(job: dict) -> tuple[str, str | None, str | None]:
    progress = job.get("progress") or {}
    target_id = str(job.get("target_id") or "")
    action = str(progress.get("action") or "").strip()
    agent_key: str | None = None
    if ":" in target_id:
        target_action, target_agent = target_id.split(":", 1)
        action = action or target_action
        if target_agent and target_agent != "all":
            agent_key = target_agent
    action = action or "start"
    return action, agent_key, _agent_control_display_name(agent_key)


def _agent_control_restart_started_after_job(
    snapshot: dict,
    job_started_at: object,
    agent_key: str | None = None,
) -> bool:
    job_started = _parse_agent_control_datetime(job_started_at)
    if not job_started:
        return False
    job_started = job_started.astimezone()
    agents = snapshot.get("agents", [])
    if agent_key:
        agents = [agent for agent in agents if agent.get("key") == agent_key]
    if not agents:
        return False
    for agent in agents:
        started_at = _parse_agent_control_datetime(agent.get("started_at"))
        if not started_at or started_at.astimezone() < job_started:
            return False
    return True


def _agent_control_running_job_reached_desired_state(job: dict, snapshot: dict) -> bool:
    action, agent_key, _display_name = _agent_control_job_context(job)
    if not _agent_control_reached_desired_state(action, snapshot, agent_key):
        return False
    if action == "restart":
        return _agent_control_restart_started_after_job(snapshot, job.get("started_at"), agent_key)
    return True


def _agent_control_elapsed_seconds(job: dict) -> float | str:
    started = _parse_agent_control_datetime(job.get("started_at"))
    if not started:
        return "-"
    return round((datetime.now().astimezone() - started.astimezone()).total_seconds(), 2)


def _complete_agent_control_running_job_from_snapshot(job_id: str, job: dict, snapshot: dict) -> None:
    action, agent_name, display_name = _agent_control_job_context(job)
    result = {
        "ok": True,
        "return_code": 0,
        "output": "Agent 상태가 목표 상태에 도달했습니다. 화면 상태를 자동으로 갱신합니다.",
        "duration_seconds": _agent_control_elapsed_seconds(job),
    }
    st.session_state["voc_command_result"] = result
    _load_agent_management_snapshot.clear()
    _load_goal_monitor_snapshot.clear()
    _store_agent_control_feedback(
        action=action,
        agent_name=agent_name,
        display_name=display_name,
        result=result,
        snapshot=snapshot,
    )
    st.session_state.pop(AGENT_CONTROL_JOB_KEY, None)
    discard_background_job(job_id)
    st.rerun(scope="app")


def _run_agent_control_command(action: str, agent_name: str | None = None) -> dict:
    """Agent 제어 명령을 실행합니다.

    전체 시작은 이미 실행 중인 Agent가 있으면 기존 bulk start 스크립트가
    포트 점유로 실패할 수 있어, 부분 RUNNING 상태에서는 중지된 Agent만
    개별 시작합니다.
    """
    if agent_name or action != "start":
        return run_agent_action(action, agent_name)

    before = agent_status_snapshot()
    agents = before.get("agents", [])
    total = int(before.get("total") or len(agents) or 0)
    running = int(before.get("running") or 0)
    if total > 0 and running == total:
        return {
            "ok": True,
            "return_code": 0,
            "output": "이미 모든 Agent가 실행 중입니다.",
            "duration_seconds": 0,
        }
    if running == 0:
        return run_agent_action("start")

    started = time.perf_counter()
    outputs = [f"이미 실행 중인 Agent {running}건은 유지하고, 중지된 Agent만 시작합니다."]
    ok = True
    return_code = 0
    for agent in agents:
        status = str(agent.get("status") or "")
        if status == "RUNNING":
            outputs.append(f"[건너뜀] {agent.get('key')} 이미 실행 중")
            continue
        result = run_agent_action("start", str(agent.get("key")))
        outputs.append(result.get("output") or "출력 없음")
        if not result.get("ok"):
            ok = False
            return_code = int(result.get("return_code") or 1)

    return {
        "ok": ok,
        "return_code": return_code,
        "output": "\n".join(outputs),
        "duration_seconds": round(time.perf_counter() - started, 2),
    }


def _store_agent_control_feedback(
    *,
    action: str,
    agent_name: str | None,
    display_name: str | None,
    result: dict,
    snapshot: dict,
) -> None:
    reached = _agent_control_reached_desired_state(action, snapshot, agent_name)
    ok = bool(result.get("ok")) and reached
    target = _agent_control_target_label(agent_name, display_name)
    action_label = _agent_control_action_label(action)
    if ok:
        title = f"{target} {action_label} 완료"
    elif result.get("ok"):
        title = f"{target} {action_label} 처리 후 상태 확인 필요"
    else:
        title = f"{target} {action_label} 실패"

    st.session_state["agent_control_feedback"] = {
        "ok": ok,
        "command_ok": bool(result.get("ok")),
        "title": title,
        "action": action,
        "target": target,
        "duration_seconds": result.get("duration_seconds", "-"),
        "running": snapshot.get("running", 0),
        "total": snapshot.get("total", 0),
        "checked_at": snapshot.get("checked_at", "-"),
        "error": "" if result.get("ok") else str(result.get("output") or ""),
    }
    _store_agent_control_log(
        action=action,
        target=target,
        result=result,
        snapshot=snapshot,
        ok=ok,
    )
    st.session_state["agent_control_latest_snapshot"] = snapshot


def _render_agent_control_feedback() -> None:
    feedback = st.session_state.get("agent_control_feedback")
    if not feedback:
        return
    icon = ":material/check_circle:" if feedback.get("ok") else ":material/error:"
    badge_color = "green" if feedback.get("ok") else "orange"
    if not feedback.get("command_ok"):
        badge_color = "red"

    with st.container(border=True, height=94):
        cols = st.columns([1.8, 0.9, 0.9, 1.4], vertical_alignment="center")
        with cols[0]:
            st.markdown(f"{icon} **{feedback.get('title', 'Agent 제어 결과')}**")
            st.caption(f"확인 시각 · {feedback.get('checked_at', '-')}")
        with cols[1]:
            st.badge(
                "완료" if feedback.get("ok") else "확인 필요",
                color=badge_color,
                icon=icon,
            )
        with cols[2]:
            st.metric("실행 중", f"{feedback.get('running', 0)} / {feedback.get('total', 0)}")
        with cols[3]:
            st.metric("처리 시간", f"{feedback.get('duration_seconds', '-')}초")


def _render_agent_control_log() -> None:
    log = st.session_state.get("agent_control_log")
    if not log:
        return
    tone = "good" if log.get("ok") else "bad" if not log.get("command_ok") else "warn"
    icon = "check_circle" if log.get("ok") else "error" if not log.get("command_ok") else "warning"
    icon_label = _html_status_chip_label(icon)
    lines = "".join(
        f"<li>{escape(str(line))}</li>"
        for line in log.get("lines", [])
    )
    st.markdown(
        f"""
        <section class="vqa-agent-log {tone}">
            <div class="vqa-agent-log-head">
                <span>{escape(icon_label)}</span>
                <div>
                    <strong>{escape(str(log.get('title') or 'Agent 제어 처리 로그'))}</strong>
                    <small>처리 시간 {escape(str(log.get('duration_seconds', '-')))}초 · 확인 시각 {escape(str(log.get('checked_at', '-')))}</small>
                </div>
            </div>
            <ol>{lines}</ol>
        </section>
        """,
        unsafe_allow_html=True,
    )


def _render_agent_credential_feedback() -> None:
    credential_items = [
        ("OpenAI", st.session_state.get("agent_openai_credential_result")),
        ("Anthropic", st.session_state.get("agent_anthropic_credential_result")),
        ("Gemini", st.session_state.get("agent_gemini_credential_result")),
    ]
    credential_items = [(label, result) for label, result in credential_items if result]
    if not credential_items:
        return
    columns = st.columns(len(credential_items), gap="small")
    for column, (label, credential_result) in zip(columns, credential_items):
        with column.container(border=True, height=112):
            if credential_result.get("ok"):
                st.success(credential_result.get("message", f"{label} 인증 점검 성공"))
            else:
                st.error(credential_result.get("message", f"{label} 인증 점검 실패"))
            meta = [
                f"상태: {_voc_status_label(credential_result.get('status', '-'))}",
                f"설정 파일: {credential_result.get('source', '미확인')}",
            ]
            if credential_result.get("env_name"):
                meta.append(f"환경변수: {credential_result.get('env_name')}")
            if credential_result.get("model"):
                meta.append(f"모델: {credential_result.get('model')}")
            meta.append(f"점검 시각: {credential_result.get('checked_at', '-')}")
            st.caption(" · ".join(meta))


def _render_agent_management_messages(snapshot: dict) -> None:
    feedback = st.session_state.get("agent_control_feedback")
    log = st.session_state.get("agent_control_log")
    current_all_running = _agent_management_all_running(snapshot)
    show_feedback = False
    if feedback:
        action = feedback.get("action")
        if current_all_running and action in {"start", "restart"}:
            show_feedback = False
        else:
            show_feedback = (
                not feedback.get("ok")
                or not current_all_running
                or action == "stop"
            )
    show_log = False
    if log:
        action = log.get("action")
        show_log = bool(log.get("command_ok")) or not current_all_running or action == "stop"

    if (
        not show_feedback
        and not show_log
        and not _has_agent_credential_feedback()
    ):
        return

    st.markdown("#### 최근 처리 상태")
    if show_feedback:
        _render_agent_control_feedback()
    if show_log:
        _render_agent_control_log()
    if _has_agent_credential_feedback():
        _render_agent_credential_feedback()


@st.fragment
def _render_agent_quick_test_fragment(agent: dict) -> None:
    test_result_key = f"agent_quick_test_result_{agent['key']}"
    quick_test_requested = st.button(
        "간편 테스트",
        key=f"quick_test_agent_{agent['key']}",
        icon=":material/network_check:",
        disabled=agent["status"] != "RUNNING",
        width="stretch",
    )

    with st.container(height=148, border=False, key=f"agent_quick_test_result_{agent['key']}"):
        if quick_test_requested:
            with st.spinner(f"{agent['name']} Agent를 실제 호출하고 있습니다..."):
                st.session_state[test_result_key] = test_agent_rpc(
                    agent["name"],
                    int(agent["port"]),
                    timeout=12.0,
                )
        test_result = st.session_state.get(test_result_key)
        if not test_result:
            st.caption("실행 중 상태에서 간편 테스트로 실제 RPC 응답을 확인할 수 있습니다.")
        elif test_result.get("ok"):
            st.success(
                f"{test_result.get('rpc', '-')} 호출 성공 · "
                f"{test_result.get('duration_seconds', '-')}초"
            )
            st.caption(f"IN: {test_result.get('input', '-')}")
            st.caption(f"OUT: {test_result.get('summary', '-')}")
        else:
            st.error(
                f"호출 실패 · {test_result.get('duration_seconds', '-')}초"
            )
            st.caption(test_result.get("summary", "-"))


def _execute_agent_control_background(
    job_id: str,
    action: str,
    agent_name: str | None = None,
    display_name: str | None = None,
) -> dict:
    target = _agent_control_target_label(agent_name, display_name)
    action_label = _agent_control_action_label(action)
    update_background_job(
        job_id,
        progress={
            "action": action,
            "target": target,
            "message": f"{target} {action_label} 명령을 실행하고 있습니다.",
            "lines": [
                f"요청 접수 · {target} {action_label}",
                "Agent 제어 명령 실행 중",
            ],
        },
    )
    result = _run_agent_control_command(action, agent_name)
    update_background_job(
        job_id,
        progress={
            "message": "Agent 상태를 다시 확인하고 있습니다.",
            "lines": [
                f"요청 접수 · {target} {action_label}",
                "Agent 제어 명령 완료",
                "최신 Agent 상태 확인 중",
            ],
        },
    )
    snapshot = agent_status_snapshot()
    return {
        "action": action,
        "agent_name": agent_name,
        "display_name": display_name,
        "result": result,
        "snapshot": snapshot,
    }


def _start_agent_control_background(
    action: str,
    agent_name: str | None = None,
    display_name: str | None = None,
) -> None:
    target = _agent_control_target_label(agent_name, display_name)
    action_label = _agent_control_action_label(action)
    _clear_agent_control_messages()
    _clear_agent_quick_test_results(agent_name)
    _reset_agent_control_confirmation()
    _load_agent_management_snapshot.clear()
    _load_goal_monitor_snapshot.clear()
    job_id = start_background_job(
        "agent-control",
        f"{action}:{agent_name or 'all'}",
        _execute_agent_control_background,
        action,
        agent_name,
        display_name,
        progress={
            "action": action,
            "target": target,
            "message": f"{target} {action_label} 작업을 준비하고 있습니다.",
            "lines": [f"요청 접수 · {target} {action_label}"],
        },
    )
    st.session_state[AGENT_CONTROL_JOB_KEY] = job_id


def _render_agent_control_running_panel(job: dict) -> None:
    progress = job.get("progress") or {}
    target = str(progress.get("target") or "전체 Agent")
    action = str(progress.get("action") or "")
    action_label = _agent_control_action_label(action)
    message = str(progress.get("message") or f"{target} {action_label} 작업을 처리하고 있습니다.")
    started_at = str(job.get("started_at") or "").replace("T", " ")[:19] or "-"
    lines = "".join(
        f"<li>{escape(str(line))}</li>"
        for line in progress.get("lines", [])
    )
    st.markdown(
        f"""
        <section class="vqa-agent-log warn">
            <div class="vqa-agent-log-head">
                <span>진행</span>
                <div>
                    <strong>{escape(target)} {escape(action_label)} 처리 중</strong>
                    <small>시작 시각 {escape(started_at)} · 완료되면 Agent 상태를 자동으로 갱신합니다.</small>
                </div>
            </div>
            <p class="vqa-agent-log-message">{escape(message)}</p>
            <ol>{lines}</ol>
        </section>
        """,
        unsafe_allow_html=True,
    )


@st.fragment(run_every="1s")
def _render_agent_control_job_monitor() -> None:
    job_id = st.session_state.get(AGENT_CONTROL_JOB_KEY)
    if not job_id:
        return
    job = background_job_snapshot(job_id)
    if not job:
        st.session_state.pop(AGENT_CONTROL_JOB_KEY, None)
        return
    if job.get("status") == "RUNNING":
        snapshot = agent_status_snapshot()
        if _agent_control_running_job_reached_desired_state(job, snapshot):
            _complete_agent_control_running_job_from_snapshot(job_id, job, snapshot)
            return
        _render_agent_control_running_panel(job)
        return

    payload = job.get("result") or {}
    action = str(payload.get("action") or (job.get("progress") or {}).get("action") or "start")
    agent_name = payload.get("agent_name")
    display_name = payload.get("display_name")
    if job.get("status") == "COMPLETED":
        result = payload.get("result") or {
            "ok": False,
            "return_code": -1,
            "output": "Agent 제어 작업 결과를 찾을 수 없습니다.",
            "duration_seconds": "-",
        }
        snapshot = payload.get("snapshot") or agent_status_snapshot()
    else:
        result = {
            "ok": False,
            "return_code": -1,
            "output": job.get("error") or "Agent 제어 작업이 실패했습니다.",
            "duration_seconds": "-",
        }
        snapshot = agent_status_snapshot()

    st.session_state["voc_command_result"] = result
    _load_agent_management_snapshot.clear()
    _load_goal_monitor_snapshot.clear()
    _store_agent_control_feedback(
        action=action,
        agent_name=agent_name,
        display_name=display_name,
        result=result,
        snapshot=snapshot,
    )
    st.session_state.pop(AGENT_CONTROL_JOB_KEY, None)
    discard_background_job(job_id)
    st.rerun(scope="app")


def _run_agent_control_and_refresh(
    action: str,
    agent_name: str | None = None,
    display_name: str | None = None,
    *,
    rerun_after: bool = True,
):
    try:
        with st.spinner(_agent_control_progress_message(action, display_name or agent_name)):
            result = _run_agent_control_command(action, agent_name)
    except Exception as exc:
        result = {
            "ok": False,
            "return_code": -1,
            "output": f"{type(exc).__name__}: {exc}",
        }
    st.session_state["voc_command_result"] = result
    _clear_agent_quick_test_results(agent_name)
    _reset_agent_control_confirmation()
    _load_agent_management_snapshot.clear()
    _load_goal_monitor_snapshot.clear()
    snapshot = agent_status_snapshot()
    _store_agent_control_feedback(
        action=action,
        agent_name=agent_name,
        display_name=display_name,
        result=result,
        snapshot=snapshot,
    )
    if rerun_after:
        st.rerun()


@st.cache_data(ttl=5, max_entries=1, show_spinner=False)
def _load_voc_dashboard_snapshot():
    return {
        "runtime": runtime_health(),
        "agents": agent_status_snapshot(),
        "testcases": test_case_summary(),
        "runs": list_voc_run_history(),
        "defects": list_voc_defects(),
        "validity_candidates": list_improvement_validity_candidates(),
        "a2a": a2a_trace_snapshot(),
    }


def _dashboard_timestamp(value: str) -> str:
    if not value:
        return "-"
    try:
        return datetime.fromisoformat(value).strftime("%Y-%m-%d %H:%M")
    except (TypeError, ValueError):
        return str(value)


def _history_table_timestamp(value: str) -> str:
    if not value:
        return "-"
    try:
        return datetime.fromisoformat(value).strftime("%m-%d %H:%M:%S")
    except (TypeError, ValueError):
        return str(value).replace("T", " ")[:19]


def _dashboard_date_range(value, today: date) -> tuple[date, date]:
    if isinstance(value, (tuple, list)) and len(value) == 2:
        return value[0], value[1]
    if isinstance(value, date):
        return value, value
    return today - timedelta(days=6), today


def _dashboard_in_period(value: str, start_date: date, end_date: date) -> bool:
    try:
        observed = datetime.fromisoformat(value).date()
    except (TypeError, ValueError):
        return False
    return start_date <= observed <= end_date


def _dashboard_status_card(icon: str, label: str, value: str, detail: str, tone: str) -> str:
    return (
        f"<article class='vqd-status-card {tone}'>"
        f"<span class='vqd-status-icon'>{_dashboard_svg_icon(icon)}</span>"
        f"<span class='vqd-status-label'>{escape(label)}</span>"
        f"<strong>{escape(value)}</strong>"
        f"<small>{escape(detail)}</small>"
        "</article>"
    )


def _dashboard_a2a_status_panel(a2a: dict) -> str:
    definitions = {
        "PASS": "최근 완전 실행 Trace에서 필수 Agent 연결이 모두 성공하고 실패 이벤트가 없습니다.",
        "FAIL": "최근 실행 Trace에 Agent 간 호출 또는 데이터 전달 실패가 기록됐습니다.",
        "NOT_VERIFIED": "최근 확인 구간에 전체 필수 연결을 통과한 완전 실행 Trace가 없습니다.",
    }
    decision = str(a2a.get("decision") or "NOT_VERIFIED").upper()
    if decision not in definitions:
        decision = "NOT_VERIFIED"
    tone = {"PASS": "pass", "FAIL": "fail", "NOT_VERIFIED": "not-verified"}[decision]
    options = "".join(
        (
            f"<span class='vqd-connection-option {'active ' + tone if status == decision else 'inactive'}' "
            f"aria-current='{'true' if status == decision else 'false'}' "
            f"title='{escape(description, quote=True)}'>{status}</span>"
        )
        for status, description in definitions.items()
    )
    reason = str(a2a.get("reason") or definitions[decision])
    recent_minutes = int(a2a.get("recent_minutes") or 30)
    return (
        f"<section class='vqd-connection-panel {tone}'>"
        f"<span class='vqd-connection-icon'>{_dashboard_svg_icon('trace')}</span>"
        "<div class='vqd-connection-heading'><b>최근 연결 판정</b><small>Agent 파이프라인 Trace</small></div>"
        f"<div class='vqd-connection-options' role='list' aria-label='연결 판정 상태'>{options}</div>"
        f"<p>{escape(reason)} <small>· 최근 {recent_minutes}분 기준</small></p>"
        "</section>"
    )


def _dashboard_agent_cards(agents: dict) -> str:
    cards = []
    for item in agents.get("agents", []):
        name = str(item.get("name") or "Agent")
        role = next((role for agent, role, _port in AGENT_PIPELINE if agent == name), "역할 정보 없음")
        healthy = bool(item.get("healthy"))
        status = "정상" if healthy else str(item.get("status") or "STOPPED")
        cards.append(
            f"<article class='vqd-agent-card {'good' if healthy else 'bad'}'>"
            f"<span class='vqd-agent-icon'>{_dashboard_agent_svg_icon(name)}</span>"
            f"<div><b>{escape(name)}</b><small>{escape(role)}</small></div>"
            f"<span class='vqd-agent-state'>{escape(status)}</span>"
            f"<em>:{int(item.get('port') or 0)} · PID {escape(str(item.get('pid') or '-'))}</em>"
            "</article>"
        )
    return f"<div class='vqd-agent-grid'>{''.join(cards)}</div>"


def _dashboard_agent_svg_icon(name: str) -> str:
    paths = {
        "Interpreter": "<path d='M4 5h16v11H8l-4 4z'/><path d='M8 9h8m-8 3h5'/>",
        "Retriever": "<circle cx='10' cy='10' r='6'/><path d='m15 15 5 5'/>",
        "Summarizer": "<path d='M6 3h9l4 4v14H6z'/><path d='M15 3v5h5M9 12h7m-7 4h7'/>",
        "Evaluator": "<path d='M5 4h14v16H5z'/><path d='m8 10 2 2 5-5m-7 9h8'/>",
        "Critic": "<path d='M12 3 4 6v6c0 5 3.4 8.3 8 10 4.6-1.7 8-5 8-10V6z'/><path d='M12 8v5m0 3h.01'/>",
        "Improver": "<path d='m4 17 5-5 4 3 7-8'/><path d='M15 7h5v5'/>",
    }
    path = paths.get(name, "<circle cx='12' cy='12' r='8'/><path d='M8 12h8'/>")
    return (
        "<svg viewBox='0 0 24 24' aria-hidden='true' fill='none' stroke='currentColor' "
        "stroke-width='1.8' stroke-linecap='round' stroke-linejoin='round'>"
        + path
        + "</svg>"
    )


def _agent_management_card_header(agent: dict) -> str:
    name = str(agent.get("name") or "Agent")
    role = next((role for agent_name, role, _port in AGENT_PIPELINE if agent_name == name), "역할 정보 없음")
    state_class = "good" if agent.get("healthy") else "bad"
    return (
        f"<div class='vqa-agent-head {state_class}'>"
        f"<span class='vqa-agent-icon'>{_dashboard_agent_svg_icon(name)}</span>"
        f"<span><b>{escape(name)}</b><small>{escape(role)}</small></span>"
        "</div>"
    )


def _build_voc_run_status_chart(runs: list[dict]):
    status_order = list(VOC_RUN_STATUS_COLORS)
    rows = [
        {
            "Run": _dashboard_timestamp(item.get("started_at", "")),
            "판정": status,
            "Case 수": int(item.get("counts", {}).get(status, 0)),
        }
        for item in reversed(runs[:12])
        for status in status_order
    ]
    frame = pd.DataFrame(rows)
    return (
        alt.Chart(frame)
        .mark_bar(cornerRadiusTopLeft=3, cornerRadiusTopRight=3)
        .encode(
            x=alt.X("Run:N", title=None, sort=None, axis=alt.Axis(labelAngle=-35)),
            y=alt.Y("sum(Case 수):Q", title="Case 수", scale=alt.Scale(zero=True)),
            color=alt.Color(
                "판정:N",
                title=None,
                scale=alt.Scale(domain=status_order, range=list(VOC_RUN_STATUS_COLORS.values())),
                legend=alt.Legend(
                    orient="bottom",
                    direction="horizontal",
                    columns=len(status_order),
                    gridAlign="all",
                    columnPadding=18,
                    labelLimit=110,
                ),
            ),
            order=alt.Order("판정:N", sort="ascending"),
            tooltip=["Run:N", "판정:N", alt.Tooltip("Case 수:Q", format=",d")],
        )
        .properties(height=270)
    )


def _build_voc_run_history_chart(runs: list[dict]):
    rows = []
    for item in reversed(runs[:12]):
        counts = item.get("counts", {})
        total = max(sum(int(counts.get(status, 0)) for status in VOC_RUN_STATUS_COLORS), 1)
        values = {
            "통과율": int(counts.get("PASS", 0)) / total * 100,
            "검토율": (int(counts.get("REVIEW_REQUIRED", 0)) + int(counts.get("NOT_RUN", 0))) / total * 100,
            "실패·오류율": (int(counts.get("FAIL", 0)) + int(counts.get("ERROR", 0))) / total * 100,
        }
        for metric, value in values.items():
            rows.append({
                "수행 시각": _dashboard_timestamp(item.get("started_at", "")),
                "지표": metric,
                "비율": round(value, 1),
                "Run ID": item.get("run_id", "-"),
            })
    frame = pd.DataFrame(rows)
    return (
        alt.Chart(frame)
        .mark_line(interpolate="monotone", point=alt.OverlayMarkDef(size=65), strokeWidth=3)
        .encode(
            x=alt.X("수행 시각:N", title=None, sort=None, axis=alt.Axis(labelAngle=-35)),
            y=alt.Y("비율:Q", title="비율 (%)", scale=alt.Scale(domain=[0, 100])),
            color=alt.Color(
                "지표:N",
                title=None,
                scale=alt.Scale(domain=list(VOC_HISTORY_COLORS), range=list(VOC_HISTORY_COLORS.values())),
                legend=alt.Legend(orient="bottom"),
            ),
            tooltip=["수행 시각:N", "Run ID:N", "지표:N", alt.Tooltip("비율:Q", format=".1f")],
        )
        .properties(height=245)
    )


def _dashboard_svg_icon(name: str) -> str:
    paths = {
        "runtime": "<path d='M4 5h16v11H4z'/><path d='M8 20h8m-4-4v4'/>",
        "agents": "<circle cx='8' cy='8' r='3'/><circle cx='17' cy='9' r='2.5'/><path d='M3 20c0-4 2-6 5-6s5 2 5 6m1-5c3 0 5 2 5 5'/>",
        "quality": "<circle cx='12' cy='12' r='9'/><path d='m8 12 3 3 6-7'/>",
        "judge": "<path d='M12 3v18M6 6h12M4 9l-2 5h8L8 9m8 0-2 5h8l-2-5'/>",
        "defect": "<path d='M12 3 2.8 20h18.4L12 3Z'/><path d='M12 9v5m0 3h.01'/>",
        "trace": "<circle cx='6' cy='6' r='2'/><circle cx='18' cy='18' r='2'/><path d='M8 6h4a4 4 0 0 1 4 4v6M6 8v8a2 2 0 0 0 2 2h8'/>",
    }
    return (
        "<svg viewBox='0 0 24 24' aria-hidden='true' fill='none' stroke='currentColor' "
        "stroke-width='1.8' stroke-linecap='round' stroke-linejoin='round'>"
        + paths[name]
        + "</svg>"
    )


def _dashboard_validity_action_counts(candidates: list[dict]) -> dict:
    counts = {
        "evaluation": 0,
        "rework": 0,
        "qa": 0,
        "business": 0,
        "approved": 0,
    }
    for candidate in candidates:
        readiness = validity_human_review_readiness(
            validity_status=candidate.get("validity_status", "NOT_RUN"),
            workflow_state=candidate.get("workflow_state", "DRAFT"),
            immediate_hold_count=candidate.get("immediate_hold_count", 0) or 0,
            formal_approval=bool(candidate.get("formal_approval")),
        )
        action = str(readiness.get("action") or "")
        if action == "VALIDITY_EVALUATION_REQUIRED":
            counts["evaluation"] += 1
        elif action == "REWORK_REQUIRED":
            counts["rework"] += 1
        elif action == "QA_REVIEW":
            counts["qa"] += 1
        elif action == "BUSINESS_APPROVAL":
            counts["business"] += 1
        elif action == "FORMAL_APPROVED":
            counts["approved"] += 1
    return counts


def _dashboard_action_summary_cards(
    candidates: list[dict],
    open_defects: list[dict],
    important_defects: list[dict],
) -> list[dict]:
    counts = _dashboard_validity_action_counts(candidates)
    open_defect_count = len(open_defects)
    high_defect_count = len(important_defects)
    return [
        {
            "icon": "fact_check",
            "label": "평가 필요",
            "value": f"{counts['evaluation']}건",
            "detail": "개선안 타당성 평가 미수행",
            "tone": "blue" if counts["evaluation"] else "gray",
            "badge": "평가",
        },
        {
            "icon": "edit_note",
            "label": "보완·재시험 필요",
            "value": f"{counts['rework']}건",
            "detail": "기준 미달 또는 보완 요청",
            "tone": "red" if counts["rework"] else "gray",
            "badge": "보완",
        },
        {
            "icon": "rate_review",
            "label": "QA 검토 대기",
            "value": f"{counts['qa']}건",
            "detail": "AI_PASS · 즉시 보류 없음",
            "tone": "green" if counts["qa"] else "gray",
            "badge": "QA",
        },
        {
            "icon": "verified",
            "label": "업무 승인 대기",
            "value": f"{counts['business']}건",
            "detail": "QA 검토 완료",
            "tone": "green" if counts["business"] else "gray",
            "badge": "승인",
        },
        {
            "icon": "bug_report",
            "label": "미종결 결함",
            "value": f"{open_defect_count}건",
            "detail": f"High/Critical {high_defect_count}건",
            "tone": "red" if high_defect_count else "orange" if open_defect_count else "gray",
            "badge": "결함",
        },
    ]


def _dashboard_action_detail_rows(
    candidates: list[dict],
    open_defects: list[dict],
    run_lookup: dict,
    *,
    limit: int = 12,
) -> pd.DataFrame:
    rows: list[dict] = []
    priority = {
        "REWORK_REQUIRED": 0,
        "QA_REVIEW": 1,
        "BUSINESS_APPROVAL": 2,
        "VALIDITY_EVALUATION_REQUIRED": 3,
    }
    for candidate in candidates:
        readiness = validity_human_review_readiness(
            validity_status=candidate.get("validity_status", "NOT_RUN"),
            workflow_state=candidate.get("workflow_state", "DRAFT"),
            immediate_hold_count=candidate.get("immediate_hold_count", 0) or 0,
            formal_approval=bool(candidate.get("formal_approval")),
        )
        action = str(readiness.get("action") or "")
        if action not in priority:
            continue
        question = str(candidate.get("question") or "-")
        if len(question) > 48:
            question = f"{question[:48]}..."
        rows.append({
            "sort": priority[action],
            "구분": "검증 후보",
            "대상": f"{candidate.get('case_id', '-')} · {question}",
            "다음 조치": readiness.get("action_label") or "확인 필요",
            "상태": (
                f"독립 LLM {_voc_status_label(candidate.get('judge_status', 'NOT_RUN'))} · "
                f"타당성 {_voc_status_label(candidate.get('validity_status', 'NOT_RUN'))}"
            ),
            "수행/등록": _dashboard_timestamp(candidate.get("started_at", "")),
            "Run": candidate.get("run_id", "-"),
        })
    for item in open_defects:
        run_id = (item.get("related_run_ids") or ["-"])[0]
        rows.append({
            "sort": -1 if item.get("severity") in {"HIGH", "CRITICAL"} else 4,
            "구분": "미종결 결함",
            "대상": item.get("title", "-"),
            "다음 조치": "결함 조치·종결 확인",
            "상태": f"{item.get('severity', '-')} · {_defect_status_label(item.get('status', ''))}",
            "수행/등록": _dashboard_timestamp(
                run_lookup.get(run_id, {}).get("started_at", "") or item.get("created_at", "")
            ),
            "Run": run_id,
        })
    if not rows:
        return pd.DataFrame(columns=["구분", "대상", "다음 조치", "상태", "수행/등록", "Run"])
    frame = pd.DataFrame(rows).sort_values(["sort", "수행/등록"], ascending=[True, False])
    return frame.drop(columns=["sort"]).head(limit)


def _render_voc_dashboard_styles() -> None:
    st.markdown(
        """
        <style>
        .vqd-status-row{display:grid;grid-template-columns:repeat(5,minmax(0,1fr));gap:10px;margin:2px 0 10px;font-family:'Malgun Gothic','Apple SD Gothic Neo',sans-serif}
        .vqd-integration-row{grid-template-columns:repeat(4,minmax(0,1fr))}
        .vqd-status-card{height:96px;border:1px solid #c8d9ee;border-top:4px solid #7b8797;border-radius:8px;background:linear-gradient(145deg,#fff,#f8fbff);display:grid;grid-template-columns:38px 1fr;grid-template-rows:auto auto 1fr;column-gap:10px;padding:10px 12px;box-sizing:border-box;box-shadow:0 3px 10px rgba(22,78,128,.05);min-width:0}
        .vqd-status-icon{grid-row:1/4;width:36px;align-self:center;color:#7b8797}.vqd-status-icon svg{width:100%;height:auto}.vqd-status-label{font-size:11px;font-weight:700;color:#40536d}.vqd-status-card strong{font-size:21px;line-height:1.15;color:#073b72;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}.vqd-status-card small{font-size:9px;color:#728095;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;align-self:end}
        .vqd-status-card.good{border-top-color:#299049}.vqd-status-card.good .vqd-status-icon,.vqd-status-card.good strong{color:#299049}
        .vqd-status-card.warn{border-top-color:#b36a08}.vqd-status-card.warn .vqd-status-icon,.vqd-status-card.warn strong{color:#b36a08}
        .vqd-status-card.bad{border-top-color:#d83f36}.vqd-status-card.bad .vqd-status-icon,.vqd-status-card.bad strong{color:#d83f36}
        .vqd-action-card-grid{display:grid;grid-template-columns:repeat(var(--vqd-action-cols,3),minmax(0,1fr));gap:8px;margin:0 0 8px;font-family:'Malgun Gothic','Apple SD Gothic Neo',sans-serif}
        .vqd-action-card{border:1px solid #c8d9ee;border-left:4px solid #7b8797;border-radius:8px;background:linear-gradient(145deg,#fff,#f8fbff);box-shadow:0 3px 10px rgba(22,78,128,.05);padding:9px 10px;box-sizing:border-box;display:grid;grid-template-rows:22px 1fr 17px;gap:2px;overflow:hidden;min-width:0}
        .vqd-action-card-head{display:flex;align-items:center;justify-content:space-between;gap:8px;min-width:0}
        .vqd-action-label{display:flex;align-items:center;gap:5px;min-width:0;color:#40536d;font-size:11px;font-weight:800;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
        .vqd-action-icon{display:flex;width:16px;min-width:16px;color:#155a96;flex:0 0 auto}.vqd-action-icon svg{width:16px;height:16px}
        .vqd-action-badge{flex:0 0 auto;display:inline-flex;align-items:center;justify-content:center;height:20px;padding:0 7px;border-radius:999px;background:#eef2f7;color:#64748b;border:1px solid #d8e2ee;font-size:9px;font-weight:850;white-space:nowrap}
        .vqd-action-card strong{align-self:center;color:#073b72;font-size:21px;line-height:1.12;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
        .vqd-action-card small{color:#728095;font-size:9px;line-height:1.2;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
        .vqd-action-card.blue{border-left-color:#155a96}.vqd-action-card.blue strong,.vqd-action-card.blue .vqd-action-icon{color:#155a96}.vqd-action-card.blue .vqd-action-badge{background:#eaf3fb;color:#155a96;border-color:#b9d2ec}
        .vqd-action-card.green{border-left-color:#299049}.vqd-action-card.green strong,.vqd-action-card.green .vqd-action-icon{color:#299049}.vqd-action-card.green .vqd-action-badge{background:#eaf7ef;color:#176b35;border-color:#a9d7b8}
        .vqd-action-card.orange{border-left-color:#b36a08}.vqd-action-card.orange strong,.vqd-action-card.orange .vqd-action-icon{color:#b36a08}.vqd-action-card.orange .vqd-action-badge{background:#fff7e6;color:#92550a;border-color:#e8c47b}
        .vqd-action-card.red{border-left-color:#d83f36}.vqd-action-card.red strong,.vqd-action-card.red .vqd-action-icon{color:#d83f36}.vqd-action-card.red .vqd-action-badge{background:#fff0ee;color:#b42318;border-color:#efaaa4}
        .vqd-action-card.gray{border-left-color:#9aa5b1}.vqd-action-card.gray strong,.vqd-action-card.gray .vqd-action-icon{color:#718096}.vqd-action-card.gray .vqd-action-badge{background:#f1f3f5;color:#7b8797;border-color:#d8dee5}
        .vqd-connection-panel{display:grid;grid-template-columns:38px 135px auto 1fr;align-items:center;gap:12px;min-height:64px;margin:0 0 12px;padding:9px 12px;border:1px solid #c8d9ee;border-left:4px solid #7b8797;border-radius:8px;background:linear-gradient(90deg,#f8fbff,#fff);box-sizing:border-box;font-family:'Malgun Gothic','Apple SD Gothic Neo',sans-serif;box-shadow:0 3px 10px rgba(22,78,128,.04)}
        .vqd-connection-panel.pass{border-left-color:#299049}.vqd-connection-panel.fail{border-left-color:#d83f36}.vqd-connection-panel.not-verified{border-left-color:#b36a08}
        .vqd-connection-icon{width:34px;color:#7b8797;display:flex}.vqd-connection-panel.pass .vqd-connection-icon{color:#299049}.vqd-connection-panel.fail .vqd-connection-icon{color:#d83f36}.vqd-connection-panel.not-verified .vqd-connection-icon{color:#b36a08}.vqd-connection-icon svg{width:100%;height:auto}
        .vqd-connection-heading b{display:block;font-size:12px;color:#173f68}.vqd-connection-heading small{display:block;font-size:9px;color:#718096;margin-top:2px}
        .vqd-connection-options{display:flex;align-items:center;gap:7px;flex-wrap:wrap}.vqd-connection-option{display:inline-flex;align-items:center;justify-content:center;height:28px;padding:0 10px;border-radius:6px;font-size:10px;font-weight:800;box-sizing:border-box}
        .vqd-connection-option.inactive{color:#9aa5b1;background:#f1f3f5;border:1px solid #d8dee5;filter:grayscale(1)}.vqd-connection-option.active.pass{color:#176b35;background:#eaf7ef;border:1px solid #8dcba2}.vqd-connection-option.active.fail{color:#b42318;background:#fff0ee;border:1px solid #efaaa4}.vqd-connection-option.active.not-verified{color:#92550a;background:#fff7e6;border:1px solid #e8c47b}
        .vqd-connection-panel p{margin:0;font-size:11px;line-height:1.45;color:#40536d}.vqd-connection-panel p small{color:#718096;white-space:nowrap}
        .vqd-agent-grid{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:8px}.vqd-agent-card{display:grid;grid-template-columns:34px 1fr auto;grid-template-rows:auto auto;gap:2px 9px;align-items:center;min-height:72px;padding:9px 10px;border:1px solid #d4e1ef;border-left:4px solid #7b8797;border-radius:7px;background:linear-gradient(145deg,#fff,#f8fbff);box-sizing:border-box}.vqd-agent-card.good{border-left-color:#299049}.vqd-agent-card.bad{border-left-color:#d83f36}.vqd-agent-icon{grid-row:1/3;width:31px;color:#7b8797;display:flex}.vqd-agent-card.good .vqd-agent-icon{color:#299049}.vqd-agent-card.bad .vqd-agent-icon{color:#d83f36}.vqd-agent-icon svg{width:100%;height:auto}.vqd-agent-card b{display:block;font-size:11px;color:#173f68}.vqd-agent-card small{display:block;font-size:9px;color:#718096}.vqd-agent-state{font-size:10px;font-weight:800;color:#7b8797}.vqd-agent-card.good .vqd-agent-state{color:#299049}.vqd-agent-card.bad .vqd-agent-state{color:#d83f36}.vqd-agent-card em{grid-column:2/4;font-size:9px;font-style:normal;color:#8795a8}
        div[data-testid="stForm"]{margin-bottom:0!important}div[data-testid="stForm"] [data-testid="stWidgetLabel"] p{font-size:10px!important;color:#40536d!important}div[data-testid="stForm"] [data-testid="stVerticalBlock"]{gap:.15rem!important}
        div[data-testid="stHeadingWithActionElements"] h1{font-size:29px!important;color:#0c3768!important;letter-spacing:-1px!important}div[data-testid="stHeadingWithActionElements"] h3{font-size:17px!important;color:#173f68!important}
        @media(max-width:1100px){.vqd-status-row,.vqd-integration-row{grid-template-columns:repeat(2,1fr)}.vqd-action-card-grid{grid-template-columns:repeat(2,minmax(0,1fr))}.vqd-connection-panel{grid-template-columns:38px 130px 1fr}.vqd-connection-panel p{grid-column:2/4}}
        @media(max-width:720px){.vqd-status-row,.vqd-integration-row,.vqd-action-card-grid{grid-template-columns:repeat(2,1fr)}.vqd-connection-panel{grid-template-columns:34px 1fr}.vqd-connection-options,.vqd-connection-panel p{grid-column:1/3}.vqd-connection-panel p small{white-space:normal}.vqd-agent-grid{grid-template-columns:1fr}}
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_dashboard():
    _render_voc_dashboard_styles()
    today = date.today()
    if not st.session_state.get("voc_dashboard_header_rendered"):
        with st.container(border=True, key="voc_dashboard_inline_filter_panel"):
            with st.form("voc_dashboard_inline_filters", border=False):
                filter_columns = st.columns([2.2, 0.9, 0.95], vertical_alignment="bottom")
                with filter_columns[0]:
                    inline_range = st.date_input(
                        "기간",
                        value=st.session_state.get(
                            "voc_dashboard_filter_range",
                            (today - timedelta(days=6), today),
                        ),
                        max_value=today,
                        key="voc_dashboard_filter_range",
                    )
                with filter_columns[1]:
                    inline_submitted = st.form_submit_button(
                        "조회",
                        icon=":material/search:",
                        type="primary",
                        width="stretch",
                    )
                with filter_columns[2]:
                    inline_refresh_requested = st.form_submit_button(
                        "새로고침",
                        icon=":material/refresh:",
                        width="stretch",
                    )
            st.session_state["voc_dashboard_filter_submitted"] = bool(inline_submitted)
            st.session_state["voc_dashboard_filter_refresh_requested"] = bool(inline_refresh_requested)
    selected_range = st.session_state.get(
        "voc_dashboard_filter_range",
        (today - timedelta(days=6), today),
    )
    submitted = bool(st.session_state.get("voc_dashboard_filter_submitted", False))
    refresh_requested = bool(st.session_state.get("voc_dashboard_filter_refresh_requested", False))

    if submitted or refresh_requested:
        _load_voc_dashboard_snapshot.clear()
        load_integration_status.clear()

    start_date, end_date = _dashboard_date_range(selected_range, today)
    if start_date > end_date:
        start_date, end_date = end_date, start_date

    snapshot = _load_voc_dashboard_snapshot()
    runtime = snapshot["runtime"]
    agents = snapshot["agents"]
    runs = [
        item for item in snapshot["runs"]
        if _dashboard_in_period(item.get("started_at", ""), start_date, end_date)
    ]
    defects = [
        item for item in snapshot["defects"]
        if _dashboard_in_period(item.get("created_at", ""), start_date, end_date)
    ]
    validity_candidates = [
        item for item in snapshot.get("validity_candidates", [])
        if _dashboard_in_period(item.get("started_at", ""), start_date, end_date)
    ]
    a2a = snapshot["a2a"]
    latest = runs[0] if runs else {}
    latest_counts = latest.get("counts", {})
    latest_judge_counts = latest.get("judge_counts", {})
    open_defects = [item for item in defects if item.get("status") != "CLOSED"]
    important_defects = [
        item for item in open_defects if item.get("severity") in {"HIGH", "CRITICAL"}
    ]

    runtime_ok = runtime.get("ok") and runtime.get("env_configured")
    quality_failures = int(latest_counts.get("FAIL", 0)) + int(latest_counts.get("ERROR", 0))
    quality_reviews = int(latest_counts.get("REVIEW_REQUIRED", 0)) + int(latest_counts.get("NOT_RUN", 0))
    judge_failures = int(latest_judge_counts.get("FAIL", 0)) + int(latest_judge_counts.get("ERROR", 0))
    judge_reviews = int(latest_judge_counts.get("REVIEW_REQUIRED", 0)) + int(latest_judge_counts.get("NOT_RUN", 0))
    action_counts = _dashboard_validity_action_counts(validity_candidates)
    action_total = (
        action_counts["evaluation"]
        + action_counts["rework"]
        + action_counts["qa"]
        + action_counts["business"]
        + len(open_defects)
    )
    cards = [
        _dashboard_status_card(
            "runtime", "실행 환경", "정상" if runtime_ok else "확인 필요",
            "필수 파일 · 환경 설정", "good" if runtime_ok else "bad",
        ),
        _dashboard_status_card(
            "agents", "Agent 가동", f"{agents.get('running', 0)} / {agents.get('total', 6)}",
            "전체 정상" if agents.get("all_running") else "중지 Agent 있음",
            "good" if agents.get("all_running") else "bad" if not agents.get("running") else "warn",
        ),
        _dashboard_status_card(
            "quality", "최신 Run 품질", "이력 없음" if not latest else f"통과 {latest_counts.get('PASS', 0)}",
            f"검토 {quality_reviews} · 실패/오류 {quality_failures}" if latest else "선택 기간 기준",
            "neutral" if not latest else "bad" if quality_failures else "warn" if quality_reviews else "good",
        ),
        _dashboard_status_card(
            "judge", "독립 LLM 평가", "미사용" if not latest or not latest.get("judge_enabled") else f"통과 {latest_judge_counts.get('PASS', 0)}",
            f"검토 {judge_reviews} · 실패/오류 {judge_failures}" if latest.get("judge_enabled") else "최신 Run 기준",
            "neutral" if not latest or not latest.get("judge_enabled") else "bad" if judge_failures else "warn" if judge_reviews else "good",
        ),
        _dashboard_status_card(
            "defect", "조치 필요", f"{action_total}건",
            (
                f"평가 {action_counts['evaluation']} · 보완 {action_counts['rework']} · "
                f"승인대기 {action_counts['qa'] + action_counts['business']} · 결함 {len(open_defects)}"
            ),
            "bad" if important_defects or action_counts["rework"] else "warn" if action_total else "good",
        ),
    ]
    st.markdown(f"<div class='vqd-status-row'>{''.join(cards)}</div>", unsafe_allow_html=True)
    st.markdown(_dashboard_a2a_status_panel(a2a), unsafe_allow_html=True)
    render_integration_status(load_integration_status(), context="voc")

    overview_columns = st.columns(2, gap="medium")
    with overview_columns[0].container(border=True, height=VOC_OVERVIEW_PANEL_HEIGHT):
        chart_heading = st.columns([1.05, 0.95], vertical_alignment="center")
        with chart_heading[0]:
            st.markdown("#### 기간 Run 판정 추이")
        with chart_heading[1]:
            if runs:
                st.caption(
                    f"{start_date.isoformat()} ~ {end_date.isoformat()} · Run {len(runs)}건",
                    text_alignment="right",
                )
        if runs:
            st.altair_chart(_build_voc_run_status_chart(runs), theme=None)
        else:
            st.info("선택 기간에 저장된 Run이 없습니다.")

    with overview_columns[1].container(border=True, height=VOC_OVERVIEW_PANEL_HEIGHT):
        st.markdown("#### Agent 운영 상태")
        if not agents.get("agents"):
            st.info("Agent 상태를 조회할 수 없습니다.")
        else:
            st.markdown(_dashboard_agent_cards(agents), unsafe_allow_html=True)

    detail_columns = st.columns(2, gap="medium")
    with detail_columns[0].container(border=True):
        history_heading = st.columns([1.05, 0.95], vertical_alignment="center")
        with history_heading[0]:
            st.markdown("#### 기간 수행 이력")
        with history_heading[1]:
            if runs:
                st.caption(
                    "Run별 통과·검토·실패/오류 비율 · 최근 12건",
                    text_alignment="right",
                )
        if not runs:
            st.info("선택 기간에 수행 이력이 없습니다.")
        else:
            st.altair_chart(_build_voc_run_history_chart(runs), theme=None)

    with detail_columns[1].container(border=True):
        _render_voc_section_heading(
            "조치 필요 현황",
            "",
            icon="pending_actions",
            right_caption=f"{start_date.isoformat()} ~ {end_date.isoformat()}",
        )
        _render_voc_summary_cards(
            _dashboard_action_summary_cards(validity_candidates, open_defects, important_defects),
            columns=3,
            height=98,
            header_badge=True,
        )
        run_lookup = {item.get("run_id"): item for item in snapshot["runs"]}
        action_rows = _dashboard_action_detail_rows(validity_candidates, open_defects, run_lookup)
        if action_rows.empty:
            st.success("선택 기간에 조치가 필요한 항목이 없습니다.")
        else:
            with st.expander(
                f"조치 대상 상세 · {len(action_rows)}건",
                expanded=False,
                icon=":material/list_alt:",
            ):
                st.dataframe(
                    action_rows,
                    hide_index=True,
                    width="stretch",
                    column_config={
                        "구분": st.column_config.TextColumn(width=92),
                        "대상": st.column_config.TextColumn(width=260),
                        "다음 조치": st.column_config.TextColumn(width=142),
                        "상태": st.column_config.TextColumn(width=176),
                        "수행/등록": st.column_config.TextColumn(width=118),
                        "Run": st.column_config.TextColumn(width=160),
                    },
                )

    st.caption(
        "기간 필터는 Run 시작일과 결함 생성일에 적용되며, 실행 환경·Agent 파이프라인은 현재 상태를 표시합니다. "
        "최종 품질 승인은 품질 보고서와 최종 인수·시연의 게이트 판정을 함께 확인하세요."
    )


def render_agents():
    health = runtime_health()
    if not health["env_configured"]:
        st.warning("Agent 실행에 필요한 `.env`가 없습니다. 환경 파일을 초기화한 뒤 API 키를 입력하세요.")

    st.markdown(
        """
        <style>
        .vqa-agent-head{display:flex;align-items:center;gap:7px;min-height:42px;margin:-2px 0 5px}
        .vqa-agent-icon{display:flex;flex:0 0 29px;width:29px;color:#7b8797}
        .vqa-agent-head.good .vqa-agent-icon{color:#299049}
        .vqa-agent-head.bad .vqa-agent-icon{color:#d83f36}
        .vqa-agent-icon svg{width:100%;height:auto}
        .vqa-agent-head b{display:block;color:#173f68;font-size:11px;line-height:1.2}
        .vqa-agent-head small{display:block;margin-top:2px;color:#718096;font-size:8px;line-height:1.15}
        div[class*="st-key-stop_agent_"] button{
            background:#d83f36!important;border-color:#d83f36!important;color:#fff!important;
        }
        div[class*="st-key-start_agent_"] button{
            background:#155a96!important;border-color:#155a96!important;color:#fff!important;
        }
        div[class*="st-key-cleanup_agent_"] button{
            background:#b36a08!important;border-color:#b36a08!important;color:#fff!important;
        }
        div[class*="st-key-stop_agent_"] button p,
        div[class*="st-key-start_agent_"] button p,
        div[class*="st-key-cleanup_agent_"] button p{
            color:#fff!important;font-weight:800!important;white-space:nowrap!important;
        }
        div[class*="st-key-agent_control_header_"] [data-testid="column"]:nth-child(2) button{
            min-height:30px!important;padding:4px 8px!important;font-size:12px!important;
        }
        div[class*="st-key-agent_quick_test_result_"]{
            overflow:hidden!important;
        }
        .vqa-agent-log{
            margin:8px 0 2px;padding:12px 14px;border:1px solid #c8d9ee;border-left:4px solid #155a96;
            border-radius:10px;background:linear-gradient(135deg,#f8fbff,#ffffff);
            font-family:'Segoe UI','Malgun Gothic',sans-serif;box-shadow:0 4px 12px rgba(22,78,128,.05);
        }
        .vqa-agent-log.good{border-left-color:#299049}
        .vqa-agent-log.warn{border-left-color:#b36a08}
        .vqa-agent-log.bad{border-left-color:#d83f36}
        .vqa-agent-log-head{display:flex;align-items:center;gap:10px;margin-bottom:8px}
        .vqa-agent-log-head span{
            flex:0 0 38px;height:24px;border-radius:999px;display:flex;align-items:center;justify-content:center;
            font-size:11px;font-weight:900;color:#155a96;background:#e7f1fb;border:1px solid #bfd7ef;
        }
        .vqa-agent-log.good .vqa-agent-log-head span{color:#1f7f43;background:#edf8f0;border-color:#bfdfca}
        .vqa-agent-log.warn .vqa-agent-log-head span{color:#9b5c07;background:#fff2dc;border-color:#e6c383}
        .vqa-agent-log.bad .vqa-agent-log-head span{color:#bd3029;background:#fff0ee;border-color:#efbbb6}
        .vqa-agent-log-head strong{display:block;color:#173f68;font-size:13px}
        .vqa-agent-log-head small{display:block;margin-top:2px;color:#718096;font-size:10px}
        .vqa-agent-log ol{
            margin:0;padding:8px 8px 8px 27px;max-height:190px;overflow:auto;border-radius:8px;background:#f3f7fc;
            color:#30465f;font-size:11px;line-height:1.55;
        }
        .vqa-agent-log-message{margin:0 0 8px;color:#40536d;font-size:12px;line-height:1.35}
        .vqa-agent-log li{padding:1px 0}
        </style>
        """,
        unsafe_allow_html=True,
    )
    with st.container(border=True):
        confirmed = st.checkbox(
            "Agent 프로세스 상태 변경",
            key=_agent_control_confirmation_key(),
        )
        agent_control_job_active = bool(st.session_state.get(AGENT_CONTROL_JOB_KEY))
        with st.container(horizontal=True):
            if st.button(
                "전체 시작",
                disabled=not confirmed or agent_control_job_active,
                width="stretch",
                icon=":material/play_arrow:",
            ):
                _start_agent_control_background("start")
                st.rerun()
            if st.button(
                "전체 재시작",
                disabled=not confirmed or agent_control_job_active,
                width="stretch",
                icon=":material/restart_alt:",
            ):
                _start_agent_control_background("restart")
                st.rerun()
            if st.button(
                "전체 중지",
                disabled=not confirmed or agent_control_job_active,
                width="stretch",
                icon=":material/stop:",
            ):
                _start_agent_control_background("stop")
                st.rerun()
            if st.button(
                "OpenAI 인증 점검",
                width="stretch",
                icon=":material/key:",
                key="check_agent_openai_credential",
            ):
                with st.spinner("Agent가 사용할 OpenAI 자격 증명을 점검하고 있습니다..."):
                    st.session_state["agent_openai_credential_result"] = (
                        check_openai_agent_credential()
                    )
            if st.button(
                "Anthropic 인증 점검",
                width="stretch",
                icon=":material/vpn_key:",
                key="check_agent_anthropic_credential",
            ):
                with st.spinner("Agent가 사용할 Anthropic 자격 증명을 실제 호출로 점검하고 있습니다..."):
                    st.session_state["agent_anthropic_credential_result"] = (
                        check_anthropic_agent_credential()
                    )
            if st.button(
                "Gemini 인증 점검",
                width="stretch",
                icon=":material/travel_explore:",
                key="check_agent_gemini_credential",
            ):
                with st.spinner("Agent가 사용할 Gemini 자격 증명을 실제 호출로 점검하고 있습니다..."):
                    st.session_state["agent_gemini_credential_result"] = (
                        check_gemini_agent_credential()
                    )

    snapshot = st.session_state.pop("agent_control_latest_snapshot", None)
    if not snapshot:
        snapshot = _load_agent_management_snapshot()
    stop_impacts = {
        "interpreter": "질문 의도와 검색 조건을 해석할 수 없어 VOC 파이프라인을 시작할 수 없습니다.",
        "retriever": "관련 VOC 근거를 검색할 수 없어 요약과 개선안 생성을 진행할 수 없습니다.",
        "summarizer": "요약 후보 생성과 전체 Agent 조정이 중단되어 최종 응답을 만들 수 없습니다.",
        "evaluator": "요약 후보를 비교·선정할 수 없어 최종 요약을 결정할 수 없습니다.",
        "critic": "요약과 개선안의 품질 검토가 누락되어 보완된 결과를 확정할 수 없습니다.",
        "improver": "정책 개선안을 생성·보완할 수 없어 최종 개선안 산출이 실패합니다.",
    }
    agent_columns = st.columns(6, gap="small")
    for index, agent in enumerate(snapshot["agents"]):
        with agent_columns[index].container(border=True, height=360):
            with st.container(
                horizontal=True,
                horizontal_alignment="distribute",
                vertical_alignment="center",
                key=f"agent_control_header_{agent['key']}",
            ):
                st.markdown(_agent_management_card_header(agent), unsafe_allow_html=True)
                if agent["status"] == "RUNNING":
                    if st.button(
                        "중지",
                        key=f"stop_agent_{agent['key']}",
                        icon=":material/stop_circle:",
                        width="content",
                    ):
                        _confirm_agent_action(agent, "stop")
                elif agent["status"] in {"STOPPED", "UNKNOWN"}:
                    if st.button(
                        "시작",
                        key=f"start_agent_{agent['key']}",
                        icon=":material/play_circle:",
                        width="content",
                    ):
                        _confirm_agent_action(agent, "start")
                elif agent["status"] == "STARTING/FAILED":
                    if st.button(
                        "정리",
                        key=f"cleanup_agent_{agent['key']}",
                        icon=":material/cleaning_services:",
                        width="content",
                    ):
                        _confirm_agent_action(agent, "stop")
                else:
                    st.button(
                        "제어 불가",
                        key=f"unmanaged_agent_{agent['key']}",
                        disabled=True,
                        width="content",
                    )
            _status_badge(
                agent["status"],
                "PASS" if agent["healthy"] else "FAIL",
                f"TCP {agent['port']} · PID {agent['pid']}",
            )
            st.caption(f"포트 {agent['port']} · PID {agent['pid']}")
            started_at = str(agent.get("started_at") or "").strip()
            if agent["status"] == "RUNNING" and started_at:
                try:
                    started_at = datetime.fromisoformat(started_at).astimezone().strftime("%Y-%m-%d %H:%M:%S")
                except ValueError:
                    pass
                st.caption(f"기동 시간 · {started_at}")
            elif agent["status"] == "RUNNING":
                st.caption("기동 시간 · 확인 불가")
            else:
                st.caption("기동 시간 · -")
            if agent["status"] == "STOPPED":
                st.markdown(f":red-badge[중지 영향] {stop_impacts[agent['key']]}")

            _render_agent_quick_test_fragment(agent)

    if st.session_state.get(AGENT_CONTROL_JOB_KEY):
        st.markdown("#### 최근 처리 상태")
        _render_agent_control_job_monitor()
    else:
        _render_agent_management_messages(snapshot)


def _status_badge(label: str, decision: str, help_text: str = ""):
    settings = {
        "PASS": ("green", ":material/check_circle:"),
        "FAIL": ("red", ":material/error:"),
        "NOT_VERIFIED": ("orange", ":material/pending:"),
    }
    color, icon = settings.get(decision, ("gray", ":material/help:"))
    st.badge(label, color=color, icon=icon, help=help_text or None)


@st.dialog("Agent 상태 변경 확인")
def _confirm_agent_action(agent: dict, action: str):
    action_label = "시작" if action == "start" else "중지"
    st.markdown(f"**{agent['name']} Agent를 {action_label}하시겠습니까?**")
    st.caption(f"대상: {agent['name']} · TCP {agent['port']} · 현재 상태 {agent['status']}")
    if action == "stop":
        st.warning(
            "이 Agent를 중지하면 해당 Agent를 사용하는 VOC 테스트가 실패할 수 있습니다. "
            "장애 시연 후 다시 시작하세요.",
            icon=":material/warning:",
        )
    else:
        st.info("시작하려면 런타임 `.env`에 필요한 API 키가 설정돼 있어야 합니다.", icon=":material/info:")

    with st.container(horizontal=True, horizontal_alignment="right"):
        if st.button("취소", key=f"cancel_{action}_{agent['key']}"):
            st.rerun()
        if st.button(
            f"{action_label} 실행",
            key=f"confirm_{action}_{agent['key']}",
            type="primary",
        ):
            _run_agent_control_and_refresh(action, agent["key"], agent["name"], rerun_after=True)


def _set_goal_testcase_selection(selected_case_id: str):
    previous_case_id = st.session_state.get("goal_testcase_selected_case_id")
    if previous_case_id == selected_case_id:
        return
    st.session_state["goal_testcase_selected_case_id"] = selected_case_id
    st.session_state.pop("goal_testcase_result", None)
    st.session_state.pop("goal_testcase_focus_result", None)
    if not st.session_state.get("goal_testcase_job_id"):
        st.session_state.pop("goal_testcase_started_at", None)
        st.session_state.pop("goal_testcase_completed_at", None)
        st.session_state.pop("goal_testcase_agent_snapshot", None)


def _table_selected_row_index(
    table_state: dict | None,
    row_count: int,
) -> int | None:
    table_state = table_state or {}
    selection = table_state.get("selection", {}) if hasattr(table_state, "get") else {}
    selected_rows = selection.get("rows", []) if hasattr(selection, "get") else []
    selected_cells = selection.get("cells", []) if hasattr(selection, "get") else []

    selected_row = None
    if selected_cells and isinstance(selected_cells[-1], (list, tuple)) and selected_cells[-1]:
        selected_row = selected_cells[-1][0]
    elif selected_rows:
        selected_row = selected_rows[0]

    if not isinstance(selected_row, int) or not 0 <= selected_row < row_count:
        return None
    return selected_row


def _promote_table_cell_to_row_selection(
    table_key: str,
    row_count: int,
) -> int | None:
    table_state = st.session_state.get(table_key, {})
    selected_row = _table_selected_row_index(table_state, row_count)
    if selected_row is None:
        return None
    selection = table_state.get("selection", {}) if hasattr(table_state, "get") else {}
    selected_cells = selection.get("cells", []) if hasattr(selection, "get") else []
    if selected_cells:
        st.session_state[table_key] = {
            "selection": {"rows": [selected_row], "columns": [], "cells": []}
        }
    return selected_row


def _remember_goal_testcase_selection(table_key: str, page_case_ids: list[str]):
    selected_row = _promote_table_cell_to_row_selection(
        table_key,
        len(page_case_ids),
    )
    if selected_row is not None:
        selected_case_id = page_case_ids[selected_row]
        previous_case_id = st.session_state.get("goal_testcase_selected_case_id")
        _set_goal_testcase_selection(selected_case_id)
        if previous_case_id and previous_case_id != selected_case_id:
            st.session_state["goal_testcase_selection_changed"] = True


def _remember_catalog_case_selection(table_key: str, case_ids: list[str]):
    selected_row = _promote_table_cell_to_row_selection(
        table_key,
        len(case_ids),
    )
    if selected_row is not None:
        st.session_state["voc_testcase_selected_case_id"] = case_ids[selected_row]


SUPPRESSED_PIPELINE_MESSAGES = {
    "Pipeline completed via agent-to-agent calls",
}


def _selected_result_matches_case(selected_case_id: str) -> bool:
    test_execution = st.session_state.get("goal_testcase_result") or {}
    return test_execution.get("case", {}).get("case_id") == selected_case_id


def _pipeline_result_message(result: dict) -> str:
    message = str(result.get("message") or "").strip()
    if not message or message in SUPPRESSED_PIPELINE_MESSAGES:
        return "VOC 테스트 실행 완료"
    return message


def _render_manual_pipeline_result_summary(result: dict, pipeline_snapshot: dict) -> None:
    ok = bool(result.get("ok"))
    generation = _manual_pipeline_generation_info(pipeline_snapshot)
    status_tone = "good" if ok else "warn"
    status_label = "Agent 파이프라인 완료" if ok else "Agent 파이프라인 확인 필요"
    detail = (
        "VOC 요약과 정책 개선안이 아래에 표시됩니다."
        if ok
        else str(result.get("message") or result.get("error") or "실행 결과 확인 필요")
    )
    provider_label = str(generation.get("provider") or "내부 Agent 파이프라인")
    model_label = str(generation.get("model") or "모델 확인 필요")
    source_label = str(generation.get("source") or "설정 기준")
    st.html(
        f"""
        <div class="vqa-pipeline-compact-summary {status_tone}">
            <div class="vqa-pipeline-compact-main">
                <span>{escape(status_label)}</span>
                <strong>{escape(detail)}</strong>
            </div>
            <div class="vqa-pipeline-compact-meta">
                <small>파이프라인 LLM</small>
                <b>{escape(provider_label)}</b>
                <em>{escape(model_label)} · {escape(source_label)}</em>
            </div>
        </div>
        <style>
        .vqa-pipeline-compact-summary{{
            display:flex;align-items:center;justify-content:space-between;gap:12px;
            min-height:58px;margin:4px 0 12px;padding:10px 12px;border:1px solid #d5e2ef;
            border-left:4px solid #1f6fb2;border-radius:12px;background:#f8fbff;
            font-family:'Segoe UI','Malgun Gothic',sans-serif;
        }}
        .vqa-pipeline-compact-summary.good{{
            border-left-color:#2f9660;background:linear-gradient(135deg,#f2fbf5,#ffffff);
        }}
        .vqa-pipeline-compact-summary.warn{{
            border-left-color:#b7791f;background:linear-gradient(135deg,#fff8e8,#ffffff);
        }}
        .vqa-pipeline-compact-main{{min-width:0;display:flex;flex-direction:column;gap:3px}}
        .vqa-pipeline-compact-main span{{
            width:max-content;max-width:100%;padding:4px 8px;border-radius:999px;
            color:#17643b;background:#e0f3e7;border:1px solid #b9dfc1;
            font-size:10px;font-weight:900;line-height:1.1;
        }}
        .vqa-pipeline-compact-summary.warn .vqa-pipeline-compact-main span{{
            color:#8a540c;background:#fff0c9;border-color:#ead6a8;
        }}
        .vqa-pipeline-compact-main strong{{
            color:#24435f;font-size:12px;line-height:1.3;font-weight:700;
            white-space:nowrap;overflow:hidden;text-overflow:ellipsis;
        }}
        .vqa-pipeline-compact-meta{{
            flex:0 0 36%;min-width:230px;text-align:right;color:#63758b;
        }}
        .vqa-pipeline-compact-meta small{{
            display:block;color:#7a889a;font-size:9px;font-weight:800;
        }}
        .vqa-pipeline-compact-meta b{{
            display:block;color:#174f85;font-size:13px;line-height:1.2;
            white-space:nowrap;overflow:hidden;text-overflow:ellipsis;
        }}
        .vqa-pipeline-compact-meta em{{
            display:block;margin-top:2px;color:#758498;font-style:normal;font-size:10px;
            white-space:nowrap;overflow:hidden;text-overflow:ellipsis;
        }}
        @media(max-width:900px){{
            .vqa-pipeline-compact-summary{{align-items:flex-start;flex-direction:column}}
            .vqa-pipeline-compact-meta{{flex:auto;min-width:0;text-align:left;width:100%}}
        }}
        </style>
        """,
    )


def _render_voc_question_highlight(question: str, *, case_id: str = "-") -> None:
    st.markdown(
        f"""
        <div class="vqa-voc-question-card">
            <div class="vqa-voc-question-label">VOC · {escape(str(case_id or "-"))}</div>
            <div class="vqa-voc-question-text">{escape(str(question or "-"))}</div>
        </div>
        <style>
        .vqa-voc-question-card{{
            margin:2px 0 10px;padding:13px 15px;border:1px solid #b9dfc1;
            border-left:5px solid #4fa96b;border-radius:12px;background:#eefaf0;
            box-shadow:0 4px 12px rgba(47,150,96,.08);
            font-family:'Segoe UI','Malgun Gothic',sans-serif;
        }}
        .vqa-voc-question-label{{
            color:#2f7b48;font-size:11px;font-weight:800;letter-spacing:.02em;margin-bottom:4px;
        }}
        .vqa-voc-question-text{{
            color:#123f26;font-size:15px;font-weight:700;line-height:1.45;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _split_policy_terms_section(policy: str) -> tuple[str, str]:
    text = str(policy or "").strip()
    if not text:
        return "", ""
    pattern = re.compile(
        r"(?ms)(^|\n)\s*(?:#{1,6}\s*)?0\.\s*용어\s*정의\s*[:：]?\s*\n?(.*?)(?=\n\s*(?:#{1,6}\s*)?1\.\s|\Z)"
    )
    match = pattern.search(text)
    if not match:
        return "", text
    terms = match.group(2).strip() or "용어 정의 내용이 없습니다."
    body = (text[: match.start()] + "\n" + text[match.end() :]).strip()
    return terms, body


def _render_policy_improvement(policy: str) -> None:
    terms, body = _split_policy_terms_section(policy)
    if terms:
        with st.expander("0. 용어 정의", icon=":material/dictionary:"):
            st.markdown(terms)
        if body:
            st.markdown(body)
        return
    st.markdown(body or "-")


def _render_manual_judge_running_panel(selected_case_id: str) -> None:
    running_case_id = st.session_state.get("goal_judge_running_case_id") or selected_case_id
    st.markdown(
        f"""
        <div class="vqa-judge-running">
            <div class="vqa-judge-running-head">
                <span>진행</span>
                <div>
                    <strong>독립 LLM 평가 진행 중</strong>
                    <small>Case {escape(str(running_case_id))} · 선택한 Provider가 저장된 Agent 파이프라인 개선안을 평가하고 있습니다.</small>
                </div>
            </div>
            <div class="vqa-judge-running-bar"><span></span></div>
            <p>이 화면을 벗어나도 평가는 백그라운드에서 계속 진행되며, 완료 후 다시 들어오면 결과가 반영됩니다.</p>
        </div>
        <style>
        .vqa-judge-running{{
            margin-top:10px;padding:13px 14px;border:1px solid #c9d8e8;border-radius:12px;
            background:linear-gradient(135deg,#f4f9ff,#ffffff);color:#234b73;
            font-family:'Segoe UI','Malgun Gothic',sans-serif;
        }}
        .vqa-judge-running-head{{display:flex;align-items:center;gap:10px;margin-bottom:10px}}
        .vqa-judge-running-head span{{
            flex:0 0 38px;height:26px;border-radius:999px;display:flex;align-items:center;justify-content:center;
            color:#1f6fb2;background:#e7f1fb;border:1px solid #bfd7ef;font-size:11px;font-weight:900;
        }}
        .vqa-judge-running-head strong{{display:block;font-size:13px}}
        .vqa-judge-running-head small{{display:block;margin-top:2px;color:#687a90;font-size:10px}}
        .vqa-judge-running-bar{{
            position:relative;overflow:hidden;height:12px;border-radius:999px;background:#dce9f5;
        }}
        .vqa-judge-running-bar span{{
            position:absolute;inset:0 auto 0 0;width:42%;border-radius:999px;
            background:linear-gradient(90deg,#7fb3e5,#1f6fb2,#7fb3e5);
            animation:vqaJudgeRunning 1.15s ease-in-out infinite;
        }}
        .vqa-judge-running p{{margin:8px 0 0;color:#5f7084;font-size:10px}}
        @keyframes vqaJudgeRunning{{
            0%{{left:-42%}} 100%{{left:100%}}
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _render_manual_judge_provider_comparison(judge_result: dict) -> None:
    rows = _manual_judge_comparison_rows(judge_result)
    if not rows:
        return

    with st.container(border=True):
        st.markdown("##### Provider별 독립 LLM 평가 비교")
        st.caption(
            "같은 Agent 파이프라인 개선안을 OpenAI, Anthropic, Google Gemini 등 서로 다른 Provider로 재평가하면 이 표에 누적됩니다."
        )
        st.dataframe(
            pd.DataFrame(rows),
            hide_index=True,
            width="stretch",
            column_config={
                "순서": st.column_config.NumberColumn(width=52),
                "평가 Provider": st.column_config.TextColumn(width=108),
                "모델": st.column_config.TextColumn(width=132),
                "판정": st.column_config.TextColumn(width=78),
                "점수": st.column_config.ProgressColumn(
                    "점수",
                    min_value=0,
                    max_value=100,
                    format="%d점",
                    width=104,
                ),
                "독립성": st.column_config.TextColumn(width=156),
                "수행 시간": st.column_config.TextColumn(width=78),
                "평가 시각": st.column_config.TextColumn(width=142),
            },
        )
        if len(rows) == 1:
            st.caption("다른 Provider를 선택해 독립 LLM 평가를 다시 실행하면 비교 행이 추가됩니다.")


def _render_goal_judge_result(selected_case_id: str):
    test_execution = st.session_state.get("goal_testcase_result")
    executed_case_id = (test_execution or {}).get("case", {}).get("case_id")
    if not test_execution or executed_case_id != selected_case_id:
        return
    judge_result = test_execution.get("judge_result", {})
    if not judge_result or judge_result.get("decision") == "NOT_RUN":
        return

    _render_goal_judge_result_focus_anchor_once()
    with st.container(border=True, key="goal_manual_judge_result"):
        st.markdown("#### 독립 LLM 평가 결과")
        st.caption(
            f"{judge_result.get('provider', '-')} · {judge_result.get('model', '-')} · "
            f"Rubric {judge_result.get('rubric_version', '-')}"
        )
        metrics = st.columns(4)
        metrics[0].metric("판정", _voc_status_label(judge_result.get("decision", "NOT_RUN")))
        metrics[1].metric(
            "총점",
            f"{judge_result['total_score']}점" if judge_result.get("total_score") is not None else "-",
        )
        metrics[2].metric("독립성", _judge_independence_grade_label(judge_result.get("independence_grade", "-")))
        metrics[3].metric(
            "수행 시간",
            f"{float(judge_result.get('duration_seconds') or 0):g}초",
        )
        _render_manual_judge_provider_comparison(judge_result)

        if judge_result.get("error"):
            st.error(judge_result["error"])
        elif judge_result.get("decision") == "NOT_RUN":
            st.warning(judge_result.get("message", "독립 LLM 평가가 실행되지 않았습니다."))
        elif judge_result.get("independence_hold"):
            st.warning(
                f"점수 기준 판정은 {_voc_status_label(judge_result.get('rubric_decision'))}이지만 "
                f"{judge_result.get('independence_hold_reason')}"
            )

        dimension_rows = []
        for dimension, detail in judge_result.get("dimension_scores", {}).items():
            if isinstance(detail, dict):
                dimension_rows.append(
                    {
                        "평가 차원": dimension,
                        "점수": detail.get("score", detail.get("points", "-")),
                        "판정 근거": detail.get("reason", "-"),
                    }
                )
            else:
                dimension_rows.append({"평가 차원": dimension, "점수": detail, "판정 근거": "-"})
        if dimension_rows:
            st.dataframe(
                pd.DataFrame(dimension_rows),
                hide_index=True,
                width="stretch",
                column_order=("평가 차원", "점수", "판정 근거"),
                column_config={
                    "평가 차원": st.column_config.TextColumn(width=92),
                    "점수": st.column_config.TextColumn(width=58),
                    "판정 근거": st.column_config.TextColumn(width="large"),
                },
            )

        details = st.columns(3, gap="medium")
        for column, title, values in (
            (details[0], "확인 근거", judge_result.get("evidence", [])),
            (details[1], "잔여 위험", judge_result.get("risks", [])),
            (details[2], "보완 권고", judge_result.get("recommendations", [])),
        ):
            with column.container(border=True, height=145):
                st.markdown(f"**{title}**")
                if values:
                    for value in values:
                        st.write(f"- {value}")
                else:
                    st.caption("표시할 내용 없음")


def _manual_followup_candidate(run_id: str, case_id: str) -> dict:
    if not run_id or not case_id:
        return {}
    try:
        return next(
            (
                candidate
                for candidate in _load_validity_candidates()
                if candidate.get("run_id") == run_id and candidate.get("case_id") == case_id
            ),
            {},
        )
    except Exception:
        return {}


def _manual_followup_action_target(run_id: str, case_id: str, action: dict) -> dict:
    action_code = str(action.get("code") or "")
    button_labels = {
        "RUN_VALIDITY": "타당성 평가로 이동",
        "REWORK_AND_RETEST": "보완·재평가로 이동",
        "QA_REVIEW": "QA 검토로 이동",
        "BUSINESS_APPROVAL": "업무 승인으로 이동",
        "CHECK_REMAINING_CASES": "잔여 Case 검토로 이동",
        "REPORT_READY": "품질 보고서로 이동",
    }
    if action_code in {"CHECK_PIPELINE_ERROR", "REVIEW_PIPELINE_RESULT"}:
        return {
            "enabled": True,
            "page": "history_detail",
            "run_id": run_id,
            "case_id": case_id,
            "action_code": action_code,
            "button_label": "Run 증적 확인",
            "detail": (
                f"{case_id}의 독립 LLM 평가가 검토 필요 상태입니다. "
                "바로 타당성 검증으로 넘기지 않고 수행 이력 상세에서 판정 근거와 Agent 결과를 먼저 확인합니다."
            ),
        }
    if action_code in VOC_HISTORY_NAVIGABLE_VALIDITY_ACTIONS:
        return {
            "enabled": True,
            "page": VOC_VALIDITY_PAGE_NAME,
            "run_id": run_id,
            "case_id": case_id,
            "action_code": action_code,
            "button_label": button_labels.get(action_code, "다음 액션으로 이동"),
            "detail": f"{case_id}의 {action.get('label', '다음 액션')} 흐름을 이어서 진행합니다.",
        }
    if action_code == "REPORT_READY":
        return {
            "enabled": True,
            "page": VOC_REPORT_PAGE_NAME,
            "run_id": run_id,
            "case_id": case_id,
            "action_code": "REPORT_READY",
            "button_label": button_labels["REPORT_READY"],
            "detail": "업무 승인 완료 증적을 품질 보고서와 최종 인수·시연 대상으로 연결합니다.",
        }
    if action_code == "RUN_JUDGE":
        return {
            "enabled": False,
            "button_label": "위에서 독립 LLM 평가 실행",
            "detail": "Agent 파이프라인 산출물을 먼저 독립 LLM으로 평가하면 타당성 검증으로 이어갈 수 있습니다.",
        }
    return {
        "enabled": False,
        "button_label": "이동 액션 없음",
        "detail": action.get("detail") or "현재 상태에서 바로 이동할 후속 화면이 없습니다.",
    }


def _manual_followup_flow_model(
    test_execution: dict | None,
    selected_case_id: str,
    *,
    candidate: dict | None = None,
) -> dict:
    test_execution = test_execution or {}
    executed_case = test_execution.get("case") or {}
    case_id = str(executed_case.get("case_id") or selected_case_id or "")
    run_id = str(test_execution.get("run_id") or "")
    if not run_id or not case_id:
        return {"visible": False}

    execution = test_execution.get("execution") if isinstance(test_execution.get("execution"), dict) else {}
    result = execution.get("result") if isinstance(execution.get("result"), dict) else {}
    pipeline_ok = bool(result.get("ok") or execution.get("ok"))
    if test_execution.get("mode") == "voc":
        pipeline_ok = pipeline_ok and bool(result.get("ok"))
    evidence_status = str(test_execution.get("evidence_status") or ("PASS" if pipeline_ok else "ERROR"))

    judge_result = test_execution.get("judge_result") if isinstance(test_execution.get("judge_result"), dict) else {}
    judge_status = str(judge_result.get("decision") or "NOT_RUN")
    if candidate is None and judge_status not in {"", "NOT_RUN"}:
        candidate = _manual_followup_candidate(run_id, case_id)
    candidate = candidate or {}

    validity_status = str(candidate.get("validity_status") or "NOT_RUN")
    workflow_state = str(candidate.get("workflow_state") or "DRAFT")
    immediate_hold_count = int(candidate.get("immediate_hold_count") or 0)
    formal_approval = bool(candidate.get("formal_approval"))
    case_state = {
        "case_id": case_id,
        "status": evidence_status,
        "message": result.get("message") or result.get("error") or "",
        "judge_status": str(candidate.get("judge_status") or judge_status),
        "validity_status": validity_status,
        "approval_state": workflow_state,
        "immediate_hold_count": immediate_hold_count,
        "formal_approval": formal_approval,
    }
    action = voc_case_next_action(case_state)
    readiness = validity_human_review_readiness(
        validity_status=validity_status,
        workflow_state=workflow_state,
        immediate_hold_count=immediate_hold_count,
        formal_approval=formal_approval,
    )
    target = _manual_followup_action_target(run_id, case_id, action)

    judge_score = candidate.get("judge_score", judge_result.get("total_score"))
    validity_score = candidate.get("validity_score")
    cards = [
        {
            "icon": "account_tree",
            "label": "Agent 파이프라인",
            "value": "완료" if pipeline_ok else "확인 필요",
            "detail": f"증적 상태: {_voc_status_label(evidence_status)}",
            "tone": "green" if pipeline_ok and evidence_status == "PASS" else ("orange" if pipeline_ok else "red"),
        },
        {
            "icon": "rule",
            "label": "독립 LLM 평가",
            "value": _voc_status_label(case_state["judge_status"]),
            "detail": f"{judge_score}점" if judge_score is not None else "평가 후 타당성 검증 가능",
            "tone": "green" if case_state["judge_status"] == "PASS" else ("gray" if case_state["judge_status"] == "NOT_RUN" else "orange"),
        },
        {
            "icon": "fact_check",
            "label": "개선안 타당성 평가",
            "value": _voc_status_label(validity_status),
            "detail": f"{validity_score}점" if validity_score is not None else "타당성 평가 전",
            "tone": "green" if validity_status == "AI_PASS" else ("gray" if validity_status == "NOT_RUN" else "orange"),
        },
        {
            "icon": "verified",
            "label": "QA·업무 승인",
            "value": readiness["action_label"],
            "detail": f"보류 규칙 {immediate_hold_count}건 · {_voc_status_label(workflow_state)}",
            "tone": "green" if readiness["action"] in {"QA_REVIEW", "BUSINESS_APPROVAL", "FORMAL_APPROVED"} else "gray",
        },
    ]
    return {
        "visible": True,
        "run_id": run_id,
        "case_id": case_id,
        "question": executed_case.get("question") or executed_case.get("name") or "-",
        "action": action,
        "action_code": action["code"],
        "target": target,
        "cards": cards,
    }


def _render_manual_demo_flow_overview(selected_case: dict) -> None:
    """Render the compact demo flow for the manual TC page."""
    selected_case_id = str(selected_case.get("case_id") or "-")
    selected_question = (
        selected_case.get("question")
        or selected_case.get("name")
        or "-"
    )
    test_execution = st.session_state.get("goal_testcase_result") or {}
    executed_case = test_execution.get("case", {}) if isinstance(test_execution.get("case"), dict) else {}
    if str(executed_case.get("case_id") or selected_case_id) != selected_case_id:
        test_execution = {}

    execution = test_execution.get("execution", {}) if isinstance(test_execution.get("execution"), dict) else {}
    result = execution.get("result", {}) if isinstance(execution.get("result"), dict) else {}
    pipeline_ok = bool(execution.get("ok") or result.get("ok"))
    if test_execution.get("mode") == "voc":
        pipeline_ok = pipeline_ok and bool(result.get("ok"))
    pipeline_has_result = bool(test_execution)
    pipeline_running = bool(st.session_state.get("goal_testcase_job_id"))
    judge_running = bool(st.session_state.get("goal_judge_job_id"))
    is_fault_case = (
        selected_case.get("execution_type") in {"fault_proxy", "isolated_fault"}
        or selected_case.get("category") == "fault_condition"
    )

    judge_result = test_execution.get("judge_result", {}) if isinstance(test_execution.get("judge_result"), dict) else {}
    judge_decision = str(judge_result.get("decision") or "NOT_RUN")
    followup_model = _manual_followup_flow_model(test_execution, selected_case_id) if pipeline_has_result else {"visible": False}
    action = followup_model.get("action", {}) if followup_model.get("visible") else {}
    target = followup_model.get("target", {}) if followup_model.get("visible") else {}

    case_implemented = selected_case.get("implementation_status") == "IMPLEMENTED"
    if pipeline_running:
        pipeline_value, pipeline_tone, pipeline_detail = "진행 중", "running", "Agent 호출 로그 갱신 중"
        a2a_value, a2a_tone, a2a_detail = "대기", "waiting", "완료 후 산출물 표시"
    elif pipeline_ok:
        pipeline_value, pipeline_tone, pipeline_detail = "완료", "done", "Agent 실행 Trace 저장"
        a2a_value, a2a_tone, a2a_detail = (
            ("격리 시험 결과", "done", "격리 장애 Case 결과 확인")
            if is_fault_case
            else ("결과 생성", "done", "VOC 요약·정책 개선안 확인")
        )
    elif pipeline_has_result:
        pipeline_value, pipeline_tone, pipeline_detail = "확인 필요", "warn", "오류 또는 보류 사유 확인"
        a2a_value, a2a_tone, a2a_detail = "확인 필요", "warn", "결과 영역에서 사유 확인"
    else:
        pipeline_value, pipeline_tone, pipeline_detail = "대기", "waiting", "선택 카드의 실행 버튼으로 시작"
        a2a_value, a2a_tone, a2a_detail = "대기", "waiting", "Agent 파이프라인 완료 후 표시"

    if judge_running:
        judge_value, judge_tone, judge_detail = "진행 중", "running", "선택 Provider가 평가 중"
    elif judge_decision and judge_decision != "NOT_RUN":
        judge_value = _voc_status_label(judge_decision)
        judge_tone = "done" if judge_decision == "PASS" else "warn"
        judge_detail = f"{judge_result.get('total_score', '-')}점 · 결과 아래 표시"
    elif pipeline_ok and not is_fault_case:
        judge_value, judge_tone, judge_detail = "실행 가능", "action", "외부 Provider 교차 평가"
    elif is_fault_case:
        judge_value, judge_tone, judge_detail = "대상 아님", "waiting", "개선안이 없는 격리 시험"
    else:
        judge_value, judge_tone, judge_detail = "대기", "waiting", "A2A 결과 생성 후 실행"

    if action:
        next_value = str(action.get("label") or target.get("button_label") or "다음 액션 확인")
        next_detail = str(target.get("detail") or action.get("detail") or "다음 업무 단계로 이동")
        next_tone = "action" if target.get("enabled") else "waiting"
    elif pipeline_running:
        next_value, next_tone, next_detail = "실행 완료 대기", "waiting", "완료 후 자동 갱신"
    elif pipeline_has_result:
        next_value, next_tone, next_detail = "독립 LLM 평가", "action", "아래 Provider 선택 후 실행"
    else:
        next_value, next_tone, next_detail = "Agent 파이프라인 실행", "action", "선택 Case로 시연 시작"

    stages = [
        {
            "no": "1",
            "label": "Test Case 선택",
            "value": selected_case_id,
            "detail": _manual_pipeline_compact_text(str(selected_question), 54),
            "tone": "done" if case_implemented else "waiting",
        },
        {
            "no": "2",
            "label": "Agent 파이프라인",
            "value": pipeline_value,
            "detail": pipeline_detail,
            "tone": pipeline_tone,
        },
        {
            "no": "3",
            "label": "A2A 결과",
            "value": a2a_value,
            "detail": a2a_detail,
            "tone": a2a_tone,
        },
        {
            "no": "4",
            "label": "독립 LLM 평가",
            "value": judge_value,
            "detail": judge_detail,
            "tone": judge_tone,
        },
        {
            "no": "5",
            "label": "다음 액션",
            "value": _manual_pipeline_compact_text(next_value, 28),
            "detail": _manual_pipeline_compact_text(next_detail, 58),
            "tone": next_tone,
        },
    ]
    stage_html = "".join(
        f"""
        <article class="vqa-manual-flow-card {escape(stage['tone'])}">
            <span>{escape(stage['no'])}</span>
            <small>{escape(stage['label'])}</small>
            <strong>{escape(stage['value'])}</strong>
            <p>{escape(stage['detail'])}</p>
        </article>
        """
        for stage in stages
    )
    st.html(
        f"""
        <div class="vqa-manual-flow-wrap">
            <div class="vqa-manual-flow-head">
                <div>
                    <strong>시연 진행 흐름</strong>
                    <small>선택 Case 기준으로 지금 봐야 할 단계만 표시합니다.</small>
                </div>
                <em>{escape(selected_case_id)} 기준</em>
            </div>
            <div class="vqa-manual-flow-grid">{stage_html}</div>
        </div>
        <style>
        .vqa-manual-flow-wrap{{
            margin:6px 0 14px;padding:14px;border:1px solid #d5e2ef;border-radius:16px;
            background:linear-gradient(180deg,#f8fbff,#ffffff);
            font-family:'Segoe UI','Malgun Gothic',sans-serif;
        }}
        .vqa-manual-flow-head{{
            display:flex;align-items:center;justify-content:space-between;gap:12px;margin-bottom:10px;
        }}
        .vqa-manual-flow-head strong{{display:block;color:#173f68;font-size:15px}}
        .vqa-manual-flow-head small{{display:block;margin-top:2px;color:#718299;font-size:10px}}
        .vqa-manual-flow-head em{{
            font-style:normal;border-radius:999px;background:#e7f1fb;color:#1f6fb2;
            padding:5px 9px;font-size:10px;font-weight:800;white-space:nowrap;
        }}
        .vqa-manual-flow-grid{{
            display:grid;grid-template-columns:repeat(5,minmax(0,1fr));gap:9px;
        }}
        .vqa-manual-flow-card{{
            position:relative;height:112px;box-sizing:border-box;padding:11px 12px 10px;
            border:1px solid #d8e1ec;border-radius:13px;background:#f8fafc;color:#516274;
            overflow:hidden;
        }}
        .vqa-manual-flow-card:not(:last-child):after{{
            content:'→';position:absolute;right:-8px;top:42%;color:#86a5c5;font-weight:900;font-size:17px;
        }}
        .vqa-manual-flow-card span{{
            display:grid;place-items:center;width:25px;height:25px;border-radius:9px;
            background:#e7edf3;color:#627285;font-size:11px;font-weight:900;margin-bottom:7px;
        }}
        .vqa-manual-flow-card small{{display:block;color:#728095;font-size:10px;font-weight:800}}
        .vqa-manual-flow-card strong{{
            display:block;margin-top:3px;color:#263f5d;font-size:15px;line-height:1.2;
            white-space:nowrap;overflow:hidden;text-overflow:ellipsis;
        }}
        .vqa-manual-flow-card p{{
            margin:5px 0 0;color:#6d7c90;font-size:10px;line-height:1.28;
            display:-webkit-box;-webkit-line-clamp:2;-webkit-box-orient:vertical;overflow:hidden;
        }}
        .vqa-manual-flow-card.done{{background:#eefaf2;border-color:#b9dfc1;color:#17643b}}
        .vqa-manual-flow-card.done span{{background:#d9f1e3;color:#247147}}
        .vqa-manual-flow-card.done strong{{color:#17643b}}
        .vqa-manual-flow-card.running{{background:#edf6ff;border-color:#a9cceb;color:#175b9d}}
        .vqa-manual-flow-card.running span{{background:#d9ebff;color:#175b9d}}
        .vqa-manual-flow-card.running strong{{color:#174f85}}
        .vqa-manual-flow-card.warn{{background:#fff8e8;border-color:#ead6a8;color:#8a540c}}
        .vqa-manual-flow-card.warn span{{background:#fff0c9;color:#8a540c}}
        .vqa-manual-flow-card.warn strong{{color:#8a540c}}
        .vqa-manual-flow-card.action{{background:#f4f9ff;border-color:#b9d2ed;color:#174f85}}
        .vqa-manual-flow-card.action span{{background:#e0efff;color:#1f6fb2}}
        .vqa-manual-flow-card.action strong{{color:#174f85}}
        @media(max-width:1100px){{
            .vqa-manual-flow-grid{{grid-template-columns:repeat(2,minmax(0,1fr));}}
            .vqa-manual-flow-card:not(:last-child):after{{display:none}}
        }}
        </style>
        """,
    )


def _render_goal_followup_focus_anchor_once():
    _render_goal_scroll_anchor(
        "goal-manual-followup-scroll-anchor",
        scroll_flag_key="goal_manual_followup_focus_once",
        block="center",
    )


def _render_manual_followup_flow(selected_case_id: str):
    model = _manual_followup_flow_model(st.session_state.get("goal_testcase_result"), selected_case_id)
    if not model.get("visible"):
        return

    action = model["action"]
    target = model["target"]
    badge_tone = {
        "green": "green",
        "blue": "blue",
        "orange": "orange",
        "red": "red",
        "gray": "gray",
    }.get(str(action.get("tone") or "gray"), "gray")
    _render_goal_followup_focus_anchor_once()
    with st.container(border=True):
        heading, state = st.columns([2.8, 1], gap="small", vertical_alignment="center")
        with heading:
            st.markdown(f"#### 다음 액션 흐름 · {model['case_id']}")
            st.caption(
                f"현재 선택한 Test Case의 최근 저장 Run 기준입니다. "
                f"Run {model['run_id']} · {model['question']}"
            )
        with state:
            st.markdown(f":{badge_tone}-badge[{action.get('label', '다음 액션 확인')}]", text_alignment="right")

        columns = st.columns(4, gap="small")
        for column, card in zip(columns, model["cards"], strict=False):
            tone = {
                "green": "green",
                "blue": "blue",
                "orange": "orange",
                "red": "red",
                "gray": "gray",
            }.get(card["tone"], "gray")
            with column.container(border=True, height=118):
                st.caption(f":material/{card['icon']}: {card['label']}")
                st.markdown(f":{tone}-badge[{card['value']}]")
                st.caption(card["detail"])

        action_col, detail_col = st.columns([0.9, 2.7], gap="medium", vertical_alignment="center")
        with action_col:
            if st.button(
                target["button_label"],
                type="primary" if target.get("enabled") else "secondary",
                icon=":material/arrow_forward:",
                disabled=not target.get("enabled"),
                width="stretch",
                key=f"goal_manual_next_action_{model['run_id']}_{model['case_id']}",
            ):
                _apply_history_next_action_target(target)
                st.rerun()
        with detail_col:
            st.caption(target.get("detail", action.get("detail", "선택 대상의 다음 업무 단계로 이동합니다.")))


def _render_goal_testcase_result(selected_case_id: str):
    test_execution = st.session_state.get("goal_testcase_result")
    executed_case_id = (test_execution or {}).get("case", {}).get("case_id")
    if not test_execution or executed_case_id != selected_case_id:
        return

    executed_case = test_execution.get("case", {})
    selected_case = _selected_goal_testcase() or {}
    result_title = (
        executed_case.get("question")
        or selected_case.get("question")
        or executed_case.get("name")
        or selected_case.get("name")
        or executed_case.get("case_id")
        or "-"
    )
    _render_voc_section_heading(
        "A2A 수행 결과",
        "Agent 파이프라인이 생성한 VOC 요약과 정책 개선안입니다. 상세 Trace는 접어서 확인합니다.",
        icon="summarize",
        badges=((executed_case.get("case_id") or selected_case_id, "blue"),),
    )
    _render_voc_question_highlight(result_title, case_id=executed_case.get("case_id") or selected_case_id)
    if test_execution.get("run_id"):
        st.caption(
            f"Run ID: {test_execution['run_id']} · 증적 상태: "
            f"{test_execution.get('evidence_status', '-')} · 저장 위치: "
            f"reports/voc_quality_runs/{test_execution['run_id']}"
        )
    if test_execution.get("mode") == "fault":
        execution = test_execution.get("execution", {})
        if execution.get("ok"):
            st.success(f"격리 장애 시험 {test_execution.get('fault_id')} 통과")
        else:
            st.error(f"격리 장애 시험 {test_execution.get('fault_id')} 실패")
        st.code(execution.get("output", "출력 없음"), language="text")
        return

    payload = test_execution.get("execution", {})
    result = payload.get("result", {})
    if not result.get("ok"):
        st.warning(result.get("message") or result.get("error") or "VOC 테스트 결과가 없습니다.")
    result_columns = st.columns(2, gap="medium")
    with result_columns[0].container(border=True):
        st.markdown("**1-1. VOC 요약**")
        st.write(result.get("summary", "-") or "-")
    with result_columns[1].container(border=True):
        st.markdown("**1-2. 정책 개선안**")
        _render_policy_improvement(result.get("policy", "-") or "-")
    with st.expander("1-3. 판정 근거 및 Agent 실행 Trace", icon=":material/account_tree:"):
        intent = _parse_json_mapping(result.get("intent_json"))
        evaluator = _parse_json_mapping(result.get("eval_json"))
        critic = _parse_json_mapping(result.get("summary_critic_json"))
        trace_summary = _parse_pipeline_trace_summary(result.get("trace"))
        trace = test_execution.get("trace") if isinstance(test_execution.get("trace"), dict) else {}
        trace_rows = _pipeline_trace_event_rows(trace)
        trace_values = trace_summary["values"]
        trace_flags = set(trace_summary["flags"])
        trace_id = trace.get("trace_id") or trace_values.get("audit_trace_id") or "-"
        failed_rows = [row for row in trace_rows if row["결과"] == "실패"]
        successful_rows = [row for row in trace_rows if row["결과"] == "성공"]

        st.info(
            "이 정보는 질문 해석, VOC 검색, 후보 평가, Critic 보완 및 Agent 연결 상태를 보여주는 "
            "판정 근거입니다. 다만 이 실행 Trace만으로 최종 통과를 확정하지 않으며, 독립 LLM 평가·품질 규칙·사람 검토와 함께 판단합니다.",
            icon=":material/fact_check:",
        )
        if not result.get("ok"):
            st.error(
                "판단: 파이프라인이 완료되지 않아 결과 품질을 평가할 수 없습니다. 아래 실패 단계와 오류를 장애 원인 근거로 확인하세요."
            )
        elif failed_rows:
            st.warning(
                f"판단: 결과는 생성됐지만 Agent 연결 실패 {len(failed_rows)}건이 있어 품질 판정을 보류하고 추가 검토해야 합니다."
            )
        elif trace_rows:
            st.success(
                "판단: 질문 해석부터 결과 생성까지의 실행 이력이 확인됩니다. 아래 평가점수와 Critic 보완 내용을 최종 결과의 타당성 검토에 활용할 수 있습니다."
            )
        else:
            st.warning(
                "판단: 요약 실행 Trace만 있어 처리 흐름은 확인할 수 있지만 Agent별 성공·실패 증적은 제한적입니다. 최종 판정에는 저장된 Run 증적을 함께 확인하세요."
            )

        task_labels = {"summary": "VOC 요약", "policy": "정책 개선안", "both": "VOC 요약 + 정책 개선안"}
        winner = trace_values.get("winner")
        numeric_scores = {
            str(key): float(value)
            for key, value in evaluator.items()
            if isinstance(value, (int, float))
        }
        if not winner and numeric_scores:
            winner = max(numeric_scores, key=numeric_scores.get)
        retrieved = trace_values.get("retrieved", "-")
        critic_status = (
            "보완 반영"
            if critic.get("need_refine") and "summary_refined" in trace_flags
            else "보완 요청"
            if critic.get("need_refine")
            else "원안 유지"
            if critic
            else "평가 없음"
        )
        with st.container(horizontal=True):
            st.metric("실행 결과", "완료" if result.get("ok") else "실패", border=True)
            st.metric("검색 VOC", f"{retrieved}건" if str(retrieved).isdigit() else retrieved, border=True)
            st.metric(
                "선택 후보",
                f"{winner} · {numeric_scores[winner]:.1f}점" if winner in numeric_scores else winner or "-",
                border=True,
            )
            st.metric("Critic 검토", critic_status, border=True)

        st.markdown("##### 세부 1 · 질문 해석과 검색 범위")
        filters = [str(value) for value in intent.get("filters", []) if str(value).strip()]
        st.write(f"- 수행 목적: **{task_labels.get(intent.get('task'), intent.get('task') or '확인 불가')}**")
        st.write(f"- 검색 키워드: **{', '.join(filters) if filters else '확인 불가'}**")
        st.write(f"- 최대 검색 범위: **{intent.get('max_items', '-')}건**")
        st.caption("질문 의도와 검색 키워드는 최종 답변이 사용자의 VOC와 관련 있는지 판단하는 근거입니다.")

        st.markdown("##### 세부 2 · 요약 후보 평가와 선택 근거")
        if numeric_scores:
            score_rows = pd.DataFrame([
                {
                    "후보": candidate,
                    "평가점수": score,
                    "선택 여부": "최종 선택" if candidate == winner else "비선택",
                }
                for candidate, score in sorted(numeric_scores.items(), key=lambda item: item[1], reverse=True)
            ])
            st.table(score_rows)
            st.caption("Evaluator 점수는 후보 간 상대 비교 근거이며, 점수 자체가 최종 품질 PASS를 의미하지는 않습니다.")
        else:
            st.warning("Evaluator 후보 점수가 없어 선택 근거를 확인할 수 없습니다.")

        st.markdown("##### 세부 3 · Critic 검토와 반영 결과")
        edits = [str(value) for value in critic.get("edits", []) if str(value).strip()]
        if critic:
            st.write(f"- 보완 필요 판단: **{'예' if critic.get('need_refine') else '아니오'}**")
            st.write(f"- 추가 VOC 표본 권고: **{'예' if critic.get('ask_more_samples') else '아니오'}**")
            st.write(f"- 실제 요약 보완: **{'반영됨' if 'summary_refined' in trace_flags else '반영 기록 없음'}**")
            st.write(f"- 실제 개선안 보완: **{'반영됨' if 'policy_refined' in trace_flags else '반영 기록 없음'}**")
            if edits:
                st.markdown("**주요 보완 의견**")
                for edit in edits:
                    st.write(f"- {edit}")
        else:
            st.warning("Critic 검토 결과가 없어 요약·개선안의 보완 여부를 확인할 수 없습니다.")

        st.markdown("##### 세부 4 · Agent 실행 이력")
        st.caption(
            f"실행 Trace ID: {trace_id} · 완료 연결 {len(trace_rows)}건 · 성공 {len(successful_rows)}건 · 실패 {len(failed_rows)}건"
        )
        if trace_rows:
            st.dataframe(
                pd.DataFrame(trace_rows),
                hide_index=True,
                width="stretch",
                column_config={
                    "순서": st.column_config.NumberColumn(width="small"),
                    "Agent 연결": st.column_config.TextColumn(width="medium"),
                    "처리 내용": st.column_config.TextColumn(width="medium"),
                    "결과": st.column_config.TextColumn(width="small"),
                    "처리시간(ms)": st.column_config.NumberColumn(format="%.2f", width="small"),
                    "판단 단서": st.column_config.TextColumn(width="large"),
                },
            )
        else:
            completed_steps = {
                "summary_refined": "Critic 의견을 반영해 요약 보완",
                "policy_refined": "Critic 의견을 반영해 개선안 보완",
                "policy_received": "최종 개선안 수신",
            }
            for flag in trace_summary["flags"]:
                st.write(f"- {completed_steps.get(flag, flag)}")
        if result.get("error"):
            st.error(f"실행 오류: {result['error']}")


@st.fragment
def _goal_testcase_selector():
    with st.container(
        horizontal=True,
        horizontal_alignment="distribute",
        vertical_alignment="center",
        gap="small",
        key="goal_testcase_compact_header",
    ):
        st.markdown("### Test Case 선택 실행")
    cases = load_unified_quality_cases().get("cases", [])
    if not cases:
        st.warning("test_cases.json에 실행할 테스트케이스가 없습니다.")
        return

    agent_snapshot = _load_agent_management_snapshot()
    page_size = 4
    total_pages = max(1, (len(cases) + page_size - 1) // page_size)
    page = st.session_state.get("goal_testcase_page", 1)
    start = (page - 1) * page_size
    page_cases = cases[start:start + page_size]
    page_case_ids = [case["case_id"] for case in page_cases]
    remembered_case_id = st.session_state.get("goal_testcase_selected_case_id")
    default_index = page_case_ids.index(remembered_case_id) if remembered_case_id in page_case_ids else 0
    rows = pd.DataFrame([
        {
            "ID": case.get("case_id", "-"),
            "분류": case.get("category", "-"),
            "질문": case.get("question", "-"),
        }
        for case in page_cases
    ])
    table_key = f"goal_testcase_table_{page}"

    table_column, detail_column = st.columns([1.75, 1], gap="medium")
    with table_column:
        selection = st.dataframe(
            rows,
            hide_index=True,
            width="stretch",
            height=245,
            row_height=48,
            key=table_key,
            on_select=partial(_remember_goal_testcase_selection, table_key, page_case_ids),
            selection_mode=["single-row-required", "single-cell"],
            selection_default={"selection": {"rows": [default_index]}},
            column_config={
                "ID": st.column_config.TextColumn("ID", width="small", pinned=True),
                "분류": st.column_config.TextColumn("분류", width="medium"),
                "질문": st.column_config.TextColumn("질문", width="large"),
            },
        )
        with st.container(horizontal_alignment="right"):
            selected_page = st.pagination(
                num_pages=total_pages,
                key="goal_testcase_page",
                persist_state="session",
            )
        if selected_page != page:
            st.rerun(scope="fragment")

    selected_rows = selection.selection.rows
    selected_index = selected_rows[0] if selected_rows else default_index
    selected_index = min(max(selected_index, 0), len(page_cases) - 1)
    selected_case = page_cases[selected_index]
    selected_case_id = selected_case["case_id"]
    _set_goal_testcase_selection(selected_case_id)
    is_fault_case = selected_case.get("execution_type") in {"fault_proxy", "isolated_fault"} or selected_case.get("category") == "fault_condition"
    case_implemented = selected_case.get("implementation_status") == "IMPLEMENTED"
    test_running = bool(st.session_state.get("goal_testcase_job_id"))

    with detail_column.container(border=True):
        st.markdown(f"**선택: {selected_case.get('case_id')}**")
        st.write(selected_case.get("question", "-"))
        st.caption(
            f"기대 의도: {selected_case.get('expected_intent', '-')}  \n"
            f"기대 작업: {selected_case.get('expected_task', '-')}"
        )
        with st.expander("판정 기준", icon=":material/rule:"):
            st.markdown(
                f"필수 출력: {', '.join(selected_case.get('required_output', [])) or '-'}  \n"
                f"금지 출력: {', '.join(selected_case.get('prohibited_output', [])) or '-'}"
            )
        if st.button(
            "Agent 파이프라인 실행",
            icon=":material/play_arrow:",
            type="primary",
            disabled=(
                not case_implemented
                or test_running
                or bool(st.session_state.get("goal_judge_job_id"))
            ),
            width="stretch",
            key=f"goal_execute_{selected_case_id}",
        ):
            _start_goal_testcase_pipeline_and_rerun(selected_case_id)
        if not case_implemented:
            st.caption("정의된 Case입니다. 자동 실행기는 후속 단계에서 구현합니다.")
        if is_fault_case:
            st.info(
                "이 케이스는 운영 Agent를 변경하지 않는 격리 장애 시험으로 실행합니다: "
                + " / ".join(selected_case.get("setup", [])),
                icon=":material/health_and_safety:",
            )
        elif not agent_snapshot["all_running"]:
            st.warning(
                "일부 Agent가 중지돼 있습니다. 장애 증상 확인을 위해 실행은 허용하며, "
                "장시간 대기를 막기 위해 20초 제한을 적용합니다."
            )

    if st.session_state.pop("goal_testcase_selection_changed", False):
        st.rerun(scope="app")


def _selected_goal_testcase() -> dict | None:
    selected_case_id = st.session_state.get("goal_testcase_selected_case_id")
    return next(
        (
            case
            for case in load_unified_quality_cases().get("cases", [])
            if case.get("case_id") == selected_case_id
        ),
        None,
    )


def _latest_goal_testcase_artifacts(selected_case: dict) -> dict:
    case_id = selected_case.get("case_id", "")
    if not case_id:
        return {}
    for run in list_voc_run_history():
        if case_id not in run.get("selected_case_ids", []):
            continue
        artifacts = load_voc_case_history_detail(run.get("run_id", ""), case_id)
        pipeline_result = artifacts.get("pipeline_result")
        if not pipeline_result:
            continue
        trace = artifacts.get("trace") or {}
        case_result = next(
            (
                item
                for item in run.get("case_results", [])
                if item.get("case_id") == case_id
            ),
            {},
        )
        restored = {
            **pipeline_result,
            "case": {**selected_case, **pipeline_result.get("case", {})},
            "trace": trace,
            "run_id": pipeline_result.get("run_id") or run.get("run_id", ""),
            "run_dir": run.get("run_dir", ""),
            "evidence_status": case_result.get("status")
            or artifacts.get("rule_result", {}).get("status")
            or pipeline_result.get("evidence_status", "-"),
        }
        if artifacts.get("judge_result"):
            restored["judge_result"] = artifacts["judge_result"]
        return {
            "result": restored,
            "trace_id": trace.get("trace_id", ""),
            "started_at": case_result.get("started_at", ""),
            "completed_at": case_result.get("finished_at", ""),
            "run_id": run.get("run_id", ""),
        }
    return {}


def _sync_goal_testcase_recent_artifacts(selected_case: dict) -> None:
    if st.session_state.get("goal_testcase_job_id"):
        return
    selected_case_id = selected_case.get("case_id", "")
    current_result = st.session_state.get("goal_testcase_result") or {}
    if current_result.get("case", {}).get("case_id") == selected_case_id:
        return
    latest = _latest_goal_testcase_artifacts(selected_case)
    if not latest:
        st.session_state.pop("goal_testcase_result", None)
        st.session_state.pop("goal_testcase_started_at", None)
        st.session_state.pop("goal_testcase_completed_at", None)
        st.session_state.pop("goal_testcase_trace_id", None)
        return
    st.session_state["goal_testcase_result"] = latest["result"]
    st.session_state["goal_testcase_trace_id"] = latest.get("trace_id", "")
    st.session_state["goal_testcase_started_at"] = latest.get("started_at", "")
    st.session_state["goal_testcase_completed_at"] = latest.get("completed_at", "")


def _ensure_goal_testcase_selection() -> dict | None:
    cases = load_unified_quality_cases().get("cases", [])
    if not cases:
        return None
    selected_case_id = st.session_state.get("goal_testcase_selected_case_id")
    selected_case = next(
        (case for case in cases if case.get("case_id") == selected_case_id),
        None,
    )
    if selected_case is None:
        selected_case = cases[0]
        _set_goal_testcase_selection(selected_case["case_id"])
    return selected_case


def _start_goal_testcase_pipeline(selected_case_id: str):
    preparation = _new_manual_preparation_progress()
    selected_case = _selected_goal_testcase() or {}
    st.session_state.goal_testcase_started_at = datetime.now().astimezone().isoformat()
    st.session_state.goal_testcase_running_case_id = selected_case_id
    st.session_state.goal_testcase_running_question = (
        selected_case.get("question")
        or selected_case.get("name")
        or ""
    )
    st.session_state.pop("goal_testcase_agent_snapshot", None)
    st.session_state.pop("goal_testcase_result", None)
    st.session_state.pop("goal_testcase_completed_at", None)
    st.session_state.goal_testcase_preparation = preparation
    st.session_state.pop("goal_testcase_trace_id", None)
    st.session_state.pop("goal_testcase_focus_result", None)
    st.session_state.pop("goal_judge_result_focus_once", None)
    st.session_state.pop("goal_manual_followup_focus_once", None)
    st.session_state.goal_testcase_focus_pipeline_once = True
    st.session_state.pop("goal_judge_error", None)
    _load_goal_monitor_snapshot.clear()
    st.session_state.goal_testcase_job_id = start_background_job(
        "manual-pipeline",
        selected_case_id,
        _execute_goal_testcase,
        selected_case_id,
        progress={"preparation": preparation},
    )


def _start_goal_testcase_pipeline_and_rerun(selected_case_id: str):
    _start_goal_testcase_pipeline(selected_case_id)
    st.rerun(scope="app")


def _start_goal_testcase_pipeline_from_callback(selected_case_id: str):
    """기존 테스트/호출부 호환용 래퍼입니다. 새 버튼은 콜백 대신 직접 rerun 흐름을 사용합니다."""
    _start_goal_testcase_pipeline(selected_case_id)


def _render_goal_scroll_anchor(
    anchor_id: str,
    *,
    scroll_flag_key: str = "",
    block: str = "start",
) -> None:
    should_scroll = bool(scroll_flag_key and st.session_state.pop(scroll_flag_key, False))
    anchor_html = f'<div id="{escape(anchor_id)}" style="height:1px"></div>'
    script_html = (
        f"""
        <script>
        (() => {{
            const focusTarget = () => {{
                const anchor = document.getElementById("{escape(anchor_id)}");
                if (anchor) {{
                    anchor.scrollIntoView({{behavior: "smooth", block: "{escape(block)}"}});
                }}
            }};
            window.setTimeout(focusTarget, 80);
            window.setTimeout(focusTarget, 350);
            window.setTimeout(focusTarget, 900);
        }})();
        </script>
        """
        if should_scroll
        else ""
    )
    st.html(
        anchor_html + script_html,
        unsafe_allow_javascript=should_scroll,
    )


def _render_goal_pipeline_focus_anchor_once():
    _render_goal_scroll_anchor(
        "goal-pipeline-scroll-anchor",
        scroll_flag_key="goal_testcase_focus_pipeline_once",
    )


def _render_goal_result_focus_anchor_once():
    _render_goal_scroll_anchor(
        "goal-result-scroll-anchor",
        scroll_flag_key="goal_testcase_focus_result",
    )


def _render_goal_judge_focus_anchor_once():
    _render_goal_scroll_anchor(
        "goal-judge-running-scroll-anchor",
        scroll_flag_key="goal_judge_focus_running_once",
    )


def _render_goal_judge_result_focus_anchor_once():
    _render_goal_scroll_anchor(
        "goal-judge-result-scroll-anchor",
        scroll_flag_key="goal_judge_result_focus_once",
        block="center",
    )


def _render_goal_execution_step(selected_case: dict):
    selected_case_id = selected_case["case_id"]
    test_running = bool(st.session_state.get("goal_testcase_job_id"))
    with st.container(border=True):
        st.markdown("#### 1단계 · Agent 파이프라인 실행")
        st.caption(
            f"{selected_case_id} · 내부 Agent 파이프라인만 먼저 수행합니다. 완료 후 독립 LLM 평가는 선택적으로 실행할 수 있습니다."
        )
        if st.button(
            "Agent 파이프라인 실행",
            icon=":material/play_arrow:",
            type="primary",
            disabled=test_running or bool(st.session_state.get("goal_judge_job_id")),
            width="stretch",
            key=f"goal_execute_{selected_case_id}",
        ):
            _start_goal_testcase_pipeline_and_rerun(selected_case_id)


def _render_goal_judge_step(selected_case: dict):
    selected_case_id = selected_case["case_id"]
    test_execution = st.session_state.get("goal_testcase_result") or {}
    if test_execution.get("case", {}).get("case_id") != selected_case_id:
        return

    is_fault_case = selected_case.get("category") == "fault_condition"
    execution = test_execution.get("execution", {})
    pipeline_ok = bool(execution.get("ok"))
    if test_execution.get("mode") == "voc":
        pipeline_ok = pipeline_ok and bool(execution.get("result", {}).get("ok"))

    _render_voc_section_heading(
        "2단계 · 독립 LLM 교차 평가",
        "Agent 파이프라인이 만든 동일 개선안을 재생성하지 않고, 선택한 외부 Provider가 독립적으로 평가합니다.",
        icon="fact_check",
    )
    if is_fault_case:
        st.info("격리 장애 Test Case는 개선안이 없어 독립 LLM 평가 대상이 아닙니다.", icon=":material/info:")
        return
    if not pipeline_ok:
        st.warning("Agent 파이프라인이 정상 완료되지 않아 독립 LLM 평가를 시작할 수 없습니다.", icon=":material/block:")
        return

    pipeline_snapshot = _manual_pipeline_llm_snapshot(test_execution)
    judge_config = _manual_judge_config_controls(
        f"goal_{selected_case_id}",
        pipeline_snapshot=pipeline_snapshot,
    )
    judge_running = bool(st.session_state.get("goal_judge_job_id"))
    with st.container(border=True):
        st.markdown("#### 평가 실행")
        if st.button(
            "독립 LLM 평가 실행",
            icon=":material/fact_check:",
            type="primary",
            disabled=judge_running or not judge_config["credential_configured"],
            width="stretch",
            key=f"goal_judge_execute_{selected_case_id}",
        ):
            st.session_state.pop("goal_judge_error", None)
            st.session_state.pop("goal_judge_result_focus_once", None)
            st.session_state.pop("goal_manual_followup_focus_once", None)
            st.session_state.goal_judge_running_case_id = selected_case_id
            st.session_state.goal_judge_focus_running_once = True
            st.session_state.goal_judge_job_id = start_background_job(
                "manual-judge",
                f"{test_execution['run_id']}:{selected_case_id}",
                _execute_goal_judge,
                test_execution["run_id"],
                selected_case_id,
                judge_config,
            )
            st.rerun()
        if judge_running:
            _render_goal_judge_focus_anchor_once()
            _render_manual_judge_running_panel(selected_case_id)
        if st.session_state.get("goal_judge_error"):
            st.error(st.session_state["goal_judge_error"], icon=":material/error:")


def render_goal_monitor():
    _ensure_goal_testcase_selection()
    _goal_testcase_selector()

    selected_case = _selected_goal_testcase()
    if selected_case:
        _sync_goal_testcase_recent_artifacts(selected_case)

    _render_goal_pipeline_focus_anchor_once()
    with st.container(
        horizontal=True,
        horizontal_alignment="distribute",
        vertical_alignment="center",
        gap="small",
        key="goal_pipeline_compact_header",
    ):
        st.markdown("### 실시간 Agent 파이프라인")
    _live_testcase_pipeline()

    if st.session_state.get("goal_judge_job_id"):
        _live_manual_judge()

    if selected_case:
        selected_case_id = selected_case["case_id"]
        _render_goal_result_focus_anchor_once()
        _render_goal_testcase_result(selected_case_id)
        _render_goal_judge_step(selected_case)
        _render_goal_judge_result(selected_case_id)
        _render_manual_followup_flow(selected_case_id)


@st.cache_data(ttl=5, max_entries=20, show_spinner=False)
def _load_batch_preflight(case_ids: tuple[str, ...]):
    return batch_preflight(list(case_ids))


def _launch_batch(
    case_ids: list[str],
    *,
    parent_run_id: str = "",
    judge_config: dict | None = None,
    rework_instruction: str = "",
):
    run = start_batch_run(
        case_ids,
        timeout_seconds=180,
        max_retries=2,
        parent_run_id=parent_run_id,
        judge_config=judge_config,
        rework_instruction=rework_instruction,
    )
    run_id = run["run_id"]
    st.session_state[BATCH_RUN_ID_KEY] = run_id
    st.session_state[BATCH_CASE_IDS_KEY] = case_ids
    st.session_state[BATCH_DIALOG_RUN_ID_KEY] = run_id
    st.session_state[f"voc_batch_initial_estimate_{run_id}"] = max(
        int(run.get("estimated_total_seconds") or 0),
        5,
    )
    st.session_state[BATCH_FUTURE_KEY] = _batch_executor().submit(
        execute_batch_run,
        run_id,
        case_ids,
        timeout_seconds=run["timeout_seconds"],
        max_retries=run["max_retries"],
        judge_config=run["judge_config"],
    )


def _parse_batch_timestamp(value: str) -> datetime | None:
    try:
        return datetime.fromisoformat(str(value or ""))
    except ValueError:
        return None


def _format_batch_duration(seconds: float) -> str:
    total_seconds = max(int(round(float(seconds or 0))), 0)
    minutes, seconds = divmod(total_seconds, 60)
    hours, minutes = divmod(minutes, 60)
    if hours:
        return f"{hours}시간 {minutes}분"
    if minutes:
        return f"{minutes}분 {seconds}초"
    return f"{seconds}초"


BATCH_RUN_ID_KEY = "voc_batch_run_id"
BATCH_CASE_IDS_KEY = "voc_batch_case_ids"
BATCH_DIALOG_RUN_ID_KEY = "voc_batch_dialog_run_id"
BATCH_FUTURE_KEY = "voc_batch_future"
BATCH_PROGRESS_RUN_TYPES = {"BATCH", "RETEST"}
BATCH_SELECTED_CASE_IDS_KEY = "voc_batch_selected_case_ids"
BATCH_ACTIVE_GROUP_KEY = "voc_batch_active_group"
BATCH_GROUP_TABLE_KEY = "voc_batch_group_table"
BATCH_GROUP_TABLE_NONCE_KEY = "voc_batch_group_table_nonce"
BATCH_GROUP_TOGGLE_KEY = "voc_batch_group_toggle"
BATCH_COMBINED_EDITOR_NONCE_KEY = "voc_batch_combined_editor_nonce"
BATCH_CASE_EDITOR_NONCE_KEY = "voc_batch_case_editor_nonce"


def _batch_progress_snapshot(run_id: str) -> dict:
    run_id = str(run_id or "").strip()
    if not run_id:
        return {}
    try:
        progress = get_batch_run_progress(run_id)
    except Exception:
        return {}
    return progress if isinstance(progress, dict) else {}


def _remember_batch_run_session(
    run_id: str,
    progress: dict | None = None,
    history_row: dict | None = None,
):
    run_id = str(run_id or "").strip()
    if not run_id:
        return
    progress = progress or {}
    history_row = history_row or {}
    st.session_state[BATCH_RUN_ID_KEY] = run_id

    verification_scope = progress.get("verification_scope", {})
    selected_ids = []
    if isinstance(verification_scope, dict):
        selected_ids = verification_scope.get("selected_case_ids") or []
    if not selected_ids:
        selected_ids = history_row.get("selected_case_ids") or []
    if selected_ids:
        st.session_state[BATCH_CASE_IDS_KEY] = [str(case_id) for case_id in selected_ids]

    estimate = progress.get("estimated_total_seconds") or 0
    try:
        estimate_seconds = int(estimate)
    except (TypeError, ValueError):
        estimate_seconds = 0
    estimate_key = f"voc_batch_initial_estimate_{run_id}"
    if estimate_seconds and not st.session_state.get(estimate_key):
        st.session_state[estimate_key] = max(estimate_seconds, 5)


def _latest_running_batch_state_from_history() -> dict:
    try:
        rows = list_voc_run_history()
    except Exception:
        return {"run_id": "", "active": False, "progress": {}, "restored": False}

    for row in sorted(rows, key=lambda item: str(item.get("started_at") or ""), reverse=True):
        if row.get("status") != "RUNNING":
            continue
        if row.get("run_type") not in BATCH_PROGRESS_RUN_TYPES:
            continue
        run_id = str(row.get("run_id") or "")
        progress = _batch_progress_snapshot(run_id)
        if progress.get("status") != "RUNNING":
            continue
        _remember_batch_run_session(run_id, progress, row)
        return {"run_id": run_id, "active": True, "progress": progress, "restored": True}
    return {"run_id": "", "active": False, "progress": {}, "restored": False}


def _active_batch_run_state(*, restore: bool = True) -> dict:
    run_id = str(st.session_state.get(BATCH_RUN_ID_KEY) or "").strip()
    progress = _batch_progress_snapshot(run_id)
    if progress.get("status") == "RUNNING":
        _remember_batch_run_session(run_id, progress)
        return {"run_id": run_id, "active": True, "progress": progress, "restored": False}
    if restore:
        restored = _latest_running_batch_state_from_history()
        if restored.get("active"):
            return restored
    return {"run_id": run_id, "active": False, "progress": progress, "restored": False}


def _batch_case_id_order(cases: list[dict]) -> tuple[str, ...]:
    return tuple(str(item.get("case_id")) for item in cases if item.get("case_id"))


def _batch_group_keys(cases: list[dict], groups: dict) -> tuple[str, ...]:
    case_group_keys = [str(item.get("group") or "ungrouped") for item in cases]
    ordered = [key for key in groups if key in case_group_keys]
    ordered.extend(key for key in case_group_keys if key not in ordered)
    return tuple(dict.fromkeys(ordered))


def _batch_cases_by_group(cases: list[dict], group_keys: tuple[str, ...]) -> dict[str, list[dict]]:
    grouped = {key: [] for key in group_keys}
    for item in cases:
        group_key = str(item.get("group") or "ungrouped")
        grouped.setdefault(group_key, []).append(item)
    return grouped


def _ordered_batch_selected_ids(selected_ids: set[str], all_case_ids: tuple[str, ...]) -> list[str]:
    return [case_id for case_id in all_case_ids if case_id in selected_ids]


def _ensure_batch_selection_state(cases: list[dict], groups: dict) -> dict:
    all_case_ids = _batch_case_id_order(cases)
    all_case_id_set = set(all_case_ids)
    group_keys = _batch_group_keys(cases, groups)
    cases_by_group = _batch_cases_by_group(cases, group_keys)

    stored_selected = st.session_state.get(BATCH_SELECTED_CASE_IDS_KEY)
    if stored_selected is None:
        selected_ids = list(all_case_ids)
    else:
        selected_ids = _ordered_batch_selected_ids(
            {str(case_id) for case_id in stored_selected if str(case_id) in all_case_id_set},
            all_case_ids,
        )
    st.session_state[BATCH_SELECTED_CASE_IDS_KEY] = selected_ids

    active_group = st.session_state.get(BATCH_ACTIVE_GROUP_KEY)
    if active_group not in group_keys:
        active_group = group_keys[0] if group_keys else ""
        st.session_state[BATCH_ACTIVE_GROUP_KEY] = active_group
    st.session_state.setdefault(BATCH_GROUP_TABLE_NONCE_KEY, 0)

    return {
        "all_case_ids": all_case_ids,
        "group_keys": group_keys,
        "cases_by_group": cases_by_group,
        "active_group": active_group,
        "selected_ids": selected_ids,
    }


def _batch_group_selection_state(group_case_ids: tuple[str, ...], selected_set: set[str]) -> dict:
    total = len(group_case_ids)
    selected_count = sum(case_id in selected_set for case_id in group_case_ids)
    if total and selected_count == total:
        return {"button": "✓", "state": "전체 선택", "count": selected_count, "all_selected": True}
    if selected_count:
        return {"button": "◩", "state": "부분 선택", "count": selected_count, "all_selected": False}
    return {"button": "□", "state": "미선택", "count": selected_count, "all_selected": False}


def _batch_group_table_rows(
    group_keys: tuple[str, ...],
    cases_by_group: dict[str, list[dict]],
    groups: dict,
    selected_ids: list[str],
) -> pd.DataFrame:
    selected_set = set(selected_ids)
    rows = []
    for group_key in group_keys:
        group_case_ids = tuple(
            str(item.get("case_id"))
            for item in cases_by_group.get(group_key, [])
            if item.get("case_id")
        )
        state = _batch_group_selection_state(group_case_ids, selected_set)
        implemented_count = sum(
            str(item.get("implementation_status") or "").upper() == "IMPLEMENTED"
            for item in cases_by_group.get(group_key, [])
        )
        rows.append({
            "그룹": groups.get(group_key, {}).get("label", group_key),
            "상태": state["state"],
            "현황": f"{state['count']} / {len(group_case_ids)}건",
            "실행 가능": f"{implemented_count}건",
            "_group": group_key,
        })
    return pd.DataFrame(rows)


def _render_batch_selector_styles() -> None:
    st.markdown(
        """
        <style>
        .vqa-batch-group-info{
            min-height:44px;display:flex;flex-direction:column;justify-content:center;
            font-family:'Segoe UI','Malgun Gothic',sans-serif;
        }
        .vqa-batch-group-info strong{
            color:#173f68;font-size:13px;line-height:1.2;font-weight:800;
            white-space:normal;word-break:keep-all;
        }
        .vqa-batch-group-info small{
            margin-top:3px;color:#6d7c8e;font-size:10px;line-height:1.25;
            white-space:normal;word-break:keep-all;
        }
        .vqa-batch-group-info.active strong{color:#0f5f9f}
        .vqa-batch-case-hint{
            display:flex;align-items:center;justify-content:space-between;gap:8px;
            margin:0 0 8px;padding:8px 10px;border:1px solid #d8e4f1;border-radius:11px;
            background:#f8fbff;font-family:'Segoe UI','Malgun Gothic',sans-serif;
        }
        .vqa-batch-case-hint strong{color:#173f68;font-size:13px}
        .vqa-batch-case-hint span{color:#6e7d8f;font-size:11px}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _render_batch_selection_mini_summary(
    *,
    selected_count: int,
    implemented_count: int,
    pending_count: int,
    selected_group_count: int,
    total_group_count: int,
) -> None:
    total_count = max(selected_count, 1)
    implemented_ratio = min(100, max(0, round(implemented_count / total_count * 100)))
    pending_ratio = min(100, max(0, round(pending_count / total_count * 100)))
    group_text = (
        f"{selected_group_count}/{total_group_count}개"
        if total_group_count
        else "-"
    )
    cards = [
        ("선택", f"{selected_count}건", "이번 일괄 Run 대상", 100 if selected_count else 0, "blue"),
        ("실행 가능", f"{implemented_count}건", "Agent 파이프라인 수행", implemented_ratio, "green"),
        ("후속 구현", f"{pending_count}건", "NOT_RUN 관리 대상", pending_ratio, "orange" if pending_count else "gray"),
        ("선택 그룹", group_text, "그룹 기준 선택 범위", 0, "gray"),
    ]
    card_html = "".join(
        f"""
        <article class="vqa-batch-mini-card {escape(tone)}">
            <small>{escape(label)}</small>
            <strong>{escape(value)}</strong>
            <p>{escape(detail)}</p>
            <i style="--value:{percent}%"></i>
        </article>
        """
        for label, value, detail, percent, tone in cards
    )
    st.markdown(
        f"""
        <div class="vqa-batch-mini-summary">
            {card_html}
        </div>
        <style>
        .vqa-batch-mini-summary{{
            display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:8px;margin:6px 0 10px;
            font-family:'Segoe UI','Malgun Gothic',sans-serif;
        }}
        .vqa-batch-mini-card{{
            position:relative;min-height:82px;padding:9px 10px;border:1px solid #d8e2ee;
            border-radius:12px;background:#f8fafc;overflow:hidden;box-sizing:border-box;
        }}
        .vqa-batch-mini-card small{{display:block;color:#748196;font-size:10px;font-weight:800}}
        .vqa-batch-mini-card strong{{display:block;margin-top:3px;color:#173f68;font-size:19px;line-height:1.1}}
        .vqa-batch-mini-card p{{
            margin:4px 0 0;color:#6e7d8f;font-size:9px;line-height:1.25;
            white-space:nowrap;overflow:hidden;text-overflow:ellipsis;
        }}
        .vqa-batch-mini-card i{{
            position:absolute;left:0;right:0;bottom:0;height:4px;background:#d8e2ee;
        }}
        .vqa-batch-mini-card i:after{{
            content:'';display:block;height:100%;width:var(--value);border-radius:999px;background:#1f6fb2;
        }}
        .vqa-batch-mini-card.green{{background:#eefaf2;border-color:#b9dfc1}}
        .vqa-batch-mini-card.green strong{{color:#17643b}}
        .vqa-batch-mini-card.green i:after{{background:#2f9660}}
        .vqa-batch-mini-card.orange{{background:#fff8e8;border-color:#ead6a8}}
        .vqa-batch-mini-card.orange strong{{color:#8a540c}}
        .vqa-batch-mini-card.orange i:after{{background:#b7791f}}
        .vqa-batch-mini-card.gray{{background:#f4f6f8;border-color:#d8e0e9}}
        .vqa-batch-mini-card.gray strong{{color:#4f5e6f}}
        .vqa-batch-mini-card.gray i:after{{background:#8a98aa}}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _set_batch_selected_ids(selected_ids: set[str], all_case_ids: tuple[str, ...]) -> list[str]:
    ordered = _ordered_batch_selected_ids(selected_ids, all_case_ids)
    st.session_state[BATCH_SELECTED_CASE_IDS_KEY] = ordered
    return ordered


def _toggle_batch_group_selection_from_click(
    group_keys: tuple[str, ...],
    group_case_ids: dict[str, tuple[str, ...]],
    all_case_ids: tuple[str, ...],
) -> None:
    click = st.session_state.get(BATCH_GROUP_TOGGLE_KEY)
    if not click:
        return
    try:
        group_key = group_keys[int(click["row"])]
    except (IndexError, TypeError, ValueError, KeyError):
        return

    selected_set = set(st.session_state.get(BATCH_SELECTED_CASE_IDS_KEY, all_case_ids))
    case_ids = set(group_case_ids.get(group_key, ()))
    if case_ids and case_ids <= selected_set:
        selected_set -= case_ids
    else:
        selected_set |= case_ids
    _set_batch_selected_ids(selected_set, all_case_ids)
    st.session_state[BATCH_ACTIVE_GROUP_KEY] = group_key
    st.session_state[BATCH_GROUP_TABLE_NONCE_KEY] = (
        int(st.session_state.get(BATCH_GROUP_TABLE_NONCE_KEY, 0) or 0) + 1
    )


def _remember_batch_group_row_selection(table_key: str, group_keys: tuple[str, ...]):
    selected_row = _promote_table_cell_to_row_selection(table_key, len(group_keys))
    if selected_row is None:
        return
    try:
        selected_group = group_keys[selected_row]
    except IndexError:
        return
    if selected_group != st.session_state.get(BATCH_ACTIVE_GROUP_KEY):
        st.session_state[BATCH_ACTIVE_GROUP_KEY] = selected_group
        st.session_state[BATCH_GROUP_TABLE_NONCE_KEY] = (
            int(st.session_state.get(BATCH_GROUP_TABLE_NONCE_KEY, 0) or 0) + 1
        )


def _apply_batch_case_row_selection(
    visible_case_ids: tuple[str, ...],
    selected_rows: list[int],
    all_case_ids: tuple[str, ...],
) -> list[str]:
    visible_set = set(visible_case_ids)
    selected_set = set(st.session_state.get(BATCH_SELECTED_CASE_IDS_KEY, all_case_ids))
    selected_set -= visible_set
    selected_set |= {
        visible_case_ids[index]
        for index in selected_rows
        if 0 <= index < len(visible_case_ids)
    }
    return _set_batch_selected_ids(selected_set, all_case_ids)


def _batch_case_selection_rows(
    group_cases: list[dict],
    selected_ids: list[str],
) -> pd.DataFrame:
    selected_set = set(selected_ids)
    rows: list[dict] = []
    for item in group_cases:
        case_id = str(item.get("case_id") or "")
        if not case_id:
            continue
        rows.append(
            {
                "체크": case_id in selected_set,
                "케이스 ID": case_id,
                "상태": _voc_status_label(
                    str(item.get("implementation_status")).upper()
                    if item.get("implementation_status") else None
                ),
                "이름": item.get("name") or item.get("question", "-"),
                "_case_id": case_id,
            }
        )
    return pd.DataFrame(rows)


def _apply_batch_case_editor_selection(
    original_rows: pd.DataFrame,
    edited_rows: pd.DataFrame,
    all_case_ids: tuple[str, ...],
) -> tuple[list[str], bool]:
    if original_rows.empty or edited_rows.empty or "체크" not in edited_rows:
        return list(st.session_state.get(BATCH_SELECTED_CASE_IDS_KEY, all_case_ids)), False

    visible_case_ids = {
        str(case_id)
        for case_id in original_rows.get("_case_id", pd.Series(dtype="string")).tolist()
        if str(case_id)
    }
    selected_set = set(st.session_state.get(BATCH_SELECTED_CASE_IDS_KEY, all_case_ids))
    before = set(selected_set)
    selected_set -= visible_case_ids

    row_count = min(len(original_rows), len(edited_rows))
    for index in range(row_count):
        case_id = str(original_rows.iloc[index].get("_case_id") or "")
        if case_id and bool(edited_rows.iloc[index].get("체크")):
            selected_set.add(case_id)

    selected = _set_batch_selected_ids(selected_set, all_case_ids)
    return selected, selected_set != before


def _batch_combined_selection_rows(
    group_keys: tuple[str, ...],
    cases_by_group: dict[str, list[dict]],
    groups: dict,
    selected_ids: list[str],
) -> pd.DataFrame:
    selected_set = set(selected_ids)
    rows: list[dict] = []
    for group_key in group_keys:
        group_case_ids = tuple(
            str(item.get("case_id"))
            for item in cases_by_group.get(group_key, [])
            if item.get("case_id")
        )
        state = _batch_group_selection_state(group_case_ids, selected_set)
        group_label = groups.get(group_key, {}).get("label", group_key)
        rows.append({
            "체크": bool(state["all_selected"]),
            "구분": "그룹",
            "대상": group_label,
            "이름": "그룹 전체 선택",
            "구현 상태": state["state"],
            "선택 현황": f"{state['count']} / {len(group_case_ids)}건",
            "_kind": "group",
            "_group": group_key,
            "_case_id": "",
        })
        for item in cases_by_group.get(group_key, []):
            case_id = str(item.get("case_id") or "")
            if not case_id:
                continue
            rows.append({
                "체크": case_id in selected_set,
                "구분": "Case",
                "대상": case_id,
                "이름": item.get("name") or item.get("question", "-"),
                "구현 상태": _voc_status_label(
                    str(item.get("implementation_status")).upper()
                    if item.get("implementation_status") else None
                ),
                "선택 현황": group_label,
                "_kind": "case",
                "_group": group_key,
                "_case_id": case_id,
            })
    return pd.DataFrame(rows)


def _apply_batch_combined_editor_selection(
    original_rows: pd.DataFrame,
    edited_rows: pd.DataFrame,
    all_case_ids: tuple[str, ...],
    cases_by_group: dict[str, list[dict]],
) -> tuple[list[str], bool]:
    if original_rows.empty or edited_rows.empty or "체크" not in edited_rows:
        return list(st.session_state.get(BATCH_SELECTED_CASE_IDS_KEY, all_case_ids)), False

    selected_set = set(st.session_state.get(BATCH_SELECTED_CASE_IDS_KEY, all_case_ids))
    changed = False
    row_count = min(len(original_rows), len(edited_rows))
    for index in range(row_count):
        original_checked = bool(original_rows.iloc[index].get("체크"))
        edited_checked = bool(edited_rows.iloc[index].get("체크"))
        if original_checked == edited_checked:
            continue
        changed = True
        row = original_rows.iloc[index]
        if row.get("_kind") == "group":
            group_key = str(row.get("_group") or "")
            group_case_ids = {
                str(item.get("case_id"))
                for item in cases_by_group.get(group_key, [])
                if item.get("case_id")
            }
            if edited_checked:
                selected_set |= group_case_ids
            else:
                selected_set -= group_case_ids
        else:
            case_id = str(row.get("_case_id") or "")
            if not case_id:
                continue
            if edited_checked:
                selected_set.add(case_id)
            else:
                selected_set.discard(case_id)
    return _set_batch_selected_ids(selected_set, all_case_ids), changed


def _render_batch_case_selector(cases: list[dict], groups: dict) -> dict:
    state = _ensure_batch_selection_state(cases, groups)
    all_case_ids = state["all_case_ids"]
    group_keys = state["group_keys"]
    cases_by_group = state["cases_by_group"]
    selected_ids = state["selected_ids"]
    group_table_nonce = int(st.session_state.get(BATCH_GROUP_TABLE_NONCE_KEY, 0) or 0)
    case_editor_nonce = int(st.session_state.get(BATCH_CASE_EDITOR_NONCE_KEY, 0) or 0)
    _render_batch_selector_styles()

    implemented_ids = {
        str(item.get("case_id"))
        for item in cases
        if item.get("case_id") and str(item.get("implementation_status") or "").upper() == "IMPLEMENTED"
    }
    selector_columns = st.columns([1.18, 1.88, 0.94], gap="medium")
    judge_config: dict = {"enabled": False, "provider": "", "model": ""}
    preflight: dict = {
        "ok": False,
        "selected_count": 0,
        "implemented_count": 0,
        "pending_count": 0,
        "agents": {"running": 0},
        "warnings": [],
        "blockers": ["선택된 테스트케이스가 없습니다."],
    }
    active_run_id = ""
    active = False

    with selector_columns[0].container(border=True, height=410):
        if not group_keys:
            st.info("표시할 테스트케이스가 없습니다.")
        else:
            selected_set = set(st.session_state.get(BATCH_SELECTED_CASE_IDS_KEY, selected_ids))
            active_group = str(st.session_state.get(BATCH_ACTIVE_GROUP_KEY) or state["active_group"])
            active_group = active_group if active_group in group_keys else group_keys[0]
            _render_voc_section_heading(
                "검증 그룹",
                "",
                icon="folder",
                badges=((f"{len(group_keys)}개", "blue"),),
            )
            action_columns = st.columns(3, gap="small")
            with action_columns[0]:
                if st.button(
                    "전체선택",
                    icon=":material/done_all:",
                    disabled=len(selected_set) == len(all_case_ids),
                    key="voc_batch_select_all_cases",
                    width="stretch",
                ):
                    selected_ids = _set_batch_selected_ids(set(all_case_ids), all_case_ids)
                    st.session_state[BATCH_CASE_EDITOR_NONCE_KEY] = case_editor_nonce + 1
                    st.session_state[BATCH_GROUP_TABLE_NONCE_KEY] = group_table_nonce + 1
                    case_editor_nonce = int(st.session_state.get(BATCH_CASE_EDITOR_NONCE_KEY, case_editor_nonce) or 0)
                    group_table_nonce = int(st.session_state.get(BATCH_GROUP_TABLE_NONCE_KEY, group_table_nonce) or 0)
                    selected_set = set(selected_ids)
            with action_columns[1]:
                if st.button(
                    "실행가능",
                    icon=":material/rule:",
                    disabled=selected_set == implemented_ids,
                    key="voc_batch_select_implemented_cases",
                    width="stretch",
                ):
                    selected_ids = _set_batch_selected_ids(implemented_ids, all_case_ids)
                    st.session_state[BATCH_CASE_EDITOR_NONCE_KEY] = case_editor_nonce + 1
                    st.session_state[BATCH_GROUP_TABLE_NONCE_KEY] = group_table_nonce + 1
                    case_editor_nonce = int(st.session_state.get(BATCH_CASE_EDITOR_NONCE_KEY, case_editor_nonce) or 0)
                    group_table_nonce = int(st.session_state.get(BATCH_GROUP_TABLE_NONCE_KEY, group_table_nonce) or 0)
                    selected_set = set(selected_ids)
            with action_columns[2]:
                if st.button(
                    "선택해제",
                    icon=":material/remove_done:",
                    disabled=not selected_set,
                    key="voc_batch_clear_all_cases",
                    width="stretch",
                ):
                    selected_ids = _set_batch_selected_ids(set(), all_case_ids)
                    st.session_state[BATCH_CASE_EDITOR_NONCE_KEY] = case_editor_nonce + 1
                    st.session_state[BATCH_GROUP_TABLE_NONCE_KEY] = group_table_nonce + 1
                    case_editor_nonce = int(st.session_state.get(BATCH_CASE_EDITOR_NONCE_KEY, case_editor_nonce) or 0)
                    group_table_nonce = int(st.session_state.get(BATCH_GROUP_TABLE_NONCE_KEY, group_table_nonce) or 0)
                    selected_set = set(selected_ids)

            group_list = st.container(height=286)
            for group_key in group_keys:
                group_cases = cases_by_group.get(group_key, [])
                group_case_ids = tuple(
                    str(item.get("case_id"))
                    for item in group_cases
                    if item.get("case_id")
                )
                group_state = _batch_group_selection_state(group_case_ids, selected_set)
                implemented_count = sum(
                    str(item.get("implementation_status") or "").upper() == "IMPLEMENTED"
                    for item in group_cases
                )
                group_label = groups.get(group_key, {}).get("label", group_key)
                is_active_group = group_key == active_group
                tone = "green" if group_state["all_selected"] else "orange" if group_state["count"] else "gray"
                with group_list.container(border=True):
                    row_columns = st.columns([0.13, 0.58, 0.29], gap="small", vertical_alignment="center")
                    checked = row_columns[0].checkbox(
                        f"{group_label} 전체 선택",
                        value=group_state["all_selected"],
                        key=f"voc_batch_group_check_{group_key}_{group_table_nonce}",
                        label_visibility="collapsed",
                    )
                    if checked != group_state["all_selected"]:
                        next_selected = set(selected_set)
                        if checked:
                            next_selected |= set(group_case_ids)
                        else:
                            next_selected -= set(group_case_ids)
                        selected_ids = _set_batch_selected_ids(next_selected, all_case_ids)
                        selected_set = set(selected_ids)
                        st.session_state[BATCH_ACTIVE_GROUP_KEY] = group_key
                        st.session_state[BATCH_CASE_EDITOR_NONCE_KEY] = case_editor_nonce + 1
                        st.session_state[BATCH_GROUP_TABLE_NONCE_KEY] = group_table_nonce + 1
                        case_editor_nonce = int(st.session_state.get(BATCH_CASE_EDITOR_NONCE_KEY, case_editor_nonce) or 0)
                        group_table_nonce = int(st.session_state.get(BATCH_GROUP_TABLE_NONCE_KEY, group_table_nonce) or 0)
                        active_group = group_key
                        group_state = _batch_group_selection_state(group_case_ids, selected_set)
                        tone = "green" if group_state["all_selected"] else "orange" if group_state["count"] else "gray"
                        is_active_group = True
                    row_columns[1].markdown(
                        f"""
                        <div class="vqa-batch-group-info {'active' if is_active_group else ''}">
                            <strong>{escape(group_label)}</strong>
                            <small>{escape(group_state['state'])} · {group_state['count']} / {len(group_case_ids)}건 선택 · 실행 가능 {implemented_count}건</small>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                    with row_columns[2]:
                        if st.button(
                            "보기",
                            key=f"voc_batch_group_view_{group_key}",
                            type="primary" if is_active_group else "secondary",
                            width="stretch",
                            help="오른쪽 Case 목록에 이 그룹을 표시합니다.",
                        ):
                            st.session_state[BATCH_ACTIVE_GROUP_KEY] = group_key
                            active_group = group_key
                            is_active_group = True
                    row_columns[2].markdown(_voc_ui_badge(group_state["state"], tone))

    with selector_columns[1].container(border=True, height=410):
        if not group_keys:
            st.info("선택할 Case가 없습니다.")
        else:
            active_group = str(st.session_state.get(BATCH_ACTIVE_GROUP_KEY) or state["active_group"])
            active_group = active_group if active_group in group_keys else group_keys[0]
            group_label = groups.get(active_group, {}).get("label", active_group)
            selected_ids = list(st.session_state.get(BATCH_SELECTED_CASE_IDS_KEY, selected_ids))
            selected_set = set(selected_ids)
            group_cases = cases_by_group.get(active_group, [])
            group_case_ids = tuple(
                str(item.get("case_id"))
                for item in group_cases
                if item.get("case_id")
            )
            group_state = _batch_group_selection_state(group_case_ids, selected_set)
            group_tone = "green" if group_state["all_selected"] else "orange" if group_state["count"] else "gray"
            _render_voc_section_heading(
                "Case 선택",
                "",
                icon="checklist",
                badges=((group_label, "blue"), (group_state["state"], group_tone)),
            )
            st.markdown(
                f"""
                <div class="vqa-batch-case-hint">
                    <strong>{escape(group_label)}</strong>
                    <span>{group_state['count']} / {len(group_case_ids)}건 선택</span>
                </div>
                """,
                unsafe_allow_html=True,
            )
            case_rows = _batch_case_selection_rows(
                group_cases,
                selected_ids,
            )
            case_table_height = max(142, min(276, 54 + (len(case_rows) * 31)))
            edited_rows = st.data_editor(
                case_rows,
                hide_index=True,
                width="stretch",
                height=case_table_height,
                row_height=30,
                key=f"voc_batch_case_editor_{active_group}_{case_editor_nonce}",
                disabled=["케이스 ID", "상태", "이름", "_case_id"],
                column_order=["체크", "케이스 ID", "상태", "이름"],
                column_config={
                    "체크": st.column_config.CheckboxColumn("체크", width=58),
                    "케이스 ID": st.column_config.TextColumn("Case ID", width=82, pinned=True),
                    "상태": st.column_config.TextColumn("구현 상태", width=132),
                    "이름": st.column_config.TextColumn("이름 / 질문", width="large"),
                    "_case_id": None,
                },
            )
            selected_ids, changed = _apply_batch_case_editor_selection(
                case_rows,
                edited_rows,
                all_case_ids,
            )
            if changed:
                st.session_state[BATCH_CASE_EDITOR_NONCE_KEY] = case_editor_nonce + 1

    with selector_columns[2].container(border=True, height=410):
        selected_ids = list(st.session_state.get(BATCH_SELECTED_CASE_IDS_KEY, selected_ids))
        selected_set = set(selected_ids)
        selected_cases = [
            item
            for item in cases
            if str(item.get("case_id") or "") in selected_set
        ]
        selected_group_count = sum(
            any(
                str(item.get("case_id")) in selected_set
                for item in cases_by_group.get(group_key, [])
                if item.get("case_id")
            )
            for group_key in group_keys
        )
        implemented_count = sum(
            str(item.get("implementation_status") or "").upper() == "IMPLEMENTED"
            for item in selected_cases
        )
        pending_count = len(selected_ids) - implemented_count
        _render_voc_section_heading(
            "선택 요약",
            "",
            icon="target",
            badges=((f"{len(selected_ids)}건", "blue"),),
        )
        _render_batch_selection_mini_summary(
            selected_count=len(selected_ids),
            implemented_count=implemented_count,
            pending_count=pending_count,
            selected_group_count=selected_group_count,
            total_group_count=len(group_keys),
        )
        with st.popover(
            "독립 LLM 평가 옵션",
            icon=":material/fact_check:",
            width="stretch",
            key="voc_batch_judge_options",
        ):
            judge_config = _judge_config_controls(
                "voc_batch",
                fault_only=bool(selected_ids) and all(case_id.startswith("FT-") for case_id in selected_ids),
            )
        _render_batch_judge_selection_badge(judge_config)
        preflight = _load_batch_preflight(tuple(selected_ids)) if selected_ids else {
            "ok": False,
            "selected_count": 0,
            "implemented_count": 0,
            "pending_count": 0,
            "agents": {"running": 0},
            "warnings": [],
            "blockers": ["선택된 테스트케이스가 없습니다."],
        }

        active_state = _active_batch_run_state()
        active_run_id = active_state["run_id"]
        active = active_state["active"]
        run_button_label = (
            "진행 화면 열기"
            if active
            else f"{len(selected_ids)}건 일괄 실행"
        )
        if st.button(
            run_button_label,
            icon=":material/open_in_new:" if active else ":material/play_arrow:",
            type="primary",
            disabled=not active and (not selected_ids or not preflight.get("ok")),
            width="stretch",
            key="voc_batch_run_selected",
        ):
            if active:
                _open_batch_progress_dialog(active_run_id)
            else:
                _launch_batch(selected_ids, judge_config=judge_config)
                st.rerun()

    return {
        "selected_ids": selected_ids,
        "judge_config": judge_config,
        "preflight": preflight,
        "active_run_id": active_run_id,
        "active": active,
    }


def _batch_case_results_for_display(case_results: list[dict]) -> pd.DataFrame:
    rows = pd.DataFrame(case_results)
    if rows.empty:
        return rows
    def label_value(value) -> str:
        if value is None or pd.isna(value):
            return "-"
        return _voc_status_label(str(value).upper())

    display_rows = pd.DataFrame()
    display_rows["케이스 ID"] = rows.get("case_id", "-")
    if "status" in rows:
        display_rows["상태"] = rows["status"].map(label_value)
    if "mode" in rows:
        display_rows["수행 유형"] = rows["mode"].map(label_value)
    if "attempt_count" in rows:
        display_rows["시도"] = rows["attempt_count"]
    if "judge_status" in rows:
        display_rows["독립 LLM 평가 상태"] = rows["judge_status"].map(label_value)
    if "judge_score" in rows:
        display_rows["독립 LLM 평가 점수"] = rows["judge_score"]
    if "judge_independence_grade" in rows:
        display_rows["독립성"] = rows["judge_independence_grade"].map(_judge_independence_grade_label)
    if "message" in rows:
        display_rows["처리 내용"] = rows["message"]
    if "finished_at" in rows:
        display_rows["완료 시각"] = rows["finished_at"]
    return display_rows


def _batch_timing(progress: dict, *, now: datetime | None = None) -> dict:
    started_at = _parse_batch_timestamp(progress.get("started_at", ""))
    finished_at = _parse_batch_timestamp(progress.get("finished_at", ""))
    now = now or datetime.now().astimezone()
    end_at = finished_at or now
    elapsed = max((end_at - started_at).total_seconds(), 0.0) if started_at else 0.0
    total = max(int(progress.get("total") or 0), 1)
    completed = max(int(progress.get("completed") or 0), 0)
    persisted_estimate = float(progress.get("estimated_total_seconds") or 0)
    initial_estimate = float(
        st.session_state.get(
            f"voc_batch_initial_estimate_{progress.get('run_id', '')}",
            persisted_estimate
            or total * (75 if progress.get("judge_config", {}).get("enabled") else 45),
        )
    )
    if completed > 0 and progress.get("status") == "RUNNING":
        observed_estimate = elapsed / completed * total
        estimated_total = max(elapsed, observed_estimate)
    elif progress.get("status") == "RUNNING":
        estimated_total = max(initial_estimate, elapsed)
    else:
        estimated_total = elapsed
    return {
        "elapsed_seconds": elapsed,
        "estimated_total_seconds": estimated_total,
        "remaining_seconds": max(estimated_total - elapsed, 0.0),
    }


def _batch_progress_fraction(
    progress: dict,
    timing: dict,
    *,
    now: datetime | None = None,
) -> float:
    status = progress.get("status")
    if status != "RUNNING":
        return 1.0
    runtime = progress.get("runtime_progress", {})
    phase = runtime.get("phase", "")
    total = max(int(progress.get("total") or 0), 1)
    completed = max(int(progress.get("completed") or 0), 0)
    now = now or datetime.now().astimezone()
    if phase == "PREFLIGHT":
        return 0.02
    if phase == "PREPARING":
        phase_started = _parse_batch_timestamp(runtime.get("phase_started_at", ""))
        phase_elapsed = max((now - phase_started).total_seconds(), 0.0) if phase_started else 0.0
        return min(0.02 + phase_elapsed / 15 * 0.03, 0.05)
    if phase == "FINALIZING":
        return 0.98
    if phase == "RUNNING":
        case_started = _parse_batch_timestamp(runtime.get("current_case_started_at", ""))
        case_elapsed = max((now - case_started).total_seconds(), 0.0) if case_started else 0.0
        case_estimate = max((timing["estimated_total_seconds"] - 15) / total, 1.0)
        current_fraction = min(case_elapsed / case_estimate, 0.9)
        case_progress = min((completed + current_fraction) / total, 1.0)
        return min(0.05 + case_progress * 0.90, 0.95)
    return min(completed / total, 0.99)


def _render_batch_stage_flow(progress: dict):
    runtime = progress.get("runtime_progress", {})
    phase = runtime.get("phase", "")
    stages = (
        ("PREFLIGHT", "사전 점검", "환경·에이전트·대상 검증"),
        ("PREPARING", "처리 준비", "실행 폴더·카탈로그 준비"),
        ("RUNNING", "테스트케이스 수행", "Agent 파이프라인·독립 LLM 평가 실행"),
        ("FINALIZING", "결과 정리", "증적·집계 저장"),
    )
    phase_index = next((index for index, stage in enumerate(stages) if stage[0] == phase), -1)
    if progress.get("status") != "RUNNING":
        phase_index = len(stages)
    cards = []
    for index, (_, label, description) in enumerate(stages):
        state = "done" if index < phase_index else "active" if index == phase_index else "waiting"
        icon = "✓" if state == "done" else str(index + 1)
        cards.append(
            f'<div class="vqb-stage {state}"><b>{icon}</b><span><strong>{label}</strong><small>{description}</small></span></div>'
        )
    st.html(
        """
        <style>
        .vqb-stage-flow{display:grid;grid-template-columns:repeat(4,1fr);gap:9px;margin:4px 0 12px}
        .vqb-stage{display:flex;align-items:center;gap:9px;padding:10px 11px;border:1px solid #d4dde7;border-radius:11px;background:#f2f4f6;color:#87919d}
        .vqb-stage>b{display:grid;place-items:center;min-width:27px;height:27px;border-radius:50%;background:#aeb6bf;color:#fff;font-size:11px}
        .vqb-stage span{display:block}.vqb-stage strong{display:block;font-size:12px}.vqb-stage small{display:block;font-size:9px;margin-top:2px}
        .vqb-stage.done{background:#f3f8fd;border-color:#b9cee2;color:#315b82}.vqb-stage.done>b{background:#2e6d9f}
        .vqb-stage.active{background:#edf6ff;border:2px solid #1767a5;color:#124b79;box-shadow:0 4px 13px rgba(23,103,165,.13)}
        .vqb-stage.active>b{background:#1767a5;animation:vqb-pulse 1.2s infinite}
        @keyframes vqb-pulse{50%{box-shadow:0 0 0 6px rgba(23,103,165,.14)}}
        @media(max-width:800px){.vqb-stage-flow{grid-template-columns:repeat(2,1fr)}}
        </style>
        <div class="vqb-stage-flow">"""
        + "".join(cards)
        + "</div>"
    )


def _render_batch_progress_styles():
    st.markdown(
        """
        <style>
        div[data-testid="stDialog"] div[data-testid="stProgress"] [role="progressbar"] {
            height: 24px !important;
            border-radius: 12px !important;
        }
        div[data-testid="stDialog"] div[data-testid="stProgress"] [role="progressbar"] > div {
            height: 100% !important;
            border-radius: 12px !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def _render_batch_progress_content(
    run_id: str,
    progress: dict,
    timing: dict | None = None,
):
    total = max(int(progress.get("total") or 0), 1)
    completed = int(progress.get("completed") or 0)
    timing = timing or _batch_timing(progress)
    progress_fraction = _batch_progress_fraction(progress, timing)
    runtime = progress.get("runtime_progress", {})
    phase_label = runtime.get("phase_label") or (
        "수행 완료" if progress.get("status") != "RUNNING" else "실행 상태 확인 중"
    )
    phase_message = runtime.get("message", "Run 진행 정보를 불러오고 있습니다.")
    if progress.get("status") == "RUNNING":
        st.info(f"**{phase_label}** · {phase_message}", icon=":material/pending:")
    st.progress(
        progress_fraction,
        text=(
            f"**전체 예상 진행률 {progress_fraction * 100:.0f}% · "
            f"완료 {completed} / {total}건**"
        ),
    )

    counts = progress.get("counts", {})
    with st.container(horizontal=True):
        st.metric("검토 필요", counts.get("REVIEW_REQUIRED", 0))
        st.metric("실패", counts.get("FAIL", 0))
        st.metric("오류", counts.get("ERROR", 0))
        st.metric("미실행", counts.get("NOT_RUN", 0))
    if progress.get("judge_config", {}).get("enabled"):
        judge_counts = progress.get("judge_counts", {})
        with st.container(horizontal=True):
            st.metric("독립 LLM 평가 통과", judge_counts.get("PASS", 0))
            st.metric("독립 LLM 평가 검토", judge_counts.get("REVIEW_REQUIRED", 0))
            st.metric("독립 LLM 평가 실패", judge_counts.get("FAIL", 0))
            st.metric("독립 LLM 평가 오류", judge_counts.get("ERROR", 0))

    rows = _batch_case_results_for_display(progress.get("case_results", []))
    if not rows.empty:
        st.dataframe(rows, hide_index=True, width="stretch")

    status = progress.get("status")
    if status == "RUNNING":
        st.caption(f"Run ID: {run_id}")
        if progress.get("stop_requested"):
            st.warning("현재 테스트케이스가 끝난 뒤 나머지는 미실행으로 기록하고 중지합니다.")
        elif st.button("실행 중지", icon=":material/stop_circle:", key=f"stop_{run_id}"):
            request_batch_stop(run_id)
            st.rerun(scope="fragment")
        return

    st.session_state.pop(BATCH_FUTURE_KEY, None)
    if status == "COMPLETED":
        st.success(f"일괄 실행이 완료되었습니다. · Run ID: {run_id}")
    elif status == "INTERRUPTED":
        st.warning(f"중지 요청에 따라 실행을 종료했습니다. · Run ID: {run_id}")
    else:
        st.error(f"일괄 실행 엔진 오류로 종료되었습니다. · Run ID: {run_id}")
    st.caption(f"증적 위치: {progress.get('run_dir', '-')}")

    retry_ids = [
        item.get("case_id") for item in progress.get("case_results", [])
        if item.get("status") in {"FAIL", "ERROR"}
    ]
    if retry_ids and st.button(
        f"실패·오류 {len(retry_ids)}건 재실행",
        icon=":material/replay:",
        key=f"retry_{run_id}",
    ):
        _launch_batch(
            retry_ids,
            parent_run_id=run_id,
            judge_config=progress.get("judge_config"),
        )
        st.rerun()


@st.fragment(run_every="1s")
def _live_batch_progress():
    run_id = _active_batch_run_state()["run_id"]
    if not run_id:
        return
    _render_batch_progress_content(run_id, get_batch_run_progress(run_id))


def _open_batch_progress_dialog(run_id: str):
    if run_id:
        st.session_state[BATCH_DIALOG_RUN_ID_KEY] = str(run_id)
        st.rerun(scope="app")


def _close_batch_progress_dialog():
    st.session_state.pop(BATCH_DIALOG_RUN_ID_KEY, None)


@st.dialog(
    "일괄 테스트케이스 수행 진행 상황",
    width="large",
    icon=":material/pending_actions:",
    on_dismiss=_close_batch_progress_dialog,
)
def _render_batch_progress_dialog(run_id: str):
    _render_batch_progress_dialog_body(run_id)


@st.fragment(run_every="1s")
def _render_batch_progress_dialog_body(run_id: str):
    progress = get_batch_run_progress(run_id)
    timing = _batch_timing(progress)
    _render_batch_progress_styles()

    status = progress.get("status", "ERROR")
    header = st.columns([1.0, 1.15, 1.25, 1.15, 0.7], vertical_alignment="center")
    header[0].metric("상태", _voc_status_label(status))
    progress_fraction = _batch_progress_fraction(progress, timing)
    header[1].metric("예상 진행률", f"{progress_fraction * 100:.0f}%")
    header[2].metric("예상 소요시간", f"약 {_format_batch_duration(timing['estimated_total_seconds'])}")
    header[3].metric("예상 남은 시간", _format_batch_duration(timing["remaining_seconds"]))
    if header[4].button(
        ":material/close: 닫기",
        key=f"voc_batch_dialog_close_{run_id}",
        width="stretch",
        help="팝업만 닫으며 일괄 수행은 백그라운드에서 계속됩니다.",
    ):
        _close_batch_progress_dialog()
        st.rerun(scope="app")

    st.caption(
        f"경과 시간 {_format_batch_duration(timing['elapsed_seconds'])} · "
        "예상 시간은 완료된 Case의 평균 처리시간으로 실시간 보정됩니다."
    )
    _render_batch_stage_flow(progress)
    _render_batch_progress_content(run_id, progress, timing)


def _render_batch_running_summary(active_state: dict):
    run_id = active_state.get("run_id", "")
    progress = active_state.get("progress") or _batch_progress_snapshot(run_id)
    if not run_id or progress.get("status") != "RUNNING":
        return

    timing = _batch_timing(progress)
    progress_fraction = _batch_progress_fraction(progress, timing)
    runtime = progress.get("runtime_progress", {})
    phase_label = runtime.get("phase_label") or "실행 상태 확인 중"
    current_case = runtime.get("current_case_id") or "준비 중"
    total = max(int(progress.get("total") or 0), 1)
    completed = int(progress.get("completed") or 0)

    with st.container(border=True):
        header = st.columns([1.25, 0.8, 0.85, 0.85, 0.9], vertical_alignment="center")
        with header[0]:
            st.markdown("#### 진행 중인 일괄 수행")
            st.caption("팝업을 닫거나 다른 페이지로 이동해도 백그라운드 실행은 계속됩니다.")
        header[1].metric("현재 단계", phase_label)
        header[2].metric("현재 테스트케이스", current_case)
        header[3].metric("진행", f"{completed}/{total}건")
        with header[4]:
            if st.button(
                "진행 화면 열기",
                icon=":material/open_in_new:",
                type="primary",
                width="stretch",
                key=f"voc_batch_resume_progress_{run_id}",
            ):
                _open_batch_progress_dialog(run_id)
        st.progress(
            progress_fraction,
            text=f"전체 예상 진행률 {progress_fraction * 100:.0f}% · 완료 {completed} / {total}건",
        )


def _render_batch_execution_safety_notice() -> None:
    with st.expander(
        "일괄 수행 중 화면을 닫으면?",
        expanded=False,
        icon=":material/help:",
    ):
        cols = st.columns(3, gap="small", vertical_alignment="top")
        with cols[0]:
            st.badge("계속 실행", color="green", icon=":material/check_circle:")
            st.markdown("**팝업·브라우저 탭 닫기**")
            st.caption("Streamlit 서버가 살아 있으면 일괄 테스트케이스는 백그라운드에서 계속 수행됩니다.")
        with cols[1]:
            st.badge("중단 가능", color="orange", icon=":material/warning:")
            st.markdown("**Streamlit 서버 종료**")
            st.caption("서버 프로세스를 끄면 실행 중인 백그라운드 작업은 이어서 수행되지 않을 수 있습니다.")
        with cols[2]:
            st.badge("증적 보존", color="blue", icon=":material/folder_open:")
            st.markdown("**다시 확인하는 방법**")
            st.caption("완료된 Case 결과는 즉시 저장되며, 수행 이력에서 Run ID 기준으로 확인합니다.")
        st.caption(
            "안전하게 중단하려면 진행 화면의 `중지 요청`을 사용하세요. "
            "강제 종료된 Run은 완료된 Case까지만 증적이 남고, 남은 Case는 재실행 또는 재시험 대상으로 확인합니다."
        )


def _batch_preflight_display_state(preflight: dict) -> dict:
    blockers = list(preflight.get("blockers", []))
    warnings = list(preflight.get("warnings", []))
    pending_count = int(preflight.get("pending_count") or 0)
    selected_count = int(preflight.get("selected_count") or 0)
    implemented_count = int(preflight.get("implemented_count") or 0)

    if blockers:
        return {
            "tone": "bad",
            "icon": "error",
            "title": "실행 차단",
            "message": blockers[0],
            "items": blockers[1:] + warnings,
        }
    if selected_count <= 0:
        return {
            "tone": "idle",
            "icon": "touch_app",
            "title": "대상 선택 필요",
            "message": "왼쪽 목록에서 일괄 수행할 테스트케이스를 선택하세요.",
            "items": [],
        }
    if warnings:
        return {
            "tone": "warn",
            "icon": "warning",
            "title": "실행 가능 · 확인 필요",
            "message": warnings[0],
            "items": warnings[1:],
        }
    if pending_count > 0:
        return {
            "tone": "warn",
            "icon": "rule",
            "title": "실행 가능 · 후속 구현 포함",
            "message": "선택 대상 중 아직 구현되지 않은 Case는 이번 Run에서 미실행으로 기록됩니다.",
            "items": [],
        }
    return {
        "tone": "good",
        "icon": "check_circle",
        "title": "실행 준비 완료",
        "message": f"선택한 {implemented_count}건은 현재 조건에서 일괄 수행할 수 있습니다.",
        "items": [],
    }


def _render_batch_preflight_readiness(preflight: dict) -> None:
    state = _batch_preflight_display_state(preflight)
    icon_label = _html_status_chip_label(str(state.get("icon") or ""))
    detail_items = "".join(
        f"<li>{escape(str(item))}</li>"
        for item in state.get("items", [])
    )
    detail_markup = f"<ul>{detail_items}</ul>" if detail_items else ""
    checked_at = str(preflight.get("checked_at") or "").replace("T", " ")[:19] or "-"
    st.markdown(
        f"""
        <section class="vqa-batch-preflight {state['tone']}">
            <span>{escape(icon_label)}</span>
            <div>
                <strong>{escape(state["title"])}</strong>
                <p>{escape(state["message"])}</p>
                {detail_markup}
                <small>점검 시각 · {escape(checked_at)}</small>
            </div>
        </section>
        <style>
        .vqa-batch-preflight{{
            min-height:82px;margin:0;padding:12px 14px;border-radius:12px;
            border:1px solid #c8d9ee;border-left:4px solid #155a96;
            background:linear-gradient(135deg,#f8fbff,#fff);display:flex;gap:12px;
            align-items:flex-start;box-sizing:border-box;font-family:'Segoe UI','Malgun Gothic',sans-serif;
        }}
        .vqa-batch-preflight.good{{border-left-color:#299049;background:linear-gradient(135deg,#f3fbf5,#fff)}}
        .vqa-batch-preflight.warn{{border-left-color:#b36a08;background:linear-gradient(135deg,#fff8ec,#fff)}}
        .vqa-batch-preflight.bad{{border-left-color:#d83f36;background:linear-gradient(135deg,#fff4f2,#fff)}}
        .vqa-batch-preflight.idle{{border-left-color:#8a98aa;background:linear-gradient(135deg,#f6f8fb,#fff)}}
        .vqa-batch-preflight>span{{
            flex:0 0 38px;height:26px;border-radius:999px;display:flex;align-items:center;justify-content:center;
            font-size:11px;font-weight:900;color:#155a96;background:#e7f1fb;border:1px solid #bfd7ef;margin-top:1px;
        }}
        .vqa-batch-preflight.good>span{{color:#1f7f43;background:#edf8f0;border-color:#bfdfca}}
        .vqa-batch-preflight.warn>span{{color:#9b5c07;background:#fff2dc;border-color:#e6c383}}
        .vqa-batch-preflight.bad>span{{color:#bd3029;background:#fff0ee;border-color:#efbbb6}}
        .vqa-batch-preflight.idle>span{{color:#6f7c8c;background:#eef1f4;border-color:#d3d9e0}}
        .vqa-batch-preflight strong{{display:block;color:#173f68;font-size:14px;line-height:1.2}}
        .vqa-batch-preflight p{{margin:4px 0 0;color:#40536d;font-size:12px;line-height:1.35}}
        .vqa-batch-preflight ul{{margin:6px 0 0;padding-left:16px;color:#40536d;font-size:11px;line-height:1.35}}
        .vqa-batch-preflight small{{display:block;margin-top:6px;color:#7a889a;font-size:10px}}
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_batch_execution():
    focus_notice = st.session_state.pop("voc_batch_focus_notice", None)
    if focus_notice:
        st.info(focus_notice, icon=":material/replay:")
    st.markdown("### 실행 대상 선택")
    st.caption(
        "기본은 순차 실행입니다. 실행 결과와 재시도 내역은 테스트케이스별 증적으로 즉시 저장됩니다."
    )
    _render_batch_execution_safety_notice()
    catalog = load_quality_test_catalog()
    cases = catalog.get("cases", [])
    groups = catalog.get("groups", {})
    if not cases:
        st.warning("quality_test_catalog.json에 실행할 Case가 없습니다.")
        return

    selection = _render_batch_case_selector(cases, groups)
    selected_ids = selection["selected_ids"]
    preflight = selection["preflight"]
    with st.container(border=True):
        st.markdown("#### 사전 점검")
        _render_batch_preflight_readiness(preflight)

    active_state = _active_batch_run_state()
    active_run_id = active_state["run_id"] or selection["active_run_id"]
    active = active_state["active"] or selection["active"]
    if active_run_id:
        if active:
            _render_batch_running_summary(active_state)

    dialog_run_id = st.session_state.get(BATCH_DIALOG_RUN_ID_KEY)
    if dialog_run_id:
        _render_batch_progress_dialog(dialog_run_id)


@st.cache_data(ttl=3, max_entries=1, show_spinner=False)
def _load_voc_history_rows():
    return list_voc_run_history()


@st.dialog("수행 이력 삭제 확인")
def _confirm_delete_voc_runs(run_ids: list[str]):
    st.warning("선택한 Run 폴더와 중앙 index 항목을 함께 삭제합니다. 이 작업은 되돌릴 수 없습니다.")
    st.code("\n".join(run_ids), language="text")
    if st.button("선택 Run 영구 삭제", type="primary", icon=":material/delete_forever:"):
        result = delete_voc_run_history(run_ids)
        _load_voc_history_rows.clear()
        st.session_state.voc_history_delete_result = result
        st.rerun()


def _history_optional_int(value, default=None):
    try:
        if value is None or value == "":
            return default
        return int(value)
    except (TypeError, ValueError):
        return default


def _history_verification_scope_model(manifest: dict, summary: dict) -> dict:
    metadata = manifest.get("run_metadata", {}) if isinstance(manifest.get("run_metadata"), dict) else {}
    scope = metadata.get("verification_scope", {}) if isinstance(metadata.get("verification_scope"), dict) else {}
    selected_ids = manifest.get("selected_case_ids", [])
    selected_count = _history_optional_int(
        scope.get("selected_count"),
        len(selected_ids) or _history_optional_int(summary.get("total"), 0),
    )
    catalog_total = _history_optional_int(scope.get("catalog_total_cases"))
    executable_count = _history_optional_int(scope.get("executable_count"))
    pending_count = _history_optional_int(scope.get("pending_count"))
    execution_type_counts = scope.get("execution_type_counts", {})
    if not isinstance(execution_type_counts, dict):
        execution_type_counts = {}
    parent_run_id = str(metadata.get("parent_run_id") or "")
    run_type = str(manifest.get("run_type") or "")
    return {
        "state_model_version": manifest.get("state_model_version", metadata.get("state_model_version", "")),
        "catalog_total_cases": catalog_total,
        "selected_count": selected_count,
        "executable_count": executable_count,
        "pending_count": pending_count,
        "execution_type_counts": execution_type_counts,
        "has_scope": bool(scope),
        "is_full_suite": bool(catalog_total and selected_count == catalog_total),
        "run_type": run_type,
        "is_retest": run_type == "RETEST",
        "parent_run_id": parent_run_id,
    }


def _render_history_verification_scope(manifest: dict, summary: dict):
    model = _history_verification_scope_model(manifest, summary)
    selected_count = model["selected_count"] or 0
    executable_count = model["executable_count"]
    pending_count = model["pending_count"]
    catalog_total = model["catalog_total_cases"]
    with st.container(border=True):
        scope_badge = "35건 전체 회차" if model["is_full_suite"] else "부분 회차"
        retest_badge = "재시험" if model["is_retest"] else _voc_status_label(model["run_type"], model["run_type"] or "-")
        _render_voc_section_heading(
            "검증 회차 범위",
            "이번 Run이 어느 범위의 Case를 대표하는지 확인합니다.",
            icon="fact_check",
            badges=((scope_badge, "blue"), (retest_badge, "gray")),
        )

        _render_voc_summary_cards(
            [
                {
                    "icon": "inventory_2",
                    "label": "카탈로그 전체",
                    "value": f"{catalog_total if catalog_total is not None else '-'}건",
                    "detail": "관리 기준 전체 Case",
                },
                {
                    "icon": "checklist",
                    "label": "선택 Case",
                    "value": f"{selected_count}건",
                    "detail": "이번 Run 대상",
                },
                {
                    "icon": "play_circle",
                    "label": "실행 가능",
                    "value": "-" if executable_count is None else f"{executable_count}건",
                    "detail": "현재 자동 수행 가능",
                },
                {
                    "icon": "pending_actions",
                    "label": "후속 구현",
                    "value": "-" if pending_count is None else f"{pending_count}건",
                    "detail": "승인 대상 후속 구현",
                },
            ],
            columns=4,
            height=112,
        )

        if selected_count and executable_count is not None:
            pending_value = pending_count or 0
            progress = min(max(executable_count / selected_count, 0.0), 1.0)
            st.progress(
                progress,
                text=f"실행 가능 {executable_count}건 · 후속 구현 {pending_value}건 · 선택 {selected_count}건 기준",
            )
        elif not model["has_scope"]:
            st.caption("이 Run은 검증 회차 범위 메타데이터가 없는 과거 기록입니다. 신규 Batch Run부터 실행 가능/후속 구현 범위가 표시됩니다.")

        if model["is_retest"]:
            if model["parent_run_id"]:
                st.info(
                    f"재시험 Run입니다. 원본 Run `{model['parent_run_id']}`와 연결되어 전후 비교 기준으로 사용할 수 있습니다.",
                    icon=":material/replay:",
                )
            else:
                st.warning(
                    "재시험 Run이지만 원본 Run 연결 정보가 없습니다. 전후 비교 신뢰도가 낮습니다.",
                    icon=":material/warning:",
                )

        if model["execution_type_counts"]:
            rows = pd.DataFrame(
                [
                    {
                        "구성": VOC_EXECUTION_TYPE_LABELS.get(str(key), str(key)),
                        "건수": int(value or 0),
                    }
                    for key, value in model["execution_type_counts"].items()
                ]
            )
            if not rows.empty:
                chart = (
                    alt.Chart(rows)
                    .mark_bar(color="#2F6FB0", cornerRadiusEnd=5, size=18)
                    .encode(
                        y=alt.Y("구성:N", title=None, sort="-x"),
                        x=alt.X("건수:Q", title=None, axis=alt.Axis(format="d", tickMinStep=1)),
                        tooltip=[
                            alt.Tooltip("구성:N", title="구성"),
                            alt.Tooltip("건수:Q", title="건수", format="d"),
                        ],
                    )
                    .properties(height=max(110, 36 + len(rows) * 28))
                    .configure_view(strokeWidth=0)
                )
                with st.expander("검증 영역별 Case 구성", icon=":material/bar_chart:"):
                    st.altair_chart(chart, theme=None)


HISTORY_ARTIFACT_LABELS = {
    "pipeline_result": "파이프라인 결과",
    "trace": "Agent 실행 Trace",
    "rule_result": "내부 규칙 판정",
    "judge_result": "독립 LLM 평가",
    "validity_result": "개선안 타당성 평가",
}
HISTORY_SCORE_LABELS = {
    "accuracy": "정확성",
    "groundedness": "근거성",
    "completeness": "충실성",
    "specificity": "구체성",
    "safety": "안전성",
    "feasibility": "실행 가능성",
    "measurability": "측정 가능성",
    "priority": "우선순위",
}


def _history_duration_label(started_at: str | None, finished_at: str | None) -> str:
    if not started_at or not finished_at:
        return "-"
    try:
        seconds = (datetime.fromisoformat(finished_at) - datetime.fromisoformat(started_at)).total_seconds()
    except (TypeError, ValueError):
        return "-"
    return f"{seconds:.1f}초"


def _history_mapping_rows(mapping: dict | None, labels: dict[str, str] | None = None) -> pd.DataFrame:
    labels = labels or {}
    rows = []
    for key, value in (mapping or {}).items():
        if isinstance(value, bool):
            display = "설정" if value else "미설정"
        elif isinstance(value, (dict, list)):
            display = json.dumps(value, ensure_ascii=False)
        else:
            display = "-" if value in (None, "") else str(value)
        rows.append({"항목": labels.get(key, key), "내용": display})
    return pd.DataFrame(rows)


def _render_history_execution_info(manifest: dict, summary: dict) -> None:
    st.markdown("#### :material/info: Run 기본 정보")
    basic_rows = pd.DataFrame(
        [
            {"항목": "Run ID", "내용": manifest.get("run_id", "-")},
            {"항목": "실행 유형", "내용": _voc_status_label(manifest.get("run_type", "-"))},
            {"항목": "실행 상태", "내용": _voc_status_label(manifest.get("status", "-"))},
            {"항목": "Suite", "내용": manifest.get("suite_id", "-")},
            {"항목": "Catalog 버전", "내용": manifest.get("catalog_version", "-")},
            {"항목": "시작 시각", "내용": _history_table_timestamp(manifest.get("started_at"))},
            {"항목": "종료 시각", "내용": _history_table_timestamp(manifest.get("finished_at"))},
            {
                "항목": "총 수행 시간",
                "내용": _history_duration_label(manifest.get("started_at"), manifest.get("finished_at")),
            },
            {
                "항목": "대상 Case",
                "내용": ", ".join(manifest.get("selected_case_ids", [])) or "-",
            },
            {
                "항목": "최종 판정",
                "내용": _voc_status_label(summary.get("deployment_decision", "미판정")),
            },
        ]
    )
    st.dataframe(
        basic_rows,
        hide_index=True,
        width="stretch",
        column_config={
            "항목": st.column_config.TextColumn(width="small", pinned=True),
            "내용": st.column_config.TextColumn(width="large"),
        },
    )

    rubric_versions = manifest.get("rubric_versions", {})
    if rubric_versions:
        st.markdown("#### :material/rule: 적용 Rubric")
        rubric_rows = []
        rubric_labels = {
            "internal_pipeline": "내부 파이프라인 품질",
            "independent_judge": "독립 LLM 평가",
            "improvement_validity": "개선안 타당성 평가",
        }
        for rubric_type, detail in rubric_versions.items():
            detail = detail if isinstance(detail, dict) else {"version": detail}
            rubric_rows.append(
                {
                    "평가 단계": rubric_labels.get(rubric_type, rubric_type),
                    "버전": detail.get("version", "-"),
                    "무결성 Hash": str(detail.get("sha256", "-"))[:16],
                }
            )
        st.dataframe(pd.DataFrame(rubric_rows), hide_index=True, width="stretch")

    model_snapshot = manifest.get("model_snapshot", {})
    if model_snapshot:
        st.markdown("#### :material/smart_toy: 실행 모델")
        model_labels = {
            "summary": "요약 생성",
            "policy": "개선안 생성",
            "judge": "독립 LLM 평가",
        }
        model_rows = []
        for role, detail in model_snapshot.items():
            detail = detail if isinstance(detail, dict) else {}
            model_rows.append(
                {
                    "역할": model_labels.get(role, role),
                    "Provider": detail.get("provider", "-"),
                    "모델": detail.get("model", "-"),
                    "사용 상태": (
                        "사용"
                        if detail.get("enabled", detail.get("credential_configured", True))
                        else "미사용"
                    ),
                }
            )
        st.dataframe(pd.DataFrame(model_rows), hide_index=True, width="stretch")

    environment = manifest.get("environment_fingerprint", {})
    if environment:
        with st.expander("실행 환경 정보", icon=":material/computer:"):
            st.dataframe(
                _history_mapping_rows(
                    environment,
                    {
                        "python_version": "Python",
                        "operating_system": "운영체제",
                        "platform_release": "OS 버전",
                        "runtime_root": "Runtime 경로",
                        "fingerprint_sha256": "환경 Hash",
                    },
                ),
                hide_index=True,
                width="stretch",
            )

    with st.expander("원본 실행 정보 JSON", icon=":material/data_object:"):
        st.json(manifest)


def _render_history_pipeline_artifact(pipeline: dict) -> None:
    execution = pipeline.get("execution", {}) if isinstance(pipeline.get("execution"), dict) else {}
    result = execution.get("result", {}) if isinstance(execution.get("result"), dict) else {}
    with st.container(horizontal=True):
        st.metric("수행 상태", "성공" if execution.get("ok") and result.get("ok") else "확인 필요", border=True)
        st.metric("실행 모드", pipeline.get("mode", "-"), border=True)
        st.metric("기록 시각", _history_table_timestamp(pipeline.get("recorded_at")), border=True)
        st.metric("실행 Trace", "연결" if result.get("trace") else "없음", border=True)
    st.markdown("#### :material/help: 고객 질문")
    st.write(execution.get("question", "-") or "-")
    result_columns = st.columns(2, gap="medium")
    with result_columns[0].container(border=True, height="stretch"):
        st.markdown("#### :material/summarize: VOC 분석 요약")
        st.write(result.get("summary", "-") or "-")
    with result_columns[1].container(border=True, height="stretch"):
        st.markdown("#### :material/lightbulb: 최종 개선안")
        _render_policy_improvement(result.get("policy", "-") or "-")

    diagnostic_rows = []
    for label, field in (
        ("Interpreter 해석", "intent_json"),
        ("Evaluator 점수", "eval_json"),
        ("Critic 보완", "summary_critic_json"),
    ):
        parsed = _parse_json_mapping(result.get(field))
        if parsed:
            diagnostic_rows.append({"Agent 진단": label, "결과": json.dumps(parsed, ensure_ascii=False)})
    if diagnostic_rows:
        with st.expander("Agent별 핵심 진단", icon=":material/account_tree:"):
            st.dataframe(pd.DataFrame(diagnostic_rows), hide_index=True, width="stretch")

    reports = execution.get("reports", {})
    if isinstance(reports, dict) and reports:
        st.caption(
            "생성 보고서 · "
            + " · ".join(f"{name.upper()} {Path(str(report_path)).name}" for name, report_path in reports.items())
        )


def _render_history_rule_artifact(rule: dict) -> None:
    with st.container(horizontal=True):
        st.metric("내부 판정", _voc_status_label(rule.get("status", "NOT_RUN")), border=True)
        st.metric("Rubric", rule.get("rubric_id", "-"), border=True)
        st.metric("버전", rule.get("rubric_version", "-"), border=True)
    message = rule.get("message")
    if rule.get("status") == "PASS":
        st.success(message or "내부 파이프라인 품질 기준을 통과했습니다.")
    elif rule.get("status") in {"FAIL", "ERROR"}:
        st.error(message or "내부 파이프라인 품질 기준을 통과하지 못했습니다.")
    else:
        st.warning(message or "내부 규칙 판정 결과를 확인하세요.")


def _history_score_rows(scores: dict | None) -> pd.DataFrame:
    rows = []
    for dimension, detail in (scores or {}).items():
        if isinstance(detail, dict):
            rows.append(
                {
                    "평가 항목": _voc_display_term(HISTORY_SCORE_LABELS.get(dimension, dimension)),
                    "점수": detail.get("score", detail.get("points", detail.get("value", "-"))),
                    "배점": detail.get("max_points", "-"),
                    "판정 근거": _voc_display_term(detail.get("reason", detail.get("comment", "-"))),
                }
            )
        else:
            rows.append(
                {
                    "평가 항목": _voc_display_term(HISTORY_SCORE_LABELS.get(dimension, dimension)),
                    "점수": detail,
                    "배점": "-",
                    "판정 근거": "-",
                }
            )
    return pd.DataFrame(rows)


HISTORY_REVIEW_ROLE_LABELS = {
    "QA": "QA 검토",
    "QA_REVIEW": "QA 검토",
    "QA_REVIEWED": "QA 검토",
    "BUSINESS": "업무 승인",
    "BUSINESS_APPROVAL": "업무 승인",
    "BUSINESS_APPROVED": "업무 승인",
}
HISTORY_REVIEW_DECISION_LABELS = {
    "APPROVE": "승인",
    "APPROVED": "승인",
    "REVISION_REQUIRED": "보완 요청",
    "REJECTED": "반려",
    "REJECT": "반려",
}


def _history_looks_broken_text(value: object) -> bool:
    text = "" if value is None else str(value).strip()
    if not text:
        return True
    if "�" in text:
        return True
    question_count = text.count("?") + text.count("？")
    has_korean = bool(re.search(r"[가-힣]", text))
    if question_count >= 3 and not has_korean:
        return True
    if re.search(r"[?？]{2,}", text) and question_count >= max(3, len(text) // 5):
        return True
    return bool(re.fullmatch(r"[\s?？�·.,:;_\-\\/()]+", text))


def _history_safe_display_text(value: object, *, fallback: str = "-") -> str:
    text = _voc_display_term("" if value is None else str(value).strip())
    if _history_looks_broken_text(text):
        return fallback
    return text


def _history_review_role_label(value: object, *, index: int) -> str:
    code = str(value or "").strip().upper()
    label = HISTORY_REVIEW_ROLE_LABELS.get(code) or _voc_status_label(value, "")
    return _history_safe_display_text(label, fallback=f"검토 {index}")


def _history_review_decision_label(value: object) -> str:
    code = str(value or "").strip().upper()
    label = HISTORY_REVIEW_DECISION_LABELS.get(code) or _voc_status_label(value, "")
    return _history_safe_display_text(label, fallback="결정 미확인")


def _history_review_decision_tone(decision_label: str) -> str:
    if "반려" in decision_label:
        return "red"
    if "보완" in decision_label or "검토" in decision_label:
        return "orange"
    if "승인" in decision_label:
        return "green"
    return "blue"


def _history_review_state_transition(review: dict) -> str:
    before = _history_safe_display_text(_voc_status_label(review.get("from_state"), ""), fallback="")
    after = _history_safe_display_text(_voc_status_label(review.get("to_state"), ""), fallback="")
    if before and after:
        return f"{before} → {after}"
    if after:
        return f"→ {after}"
    if before:
        return before
    return "-"


def _history_review_comment_fallback(review: dict, *, index: int) -> str:
    role = _history_review_role_label(review.get("reviewer_role"), index=index)
    decision = _history_review_decision_label(review.get("decision"))
    transition = _history_review_state_transition(review)
    if transition and transition != "-":
        return f"{role} 단계에서 {decision} 처리되었습니다. ({transition})"
    return f"{role} 단계에서 {decision} 처리되었습니다."


def _history_validity_review_rows(reviews: list | tuple | None) -> pd.DataFrame:
    rows = []
    for index, review in enumerate(reviews or [], start=1):
        if not isinstance(review, dict):
            continue
        rows.append(
            {
                "단계": _history_review_role_label(review.get("reviewer_role"), index=index),
                "결정": _history_review_decision_label(review.get("decision")),
                "검토자": _history_safe_display_text(
                    review.get("reviewer_name_or_id"),
                    fallback="검토자 미확인",
                ),
                "상태 변화": _history_review_state_transition(review),
                "검토 의견": _history_safe_display_text(
                    review.get("comment"),
                    fallback=_history_review_comment_fallback(review, index=index),
                ),
                "검토 시각": _history_table_timestamp(review.get("reviewed_at")),
            }
        )
    return pd.DataFrame(rows)


def _history_action_tone(label: object) -> str:
    text = str(label or "")
    if any(token in text for token in ("승인", "통과", "완료", "보고서", "시연")):
        return "green"
    if any(token in text for token in ("오류", "실패", "반려", "보완", "재시험")):
        return "red"
    if any(token in text for token in ("평가", "검토", "확인", "재평가")):
        return "orange"
    return "blue"


def _render_history_validity_review_timeline(reviews: list | tuple | None) -> None:
    rows = _history_validity_review_rows(reviews)
    if rows.empty:
        return
    with st.container(border=True):
        _render_voc_section_heading(
            "QA·업무 승인 이력",
            "QA 검토와 업무 승인 단계에서 남긴 결정 이력입니다.",
            icon="approval",
            badges=((f"{len(rows)}건", "blue"),),
        )
        for row in rows.to_dict("records"):
            with st.container(border=True):
                cols = st.columns([0.9, 0.85, 1.45, 2.8], gap="small", vertical_alignment="center")
                with cols[0]:
                    st.markdown(f"**{row['단계']}**")
                    st.caption(row["검토 시각"] or "-")
                with cols[1]:
                    st.markdown(
                        _voc_ui_badge(
                            row["결정"],
                            _history_review_decision_tone(row["결정"]),
                        )
                    )
                with cols[2]:
                    st.caption("상태 변화")
                    st.write(row["상태 변화"])
                with cols[3]:
                    st.caption("검토 의견")
                    st.write(row["검토 의견"])


def _history_float_value(value) -> float | None:
    try:
        if value in (None, ""):
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _history_judge_pass_threshold(rubric: dict | None) -> float:
    rubric = rubric if isinstance(rubric, dict) else {}
    decisions = rubric.get("automatic_decisions")
    if isinstance(decisions, list):
        for item in decisions:
            if isinstance(item, dict) and str(item.get("decision", "")).upper() == "PASS":
                threshold = _history_float_value(item.get("min_score"))
                if threshold is not None:
                    return threshold
    return 80.0


def _history_limited_items(values, *, limit: int = 2) -> list[str]:
    if not isinstance(values, list):
        return []
    items = [str(value).strip() for value in values if str(value or "").strip()]
    if len(items) <= limit:
        return items
    return [*items[:limit], f"외 {len(items) - limit}건"]


def _history_judge_review_blockers(judge: dict, rubric: dict | None) -> list[str]:
    judge = judge if isinstance(judge, dict) else {}
    rubric = rubric if isinstance(rubric, dict) else {}
    decision = str(judge.get("decision") or "NOT_RUN").strip().upper() or "NOT_RUN"
    threshold = _history_judge_pass_threshold(rubric)
    blockers: list[str] = []

    if decision == "NOT_RUN":
        blockers.append("독립 LLM 평가 결과가 아직 없어 재평가가 아니라 최초 평가가 필요합니다.")
        return blockers

    error = str(judge.get("error") or "").strip()
    if error:
        blockers.append(f"이전 평가 오류: {_voc_display_term(error)}")

    triggered = [
        str(rule).strip()
        for rule in (judge.get("immediate_fail_rules_triggered") or [])
        if str(rule or "").strip()
    ]
    if triggered:
        blockers.append("즉시 실패 규칙 감지: " + ", ".join(triggered[:3]))

    if judge.get("independence_hold"):
        reason = str(judge.get("independence_hold_reason") or "").strip()
        blockers.append(
            "독립성 보류: "
            + _voc_display_term(reason or "평가 Provider/모델 독립성 기준으로 추가 확인이 필요합니다.")
        )

    score = _history_float_value(judge.get("total_score"))
    if decision in {"REVIEW_REQUIRED", "FAIL"}:
        if score is None:
            blockers.append("총점이 저장되지 않아 독립 LLM 판정 근거를 확인할 수 없습니다.")
        elif score < threshold:
            blockers.append(f"총점 기준 미달: {score:g} / {threshold:g}점")

    dimension_scores = judge.get("dimension_scores")
    dimensions = rubric.get("dimensions")
    if isinstance(dimension_scores, dict) and isinstance(dimensions, dict):
        for key, spec in dimensions.items():
            if not isinstance(spec, dict):
                continue
            floor = _history_float_value(spec.get("pass_floor"))
            if floor is None:
                continue
            detail = dimension_scores.get(key)
            score_value = (
                _history_float_value(detail.get("score", detail.get("points", detail.get("value"))))
                if isinstance(detail, dict)
                else _history_float_value(detail)
            )
            label = str(spec.get("label") or HISTORY_SCORE_LABELS.get(str(key), str(key)))
            if score_value is None:
                blockers.append(f"세부 항목 미평가: {label} / 하한 {floor:g}점")
            elif score_value < floor:
                blockers.append(f"세부 항목 하한 미달: {label} {score_value:g} / {floor:g}점")

    for risk in _history_limited_items(judge.get("risks")):
        blockers.append(f"위험 지적: {_voc_display_term(risk)}")
    for recommendation in _history_limited_items(judge.get("recommendations")):
        blockers.append(f"보완 권고: {_voc_display_term(recommendation)}")

    if not blockers and decision == "PASS":
        blockers.append("PASS 상태입니다. 재평가는 Provider 교차 확인이나 Rubric 변경 영향 확인이 필요할 때만 실행합니다.")
    elif not blockers:
        blockers.append("저장된 평가 결과만으로 세부 원인을 식별하지 못했습니다. 원본 평가 JSON의 판정 근거를 확인하세요.")
    return blockers[:8]


def _history_judge_review_focus(
    *,
    decision: str,
    rubric_changed: bool,
    has_error: bool,
    has_content_blockers: bool,
) -> list[str]:
    focus = []
    if has_error:
        focus.append("이전 오류가 사라지고 정상 평가 결과가 생성됐는지 확인")
    if rubric_changed:
        focus.append("변경된 Rubric에서 총점·세부 하한·판정이 어떻게 달라졌는지 확인")
    if decision != "PASS":
        focus.append("기존 검토 필요 원인이 해소되어 PASS로 전환됐는지 확인")
    focus.append("독립성 보류가 남아 있으면 다른 Provider로 교차 평가")
    if has_content_blockers:
        focus.append("같은 원인이 남으면 재평가 반복이 아니라 개선안 보완 또는 RETEST 진행")
    else:
        focus.append("판정만 달라졌다면 Provider/모델별 평가 차이를 비교")
    return focus


def _history_judge_reevaluation_result_key(run_id: str, case_id: str) -> str:
    return f"voc_history_judge_reevaluation_result::{run_id}::{case_id}"


def _history_judge_reevaluation_job_key(run_id: str, case_id: str) -> str:
    return f"voc_history_judge_reevaluation_job::{run_id}::{case_id}"


def _history_judge_reevaluation_focus_key(run_id: str, case_id: str) -> str:
    return f"voc_history_judge_reevaluation_focus::{run_id}::{case_id}"


def _history_judge_reevaluation_error_key(run_id: str, case_id: str) -> str:
    return f"voc_history_judge_reevaluation_error::{run_id}::{case_id}"


def _history_dom_id(value: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "-", str(value or "")).strip("-")
    return safe or "target"


def _history_judge_reevaluation_progress_model(job: dict | None) -> dict:
    job = job if isinstance(job, dict) else {}
    status = str(job.get("status") or "RUNNING")
    progress = job.get("progress", {}) if isinstance(job.get("progress"), dict) else {}
    raw_percent = progress.get("percent")
    try:
        percent = int(raw_percent)
    except (TypeError, ValueError):
        percent = 8
    if job.get("done") and status == "COMPLETED":
        percent = 100
    elif job.get("done") or status == "ERROR":
        percent = 100
    else:
        percent = min(max(percent, 5), 96)
    stage = str(progress.get("stage") or "독립 LLM 재평가 진행 중")
    detail = str(progress.get("detail") or "저장된 동일 Agent 파이프라인 결과를 독립 LLM으로 다시 평가합니다.")
    if status == "ERROR":
        stage = "독립 LLM 재평가 오류"
        detail = str(job.get("error") or detail)
    elif job.get("done") and status == "COMPLETED":
        stage = "독립 LLM 재평가 완료"
        detail = "재평가 결과를 불러와 결과 영역으로 이동합니다."
    return {
        "status": status,
        "percent": percent,
        "fraction": min(max(percent / 100, 0.0), 1.0),
        "stage": stage,
        "detail": detail,
        "started_at": _history_table_timestamp(job.get("started_at")),
        "updated_at": _history_table_timestamp(job.get("updated_at")),
        "error": str(job.get("error") or ""),
    }


def _render_history_judge_reevaluation_progress(
    run_id: str,
    case_id: str,
    job: dict | None,
) -> None:
    model = _history_judge_reevaluation_progress_model(job)
    tone = "red" if model["status"] == "ERROR" else "blue"
    with st.container(border=True):
        heading, state = st.columns([2.5, 1], gap="small", vertical_alignment="center")
        with heading:
            st.markdown(f"#### 독립 LLM 재평가 진행 · {case_id}")
            st.caption(model["detail"])
        with state:
            st.markdown(f":{tone}-badge[{model['stage']}]", text_alignment="right")
        st.progress(
            model["fraction"],
            text=f"{model['percent']}% · {model['stage']}",
        )
        st.caption(
            f"Run {run_id} · 시작 {model['started_at'] or '-'} · 갱신 {model['updated_at'] or '-'} · "
            "현재 Streamlit 서버가 켜져 있으면 팝업을 닫아도 백그라운드 작업은 계속됩니다."
        )


@st.fragment(run_every="1s")
def _live_history_judge_reevaluation(run_id: str, case_id: str) -> None:
    job_key = _history_judge_reevaluation_job_key(run_id, case_id)
    job_id = str(st.session_state.get(job_key) or "")
    if not job_id:
        return
    job = background_job_snapshot(job_id)
    if not job:
        st.session_state.pop(job_key, None)
        st.session_state[_history_judge_reevaluation_error_key(run_id, case_id)] = (
            "독립 LLM 재평가 작업 상태를 찾을 수 없습니다. 다시 실행해 주세요."
        )
        st.rerun()
        return

    _render_history_judge_reevaluation_progress(run_id, case_id, job)
    if not job.get("done"):
        return

    result_key = _history_judge_reevaluation_result_key(run_id, case_id)
    if job.get("status") == "COMPLETED":
        st.session_state[result_key] = job.get("result") or {}
        st.session_state[f"voc_history_artifact_pending_{run_id}"] = "judge_result"
        st.session_state[HISTORY_SELECTED_RUN_ID_KEY] = run_id
        st.session_state[_history_judge_reevaluation_focus_key(run_id, case_id)] = True
        _load_voc_history_rows.clear()
        _load_validity_candidates.clear()
    else:
        st.session_state[_history_judge_reevaluation_error_key(run_id, case_id)] = (
            job.get("error") or "독립 LLM 재평가 중 오류가 발생했습니다."
        )
    st.session_state.pop(job_key, None)
    discard_background_job(job_id)
    st.rerun()


def _history_score_delta_label(previous_score, current_score) -> str:
    before = _history_float_value(previous_score)
    after = _history_float_value(current_score)
    if before is None or after is None:
        return "-"
    delta = round(after - before, 2)
    if delta > 0:
        return f"+{delta:g}점 상승"
    if delta < 0:
        return f"{delta:g}점 하락"
    return "점수 변화 없음"


def _history_judge_reevaluation_result_model(
    reevaluation_result: dict | None,
    *,
    current_judge_rubric: dict | None = None,
) -> dict:
    reevaluation_result = reevaluation_result if isinstance(reevaluation_result, dict) else {}
    judge = reevaluation_result.get("judge_result", {})
    if not isinstance(judge, dict):
        judge = {}
    history = judge.get("evaluation_history", [])
    previous = history[-1] if isinstance(history, list) and history and isinstance(history[-1], dict) else {}
    current_judge_rubric = current_judge_rubric or load_independent_judge_rubric()

    before_decision = str(previous.get("decision") or "NOT_RUN").strip().upper() or "NOT_RUN"
    after_decision = str(judge.get("decision") or "ERROR").strip().upper() or "ERROR"
    before_score = previous.get("total_score")
    after_score = judge.get("total_score")
    decision_changed = before_decision != after_decision
    score_delta = _history_score_delta_label(before_score, after_score)
    blockers = _history_judge_review_blockers(judge, current_judge_rubric)
    provider = _manual_provider_label(judge.get("provider"))
    model = judge.get("model") or "-"

    if after_decision == "PASS":
        next_action = {
            "tone": "green",
            "label": "개선안 타당성 평가 실행",
            "detail": "독립 LLM 평가가 PASS로 전환되었습니다. 타당성 검증 화면의 평가 설정에서 바로 평가를 실행하세요.",
            "target": {
                "enabled": True,
                "page": VOC_VALIDITY_PAGE_NAME,
                "run_id": reevaluation_result.get("run_id", ""),
                "case_id": reevaluation_result.get("case_id", ""),
                "action_code": "RUN_VALIDITY",
                "button_label": "타당성 평가로 이동",
            },
        }
    elif after_decision == "ERROR":
        next_action = {
            "tone": "red",
            "label": "평가 오류 조치",
            "detail": "Provider API 키, 모델명, 네트워크 상태를 확인한 뒤 다시 평가하세요.",
            "target": {"enabled": False, "button_label": "오류 조치 후 재시도"},
        }
    else:
        next_action = {
            "tone": "orange",
            "label": "검토 필요 원인 확인",
            "detail": (
                "판정이 여전히 검토 필요/실패이면 같은 입력 재평가를 반복하기보다 "
                "아래 원인을 기준으로 개선안 보완 또는 RETEST 여부를 판단하세요."
            ),
            "target": {"enabled": False, "button_label": "원인 확인 후 보완 판단"},
        }

    return {
        "run_id": reevaluation_result.get("run_id", ""),
        "case_id": reevaluation_result.get("case_id", ""),
        "before_decision": before_decision,
        "before_label": _voc_status_label(before_decision),
        "before_score": before_score,
        "after_decision": after_decision,
        "after_label": _voc_status_label(after_decision),
        "after_score": after_score,
        "decision_changed": decision_changed,
        "decision_change_label": "판정 변경" if decision_changed else "판정 유지",
        "score_delta": score_delta,
        "provider": provider,
        "model": model,
        "evaluated_at": _history_table_timestamp(judge.get("evaluated_at")),
        "blockers": blockers,
        "next_action": next_action,
    }


def _render_history_judge_reevaluation_result(
    run_id: str,
    case_id: str,
    reevaluation_result: dict | None,
) -> None:
    if not reevaluation_result:
        return
    _render_goal_scroll_anchor(
        (
            "history-judge-reevaluation-result-"
            f"{_history_dom_id(run_id)}-{_history_dom_id(case_id)}"
        ),
        scroll_flag_key=_history_judge_reevaluation_focus_key(run_id, case_id),
        block="start",
    )
    model = _history_judge_reevaluation_result_model(reevaluation_result)
    action = model["next_action"]
    action_tone = {
        "green": "green",
        "red": "red",
        "orange": "orange",
        "gray": "gray",
    }.get(action.get("tone", "gray"), "gray")
    result_key = _history_judge_reevaluation_result_key(run_id, case_id)
    with st.container(border=True):
        heading, state = st.columns([2.5, 1], gap="small", vertical_alignment="center")
        with heading:
            st.markdown(f"#### 독립 LLM 재평가 결과 · {case_id}")
            st.caption(
                f"{model['provider']} · {model['model']} · "
                f"{model['evaluated_at'] or '평가 시각 미확인'}"
            )
        with state:
            st.markdown(f":{action_tone}-badge[{action['label']}]", text_alignment="right")

        columns = st.columns(4, gap="small")
        card_payloads = [
            ("history", "이전 판정", f"{model['before_label']} · {model['before_score'] if model['before_score'] is not None else '-'}점", "재평가 전 저장 결과"),
            ("published_with_changes", "재평가 판정", f"{model['after_label']} · {model['after_score'] if model['after_score'] is not None else '-'}점", model["decision_change_label"]),
            ("trending_up", "점수 변화", model["score_delta"], "같은 Agent 결과를 다시 채점한 변화"),
            ("arrow_forward", "다음 액션", action["label"], action["detail"]),
        ]
        for column, (icon, label, value, detail_text) in zip(columns, card_payloads, strict=False):
            with column.container(border=True, height=126):
                st.caption(f":material/{icon}: {label}")
                st.markdown(f"##### {value}")
                st.caption(detail_text)

        if model["after_decision"] != "PASS":
            with st.container(border=True):
                st.caption(":material/troubleshoot: 재평가 후에도 확인해야 할 원인")
                for blocker in model["blockers"]:
                    st.write(f"- {blocker}")

        button_columns = st.columns([1.2, 1.2, 2.2], gap="small", vertical_alignment="center")
        target = action.get("target", {})
        with button_columns[0]:
            if st.button(
                target.get("button_label", "다음 액션으로 이동"),
                icon=":material/arrow_forward:",
                type="primary" if target.get("enabled") else "secondary",
                disabled=not target.get("enabled"),
                width="stretch",
                key=f"voc_history_judge_result_next_{run_id}_{case_id}",
            ):
                st.session_state.pop(result_key, None)
                _dismiss_history_detail_dialog()
                _apply_history_next_action_target(target)
                st.rerun()
        with button_columns[1]:
            if st.button(
                "결과 확인 완료",
                icon=":material/check_circle:",
                width="stretch",
                key=f"voc_history_judge_result_clear_{run_id}_{case_id}",
            ):
                st.session_state.pop(result_key, None)
                st.rerun()
        with button_columns[2]:
            st.caption(action["detail"])


def _render_history_judge_artifact(judge: dict) -> None:
    with st.container(horizontal=True):
        st.metric("유효 판정", _voc_status_label(judge.get("decision", "NOT_RUN")), border=True)
        st.metric(
            "총점",
            f"{judge.get('total_score')}점" if judge.get("total_score") is not None else "-",
            border=True,
        )
        st.metric("독립성", _judge_independence_grade_label(judge.get("independence_grade", "-")), border=True)
        st.metric("수행 시간", f"{float(judge.get('duration_seconds') or 0):g}초", border=True)
    st.caption(
        f"{judge.get('provider', '-')} · {judge.get('model', '-')} · "
        f"Rubric {judge.get('rubric_version', '-')} · "
        f"{_history_table_timestamp(judge.get('evaluated_at'))}"
    )
    if judge.get("error"):
        st.error(judge["error"])
    elif judge.get("independence_hold"):
        st.warning(
            f"점수 기준 판정은 {_voc_status_label(judge.get('rubric_decision', '-'))}이지만 "
            f"{judge.get('independence_hold_reason', '독립성 기준으로 검토가 필요합니다.')}"
        )
    score_rows = _history_score_rows(judge.get("dimension_scores"))
    if not score_rows.empty:
        st.dataframe(score_rows, hide_index=True, width="stretch")
    detail_columns = st.columns(3, gap="small")
    for column, title, values in (
        (detail_columns[0], "확인 근거", judge.get("evidence", [])),
        (detail_columns[1], "잔여 위험", judge.get("risks", [])),
        (detail_columns[2], "보완 권고", judge.get("recommendations", [])),
    ):
        with column.container(border=True, height=190):
            st.markdown(f"**{title}**")
            if values:
                for value in values:
                    st.write(f"- {value}")
            else:
                st.caption("표시할 내용 없음")


def _render_history_validity_artifact(validity: dict) -> None:
    with st.container(horizontal=True):
        st.metric("자동 판정", _voc_status_label(validity.get("decision", "NOT_RUN")), border=True)
        st.metric(
            "개선안 타당성 점수",
            f"{validity.get('total_score')}점" if validity.get("total_score") is not None else "-",
            border=True,
        )
        st.metric("승인 단계", _voc_status_label(validity.get("workflow_state", "DRAFT")), border=True)
        st.metric("정식 승인", "승인" if validity.get("formal_approval") else "미승인", border=True)
    st.caption(
        f"{validity.get('provider', '-')} · {validity.get('model', '-')} · "
        f"Rubric {validity.get('rubric_version', '-')} · "
        f"{_history_table_timestamp(validity.get('evaluated_at'))}"
    )
    if validity.get("error"):
        st.error(validity["error"])
    holds = validity.get("immediate_hold_rules_triggered", [])
    if holds:
        st.error("즉시 승인 보류 · " + ", ".join(_validity_hold_rule_label(rule) for rule in holds))
    score_rows = _history_score_rows(validity.get("dimension_scores"))
    if not score_rows.empty:
        st.dataframe(score_rows, hide_index=True, width="stretch")
    recommendations = validity.get("recommendations", [])
    if recommendations:
        with st.container(border=True):
            st.markdown("#### :material/edit_note: 보완 권고")
            for recommendation in recommendations:
                st.write(f"- {recommendation}")
    _render_history_validity_review_timeline(validity.get("human_reviews", []))


def _history_judge_reevaluation_context_model(
    case_id: str,
    artifacts: dict,
    detail: dict,
    *,
    current_judge_rubric: dict | None = None,
) -> dict:
    artifacts = artifacts if isinstance(artifacts, dict) else {}
    detail = detail if isinstance(detail, dict) else {}
    judge = artifacts.get("judge_result", {})
    if not isinstance(judge, dict):
        judge = {}
    plan = detail.get("rubric_reevaluation_plan", {})
    if not isinstance(plan, dict):
        plan = {}
    manifest = detail.get("manifest", {})
    if not isinstance(manifest, dict):
        manifest = {}
    current_judge_rubric = current_judge_rubric or load_independent_judge_rubric()
    stored_versions = manifest.get("rubric_versions", {})
    if not isinstance(stored_versions, dict):
        stored_versions = {}
    stored_judge = stored_versions.get("independent_judge", {})
    if not isinstance(stored_judge, dict):
        stored_judge = {}

    decision = str(judge.get("decision") or "NOT_RUN").strip().upper()
    if not decision:
        decision = "NOT_RUN"
    reasons = []
    plan_actions = [
        action for action in (plan.get("actions") or [])
        if isinstance(action, dict)
        and action.get("method") == "JUDGE_REEVALUATION"
        and (
            not action.get("target_case_ids")
            or str(case_id) in {str(item) for item in action.get("target_case_ids", [])}
        )
    ]
    if plan_actions:
        reasons.append("독립 LLM 평가 Rubric 기준 변경 대상입니다.")
    if decision == "NOT_RUN":
        reasons.append("저장된 독립 LLM 평가 결과가 없어 현재 기준 평가가 필요합니다.")
    elif decision == "ERROR":
        reasons.append("기존 독립 LLM 평가가 오류로 종료되어 재평가가 필요합니다.")
    elif decision != "PASS":
        reasons.append(f"기존 독립 LLM 판정이 {_voc_status_label(decision)} 상태입니다.")
    if not reasons:
        reasons.append("필수 재평가 대상은 아니며, Provider·모델 교차 확인이 필요할 때만 실행합니다.")

    stored_version = stored_judge.get("version") or "-"
    current_version = current_judge_rubric.get("version") or "-"
    stored_hash = stored_judge.get("sha256") or ""
    current_hash = current_judge_rubric.get("sha256") or current_judge_rubric.get("rubric_sha256") or ""
    rubric_changed = (
        stored_version != "-"
        and current_version != "-"
        and (
            stored_version != current_version
            or (bool(stored_hash) and bool(current_hash) and stored_hash != current_hash)
        )
    )
    blockers = _history_judge_review_blockers(judge, current_judge_rubric)
    has_error = bool(str(judge.get("error") or "").strip()) or decision == "ERROR"
    has_content_blockers = any(
        marker in item
        for item in blockers
        for marker in ("총점 기준 미달", "세부 항목", "위험 지적", "보완 권고", "즉시 실패")
    )
    review_focus = _history_judge_review_focus(
        decision=decision,
        rubric_changed=rubric_changed,
        has_error=has_error,
        has_content_blockers=has_content_blockers,
    )
    return {
        "decision": decision,
        "decision_label": _voc_status_label(decision),
        "score": judge.get("total_score"),
        "provider": judge.get("provider") or "-",
        "model": judge.get("model") or "-",
        "stored_rubric_version": stored_version,
        "current_rubric_version": current_version,
        "rubric_changed": rubric_changed,
        "reasons": reasons,
        "blockers": blockers,
        "reuses": [
            "저장된 Agent 파이프라인 결과",
            "동일 VOC 질문·요약·개선안",
            "저장된 Trace·Rule 증적",
        ],
        "not_changed": [
            "Agent 개선안 내용",
            "VOC 원본 데이터",
            "타당성 보완 입력",
        ],
        "review_focus": review_focus,
        "after": (
            "PASS면 개선안 타당성 평가와 QA 검토 흐름으로 이어갈 수 있고, "
            "검토 필요/실패면 Agent 개선안 보완 또는 RETEST를 검토합니다."
        ),
    }


def _render_history_judge_reevaluation_context(
    run_id: str,
    case_id: str,
    artifacts: dict,
    detail: dict,
) -> None:
    context = _history_judge_reevaluation_context_model(case_id, artifacts, detail)
    with st.container(border=True):
        heading, state = st.columns([2.6, 1], gap="small", vertical_alignment="center")
        with heading:
            st.markdown("#### 독립 LLM 재평가 기준")
            st.caption(
                "보완 입력을 자동 생성하지 않습니다. 저장된 Agent 결과를 그대로 두고 현재 독립 LLM 평가 기준과 선택 Provider로 다시 판정합니다."
            )
        with state:
            tone = "green" if context["decision"] == "PASS" else ("gray" if context["decision"] == "NOT_RUN" else "orange")
            st.markdown(f":{tone}-badge[기존 판정 {context['decision_label']}]", text_alignment="right")
            score = context["score"]
            st.caption(f"{context['provider']} · {context['model']} · {score if score is not None else '-'}점")

        columns = st.columns(3, gap="small")
        with columns[0].container(border=True, height=220):
            st.caption(":material/troubleshoot: 검토 필요 원인")
            for blocker in context["blockers"]:
                st.write(f"- {blocker}")
            st.caption(
                f"Judge Rubric {context['stored_rubric_version']} → {context['current_rubric_version']}"
            )
        with columns[1].container(border=True, height=220):
            st.caption(":material/replay: 재평가의 의미")
            for reason in context["reasons"]:
                st.write(f"- {reason}")
            st.caption("그대로 쓰는 입력")
            for item in context["reuses"]:
                st.write(f"- {item}")
        with columns[2].container(border=True, height=220):
            st.caption(":material/checklist: 재평가 후 볼 것")
            for item in context["review_focus"]:
                st.write(f"- {item}")
            st.caption("자동 보완하지 않는 것")
            for item in context["not_changed"]:
                st.write(f"- {item}")
            st.caption(context["after"])


def _render_history_case_artifact(
    artifact_name: str,
    artifacts: dict,
    *,
    run_id: str,
    case_id: str,
) -> None:
    artifact = artifacts.get(artifact_name, {})
    if not isinstance(artifact, dict):
        st.warning("선택한 증적의 데이터 형식을 읽을 수 없습니다.")
        return
    if artifact_name == "pipeline_result":
        _render_history_pipeline_artifact(artifact)
    elif artifact_name == "trace":
        _render_validity_trace_evidence({"trace": artifact})
    elif artifact_name == "rule_result":
        _render_history_rule_artifact(artifact)
    elif artifact_name == "judge_result":
        _render_history_judge_artifact(artifact)
    elif artifact_name == "validity_result":
        _render_history_validity_artifact(artifact)
    else:
        st.info("이 증적은 전용 화면이 없어 원본 형식으로 표시합니다.")

    with st.expander("원본 증적 JSON", icon=":material/data_object:"):
        st.caption(f"{run_id} · {case_id} · {artifact_name}")
        st.json(artifact)


def _history_run_detail_summary_cards(manifest: dict, summary: dict) -> list[dict]:
    total_cases = summary.get("total")
    if total_cases in (None, ""):
        total_cases = len(manifest.get("selected_case_ids", []) or [])
    counts = summary.get("counts", {}) if isinstance(summary.get("counts"), dict) else {}
    judge_counts = summary.get("judge_counts", {}) if isinstance(summary.get("judge_counts"), dict) else {}
    case_results = summary.get("case_results", []) if isinstance(summary.get("case_results"), list) else []
    approved_cases = sum(1 for item in case_results if isinstance(item, dict) and item.get("formal_approval"))
    issue_cases = int(counts.get("ERROR") or 0) + int(counts.get("FAIL") or 0)
    review_cases = int(counts.get("REVIEW_REQUIRED") or 0)
    pass_cases = int(counts.get("PASS") or 0)
    judge_pass_cases = int(judge_counts.get("PASS") or 0)
    judge_review_cases = int(judge_counts.get("REVIEW_REQUIRED") or 0)
    judge_error_cases = int(judge_counts.get("ERROR") or 0)
    run_type = _voc_status_label(manifest.get("run_type", "-"))
    try:
        action = voc_run_next_action(
            {
                **manifest,
                "counts": counts,
                "judge_counts": judge_counts,
                "selected_count": total_cases,
                "validity_state": summary.get("validity_state", "DRAFT"),
                "deployment_decision": summary.get("deployment_decision", "NOT_EVALUATED"),
            }
        )
    except Exception:
        action = {
            "label": "상세 확인",
            "detail": "선택 Run의 저장 증적을 확인합니다.",
        }
    return [
        {
            "icon": "task_alt",
            "label": "Run 상태",
            "value": _voc_status_label(manifest.get("status", "-")),
            "detail": f"유형: {run_type}",
            "tone": "green" if str(manifest.get("status", "")).upper() == "COMPLETED" else "orange",
        },
        {
            "icon": "format_list_numbered",
            "label": "대상 Case",
            "value": f"{total_cases}건",
            "detail": f"통과 {pass_cases} · 검토 {review_cases} · 실패/오류 {issue_cases}",
            "tone": "green" if total_cases and issue_cases == 0 else "orange",
        },
        {
            "icon": "psychology",
            "label": "독립 LLM 평가",
            "value": "사용" if manifest.get("judge_enabled") else "미사용",
            "detail": f"PASS {judge_pass_cases} · 검토 {judge_review_cases} · 오류 {judge_error_cases}",
            "tone": "green" if manifest.get("judge_enabled") and judge_error_cases == 0 else "gray",
        },
        {
            "icon": "approval",
            "label": "타당성·승인",
            "value": _voc_status_label(summary.get("validity_state", "DRAFT")),
            "detail": f"정식 승인 {approved_cases}건 · 개선안 타당성",
            "tone": "green" if approved_cases else "gray",
        },
        {
            "icon": "conversion_path",
            "label": "다음 액션",
            "value": action.get("label", "상세 확인"),
            "detail": action.get("detail", "선택 Run의 다음 처리 단계를 확인합니다."),
            "tone": _history_action_tone(action.get("label")),
        },
    ]


def _history_case_question_lookup() -> dict[str, str]:
    try:
        return {
            str(case.get("case_id") or ""): str(case.get("question") or case.get("name") or "")
            for case in load_unified_quality_cases().get("cases", [])
            if case.get("case_id")
        }
    except Exception:
        return {}


def _history_case_result_item(case_results: list, case_id: str) -> dict:
    return next(
        (
            item for item in case_results
            if isinstance(item, dict) and str(item.get("case_id") or "") == str(case_id)
        ),
        {},
    )


def _history_case_evidence_summary_cards(
    case_id: str,
    case_item: dict,
    artifacts: dict,
) -> list[dict]:
    pipeline = artifacts.get("pipeline_result", {}) if isinstance(artifacts.get("pipeline_result"), dict) else {}
    execution = pipeline.get("execution", {}) if isinstance(pipeline.get("execution"), dict) else {}
    pipeline_result = execution.get("result", {}) if isinstance(execution.get("result"), dict) else {}
    judge = artifacts.get("judge_result", {}) if isinstance(artifacts.get("judge_result"), dict) else {}
    validity = artifacts.get("validity_result", {}) if isinstance(artifacts.get("validity_result"), dict) else {}
    try:
        action = voc_case_next_action({**case_item, "case_id": case_id})
    except Exception:
        action = {
            "label": "Case 증적 확인",
            "detail": "저장된 Agent 결과와 평가 증적을 확인합니다.",
        }
    question_lookup = _history_case_question_lookup()
    question = (
        question_lookup.get(str(case_id))
        or execution.get("question")
        or case_item.get("message")
        or "-"
    )
    pipeline_ok = bool(pipeline) and execution.get("ok") is not False and pipeline_result.get("ok") is not False
    return [
        {
            "icon": "help",
            "label": "Case",
            "value": case_id,
            "detail": _manual_pipeline_compact_text(question, 68),
            "tone": "blue",
        },
        {
            "icon": "account_tree",
            "label": "Agent 파이프라인",
            "value": "완료" if pipeline_ok else "확인 필요",
            "detail": f"Trace {_history_safe_display_text(pipeline_result.get('trace'), fallback='미연결')}",
            "tone": "green" if pipeline_ok else "orange",
        },
        {
            "icon": "psychology",
            "label": "독립 LLM 평가",
            "value": _voc_status_label(judge.get("decision", case_item.get("judge_status", "NOT_RUN"))),
            "detail": f"{judge.get('total_score', case_item.get('judge_score', '-'))}점",
            "tone": "green" if str(judge.get("decision") or case_item.get("judge_status") or "").upper() == "PASS" else "orange",
        },
        {
            "icon": "approval",
            "label": "개선안 타당성",
            "value": _voc_status_label(validity.get("decision", case_item.get("validity_status", "NOT_RUN"))),
            "detail": f"{validity.get('total_score', case_item.get('validity_score', '-'))}점",
            "tone": "green" if str(validity.get("decision") or case_item.get("validity_status") or "").upper() == "AI_PASS" else "orange",
        },
        {
            "icon": "conversion_path",
            "label": "다음 액션",
            "value": action.get("label", "상세 확인"),
            "detail": action.get("detail", "Case 증적을 확인합니다."),
            "tone": _history_action_tone(action.get("label")),
        },
    ]


def _render_history_integrity_status(integrity: dict) -> None:
    errors = [str(error) for error in integrity.get("errors", []) if str(error or "").strip()]
    warnings = [str(warning) for warning in integrity.get("warnings", []) if str(warning or "").strip()]
    if integrity.get("ok"):
        st.caption(":material/check_circle: 증적 무결성 정상 · Run 폴더, index, Case 증적이 일치합니다.")
    else:
        st.error("증적 무결성 확인이 필요합니다.", icon=":material/error:")
        for error in errors:
            st.caption(f"- {error}")
    if warnings:
        with st.expander("무결성 경고 보기", icon=":material/warning:"):
            for warning in warnings:
                st.write(f"- {warning}")


def _render_voc_run_detail(run_id: str):
    detail = load_voc_run_history_detail(run_id)
    manifest = detail.get("manifest", {})
    summary = detail.get("summary", {})
    integrity = detail.get("integrity", {})
    with st.container(border=True):
        status_label = _voc_status_label(manifest.get("status", "-"))
        status_tone = "green" if str(manifest.get("status", "")).upper() == "COMPLETED" else "orange"
        _render_voc_section_heading(
            "실행 상세 · Run 요약",
            "상태, 대상 범위, 다음 액션만 먼저 확인합니다.",
            icon="history",
            badges=((status_label, status_tone), (_voc_status_label(manifest.get("run_type", "-")), "gray")),
            right_caption=(
                f"Run {run_id} · "
                f"{_history_table_timestamp(manifest.get('finished_at') or manifest.get('started_at'))}"
            ),
        )
        _render_voc_summary_cards(
            _history_run_detail_summary_cards(manifest, summary),
            columns=5,
            height=124,
        )
        _render_history_integrity_status(integrity)

    with st.expander("검증 회차 범위 보기", expanded=False, icon=":material/fact_check:"):
        _render_history_verification_scope(manifest, summary)

    view = st.segmented_control(
        "상세 구분",
        ["Case 결과", "Case 증적", "실행 정보"],
        default="Case 결과",
        key=f"voc_history_detail_view_{run_id}",
    )
    case_results = summary.get("case_results", [])
    if view == "Case 결과":
        question_lookup = _history_case_question_lookup()
        display_rows = []
        detailed_rows = []
        for item in case_results:
            action = voc_case_next_action(item)
            case_id = str(item.get("case_id") or "-")
            question = question_lookup.get(case_id, "")
            row = {
                "Case ID": case_id,
                "질문": _manual_pipeline_compact_text(question or item.get("message") or "-", 58),
                "상태": _voc_status_label(item.get("status", "NOT_RUN")),
                "다음 액션": action["label"],
                "독립 LLM": _voc_status_label(item.get("judge_status", "NOT_RUN")),
                "LLM 점수": item.get("judge_score"),
                "타당성": _voc_status_label(item.get("validity_status", "NOT_RUN")),
                "타당성 점수": item.get("validity_score"),
                "승인": "승인" if item.get("formal_approval") else _voc_status_label(item.get("approval_state", "DRAFT")),
                "종료": _history_table_timestamp(item.get("finished_at")),
            }
            display_rows.append(row)
            detailed_rows.append(
                {
                    **row,
                    "수행 모드": item.get("mode", "-") or "-",
                    "시도": item.get("attempt_count", 0),
                    "독립성": _judge_independence_grade_label(item.get("judge_independence_grade", "-")),
                    "메시지": item.get("message", "-") or "-",
                    "시작": _history_table_timestamp(item.get("started_at")),
                }
            )
        rows = pd.DataFrame(display_rows)
        if rows.empty:
            st.info("아직 저장된 Case 결과가 없습니다.")
        else:
            _render_voc_section_heading(
                "Case 결과",
                "Case, 질문, 현재 상태와 다음 액션만 먼저 표시합니다.",
                icon="table_chart",
                badges=((f"{len(rows)}건", "blue"),),
            )
            st.dataframe(
                rows,
                hide_index=True,
                width="stretch",
                column_order=(
                    "Case ID",
                    "질문",
                    "상태",
                    "다음 액션",
                    "독립 LLM",
                    "LLM 점수",
                    "타당성",
                    "타당성 점수",
                    "승인",
                    "종료",
                ),
                column_config={
                    "Case ID": st.column_config.TextColumn(width="small", pinned=True),
                    "질문": st.column_config.TextColumn(width="large"),
                    "상태": st.column_config.TextColumn(width="small"),
                    "다음 액션": st.column_config.TextColumn(width="medium"),
                    "독립 LLM": st.column_config.TextColumn(width="small"),
                    "LLM 점수": st.column_config.ProgressColumn(
                        min_value=0, max_value=100, format="%g점", width="small"
                    ),
                    "타당성": st.column_config.TextColumn(width="small"),
                    "타당성 점수": st.column_config.ProgressColumn(
                        min_value=0, max_value=100, format="%g점", width="small"
                    ),
                    "승인": st.column_config.TextColumn(width="small"),
                    "종료": st.column_config.TextColumn(width="small"),
                },
            )
            with st.expander("Case 결과 상세 컬럼 보기", icon=":material/view_column:"):
                st.dataframe(
                    pd.DataFrame(detailed_rows),
                    hide_index=True,
                    width="stretch",
                    column_config={
                        "Case ID": st.column_config.TextColumn(width="small", pinned=True),
                        "질문": st.column_config.TextColumn(width="large"),
                        "시도": st.column_config.NumberColumn(format="%d", width="small"),
                        "독립성": st.column_config.TextColumn(width=156),
                        "메시지": st.column_config.TextColumn(width="large"),
                    },
                )
    elif view == "실행 정보":
        _render_history_execution_info(manifest, summary)
    else:
        completed_case_ids = [item.get("case_id") for item in case_results if item.get("case_id")]
        if not completed_case_ids:
            st.info("조회할 Case 증적이 없습니다.")
        else:
            with st.container(border=True):
                _render_voc_section_heading(
                    "Case 증적",
                    "",
                    icon="folder_open",
                    badges=((f"{len(completed_case_ids)}건", "blue"),),
                )
                control_cols = st.columns([0.85, 1.15], gap="small", vertical_alignment="bottom")
                with control_cols[0]:
                    selected_case_id = st.selectbox(
                        "Case",
                        completed_case_ids,
                        key=f"voc_history_case_{run_id}",
                    )
                reevaluation_result_key = _history_judge_reevaluation_result_key(run_id, selected_case_id)
                artifacts = load_voc_case_history_detail(run_id, selected_case_id)
                artifact_names = [
                    name for name in
                    ("pipeline_result", "trace", "rule_result", "judge_result", "validity_result")
                    if name in artifacts
                ]
                with control_cols[1]:
                    artifact_name = None
                    if artifact_names:
                        artifact_key = f"voc_history_artifact_{run_id}"
                        preferred_artifact_key = f"voc_history_artifact_pending_{run_id}"
                        preferred_artifact = st.session_state.pop(preferred_artifact_key, None)
                        if preferred_artifact in artifact_names:
                            st.session_state[artifact_key] = preferred_artifact
                        if st.session_state.get(artifact_key) not in artifact_names:
                            st.session_state[artifact_key] = artifact_names[0]
                        artifact_name = st.selectbox(
                            "증적",
                            artifact_names,
                            key=artifact_key,
                            format_func=lambda value: HISTORY_ARTIFACT_LABELS.get(value, value),
                        )
                    else:
                        st.caption("증적 파일 없음")
                case_item = _history_case_result_item(case_results, selected_case_id)
                _render_voc_summary_cards(
                    _history_case_evidence_summary_cards(selected_case_id, case_item, artifacts),
                    columns=5,
                    height=124,
                )
            if artifact_names and artifact_name:
                with st.container(border=True):
                    _render_voc_section_heading(
                        HISTORY_ARTIFACT_LABELS.get(artifact_name, artifact_name),
                        "",
                        icon="description",
                        badges=((selected_case_id, "blue"),),
                    )
                    _render_history_case_artifact(
                        artifact_name,
                        artifacts,
                        run_id=run_id,
                        case_id=selected_case_id,
                    )
            else:
                st.warning("Case 증적 JSON을 읽을 수 없습니다.")
            pipeline = artifacts.get("pipeline_result", {})
            execution = pipeline.get("execution", {})
            result = execution.get("result", {}) if isinstance(execution, dict) else {}
            if pipeline.get("mode") == "voc" and execution.get("ok") and result.get("ok"):
                reevaluation_job_key = _history_judge_reevaluation_job_key(run_id, selected_case_id)
                reevaluation_running = bool(st.session_state.get(reevaluation_job_key))
                reevaluation_error = st.session_state.pop(
                    _history_judge_reevaluation_error_key(run_id, selected_case_id),
                    None,
                )
                if reevaluation_error:
                    st.error(reevaluation_error)
                reevaluation_result = st.session_state.get(reevaluation_result_key)
                with st.expander(
                    "독립 LLM 재평가",
                    expanded=reevaluation_running or bool(reevaluation_result),
                    icon=":material/replay:",
                ):
                    _live_history_judge_reevaluation(run_id, selected_case_id)
                    _render_history_judge_reevaluation_result(
                        run_id,
                        selected_case_id,
                        reevaluation_result,
                    )
                    _render_history_judge_reevaluation_context(
                        run_id,
                        selected_case_id,
                        artifacts,
                        detail,
                    )
                    reevaluation_config = _judge_config_controls(
                        f"history_{run_id}_{selected_case_id}"
                    )
                    if st.button(
                        "재평가 진행 중" if reevaluation_running else "독립 LLM 재평가",
                        icon=":material/replay:",
                        disabled=reevaluation_running or not reevaluation_config.get("enabled"),
                        key=f"reevaluate_{run_id}_{selected_case_id}",
                    ):
                        st.session_state[reevaluation_job_key] = start_background_job(
                            "history-judge-reevaluation",
                            f"{run_id}::{selected_case_id}",
                            _execute_history_judge_reevaluation,
                            run_id,
                            selected_case_id,
                            reevaluation_config,
                            progress={
                                "percent": 6,
                                "stage": "재평가 대기열 등록",
                                "detail": "독립 LLM 재평가 작업을 백그라운드로 시작했습니다.",
                            },
                        )
                        st.rerun()

    evidence = download_voc_run_evidence(run_id)
    with st.container(horizontal=True, horizontal_alignment="right"):
        st.download_button(
            "Run 전체 증적 ZIP 다운로드",
            data=evidence,
            file_name=f"{run_id}.zip",
            mime="application/zip",
            icon=":material/download:",
            key=f"download_{run_id}",
        )


def _history_retest_sort_key(item: dict) -> tuple[str, str]:
    return (str(item.get("started_at") or ""), str(item.get("run_id") or ""))


def _history_retest_pair_basis(history: list[dict], run_item: dict) -> dict:
    run_id = str(run_item.get("run_id") or "")
    by_run_id = {str(item.get("run_id") or ""): item for item in history if item.get("run_id")}
    selected = by_run_id.get(run_id, run_item)
    if not run_id:
        return {
            "found": False,
            "enabled": False,
            "state_label": "비교 대상 없음",
            "detail": "선택한 Run ID를 확인할 수 없습니다.",
        }
    if str(selected.get("status") or "") == "RUNNING":
        return {
            "found": False,
            "enabled": False,
            "state_label": "실행 중",
            "detail": "실행 중인 Run은 종료 후 재시험 전후 비교를 확인할 수 있습니다.",
        }

    if str(selected.get("run_type") or "") == "RETEST":
        parent_run_id = str(selected.get("parent_run_id") or "")
        parent = by_run_id.get(parent_run_id)
        if not parent_run_id:
            return {
                "found": False,
                "enabled": False,
                "state_label": "원본 연결 없음",
                "detail": "재시험 Run에 원본 Run 연결 정보가 없어 자동 비교할 수 없습니다.",
            }
        if not parent:
            return {
                "found": False,
                "enabled": False,
                "state_label": "원본 없음",
                "detail": f"원본 Run {parent_run_id}를 현재 이력에서 찾을 수 없습니다.",
            }
        if str(parent.get("status") or "") == "RUNNING":
            return {
                "found": False,
                "enabled": False,
                "state_label": "원본 실행 중",
                "detail": "원본 Run이 아직 실행 중이라 전후 비교를 할 수 없습니다.",
            }
        return {
            "found": True,
            "enabled": True,
            "mode": "selected_retest",
            "state_label": "비교 대상 확인",
            "detail": f"선택 Run은 재시험입니다. 원본 {parent_run_id}와 비교합니다.",
            "baseline_run_id": parent_run_id,
            "candidate_run_id": run_id,
            "candidate_count": 1,
            "candidate_started_at": selected.get("started_at", ""),
        }

    linked_retests = [
        item for item in history
        if str(item.get("run_type") or "") == "RETEST"
        and str(item.get("parent_run_id") or "") == run_id
        and str(item.get("status") or "") != "RUNNING"
    ]
    if not linked_retests:
        return {
            "found": False,
            "enabled": False,
            "state_label": "비교 대상 없음",
            "detail": "선택 Run을 원본으로 하는 완료된 재시험이 아직 없습니다.",
        }
    linked_retests.sort(key=_history_retest_sort_key, reverse=True)
    latest_retest = linked_retests[0]
    candidate_run_id = str(latest_retest.get("run_id") or "")
    return {
        "found": True,
        "enabled": True,
        "mode": "selected_baseline",
        "state_label": "비교 대상 확인",
        "detail": f"연결 재시험 {len(linked_retests)}건 중 최신 Run {candidate_run_id}와 비교합니다.",
        "baseline_run_id": run_id,
        "candidate_run_id": candidate_run_id,
        "candidate_count": len(linked_retests),
        "candidate_started_at": latest_retest.get("started_at", ""),
    }


def _history_retest_comparison_plan(history: list[dict], run_item: dict) -> dict:
    basis = _history_retest_pair_basis(history, run_item)
    if not basis.get("found"):
        return {**basis, "pair_key": "", "comparison": None}

    baseline_run_id = str(basis.get("baseline_run_id") or "")
    candidate_run_id = str(basis.get("candidate_run_id") or "")
    pair_key = f"{baseline_run_id}::{candidate_run_id}"
    try:
        comparison = compare_voc_runs(baseline_run_id, candidate_run_id)
    except Exception as exc:
        return {
            **basis,
            "enabled": False,
            "pair_key": pair_key,
            "comparison": None,
            "state_label": "비교 증적 오류",
            "detail": f"비교 증적을 읽을 수 없습니다: {type(exc).__name__}",
        }
    if not comparison.get("compatible"):
        differences = ", ".join(comparison.get("compatibility_differences") or [])
        return {
            **basis,
            "enabled": False,
            "pair_key": pair_key,
            "comparison": comparison,
            "state_label": "기준 불일치",
            "detail": f"Catalog·TC·Rubric 또는 부모 Run 조건이 달라 비교가 차단됐습니다: {differences or '-'}",
        }
    return {
        **basis,
        "enabled": True,
        "pair_key": pair_key,
        "comparison": comparison,
        "state_label": "비교 가능",
    }


def _render_retest_comparison_result(comparison: dict) -> None:
    st.success("동일 Catalog·TC·Rubric과 부모 Run 연결이 확인됐습니다.")
    count_rows = pd.DataFrame(
        [
            {
                "상태": _voc_status_label(item.get("status")),
                "원본": int(item.get("baseline") or 0),
                "재시험": int(item.get("candidate") or 0),
                "증감": int(item.get("delta") or 0),
            }
            for item in comparison.get("count_comparison", [])
        ]
    )
    st.dataframe(count_rows, hide_index=True, width="stretch")
    case_rows = [
        {
            "Case ID": item.get("case_id", "-"),
            "원본 상태": _voc_status_label(item.get("baseline_status")),
            "재시험 상태": _voc_status_label(item.get("candidate_status")),
            "변경": "변경" if item.get("changed") else "동일",
            "원본 시도": int(item.get("baseline_attempts") or 0),
            "재시험 시도": int(item.get("candidate_attempts") or 0),
        }
        for item in comparison.get("case_comparison", [])
    ]
    changed = [item for item in case_rows if item["변경"] == "변경"]
    st.markdown(f"#### 상태 변경 Case · {len(changed)}건")
    st.dataframe(pd.DataFrame(changed or case_rows), hide_index=True, width="stretch")


def _render_retest_comparison(history: list[dict], plan: dict | None = None):
    st.markdown("### 재시험 전후 비교")
    if plan:
        st.caption("선택 Run 기준으로 원본 Run과 연결 재시험 Run을 자동 매칭해 비교합니다.")
        comparison = plan.get("comparison") or {}
        if not plan.get("enabled") or not comparison.get("compatible"):
            st.warning(plan.get("detail") or "재시험 전후 비교 조건이 맞지 않습니다.")
            return
        st.caption(
            f"원본 Run `{comparison.get('baseline_run_id')}` · "
            f"재시험 Run `{comparison.get('candidate_run_id')}`"
        )
        _render_retest_comparison_result(comparison)
        return

    st.caption("원본 Run과 연결된 재시험 Run의 Case 상태 변화를 비교합니다. Catalog·TC·Rubric 버전이 다르면 비교를 차단합니다.")
    completed = [item for item in history if item.get("status") != "RUNNING"]
    retests = [item for item in completed if item.get("run_type") == "RETEST"]
    if not retests:
        st.info("비교 가능한 재시험 Run이 아직 없습니다. 실패·오류 재실행 후 사용할 수 있습니다.")
        return
    baseline_id = st.selectbox(
        "원본 Run",
        [item["run_id"] for item in completed],
        key="voc_history_baseline_run",
    )
    candidate_id = st.selectbox(
        "재시험 Run",
        [item["run_id"] for item in retests],
        key="voc_history_retest_run",
    )
    comparison = compare_voc_runs(baseline_id, candidate_id)
    if not comparison["compatible"]:
        st.error(
            "재시험 전후 비교 조건이 일치하지 않습니다: "
            + ", ".join(comparison["compatibility_differences"])
        )
        return
    _render_retest_comparison_result(comparison)


HISTORY_SELECTED_RUN_ID_KEY = "voc_history_selected_run_id"
HISTORY_TABLE_KEY = "voc_history_table"
HISTORY_TABLE_NONCE_KEY = "voc_history_table_nonce"
HISTORY_TABLE_SIGNATURE_KEY = "voc_history_table_signature"
HISTORY_DETAIL_DIALOG_RUN_ID_KEY = "voc_history_detail_dialog_run_id"
HISTORY_RETEST_COMPARE_TARGET_KEY = "voc_history_retest_compare_target"
VALIDITY_CANDIDATE_TABLE_KEY = "voc_validity_candidate_table"
VALIDITY_CANDIDATE_TABLE_NONCE_KEY = "voc_validity_candidate_table_nonce"
VALIDITY_DETAIL_DIALOG_CANDIDATE_KEY = "voc_validity_detail_dialog_candidate_key"
VOC_QUALITY_MENU_NAME = "VOC 품질진단"
VOC_VALIDITY_PAGE_NAME = "개선안 타당성 검증"
VOC_HISTORY_PAGE_NAME = "수행 이력"
VOC_BATCH_PAGE_NAME = "일괄 TC 수행"
VOC_REPORT_PAGE_NAME = "품질 보고서"
VOC_ACCEPTANCE_PAGE_NAME = "최종 인수·시연"
VOC_HISTORY_NAVIGABLE_VALIDITY_ACTIONS = {
    "RUN_VALIDITY",
    "REWORK_AND_RETEST",
    "QA_REVIEW",
    "BUSINESS_APPROVAL",
    "CHECK_REMAINING_CASES",
}


def _open_history_detail_dialog(run_id: str) -> None:
    if run_id:
        st.session_state[HISTORY_DETAIL_DIALOG_RUN_ID_KEY] = run_id


def _history_table_widget_key() -> str:
    nonce = int(st.session_state.get(HISTORY_TABLE_NONCE_KEY, 0) or 0)
    return f"{HISTORY_TABLE_KEY}_{nonce}"


def _reset_history_table_selection() -> None:
    st.session_state[HISTORY_TABLE_NONCE_KEY] = int(
        st.session_state.get(HISTORY_TABLE_NONCE_KEY, 0) or 0
    ) + 1


def _dismiss_history_detail_dialog() -> None:
    st.session_state.pop(HISTORY_DETAIL_DIALOG_RUN_ID_KEY, None)
    _reset_history_table_selection()


@st.dialog(
    "수행 이력 상세",
    width="large",
    icon=":material/history:",
    on_dismiss=_dismiss_history_detail_dialog,
)
def _render_history_detail_dialog(run_id: str) -> None:
    _render_contextual_jira_action_menu(
        area_label="수행 이력 상세",
        target_label="선택 Run",
        run_id=run_id,
        status_label="상세 팝업",
        extra_detail=(
            "수행 이력 상세 팝업에서 선택한 Run 기준입니다.\n"
            f"{_jira_context_run_detail_extra(run_id)}"
        ),
        key=f"history_dialog_{run_id}",
    )
    _render_voc_run_detail(run_id)


def _history_status_total(history: list[dict], status: str) -> int:
    return sum(int((item.get("counts") or {}).get(status, 0) or 0) for item in history)


def _history_judge_total(history: list[dict], status: str) -> int:
    return sum(int((item.get("judge_counts") or {}).get(status, 0) or 0) for item in history)


def _history_run_needs_followup(item: dict) -> bool:
    counts = item.get("counts") or {}
    judge_counts = item.get("judge_counts") or {}
    validity_state = str(item.get("validity_state") or "DRAFT")
    deployment_decision = str(item.get("deployment_decision") or "미판정")
    return (
        int(counts.get("REVIEW_REQUIRED", 0) or 0) > 0
        or int(counts.get("FAIL", 0) or 0) > 0
        or int(counts.get("ERROR", 0) or 0) > 0
        or int(judge_counts.get("ERROR", 0) or 0) > 0
        or validity_state in {"REVISION_REQUIRED", "REJECTED"}
        or deployment_decision in {"REVISION_REQUIRED", "REJECTED"}
    )


def _history_summary_cards(history: list[dict]) -> list[dict]:
    pass_count = _history_status_total(history, "PASS")
    fail_count = _history_status_total(history, "FAIL")
    error_count = _history_status_total(history, "ERROR")
    review_count = _history_status_total(history, "REVIEW_REQUIRED")
    decided_count = pass_count + fail_count
    pass_rate = round(pass_count / decided_count * 100, 1) if decided_count else None
    judge_not_run = _history_judge_total(history, "NOT_RUN")
    judge_error = _history_judge_total(history, "ERROR")
    followup_runs = sum(1 for item in history if _history_run_needs_followup(item))
    retest_runs = sum(1 for item in history if item.get("parent_run_id"))
    rubric_drift_runs = sum(1 for item in history if item.get("reevaluation_required"))
    undecided_runs = sum(
        1 for item in history
        if str(item.get("deployment_decision") or "미판정") == "미판정"
    )
    return [
        {
            "icon": "check_circle",
            "label": "통과 Case",
            "value": f"{pass_count}건",
            "detail": f"통과율 {pass_rate:.1f}%" if pass_rate is not None else "아직 통과/실패 판정 없음",
        },
        {
            "icon": "rate_review",
            "label": "검토 필요 Case",
            "value": f"{review_count}건",
            "detail": "근거 부족·사람 확인 필요",
        },
        {
            "icon": "error",
            "label": "실패·오류 Case",
            "value": f"{fail_count + error_count}건",
            "detail": f"실패 {fail_count} · 오류 {error_count}",
        },
        {
            "icon": "fact_check",
            "label": "독립 LLM 평가 필요",
            "value": f"{judge_not_run + judge_error}건",
            "detail": f"미수행 {judge_not_run} · 오류 {judge_error}",
        },
        {
            "icon": "conversion_path",
            "label": "후속 조치 Run",
            "value": f"{followup_runs}건",
            "detail": f"미판정 {undecided_runs} · 재시험 {retest_runs} · 재평가 {rubric_drift_runs}",
        },
    ]


def _render_history_summary_cards(history: list[dict]) -> None:
    cards = _history_summary_cards(history)
    columns = st.columns(len(cards), gap="small")
    for column, card in zip(columns, cards, strict=False):
        with column.container(border=True, height=136):
            st.caption(f":material/{card['icon']}: {card['label']}")
            st.markdown(f"#### {card['value']}")
            st.caption(card["detail"])


def _history_action_badge(action: dict) -> str:
    tone = str(action.get("tone") or "gray")
    color = {
        "blue": "blue",
        "green": "green",
        "red": "red",
        "orange": "orange",
        "gray": "gray",
    }.get(tone, "gray")
    return f":{color}-badge[{action.get('label', '다음 액션 확인')}]"


def _history_rubric_badge(drift: dict) -> str:
    tone = str(drift.get("tone") or "gray")
    color = {
        "green": "green",
        "red": "red",
        "orange": "orange",
        "gray": "gray",
    }.get(tone, "gray")
    return f":{color}-badge[{drift.get('status', '기준 확인')}]"


def _history_rubric_plan_label(status: str) -> str:
    labels = {
        "NOT_REQUIRED": "재평가 불필요",
        "RETEST_RECOMMENDED": "재시험 권장",
        "REEVALUATION_READY": "재평가 가능",
    }
    return labels.get(str(status or ""), "계획 없음")


def _history_rubric_plan_next_targets(run_id: str, plan: dict | None) -> list[dict]:
    plan = plan if isinstance(plan, dict) else {}
    targets: list[dict] = []
    for action in plan.get("actions") or []:
        if not isinstance(action, dict):
            continue
        method = str(action.get("method") or "")
        case_ids = [
            str(case_id)
            for case_id in (action.get("target_case_ids") or [])
            if str(case_id).strip()
        ]
        first_case_id = case_ids[0] if case_ids else ""
        target_count = int(action.get("target_count") or len(case_ids) or 0)
        label = str(action.get("label") or "Rubric")
        if method == "JUDGE_REEVALUATION":
            targets.append(
                {
                    "enabled": bool(first_case_id),
                    "page": "history_detail",
                    "run_id": run_id,
                    "case_id": first_case_id,
                    "action_code": "RUBRIC_JUDGE_REEVALUATION",
                    "button_label": "독립 LLM 재평가 열기",
                    "title": "독립 LLM 기준 변경",
                    "detail": (
                        f"{target_count}건 대상입니다. 수행 이력 상세 팝업에서 저장된 Agent 파이프라인 결과를 "
                        "보완 없이 그대로 사용해 현재 독립 LLM Rubric으로 다시 판정합니다."
                    ),
                    "icon": "rule",
                }
            )
        elif method == "VALIDITY_REEVALUATION":
            targets.append(
                {
                    "enabled": bool(first_case_id),
                    "page": VOC_VALIDITY_PAGE_NAME,
                    "run_id": run_id,
                    "case_id": first_case_id,
                    "action_code": "RUBRIC_VALIDITY_REEVALUATION",
                    "button_label": "타당성 재평가로 이동",
                    "title": "타당성 기준 변경",
                    "detail": (
                        f"{target_count}건 대상입니다. 저장된 개선안과 독립 LLM 평가 증적을 유지한 채 "
                        "현재 타당성 Rubric 기준으로 다시 평가합니다."
                    ),
                    "icon": "fact_check",
                }
            )
        elif method == "RETEST_REQUIRED":
            targets.append(
                {
                    "enabled": True,
                    "page": VOC_BATCH_PAGE_NAME,
                    "run_id": run_id,
                    "case_id": first_case_id,
                    "action_code": "RUBRIC_RETEST_REQUIRED",
                    "button_label": "RETEST 준비로 이동",
                    "title": f"{label} 기준 변경",
                    "detail": (
                        f"{target_count}건 대상입니다. Agent 파이프라인 기준 변경은 기존 답변 재점수화보다 "
                        "부모 Run을 둔 RETEST로 다시 실행하는 편이 안전합니다."
                    ),
                    "icon": "replay",
                }
            )
    return targets


def _save_history_rubric_reevaluation_plan(run_id: str) -> None:
    with st.spinner("Rubric 기준 영향 계획을 저장하고 있습니다..."):
        saved = save_voc_rubric_reevaluation_plan(run_id)
    st.session_state.voc_rubric_reevaluation_plan_result = saved.get("plan", {})
    _load_voc_history_rows.clear()


def _render_history_rubric_reevaluation_plan(
    run_id: str,
    detail: dict,
    rubric_drift: dict,
) -> None:
    plan = detail.get("rubric_reevaluation_plan", {}) if isinstance(detail, dict) else {}
    plan_required = bool(rubric_drift.get("requires_reevaluation"))
    if not plan_required and not plan:
        return

    with st.container(border=True):
        heading, state_col, action_col = st.columns(
            [2.2, 1.1, 0.95],
            gap="small",
            vertical_alignment="center",
        )
        with heading:
            st.markdown("#### Rubric 기준 영향 계획")
            st.caption("원본 Agent 파이프라인 결과는 보존하고, 현재 Rubric 기준 변경으로 필요한 후속 실행만 정리합니다.")
        with state_col:
            if plan:
                saved_at = _history_table_timestamp(plan.get("saved_at"))
                st.markdown(f":green-badge[계획 저장됨]", text_alignment="right")
                st.caption(saved_at or "-")
            else:
                st.markdown(":orange-badge[계획 미작성]", text_alignment="right")
                st.caption("기준 변경 영향 확인 필요")
        with action_col:
            if st.button(
                "기준 영향 계획 갱신" if plan else "기준 영향 계획 저장",
                icon=":material/assignment:",
                type="primary" if plan_required and not plan else "secondary",
                disabled=not plan_required,
                width="stretch",
                key=f"voc_history_save_rubric_plan_{run_id}",
            ):
                _save_history_rubric_reevaluation_plan(run_id)
                st.rerun()

        if plan:
            status_label = _history_rubric_plan_label(plan.get("status"))
            st.caption(f"{status_label} · {plan.get('recommendation', '-')}")
            actions = plan.get("actions") or []
            if actions:
                table = pd.DataFrame(
                    [
                        {
                            "기준": item.get("label", "-"),
                            "처리 방식": item.get("method_label", "-"),
                            "대상": f"{item.get('target_count', 0)}건",
                            "연결 화면": item.get("menu", "-"),
                            "설명": item.get("description", "-"),
                        }
                        for item in actions
                    ]
                )
                st.dataframe(
                    table,
                    hide_index=True,
                    width="stretch",
                    height=min(176, 46 + len(table) * 42),
                    column_config={
                        "기준": st.column_config.TextColumn(width=112, pinned=True),
                        "처리 방식": st.column_config.TextColumn(width=112),
                        "대상": st.column_config.TextColumn(width=72),
                        "연결 화면": st.column_config.TextColumn(width=148),
                        "설명": st.column_config.TextColumn(width="large"),
                        },
                    )
                targets = _history_rubric_plan_next_targets(run_id, plan)
                if targets:
                    st.markdown("##### 기준 영향 계획 저장 후 다음 실행")
                    target_columns = st.columns(min(3, len(targets)), gap="small")
                    for index, target in enumerate(targets):
                        column = target_columns[index % len(target_columns)]
                        with column.container(border=True, height=148):
                            st.caption(f":material/{target.get('icon', 'arrow_forward')}: {target.get('title', '다음 실행')}")
                            st.markdown(f"**{target['button_label']}**")
                            st.caption(target.get("detail", "계획에 따라 다음 화면으로 이동합니다."))
                            if st.button(
                                "열기",
                                icon=":material/arrow_forward:",
                                disabled=not target.get("enabled"),
                                width="stretch",
                                key=f"voc_history_rubric_target_{run_id}_{index}_{target.get('action_code')}",
                            ):
                                _apply_history_next_action_target(target)
                                st.rerun()
            skipped = plan.get("skipped_cases") or []
            if skipped:
                with st.expander(f"재평가 제외 Case {len(skipped)}건", expanded=False):
                    st.dataframe(
                        pd.DataFrame(skipped),
                        hide_index=True,
                        width="stretch",
                        column_config={
                            "case_id": st.column_config.TextColumn("Case ID", width=90),
                            "reason": st.column_config.TextColumn("제외 사유", width="large"),
                        },
                    )


def _history_case_action_groups(case_results: list[dict]) -> dict[str, dict]:
    groups: dict[str, dict] = {}
    for case in case_results:
        action = voc_case_next_action(case)
        code = action["code"]
        group = groups.setdefault(
            code,
            {
                "action": action,
                "count": 0,
                "case_ids": [],
            },
        )
        group["count"] += 1
        if case.get("case_id"):
            group["case_ids"].append(str(case["case_id"]))
    return groups


def _history_priority_case_action(case_results: list[dict]) -> dict:
    groups = _history_case_action_groups(case_results)
    priority = (
        "CHECK_PIPELINE_ERROR",
        "REVIEW_PIPELINE_RESULT",
        "RUN_JUDGE",
        "RUN_VALIDITY",
        "REWORK_AND_RETEST",
        "QA_REVIEW",
        "BUSINESS_APPROVAL",
        "CHECK_REMAINING_CASES",
        "REPORT_READY",
        "NO_ACTION",
    )
    for code in priority:
        if code in groups:
            group = groups[code]
            action = dict(group["action"])
            samples = ", ".join(group["case_ids"][:4]) or "-"
            if len(group["case_ids"]) > 4:
                samples += f" 외 {len(group['case_ids']) - 4}건"
            action["count"] = group["count"]
            action["case_ids"] = samples
            return action
    return voc_case_next_action({})


def _history_case_for_next_action(case_results: list[dict], action_code: str) -> dict | None:
    if not case_results:
        return None
    preferred_codes = [action_code]
    if action_code == "CHECK_REMAINING_CASES":
        preferred_codes = [
            "RUN_VALIDITY",
            "REWORK_AND_RETEST",
            "QA_REVIEW",
            "BUSINESS_APPROVAL",
            "CHECK_REMAINING_CASES",
        ]
    for code in preferred_codes:
        for case in case_results:
            if voc_case_next_action(case)["code"] == code:
                return case
    return case_results[0]


def _history_next_action_target(
    run_item: dict,
    case_results: list[dict],
    run_action: dict,
    case_action: dict,
) -> dict:
    run_id = str(run_item.get("run_id") or "")
    action_code = str(case_action.get("code") or run_action.get("code") or "")
    if action_code in VOC_HISTORY_NAVIGABLE_VALIDITY_ACTIONS:
        target_case = _history_case_for_next_action(case_results, action_code)
        if not target_case:
            return {
                "enabled": False,
                "button_label": "다음 액션 대상 없음",
                "detail": "선택 Run에서 이어서 처리할 Case를 찾지 못했습니다.",
            }
        case_id = str(target_case.get("case_id") or "")
        action_label = {
            "RUN_VALIDITY": "개선안 타당성 평가 진행",
            "REWORK_AND_RETEST": "보완·재시험 진행",
            "QA_REVIEW": "QA 검토 진행",
            "BUSINESS_APPROVAL": "업무 승인 진행",
            "CHECK_REMAINING_CASES": "잔여 Case 검토",
        }.get(action_code, "다음 액션 진행")
        return {
            "enabled": True,
            "page": VOC_VALIDITY_PAGE_NAME,
            "run_id": run_id,
            "case_id": case_id,
            "action_code": action_code,
            "button_label": f"{case_id} {action_label}",
            "detail": f"{VOC_VALIDITY_PAGE_NAME} 화면으로 이동해 {case_id}의 {action_label} 흐름을 이어갑니다.",
        }
    if str(run_action.get("code")) == "REPORT_READY" or action_code == "REPORT_READY":
        target_case = _history_case_for_next_action(case_results, "REPORT_READY")
        case_id = str((target_case or {}).get("case_id") or "")
        return {
            "enabled": True,
            "page": VOC_REPORT_PAGE_NAME,
            "run_id": run_id,
            "case_id": case_id,
            "action_code": "REPORT_READY",
            "button_label": "품질 보고서·최종 시연 연결",
            "detail": "업무 승인 완료 증적을 품질 보고서와 최종 인수·시연 대상으로 연결합니다.",
        }
    if str(run_action.get("code")) == "RUBRIC_REEVALUATE" or action_code == "RUBRIC_REEVALUATE":
        plan_status = str(run_item.get("rubric_reevaluation_plan_status") or "")
        return {
            "enabled": True,
            "page": "rubric_reevaluation_plan",
            "run_id": run_id,
            "case_id": "",
            "action_code": "RUBRIC_REEVALUATE",
            "button_label": "기준 영향 계획 갱신" if plan_status else "기준 영향 계획 저장",
            "detail": (
                "독립 LLM PASS 확보용 실행이 아니라, Run 저장 당시 Rubric과 현재 Rubric의 차이를 확인하고 "
                "독립 LLM 재평가·타당성 재평가·RETEST 중 필요한 후속 실행을 정리합니다."
            ),
        }
    if action_code in {"CHECK_PIPELINE_ERROR", "REVIEW_PIPELINE_RESULT", "RUN_JUDGE"}:
        return {
            "enabled": True,
            "page": "history_detail",
            "run_id": run_id,
            "case_id": "",
            "action_code": action_code,
            "button_label": "Run 증적 확인",
            "detail": "수행 이력 상세 팝업에서 Case 증적과 독립 LLM 평가 상태를 확인합니다.",
        }
    return {
        "enabled": False,
        "button_label": "추가 이동 없음",
        "detail": "현재 상태에서는 바로 이동할 후속 화면이 없습니다.",
    }


def _apply_history_next_action_target(target: dict) -> None:
    page = target.get("page")
    run_id = str(target.get("run_id") or "")
    case_id = str(target.get("case_id") or "")
    action_code = str(target.get("action_code") or "")
    if page == "rubric_reevaluation_plan":
        _save_history_rubric_reevaluation_plan(run_id)
        return
    if page == "history_detail":
        st.session_state["current_menu"] = VOC_QUALITY_MENU_NAME
        st.session_state["current_sub_menu"] = VOC_HISTORY_PAGE_NAME
        st.session_state[HISTORY_SELECTED_RUN_ID_KEY] = run_id
        if case_id:
            st.session_state[f"voc_history_case_{run_id}"] = case_id
        _open_history_detail_dialog(run_id)
        return
    if page == VOC_VALIDITY_PAGE_NAME:
        _load_validity_candidates.clear()
        st.session_state.current_menu = VOC_QUALITY_MENU_NAME
        st.session_state.current_sub_menu = VOC_VALIDITY_PAGE_NAME
        st.session_state.voc_validity_selected_key = f"{run_id}::{case_id}"
        st.session_state.voc_validity_candidate_query = run_id or case_id
        st.session_state.voc_validity_run_type = "전체"
        st.session_state.voc_validity_candidate_status = {
            "RUN_VALIDITY": "전체",
            "QA_REVIEW": "QA 검토 가능",
            "BUSINESS_APPROVAL": "업무 승인 가능",
            "REPORT_READY": "정식 승인",
        }.get(action_code, "전체")
        st.session_state.voc_validity_focus_action_code = action_code
        st.session_state.voc_validity_focus_target_key = f"{run_id}::{case_id}"
        st.session_state.voc_validity_focus_notice = (
            f"독립 LLM PASS 결과를 이어받았습니다. 대상 {run_id} / {case_id}의 "
            "개선안 타당성 평가를 실행하세요."
            if action_code == "RUN_VALIDITY"
            else f"수행 이력의 다음 액션에서 이동했습니다. 대상: {run_id} / {case_id}"
        )
        if action_code == "RUN_VALIDITY":
            st.session_state.voc_validity_evaluation_focus_once = True
        _reset_validity_candidate_table_selection()
        st.session_state.pop(VALIDITY_DETAIL_DIALOG_CANDIDATE_KEY, None)
        st.session_state.pop("voc_validity_dialog_opened_key", None)
        return
    if page == VOC_BATCH_PAGE_NAME:
        st.session_state.current_menu = VOC_QUALITY_MENU_NAME
        st.session_state.current_sub_menu = VOC_BATCH_PAGE_NAME
        st.session_state.voc_batch_focus_notice = (
            f"Rubric 기준 영향 확인 결과 RETEST가 권장되었습니다. 원본 Run: {run_id}"
        )
        return
    if page == VOC_REPORT_PAGE_NAME:
        _go_to_voc_report(
            run_id,
            case_id,
            notice=f"수행 이력의 다음 액션에서 이동했습니다. 보고서 대상 Run: {run_id}",
        )


def _go_to_voc_report(run_id: str, case_id: str = "", *, notice: str = "") -> None:
    st.session_state.current_menu = VOC_QUALITY_MENU_NAME
    st.session_state.current_sub_menu = VOC_REPORT_PAGE_NAME
    st.session_state.voc_report_run_id = run_id
    st.session_state.voc_report_focus_case_id = case_id
    st.session_state.voc_report_focus_notice = notice or (
        f"업무 승인 완료 증적을 품질 보고서 대상으로 연결했습니다. Run: {run_id}"
    )


def _go_to_voc_acceptance(run_id: str, case_id: str = "", *, notice: str = "") -> None:
    st.session_state.current_menu = VOC_QUALITY_MENU_NAME
    st.session_state.current_sub_menu = VOC_ACCEPTANCE_PAGE_NAME
    st.session_state.voc_acceptance_focus_run_id = run_id
    st.session_state.voc_acceptance_focus_case_id = case_id
    st.session_state.voc_acceptance_focus_notice = notice or (
        f"업무 승인 완료 증적을 최종 인수·시연 대상으로 연결했습니다. Run: {run_id}"
    )


def _render_history_next_action_cards(run_item: dict, *, retest_plan: dict | None = None) -> None:
    run_id = str(run_item.get("run_id") or "")
    try:
        detail = load_voc_run_history_detail(run_id)
        summary = detail.get("summary", {})
    except Exception:
        summary = {}
    case_results = summary.get("case_results", [])
    run_action = run_item.get("next_action") or voc_run_next_action({**run_item, **summary})
    case_action = _history_priority_case_action(case_results)
    action_groups = _history_case_action_groups(case_results)
    qa_waiting = int(action_groups.get("QA_REVIEW", {}).get("count", 0) or 0)
    business_waiting = int(action_groups.get("BUSINESS_APPROVAL", {}).get("count", 0) or 0)
    approved = sum(1 for case in case_results if case.get("formal_approval"))
    total_cases = len(case_results) or int(run_item.get("selected_count") or 0)

    with st.container(border=True):
        heading, run_state = st.columns([2.4, 1], vertical_alignment="center")
        with heading:
            st.markdown("#### 선택 Run 다음 액션")
            st.caption("선택한 Run과 Case 결과를 기준으로 이어서 해야 할 업무를 요약합니다.")
        with run_state:
            st.markdown(_history_action_badge(run_action), text_alignment="right")

        columns = st.columns(4, gap="small")
        card_payloads = [
            (
                run_action.get("icon", "conversion_path"),
                "Run 기준",
                run_action["label"],
                f"{run_action['menu']} · {run_action['detail']}",
            ),
            (
                case_action.get("icon", "fact_check"),
                "Case 우선 액션",
                f"{case_action.get('label', '-')} · {case_action.get('count', 0)}건",
                f"대상: {case_action.get('case_ids', '-')}",
            ),
            (
                "rate_review",
                "QA·업무 대기",
                f"QA {qa_waiting}건 · 업무 {business_waiting}건",
                "QA 검토 완료 후 업무 승인으로 이어집니다.",
            ),
            (
                "verified",
                "승인 완료",
                f"{approved}/{total_cases or '-'}건",
                "정식 승인 완료 건은 보고서와 최종 시연 대상으로 연결합니다.",
            ),
        ]
        for column, (icon, label, value, detail_text) in zip(columns, card_payloads, strict=False):
            with column.container(border=True, height=128):
                st.caption(f":material/{icon}: {label}")
                st.markdown(f"##### {value}")
                st.caption(detail_text)

        lineage = run_item.get("lineage_policy") or {}
        rubric_drift = run_item.get("rubric_drift") or {}
        changed_labels = run_item.get("rubric_changed_labels") or "-"
        parent_run_id = lineage.get("parent_run_id") or run_item.get("parent_run_id") or "-"
        lineage_detail = lineage.get("lineage_rule") or "Run 계보 정보를 확인합니다."
        if lineage.get("has_parent"):
            lineage_detail = f"원본 Run {parent_run_id} 기준으로 전후 비교합니다."
        rubric_detail = rubric_drift.get("detail") or "Rubric 기준 상태를 확인합니다."
        if changed_labels and changed_labels != "-":
            rubric_detail = f"변경 범위: {changed_labels}"
        reevaluation_value = "평가만 재확인" if rubric_drift.get("requires_reevaluation") else "추가 재평가 없음"
        reevaluation_detail = (
            "Agent 파이프라인 실행 결과는 보존하고 Rubric 기준 변경 영향만 확인합니다."
            if rubric_drift.get("requires_reevaluation")
            else "현재 기준으로 유지해도 되는 Run입니다."
        )
        retest_plan = retest_plan or {}
        retest_detail = retest_plan.get("detail") or "선택 Run의 연결 재시험 여부를 확인합니다."
        basis_columns = st.columns(4, gap="small")
        basis_payloads = [
            (
                "account_tree",
                "Run/Case 계보",
                lineage.get("label") or _voc_status_label(run_item.get("run_type")),
                lineage_detail,
            ),
            (
                "rule_settings",
                "Rubric 기준",
                rubric_drift.get("status") or "기준 확인",
                rubric_detail,
            ),
            (
                "published_with_changes",
                "재수행·재평가 판단",
                reevaluation_value,
                reevaluation_detail,
            ),
            (
                "compare_arrows",
                "재시험 비교",
                retest_plan.get("state_label") or "비교 대상 없음",
                retest_detail,
            ),
        ]
        for column, (icon, label, value, detail_text) in zip(basis_columns, basis_payloads, strict=False):
            with column.container(border=True, height=124):
                st.caption(f":material/{icon}: {label}")
                if label == "Rubric 기준":
                    st.markdown(f"##### {_history_rubric_badge(rubric_drift)}")
                else:
                    st.markdown(f"##### {value}")
                st.caption(detail_text)

        _render_history_rubric_reevaluation_plan(run_id, detail, rubric_drift)

        target = _history_next_action_target(run_item, case_results, run_action, case_action)
        action_col, guide_col = st.columns([0.9, 2.8], gap="medium", vertical_alignment="center")
        with action_col:
            if st.button(
                target["button_label"],
                type="primary" if target.get("enabled") else "secondary",
                icon=":material/arrow_forward:",
                disabled=not target.get("enabled"),
                width="stretch",
                key=f"voc_history_next_action_go_{run_id}",
            ):
                _apply_history_next_action_target(target)
                st.rerun()
        with guide_col:
            st.caption(target.get("detail", "선택 Run의 다음 업무 단계로 이동합니다."))


def _remember_history_run_selection(table_key: str, run_ids: tuple[str, ...]) -> None:
    selected_row = _promote_table_cell_to_row_selection(table_key, len(run_ids))
    if selected_row is None:
        return
    try:
        selected_run_id = run_ids[selected_row]
    except IndexError:
        return
    st.session_state[HISTORY_SELECTED_RUN_ID_KEY] = selected_run_id
    _open_history_detail_dialog(selected_run_id)


def _render_history_filter_controls(
    history: list[dict],
    min_date: date,
    max_date: date,
    *,
    section_title: str = "",
):
    if section_title:
        filter_columns = st.columns([0.9, 1.45, 1.05, 1.05, 0.8, 1.15], vertical_alignment="bottom")
        with filter_columns[0]:
            st.markdown(f"### {section_title}")
        control_columns = filter_columns[1:]
    else:
        control_columns = st.columns([1.45, 1.1, 1.1, 0.8, 1.15], vertical_alignment="bottom")
    with control_columns[0]:
        date_range = st.date_input(
            "실행 기간",
            value=(min_date, max_date),
            min_value=min_date,
            max_value=max_date,
            key="voc_history_date_range",
        )
    with control_columns[1]:
        statuses = st.multiselect(
            "Run 상태",
            sorted({item.get("status", "") for item in history}),
            key="voc_history_status_filter",
            format_func=_voc_status_label,
        )
    with control_columns[2]:
        run_types = st.multiselect(
            "실행 유형",
            sorted({item.get("run_type", "") for item in history}),
            key="voc_history_type_filter",
            format_func=_voc_status_label,
        )
    with control_columns[3]:
        judge_filter = st.selectbox(
            "독립 LLM 평가",
            ["전체", "사용", "미사용"],
            key="voc_history_judge_filter",
        )
    with control_columns[4]:
        case_query = st.text_input(
            "Case ID",
            placeholder="예: TC-01",
            key="voc_history_case_filter",
        )
    return date_range, statuses, run_types, judge_filter, case_query


def render_voc_history():
    history = _load_voc_history_rows()
    delete_result = st.session_state.pop("voc_history_delete_result", None)
    if delete_result:
        st.success(f"{delete_result['deleted_count']}개 Run을 삭제했습니다.")
    reevaluation_result = st.session_state.pop("voc_judge_reevaluation_result", None)
    if reevaluation_result:
        judge = reevaluation_result.get("judge_result", {})
        st.success(
            f"{reevaluation_result['case_id']} 독립 LLM 재평가 완료 · "
            f"{_voc_status_label(judge.get('decision', 'ERROR'))} · {judge.get('total_score', '-')}점"
        )
    rubric_plan_result = st.session_state.pop("voc_rubric_reevaluation_plan_result", None)
    if rubric_plan_result:
        changed = ", ".join(rubric_plan_result.get("changed_labels", [])) or "-"
        st.success(
            f"Rubric 기준 영향 계획 저장 완료 · {_history_rubric_plan_label(rubric_plan_result.get('status'))} · 변경 범위 {changed}"
        )
    if not history:
        st.info("저장된 VOC 실행 이력이 없습니다.")
        return

    started_dates = [
        datetime.fromisoformat(item["started_at"]).date()
        for item in history if item.get("started_at")
    ]
    min_date, max_date = min(started_dates), max(started_dates)
    with st.container(border=True):
        date_range, statuses, run_types, judge_filter, case_query = _render_history_filter_controls(
            history,
            min_date,
            max_date,
            section_title="수행 이력",
        )

        if isinstance(date_range, tuple) and len(date_range) == 2:
            start_date, end_date = date_range
        else:
            start_date = end_date = date_range
        filtered = []
        for item in history:
            started_date = datetime.fromisoformat(item["started_at"]).date() if item.get("started_at") else min_date
            if not (start_date <= started_date <= end_date):
                continue
            if statuses and item.get("status") not in statuses:
                continue
            if run_types and item.get("run_type") not in run_types:
                continue
            if judge_filter != "전체" and item.get("judge_status") != judge_filter:
                continue
            if case_query and case_query.strip().upper() not in {
                str(case_id).upper() for case_id in item.get("selected_case_ids", [])
            }:
                continue
            filtered.append(item)

        _render_history_summary_cards(filtered)
        if not filtered:
            st.info("필터 조건에 맞는 Run이 없습니다.")
            return

        run_ids = tuple(str(item.get("run_id", "")) for item in filtered)
        selected_run_id = st.session_state.get(HISTORY_SELECTED_RUN_ID_KEY)
        if selected_run_id not in run_ids:
            selected_run_id = run_ids[0]
            st.session_state[HISTORY_SELECTED_RUN_ID_KEY] = selected_run_id
        table_signature = json.dumps(run_ids, ensure_ascii=False)
        if st.session_state.get(HISTORY_TABLE_SIGNATURE_KEY) != table_signature:
            _reset_history_table_selection()
            st.session_state[HISTORY_TABLE_SIGNATURE_KEY] = table_signature
        table_key = _history_table_widget_key()

        table_payloads = []
        for item in filtered:
            next_action = item.get("next_action") or voc_run_next_action(item)
            rubric_plan_status = item.get("rubric_reevaluation_plan_status")
            rubric_plan_label = (
                _history_rubric_plan_label(rubric_plan_status)
                if item.get("reevaluation_required") or rubric_plan_status
                else "-"
            )
            table_payloads.append(
                {
                    "Run ID": item.get("run_id"),
                    "실행 시각": _history_table_timestamp(item.get("started_at")),
                    "유형": _voc_status_label(item.get("run_type")),
                    "상태": _voc_status_label(item.get("status")),
                    "대상": item.get("selected_count", 0),
                    "실행 가능": "-" if item.get("executable_count") is None else f"{item.get('executable_count')}",
                    "후속 구현": "-" if item.get("pending_count") is None else f"{item.get('pending_count')}",
                    "원본 Run": item.get("parent_run_id") or "-",
                    "계보 기준": item.get("lineage_label") or _voc_status_label(item.get("run_type")),
                    "Rubric 기준": item.get("rubric_status") or "기준 확인",
                    "변경 범위": item.get("rubric_changed_labels") or "-",
                    "재평가 계획": rubric_plan_label,
                    "완료": item.get("completed_count", 0),
                    "진행률": item.get("completion_rate", 0.0),
                    "통과": item.get("counts", {}).get("PASS", 0),
                    "실패": item.get("counts", {}).get("FAIL", 0),
                    "오류": item.get("counts", {}).get("ERROR", 0),
                    "검토 필요": item.get("counts", {}).get("REVIEW_REQUIRED", 0),
                    "독립 LLM 평가": _voc_status_label(item.get("judge_status")),
                    "독립 통과": item.get("judge_counts", {}).get("PASS", 0),
                    "독립 오류": item.get("judge_counts", {}).get("ERROR", 0),
                    "개선안 타당성": _voc_status_label(item.get("validity_state", "DRAFT")),
                    "배포 판정": _voc_status_label(item.get("deployment_decision")),
                    "다음 액션": next_action["label"],
                }
            )
        table_rows = pd.DataFrame(table_payloads)
        st.dataframe(
            table_rows,
            hide_index=True,
            width="stretch",
            on_select=partial(_remember_history_run_selection, table_key, run_ids),
            selection_mode=["single-row", "single-cell"],
            key=table_key,
            row_height=32,
            column_order=[
                "Run ID",
                "실행 시각",
                "유형",
                "상태",
                "대상",
                "실행 가능",
                "후속 구현",
                "원본 Run",
                "계보 기준",
                "Rubric 기준",
                "변경 범위",
                "재평가 계획",
                "완료",
                "진행률",
                "통과",
                "실패",
                "오류",
                "검토 필요",
                "독립 LLM 평가",
                "독립 통과",
                "독립 오류",
                "개선안 타당성",
                "배포 판정",
                "다음 액션",
            ],
            column_config={
                "Run ID": st.column_config.TextColumn("Run ID", width=210, pinned=True),
                "실행 시각": st.column_config.TextColumn("실행 시각", width=124),
                "유형": st.column_config.TextColumn("유형", width=82),
                "상태": st.column_config.TextColumn("상태", width=82),
                "대상": st.column_config.NumberColumn("대상", width=58),
                "실행 가능": st.column_config.TextColumn("실행 가능", width=74),
                "후속 구현": st.column_config.TextColumn("후속 구현", width=74),
                "원본 Run": st.column_config.TextColumn("원본 Run", width=150),
                "계보 기준": st.column_config.TextColumn("계보 기준", width=108),
                "Rubric 기준": st.column_config.TextColumn("Rubric 기준", width=104),
                "변경 범위": st.column_config.TextColumn("변경 범위", width=132),
                "재평가 계획": st.column_config.TextColumn("재평가 계획", width=104),
                "완료": st.column_config.NumberColumn("완료", width=58),
                "진행률": st.column_config.ProgressColumn("진행률", min_value=0, max_value=100, format="%.0f%%", width=92),
                "통과": st.column_config.NumberColumn("통과", width=58),
                "실패": st.column_config.NumberColumn("실패", width=58),
                "오류": st.column_config.NumberColumn("오류", width=62),
                "검토 필요": st.column_config.NumberColumn("검토 필요", width=78),
                "독립 LLM 평가": st.column_config.TextColumn("독립 LLM 평가", width=98),
                "독립 통과": st.column_config.NumberColumn("독립 통과", width=74),
                "독립 오류": st.column_config.NumberColumn("독립 오류", width=74),
                "개선안 타당성": st.column_config.TextColumn("개선안 타당성", width=104),
                "배포 판정": st.column_config.TextColumn("배포 판정", width=82),
                "다음 액션": st.column_config.TextColumn("다음 액션", width=128),
            },
        )
        selected_run_id = st.session_state.get(HISTORY_SELECTED_RUN_ID_KEY, selected_run_id)
        selected_items = [item for item in filtered if item.get("run_id") == selected_run_id]
        selected_retest_plan = (
            _history_retest_comparison_plan(history, selected_items[0])
            if selected_items
            else {"enabled": False, "pair_key": "", "detail": "선택 Run이 없습니다."}
        )
        if st.session_state.get(HISTORY_RETEST_COMPARE_TARGET_KEY) != selected_retest_plan.get("pair_key"):
            st.session_state.pop(HISTORY_RETEST_COMPARE_TARGET_KEY, None)
        if selected_items:
            _render_history_next_action_cards(selected_items[0], retest_plan=selected_retest_plan)
        with st.container(horizontal=True):
            if selected_items:
                selected_next_action = selected_items[0].get("next_action") or voc_run_next_action(selected_items[0])
                _render_contextual_jira_action_menu(
                    area_label="수행 이력",
                    target_label="선택 Run",
                    run_id=selected_run_id,
                    status_label=selected_next_action.get("label", ""),
                    extra_detail=(
                        f"유형: {_voc_status_label(selected_items[0].get('run_type'))}\n"
                        f"상태: {_voc_status_label(selected_items[0].get('status'))}\n"
                        f"대상: {selected_items[0].get('selected_count', 0)}건\n"
                        f"통과: {(selected_items[0].get('counts') or {}).get('PASS', 0)}건\n"
                        f"검토 필요: {(selected_items[0].get('counts') or {}).get('REVIEW_REQUIRED', 0)}건\n"
                        f"실패: {(selected_items[0].get('counts') or {}).get('FAIL', 0)}건\n"
                        f"오류: {(selected_items[0].get('counts') or {}).get('ERROR', 0)}건"
                    ),
                    key=f"history_{selected_run_id}",
                )
            if st.button(
                "선택 Run 상세",
                icon=":material/open_in_new:",
                disabled=not selected_items,
                key="voc_history_open_selected_run",
            ):
                _open_history_detail_dialog(selected_run_id)
                st.rerun()
            if st.button(
                "재시험 전후 비교",
                icon=":material/compare_arrows:",
                disabled=not selected_retest_plan.get("enabled"),
                key="voc_history_open_retest_comparison",
            ):
                st.session_state[HISTORY_RETEST_COMPARE_TARGET_KEY] = selected_retest_plan["pair_key"]
                st.rerun()
            if st.button(
                "선택 Run 삭제",
                icon=":material/delete:",
                disabled=not selected_items or selected_items[0].get("status") == "RUNNING",
            ):
                _confirm_delete_voc_runs([selected_items[0]["run_id"]])
            if st.button("이력 새로고침", icon=":material/refresh:"):
                _load_voc_history_rows.clear()
                st.rerun()

    if not selected_items:
        st.caption("상세 조회할 Run 행을 선택하세요.")
    else:
        st.caption("Run 행을 선택하면 상세 팝업이 열리고, 다음 액션 대상도 함께 바뀝니다.")

    dialog_run_id = st.session_state.get(HISTORY_DETAIL_DIALOG_RUN_ID_KEY)
    if dialog_run_id:
        if any(str(item.get("run_id")) == dialog_run_id for item in history):
            _render_history_detail_dialog(dialog_run_id)
        else:
            _dismiss_history_detail_dialog()

    if (
        selected_items
        and selected_retest_plan.get("enabled")
        and st.session_state.get(HISTORY_RETEST_COMPARE_TARGET_KEY) == selected_retest_plan.get("pair_key")
    ):
        _render_retest_comparison(history, selected_retest_plan)


@st.cache_data(ttl=3, max_entries=1, show_spinner=False)
def _load_validity_candidates():
    return list_improvement_validity_candidates()


def _validity_candidate_key(candidate: dict) -> str:
    return f"{candidate.get('run_id', '')}::{candidate.get('case_id', '')}"


def _validity_candidate_table_widget_key() -> str:
    nonce = int(st.session_state.get(VALIDITY_CANDIDATE_TABLE_NONCE_KEY, 0) or 0)
    return f"{VALIDITY_CANDIDATE_TABLE_KEY}_{nonce}"


def _reset_validity_candidate_table_selection() -> None:
    st.session_state[VALIDITY_CANDIDATE_TABLE_NONCE_KEY] = int(
        st.session_state.get(VALIDITY_CANDIDATE_TABLE_NONCE_KEY, 0) or 0
    ) + 1


def _open_validity_candidate_dialog(candidate: dict | str) -> None:
    candidate_key = _validity_candidate_key(candidate) if isinstance(candidate, dict) else str(candidate or "")
    if candidate_key:
        st.session_state.voc_validity_selected_key = candidate_key
        st.session_state[VALIDITY_DETAIL_DIALOG_CANDIDATE_KEY] = candidate_key


def _dismiss_validity_candidate_dialog() -> None:
    st.session_state.pop(VALIDITY_DETAIL_DIALOG_CANDIDATE_KEY, None)
    st.session_state.pop("voc_validity_dialog_opened_key", None)
    _reset_validity_candidate_table_selection()


def _remember_validity_candidate_selection(table_key: str, candidate_keys: tuple[str, ...]) -> None:
    selected_row = _promote_table_cell_to_row_selection(table_key, len(candidate_keys))
    if selected_row is None:
        return
    try:
        selected_key = candidate_keys[selected_row]
    except IndexError:
        return
    st.session_state.voc_validity_selected_key = selected_key
    st.session_state[VALIDITY_DETAIL_DIALOG_CANDIDATE_KEY] = selected_key


def _apply_pending_validity_candidate_filters() -> None:
    """Apply deferred validity filters before Streamlit widgets are instantiated."""
    pending_to_widget = {
        "voc_validity_pending_candidate_query": "voc_validity_candidate_query",
        "voc_validity_pending_candidate_status": "voc_validity_candidate_status",
        "voc_validity_pending_run_type": "voc_validity_run_type",
    }
    for pending_key, widget_key in pending_to_widget.items():
        if pending_key in st.session_state:
            st.session_state[widget_key] = st.session_state.pop(pending_key)


def _queue_validity_candidate_focus(
    candidate: dict,
    action_code: str,
    *,
    notice: str = "",
    scroll_to_approval: bool = False,
) -> None:
    """Keep the selected Run/Case visible on the next rerun without mutating live widgets."""
    run_id = str(candidate.get("run_id") or "")
    case_id = str(candidate.get("case_id") or "")
    if run_id and case_id:
        st.session_state.voc_validity_selected_key = f"{run_id}::{case_id}"
        st.session_state.voc_validity_focus_target_key = f"{run_id}::{case_id}"
    if run_id:
        st.session_state.voc_validity_pending_candidate_query = run_id
    elif case_id:
        st.session_state.voc_validity_pending_candidate_query = case_id
    st.session_state.voc_validity_pending_run_type = "전체"
    status_filter = VALIDITY_STATUS_FILTER_BY_ACTION.get(action_code)
    if status_filter:
        st.session_state.voc_validity_pending_candidate_status = status_filter
    if action_code:
        st.session_state.voc_validity_focus_action_code = action_code
    if notice:
        st.session_state.voc_validity_focus_notice = notice
    if action_code == "RUN_VALIDITY":
        st.session_state.voc_validity_evaluation_focus_once = True
    if scroll_to_approval:
        st.session_state.voc_validity_approval_focus_once = True


def _render_validity_approval_focus_anchor_once() -> None:
    _render_goal_scroll_anchor(
        "voc-validity-approval-scroll-anchor",
        scroll_flag_key="voc_validity_approval_focus_once",
        block="start",
    )


def _render_validity_evaluation_focus_anchor_once() -> None:
    _render_goal_scroll_anchor(
        "voc-validity-evaluation-scroll-anchor",
        scroll_flag_key="voc_validity_evaluation_focus_once",
        block="start",
    )


def _validity_post_evaluation_action_model(candidate: dict, result: dict | None) -> dict:
    """Return the user-facing next action after the validity result is known."""
    result = result or {}
    result_has_hold_state = "immediate_hold_rules_triggered" in result
    immediate_hold_count = (
        len(_validity_immediate_holds(result))
        if result_has_hold_state
        else int(candidate.get("immediate_hold_count") or 0)
    )
    readiness = validity_human_review_readiness(
        validity_status=result.get("decision", candidate.get("validity_status", "NOT_RUN")),
        workflow_state=result.get("workflow_state", candidate.get("workflow_state", "DRAFT")),
        immediate_hold_count=immediate_hold_count,
        formal_approval=bool(result.get("formal_approval") or candidate.get("formal_approval")),
    )
    action = readiness["action"]
    if action == "QA_REVIEW":
        return {
            "visible": True,
            "action_code": action,
            "title": "다음 단계 · QA 검토 저장",
            "detail": "AI_PASS와 즉시 보류 0건이 확인됐습니다. 이제 QA 검토 의견을 저장하면 업무 승인 단계로 이어집니다.",
            "button_label": "QA 검토 영역으로 이동",
            "target": "approval",
            "tone": "green",
            "icon": "rate_review",
        }
    if action == "BUSINESS_APPROVAL":
        return {
            "visible": True,
            "action_code": action,
            "title": "다음 단계 · 업무 승인 저장",
            "detail": "QA 검토가 완료됐습니다. 업무 관점의 최종 승인 여부를 저장하면 보고서와 최종 시연 대상이 됩니다.",
            "button_label": "업무 승인 영역으로 이동",
            "target": "approval",
            "tone": "green",
            "icon": "verified",
        }
    if action == "FORMAL_APPROVED":
        return {
            "visible": True,
            "action_code": "REPORT_READY",
            "title": "다음 단계 · 보고서/최종 시연",
            "detail": "정식 승인 완료 상태입니다. 품질 보고서와 최종 인수·시연 화면에서 증적을 확인할 수 있습니다.",
            "button_label": "품질 보고서로 이동",
            "secondary_button_label": "최종 인수·시연으로 이동",
            "target": "report",
            "tone": "green",
            "icon": "summarize",
        }
    return {"visible": False, "action_code": action, "readiness": readiness}


def _approved_demo_cases(
    *,
    focus_run_id: str = "",
    focus_case_id: str = "",
) -> list[dict]:
    approved = [
        dict(candidate)
        for candidate in _load_validity_candidates()
        if candidate.get("formal_approval")
        or str(candidate.get("workflow_state") or "") == "BUSINESS_APPROVED"
    ]
    approved.sort(key=lambda item: item.get("started_at", ""), reverse=True)
    if focus_run_id:
        run_cases = [
            item
            for item in approved
            if str(item.get("run_id") or "") == focus_run_id
        ]
        focused = [
            item
            for item in run_cases
            if not focus_case_id or str(item.get("case_id") or "") == focus_case_id
        ] or run_cases
        if focused:
            return focused + [
                item
                for item in approved
                if _validity_candidate_key(item) not in {_validity_candidate_key(focus) for focus in focused}
            ]
    return approved


def _approved_demo_case_label(candidate: dict) -> str:
    score = candidate.get("validity_score")
    score_text = "-" if score in (None, "") else f"{score}점"
    return (
        f"{candidate.get('case_id', '-')} · "
        f"{candidate.get('run_id', '-')} · "
        f"개선안 타당성 {score_text}"
    )


def _approved_demo_case_rows(candidates: list[dict]) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "Run ID": item.get("run_id", "-"),
                "Case ID": item.get("case_id", "-"),
                "수행 유형": _voc_status_label(item.get("run_type", "-")),
                "질문": item.get("question", "-") or "-",
                "독립 LLM 평가": _voc_status_label(item.get("judge_status", "NOT_RUN")),
                "개선안 타당성": _voc_status_label(item.get("validity_status", "NOT_RUN")),
                "점수": item.get("validity_score"),
                "승인 단계": _voc_status_label(item.get("workflow_state", "DRAFT")),
            }
            for item in candidates
        ]
    )


def _render_demo_flow_step_cards(candidate: dict) -> None:
    judge_status = str(candidate.get("judge_status") or "NOT_RUN")
    validity_status = str(candidate.get("validity_status") or "NOT_RUN")
    workflow_state = str(candidate.get("workflow_state") or "DRAFT")
    formal_approval = bool(candidate.get("formal_approval")) or workflow_state == "BUSINESS_APPROVED"
    qa_done = workflow_state in {"QA_REVIEWED", "BUSINESS_APPROVED"} or formal_approval
    steps = [
        {
            "icon": "account_tree",
            "label": "Agent 파이프라인",
            "status": "완료",
            "tone": "green",
            "detail": "Agent 파이프라인 증적",
        },
        {
            "icon": "psychology",
            "label": "독립 LLM 평가",
            "status": _voc_status_label(judge_status),
            "tone": "green" if judge_status not in {"NOT_RUN", "ERROR", ""} else "gray",
            "detail": f"{candidate.get('judge_score') if candidate.get('judge_score') is not None else '-'}점",
        },
        {
            "icon": "fact_check",
            "label": "개선안 타당성 평가",
            "status": _voc_status_label(validity_status),
            "tone": "green" if validity_status == "AI_PASS" else "orange",
            "detail": f"{candidate.get('validity_score') if candidate.get('validity_score') is not None else '-'}점",
        },
        {
            "icon": "rate_review",
            "label": "QA 검토",
            "status": "완료" if qa_done else "대기",
            "tone": "green" if qa_done else "gray",
            "detail": _voc_status_label(workflow_state),
        },
        {
            "icon": "verified",
            "label": "업무 승인",
            "status": "완료" if formal_approval else "대기",
            "tone": "green" if formal_approval else "gray",
            "detail": "정식 승인 증적",
        },
        {
            "icon": "approval",
            "label": "보고서·시연",
            "status": "연결 가능" if formal_approval else "대기",
            "tone": "blue" if formal_approval else "gray",
            "detail": "품질 보고서 / 최종 시연",
        },
    ]
    columns = st.columns(6, gap="small")
    for column, step in zip(columns, steps, strict=False):
        badge = {
            "green": f":green-badge[{step['status']}]",
            "blue": f":blue-badge[{step['status']}]",
            "orange": f":orange-badge[{step['status']}]",
            "gray": f":gray-badge[{step['status']}]",
        }.get(step["tone"], f":gray-badge[{step['status']}]")
        with column.container(border=True, height=124):
            st.caption(f":material/{step['icon']}: {step['label']}")
            st.markdown(badge)
            st.caption(step["detail"])


def _render_approved_demo_flow_panel(
    location: str,
    *,
    focus_run_id: str = "",
    focus_case_id: str = "",
) -> dict | None:
    approved = _approved_demo_cases(focus_run_id=focus_run_id, focus_case_id=focus_case_id)
    if not approved:
        with st.container(border=True):
            st.markdown("#### 업무 승인 완료 → 보고서/최종 시연 연결")
            st.caption("업무 승인 완료 Case가 생기면 이 영역에서 시연 가능한 한 줄 흐름을 보여줍니다.")
            st.markdown(":gray-badge[승인 완료 Case 없음]")
        return None

    keys = [_validity_candidate_key(item) for item in approved]
    target_key = f"voc_demo_target_{location}"
    if st.session_state.get(target_key) not in keys:
        st.session_state[target_key] = keys[0]
    candidate_map = {
        _validity_candidate_key(item): item
        for item in approved
    }

    with st.container(border=True):
        heading, summary = st.columns([2.4, 1], vertical_alignment="center")
        with heading:
            st.markdown("#### 업무 승인 완료 → 품질 보고서 → 최종 인수·시연")
            st.caption("업무 승인 완료된 Run·Case를 보고서 증적과 최종 시연 대상으로 연결합니다.")
        with summary:
            run_count = len({item.get("run_id") for item in approved})
            st.markdown(
                f":green-badge[승인 Case {len(approved)}건] :blue-badge[Run {run_count}개]",
                text_alignment="right",
            )

        selected_key = st.selectbox(
            "시연 대상 승인 Case",
            keys,
            index=None,
            key=target_key,
            format_func=lambda value: _approved_demo_case_label(candidate_map.get(value, {})),
            width="stretch",
        )
        selected = candidate_map.get(selected_key) or approved[0]

        st.markdown(
            f"**현재 시연 대상** · `{selected.get('run_id', '-')}` / `{selected.get('case_id', '-')}`"
        )
        st.caption(selected.get("question", "-") or "-")
        _render_demo_flow_step_cards(selected)

        action_columns = st.columns([1, 1, 2], gap="small", vertical_alignment="center")
        with action_columns[0]:
            if location != "report":
                if st.button(
                    "품질 보고서로 이동",
                    icon=":material/article:",
                    width="stretch",
                    key=f"approved_demo_to_report_{location}_{_validity_candidate_key(selected)}",
                ):
                    _go_to_voc_report(
                        str(selected.get("run_id") or ""),
                        str(selected.get("case_id") or ""),
                        notice=(
                            "업무 승인 완료 Case에서 품질 보고서로 이동했습니다. "
                            f"대상: {selected.get('run_id')} / {selected.get('case_id')}"
                        ),
                    )
                    st.rerun()
            else:
                st.markdown(":blue-badge[보고서 화면]")
        with action_columns[1]:
            if location != "acceptance":
                if st.button(
                    "최종 시연으로 이동",
                    icon=":material/approval:",
                    width="stretch",
                    key=f"approved_demo_to_acceptance_{location}_{_validity_candidate_key(selected)}",
                ):
                    _go_to_voc_acceptance(
                        str(selected.get("run_id") or ""),
                        str(selected.get("case_id") or ""),
                        notice=(
                            "업무 승인 완료 Case에서 최종 인수·시연으로 이동했습니다. "
                            f"대상: {selected.get('run_id')} / {selected.get('case_id')}"
                        ),
                    )
                    st.rerun()
            else:
                st.markdown(":blue-badge[최종 시연 화면]")
        with action_columns[2]:
            st.caption(
                "정식 배포 판정은 실행 가능 Case 승인과 후속 구현 Case 승인 상태를 함께 봅니다. "
                "이 카드는 단일 승인 Case의 시연 흐름을 빠르게 확인하는 용도입니다."
            )

        with st.expander("업무 승인 완료 Case 목록", expanded=False, icon=":material/table_view:"):
            st.dataframe(
                _approved_demo_case_rows(approved),
                hide_index=True,
                width="stretch",
                height=min(260, 58 + len(approved) * 34),
                column_config={
                    "Run ID": st.column_config.TextColumn(width=220, pinned=True),
                    "Case ID": st.column_config.TextColumn(width=80),
                    "수행 유형": st.column_config.TextColumn(width=88),
                    "질문": st.column_config.TextColumn(width=260),
                    "독립 LLM 평가": st.column_config.TextColumn(width=112),
                    "개선안 타당성": st.column_config.TextColumn(width=118),
                    "점수": st.column_config.ProgressColumn(width=82, min_value=0, max_value=100, format="%g점"),
                    "승인 단계": st.column_config.TextColumn(width=118),
                },
            )
    return selected


def _quality_case_by_id(case_id: str) -> dict:
    try:
        cases = load_unified_quality_cases().get("cases", [])
    except Exception:
        return {}
    for case in cases:
        if isinstance(case, dict) and str(case.get("case_id") or "") == case_id:
            return case
    return {}


def _go_to_voc_history_detail(run_id: str, case_id: str = "", *, notice: str = "") -> None:
    st.session_state.current_menu = VOC_QUALITY_MENU_NAME
    st.session_state.current_sub_menu = VOC_HISTORY_PAGE_NAME
    st.session_state[HISTORY_SELECTED_RUN_ID_KEY] = run_id
    if case_id:
        st.session_state[f"voc_history_case_{run_id}"] = case_id
    if notice:
        st.session_state.voc_history_focus_notice = notice
    _open_history_detail_dialog(run_id)


def _render_retest_formal_approval_flow(
    *,
    run_id: str,
    release_scope: dict,
    location: str,
) -> None:
    linked_retests = release_scope.get("linked_retest_evidence")
    if not isinstance(linked_retests, list) or not linked_retests:
        return

    primary = linked_retests[0] if isinstance(linked_retests[0], dict) else {}
    case_id = str(primary.get("case_id") or "-")
    retest_run_id = str(primary.get("retest_run_id") or "-")
    case = _quality_case_by_id(case_id)
    question = (
        case.get("question")
        or (case.get("execution") or {}).get("question")
        or case.get("name")
        or "보완 재시험 Case"
    )
    expected_behavior = case.get("expected_system_behavior") or case.get("acceptance") or "보완 재시험 결과를 원본 회차 승인 근거에 반영합니다."
    judge_score = primary.get("judge_score")
    validity_score = primary.get("validity_score")
    score_text = (
        f"독립 LLM {judge_score if judge_score is not None else '-'}점 · "
        f"타당성 {validity_score if validity_score is not None else '-'}점"
    )
    status_text = (
        f"{_voc_status_label(primary.get('status'))} · "
        f"{_voc_status_label(primary.get('judge_decision'))} · "
        f"{_voc_status_label(primary.get('validity_state'))}"
    )

    with st.container(border=True):
        heading_col, status_col = st.columns([2.4, 1.2], gap="small", vertical_alignment="center")
        with heading_col:
            st.markdown(f"#### {case_id} RETEST 반영 흐름")
            st.caption(
                "원본 35건 Run에서 보완이 필요했던 Case가 별도 RETEST로 통과했고, "
                "그 승인 증적이 원본 회차의 FORMAL_APPROVED 근거로 연결됩니다."
            )
        with status_col:
            st.markdown(":green-badge[보완 반영] :green-badge[정식 승인 연결]", text_alignment="right")
            st.caption(f"원본 Run {run_id}", text_alignment="right")

        story_cards = [
            {
                "icon": "inventory_2",
                "label": "원본 35건 Run",
                "value": "35건 회차",
                "detail": f"{case_id} 포함 · 최종 판정 대상",
                "badge": ":blue-badge[원본]",
            },
            {
                "icon": "help",
                "label": f"{case_id} 보완 포인트",
                "value": "무근거 VOC",
                "detail": "근거 없는 상품·보상 정보를 창작하지 않는지 확인",
                "badge": ":orange-badge[보완 필요]",
            },
            {
                "icon": "sync_alt",
                "label": "RETEST 승인",
                "value": "통과",
                "detail": f"{retest_run_id} · {score_text}",
                "badge": ":green-badge[승인]",
            },
            {
                "icon": "verified",
                "label": "원본 Run 반영",
                "value": "FORMAL_APPROVED",
                "detail": "RETEST 증적을 원본 35건 인수 게이트에 반영",
                "badge": ":green-badge[정식 승인]",
            },
        ]
        for column, card in zip(st.columns(4, gap="small"), story_cards, strict=False):
            with column.container(border=True, height=136):
                st.caption(f":material/{card['icon']}: {card['label']}")
                st.markdown(f"#### {card['value']}")
                st.caption(card["detail"])
                st.markdown(card["badge"])

        st.caption(f"질문: {question}")
        st.caption(f"판정 의미: {expected_behavior}")

        action_cols = st.columns([1, 1, 1, 1.6], gap="small", vertical_alignment="center")
        with action_cols[0]:
            if location != "report":
                if st.button(
                    "품질 보고서에서 확인",
                    icon=":material/article:",
                    width="stretch",
                    key=f"retest_flow_to_report_{location}_{run_id}_{case_id}_{retest_run_id}",
                ):
                    _go_to_voc_report(
                        run_id,
                        case_id,
                        notice=f"{case_id} RETEST 반영 흐름을 품질 보고서에서 확인합니다. 원본 Run: {run_id}",
                    )
                    st.rerun()
            else:
                st.markdown(":blue-badge[품질 보고서]")
        with action_cols[1]:
            if location != "acceptance":
                if st.button(
                    "최종 시연에서 확인",
                    icon=":material/approval:",
                    width="stretch",
                    key=f"retest_flow_to_acceptance_{location}_{run_id}_{case_id}_{retest_run_id}",
                ):
                    _go_to_voc_acceptance(
                        run_id,
                        case_id,
                        notice=f"{case_id} RETEST 반영 흐름을 최종 인수·시연에서 확인합니다. 원본 Run: {run_id}",
                    )
                    st.rerun()
            else:
                st.markdown(":blue-badge[최종 시연]")
        with action_cols[2]:
            if st.button(
                "RETEST 이력 보기",
                icon=":material/history:",
                width="stretch",
                key=f"retest_flow_history_{location}_{run_id}_{case_id}_{retest_run_id}",
            ):
                _go_to_voc_history_detail(
                    retest_run_id,
                    case_id,
                    notice=f"{case_id} 보완 RETEST 상세 증적을 열었습니다. 원본 Run: {run_id}",
                )
                st.rerun()
        with action_cols[3]:
            st.caption(f"RETEST 결과: {status_text} · 시작 {_history_table_timestamp(primary.get('started_at'))}")

        if len(linked_retests) > 1:
            rows = []
            for item in linked_retests:
                if not isinstance(item, dict):
                    continue
                rows.append(
                    {
                        "Case ID": item.get("case_id", "-"),
                        "RETEST Run": item.get("retest_run_id", "-"),
                        "파이프라인": _voc_status_label(item.get("status")),
                        "독립 LLM": f"{_voc_status_label(item.get('judge_decision'))} · {item.get('judge_score', '-')}점",
                        "타당성/승인": f"{_voc_status_label(item.get('validity_state'))} · {item.get('validity_score', '-')}점",
                        "수행 시각": _history_table_timestamp(item.get("started_at")),
                    }
                )
            with st.expander("연결 RETEST 전체 목록", expanded=False, icon=":material/table_view:"):
                st.dataframe(pd.DataFrame(rows), hide_index=True, width="stretch")


def _render_acceptance_formal_connection(snapshot: dict) -> None:
    scope = snapshot.get("release_scope_summary", {}) if isinstance(snapshot.get("release_scope_summary"), dict) else {}
    release_decision = str(snapshot.get("release_report_decision") or "")
    report_state = str(snapshot.get("report_state") or "")
    readiness = str(snapshot.get("decision") or "")
    is_formal = release_decision == "FORMAL_APPROVED"
    ready_for_uat = readiness == "READY_FOR_UAT"
    badge_tone = "green" if is_formal else "orange"
    readiness_tone = "green" if ready_for_uat else "orange"

    with st.container(border=True):
        _render_voc_section_heading(
            "정식 승인 결과 연결",
            "품질 보고서의 최종 판정이 최종 인수·시연 게이트에 연결된 상태입니다. Run ID보다 인수 범위와 보완 RETEST 반영 여부를 먼저 확인합니다.",
            icon="verified",
            badges=(
                (_voc_status_label(release_decision), badge_tone),
                (_voc_status_label(report_state), "blue"),
                (_voc_status_label(readiness), readiness_tone),
            ),
            right_caption=f"Run {snapshot.get('run_id', '-')}",
        )

        cards = [
            {
                "icon": "task_alt",
                "label": "VOC 개선 Case",
                "value": f"{int(scope.get('voc_count') or 0)}건",
                "detail": "PASS·독립 LLM·업무 승인 완료",
                "tone": "green" if scope.get("release_scope_ready") else "orange",
            },
            {
                "icon": "shield",
                "label": "장애 검증 Case",
                "value": f"{int(scope.get('fault_count') or 0)}건",
                "detail": "장애 보호 동작 실행 확인",
                "tone": "blue",
            },
            {
                "icon": "pending_actions",
                "label": "후속 구현 Case",
                "value": f"{int(scope.get('pending_count') or 0)}건",
                "detail": "승인된 미실행 계획으로 관리",
                "tone": "blue",
            },
            {
                "icon": "sync_alt",
                "label": "연결 RETEST",
                "value": f"{int(scope.get('linked_retest_count') or 0)}건",
                "detail": "부족 Case 보완 결과를 원본 회차에 반영",
                "tone": "green" if int(scope.get("linked_retest_count") or 0) else "gray",
                "badge": "연결",
            },
        ]
        for card in cards:
            card.setdefault(
                "badge",
                {
                    "green": "충족",
                    "blue": "확인",
                    "orange": "확인 필요",
                    "gray": "없음",
                }.get(card["tone"], "확인"),
            )
        _render_voc_summary_cards(cards, columns=4, height=124)


def _validity_candidate_rows(candidates: list[dict], selected_key: str = "") -> pd.DataFrame:
    rows = []
    for candidate in candidates:
        validity_status = candidate.get("validity_status", "NOT_RUN")
        review_readiness = _candidate_review_readiness(candidate)
        next_action = candidate.get("next_action") or {}
        judge_gate = _validity_judge_gate_model(candidate)
        review_already_started = (
            candidate.get("workflow_state") in {"QA_REVIEWED", "BUSINESS_APPROVED"}
            or bool(candidate.get("formal_approval"))
        )
        if candidate.get("judge_status") and judge_gate["blocked"] and not review_already_started:
            next_action_label = judge_gate["next_title"]
        else:
            next_action_label = (
                next_action.get("label")
                or candidate.get("review_action_label")
                or review_readiness["action_label"]
            )
        rows.append(
            {
                "Case ID": candidate.get("case_id", "-"),
                "질문": candidate.get("question", "-") or "-",
                "다음 조치": next_action_label,
                "개선안 타당성": _voc_status_label(validity_status),
                "타당성 점수": candidate.get("validity_score"),
                "승인 단계": _voc_status_label(candidate.get("workflow_state", "DRAFT")),
                "독립 LLM 평가": _voc_status_label(candidate.get("judge_status", "NOT_RUN")),
                "독립 LLM 점수": candidate.get("judge_score"),
                "수행 유형": _voc_status_label(candidate.get("run_type", "-")),
                "수행 일시": _dashboard_timestamp(candidate.get("started_at", "")),
                "정식 승인": "승인" if candidate.get("formal_approval") else "미승인",
                "Run ID": candidate.get("run_id", "-"),
            }
        )
    return pd.DataFrame(rows)


def _filter_validity_candidates(
    candidates: list[dict],
    *,
    query: str,
    status_filter: str,
    run_type_filter: str = "전체",
) -> list[dict]:
    query = str(query or "").strip().lower()
    run_type_value = VALIDITY_RUN_TYPE_FILTER_VALUES.get(run_type_filter or "전체")
    filtered = []
    for candidate in candidates:
        if query and query not in " ".join(
            str(candidate.get(key, ""))
            for key in ("run_id", "case_id", "question", "run_type")
        ).lower():
            continue
        if run_type_value and candidate.get("run_type") != run_type_value:
            continue
        validity_status = candidate.get("validity_status", "NOT_RUN")
        if status_filter == "평가 전" and validity_status != "NOT_RUN":
            continue
        if status_filter == "평가 완료" and validity_status == "NOT_RUN":
            continue
        if status_filter == "QA 검토 가능" and not _candidate_qa_review_ready(candidate):
            continue
        if status_filter == "업무 승인 가능" and not _candidate_business_approval_ready(candidate):
            continue
        if status_filter == "정식 승인" and not candidate.get("formal_approval"):
            continue
        filtered.append(candidate)
    return filtered


def _validity_immediate_holds(result: dict | None) -> list[str]:
    holds = (result or {}).get("immediate_hold_rules_triggered", []) or []
    if isinstance(holds, str):
        return [holds] if holds.strip() else []
    if isinstance(holds, list):
        return [str(item) for item in holds if str(item).strip()]
    try:
        return [str(item) for item in holds if str(item).strip()]
    except TypeError:
        return [str(holds)] if holds else []


def _candidate_qa_review_ready(candidate: dict) -> bool:
    return _candidate_review_readiness(candidate)["can_qa_review"]


def _candidate_business_approval_ready(candidate: dict) -> bool:
    return _candidate_review_readiness(candidate)["can_business_approve"]


def _candidate_review_readiness(candidate: dict) -> dict:
    return validity_human_review_readiness(
        validity_status=candidate.get("validity_status", "NOT_RUN"),
        workflow_state=candidate.get("workflow_state", "DRAFT"),
        immediate_hold_count=candidate.get("immediate_hold_count", 0) or 0,
        formal_approval=bool(candidate.get("formal_approval")),
    )


def _validity_focus_cards(candidates: list[dict]) -> list[dict]:
    readiness = [_candidate_review_readiness(candidate) for candidate in candidates]
    action_counts = pd.Series([item["action"] for item in readiness]).value_counts().to_dict()
    hold_count = sum(item["immediate_hold_count"] for item in readiness)
    total = max(1, len(candidates))
    approved_count = action_counts.get("FORMAL_APPROVED", 0)
    approved_percent = round(approved_count / total * 100)
    return [
        {
            "icon": "fact_check",
            "label": "개선안 타당성 평가 필요",
            "value": f"{action_counts.get('VALIDITY_EVALUATION_REQUIRED', 0)}건",
            "delta": "개선안 타당성 평가 미수행",
            "detail": "개선안 타당성 평가 미수행",
            "tone": "blue",
            "badge": "평가",
        },
        {
            "icon": "edit_note",
            "label": "보완·재시험 필요",
            "value": f"{action_counts.get('REWORK_REQUIRED', 0)}건",
            "delta": f"즉시 보류 {hold_count}건 포함",
            "detail": f"즉시 보류 {hold_count}건 포함",
            "tone": "red" if action_counts.get("REWORK_REQUIRED", 0) else "gray",
            "badge": "보완",
        },
        {
            "icon": "rate_review",
            "label": "QA 검토 가능",
            "value": f"{action_counts.get('QA_REVIEW', 0)}건",
            "delta": "AI 통과 · 보류 없음",
            "detail": "AI 통과 · 보류 없음",
            "tone": "green" if action_counts.get("QA_REVIEW", 0) else "gray",
            "badge": "QA",
        },
        {
            "icon": "verified",
            "label": "업무 승인 가능",
            "value": f"{action_counts.get('BUSINESS_APPROVAL', 0)}건",
            "delta": "QA 검토 완료",
            "detail": "QA 검토 완료",
            "tone": "green" if action_counts.get("BUSINESS_APPROVAL", 0) else "gray",
            "badge": "승인",
        },
        {
            "icon": "approval",
            "label": "정식 승인 완료",
            "value": f"{approved_count}건",
            "delta": f"전체 대비 {approved_percent}%",
            "detail": f"전체 대비 {approved_percent}%",
            "tone": "green" if approved_count else "gray",
            "badge": "완료",
        },
    ]


def _validity_review_queue_rows(candidates: list[dict]) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "Case ID": candidate.get("case_id", "-"),
                "Run ID": candidate.get("run_id", "-"),
                "점수": candidate.get("validity_score"),
                "다음 조치": candidate.get("review_action_label") or _candidate_review_readiness(candidate)["action_label"],
                "질문": candidate.get("question", "-") or "-",
            }
            for candidate in candidates
        ]
    )


def _validity_status_badge_tone(status: str) -> str:
    status = str(status or "NOT_RUN").upper()
    if status in {"AI_PASS", "BUSINESS_APPROVED", "QA_REVIEWED", "FORMAL_APPROVED"}:
        return "green"
    if status in {"REVISION_REQUIRED", "REJECTED", "FAIL", "ERROR"}:
        return "red"
    if status in {"NOT_RUN", "NOT_EVALUATED", "DRAFT"}:
        return "gray"
    return "orange"


def _validity_status_count_badges(candidates: list[dict]) -> tuple[tuple[str, str], ...]:
    order = ("NOT_RUN", "REVISION_REQUIRED", "AI_PASS", "QA_REVIEWED", "BUSINESS_APPROVED")
    counts = pd.Series(
        [str(item.get("validity_status") or "NOT_RUN").upper() for item in candidates]
    ).value_counts().to_dict()
    ordered_statuses = [status for status in order if counts.get(status)]
    ordered_statuses.extend(
        status for status in counts if status not in set(ordered_statuses)
    )
    badges: list[tuple[str, str]] = [(f"전체 {len(candidates)}건", "blue")]
    badges.extend(
        (
            f"{_voc_status_label(status)} {int(counts.get(status) or 0)}건",
            _validity_status_badge_tone(status),
        )
        for status in ordered_statuses
    )
    return tuple(badges)


def _validity_review_queue_preview_items(candidates: list[dict], limit: int = 3) -> list[str]:
    preview: list[str] = []
    for candidate in candidates[:limit]:
        case_id = candidate.get("case_id") or "-"
        score = candidate.get("validity_score")
        score_text = f"{score:g}점" if isinstance(score, (int, float)) else "점수 없음"
        question = str(candidate.get("question") or "-")
        if len(question) > 42:
            question = f"{question[:42]}..."
        action_label = candidate.get("review_action_label") or _candidate_review_readiness(candidate)["action_label"]
        preview.append(f"- **{case_id}** · {action_label} · {score_text}  \n  :gray[{question}]")
    return preview


def _render_validity_review_queue(candidates: list[dict]):
    qa_ready = [item for item in candidates if _candidate_qa_review_ready(item)]
    business_ready = [item for item in candidates if _candidate_business_approval_ready(item)]
    if not qa_ready and not business_ready:
        return

    with st.expander(
        f"QA/승인 대기 빠른 확인 · QA {len(qa_ready)}건 · 업무 승인 {len(business_ready)}건",
        expanded=False,
        icon=":material/rate_review:",
    ):
        st.caption("상세 처리와 저장은 아래 선택 Case 작업 영역에서 진행합니다. 전체 대상은 목록의 상태 필터로 좁혀 확인하세요.")
        columns = st.columns(2, gap="small", vertical_alignment="top")
        queues = [
            ("QA 검토 가능", qa_ready, ":material/rate_review:", "QA 검토 의견 저장 대상"),
            ("업무 승인 가능", business_ready, ":material/verified:", "QA 검토 완료 후 업무 승인 대상"),
        ]
        for column, (title, items, icon, detail) in zip(columns, queues, strict=False):
            with column.container(border=True, height="stretch"):
                st.markdown(f"{icon} **{title}**")
                st.markdown(f":{'green' if items else 'gray'}-badge[{len(items)}건]")
                st.caption(detail)
                preview_items = _validity_review_queue_preview_items(items)
                if preview_items:
                    st.markdown("\n".join(preview_items))
                    if len(items) > len(preview_items):
                        st.caption(f"외 {len(items) - len(preview_items)}건은 목록 필터에서 확인하세요.")
                else:
                    st.caption("현재 대기 대상이 없습니다.")


def _validity_selection_basis(candidate: dict, artifacts: dict | None = None) -> dict:
    artifacts = artifacts or {}
    pipeline = artifacts.get("pipeline_result", {}) if isinstance(artifacts.get("pipeline_result"), dict) else {}
    execution = pipeline.get("execution", {}) if isinstance(pipeline.get("execution"), dict) else {}
    result = execution.get("result", {}) if isinstance(execution.get("result"), dict) else {}
    pipeline_mode = str(pipeline.get("mode") or "voc").upper()
    execution_ok = execution.get("ok", True)
    result_ok = result.get("ok", True)
    pipeline_success = pipeline_mode == "VOC" and execution_ok is not False and result_ok is not False
    question = candidate.get("question") or execution.get("question") or "-"
    run_type = candidate.get("run_type") or "-"
    parent_run_id = candidate.get("parent_run_id") or ""
    review_readiness = _candidate_review_readiness(candidate)
    next_action = candidate.get("next_action") or {}
    return {
        "run_id": candidate.get("run_id", "-"),
        "case_id": candidate.get("case_id", "-"),
        "run_type": run_type,
        "run_type_label": _voc_status_label(run_type),
        "started_at": _dashboard_timestamp(candidate.get("started_at", "")),
        "parent_run_id": parent_run_id,
        "question": question,
        "pipeline_success": pipeline_success,
        "pipeline_label": "VOC 파이프라인 성공" if pipeline_success else "파이프라인 증적 확인 필요",
        "pipeline_detail": "완료 Run의 VOC 실행 결과를 기준으로 검증합니다.",
        "judge_label": _voc_status_label(candidate.get("judge_status", "NOT_RUN")),
        "judge_score": candidate.get("judge_score"),
        "validity_label": _voc_status_label(candidate.get("validity_status", "NOT_RUN")),
        "validity_score": candidate.get("validity_score"),
        "next_action_label": next_action.get("label") or review_readiness["action_label"],
        "next_action_detail": next_action.get("detail") or "선택한 Run·Case의 다음 처리 단계를 확인합니다.",
    }


def _validity_qa_gate_model(candidate: dict, result: dict | None) -> dict:
    result = result or {}
    has_result = bool(result)
    decision = result.get("decision") or candidate.get("validity_status", "NOT_RUN")
    workflow_state = result.get("workflow_state") or candidate.get("workflow_state", "DRAFT")
    formal_approval = bool(result.get("formal_approval"))
    holds = _validity_immediate_holds(result)
    if not holds and candidate.get("immediate_hold_count"):
        holds = [f"{candidate.get('immediate_hold_count')}건"]
    approval_stage_ok = workflow_state in {"AI_REVIEWED", "QA_REVIEWED", "BUSINESS_APPROVED"} or formal_approval
    checks = [
        {
            "label": "개선안 타당성 평가",
            "value": "완료" if has_result else "대기",
            "ok": has_result,
            "detail": "개선안 타당성 평가 결과가 저장되어야 합니다.",
        },
        {
            "label": "AI 판정",
            "value": _voc_status_label(decision),
            "ok": decision == "AI_PASS",
            "detail": "AI 통과일 때만 QA 검토를 시작합니다.",
        },
        {
            "label": "즉시 보류 규칙",
            "value": "없음" if not holds else f"{len(holds)}건",
            "ok": not holds,
            "detail": "보류 규칙이 있으면 보완 후 재평가합니다.",
        },
        {
            "label": "승인 단계",
            "value": _voc_status_label(workflow_state),
            "ok": approval_stage_ok,
            "detail": "AI 검토 완료 이후 QA 검토와 업무 승인이 순차로 열립니다.",
        },
    ]
    ready = all(item["ok"] for item in checks)
    completed = workflow_state == "BUSINESS_APPROVED" or formal_approval
    if completed:
        summary = "정식 승인 완료"
    elif workflow_state == "QA_REVIEWED":
        summary = "QA 검토 완료 · 업무 승인 가능"
    elif ready:
        summary = "QA 검토 가능"
    else:
        summary = "QA 검토 대기"
    blocked_reasons = [item["label"] for item in checks if not item["ok"]]
    return {
        "ready": ready,
        "completed": completed,
        "summary": summary,
        "blocked_reasons": blocked_reasons,
        "checks": checks,
        "holds": holds,
        "workflow_state": workflow_state,
        "decision": decision,
    }


def _render_validity_selection_basis(
    candidate: dict,
    artifacts: dict | None = None,
    *,
    compact: bool = False,
):
    basis = _validity_selection_basis(candidate, artifacts)
    with st.container(border=True):
        readiness = _candidate_review_readiness(candidate)
        action_tone = {
            "FORMAL_APPROVED": "green",
            "BUSINESS_APPROVAL": "green",
            "QA_REVIEW": "green",
            "REWORK_REQUIRED": "red",
            "VALIDITY_EVALUATION_REQUIRED": "blue",
        }.get(readiness["action"], "gray")
        _render_voc_section_heading(
            f"선택 Case 판단 · {basis['case_id']}",
            "",
            icon="ads_click",
            badges=((basis["next_action_label"], action_tone),),
            right_caption=basis["started_at"],
        )
        if compact:
            _render_voc_summary_cards(
                [
                    {
                        "icon": "fact_check",
                        "label": "Case",
                        "value": basis["case_id"],
                        "detail": _manual_pipeline_compact_text(basis["question"], 64),
                        "tone": "blue",
                        "badge": basis["run_type_label"],
                    },
                    {
                        "icon": "conversion_path",
                        "label": "다음 액션",
                        "value": basis["next_action_label"],
                        "detail": _manual_pipeline_compact_text(basis["next_action_detail"], 58),
                        "tone": action_tone,
                        "badge": "지금 할 일",
                    },
                    {
                        "icon": "approval",
                        "label": "타당성·승인",
                        "value": basis["validity_label"],
                        "detail": (
                            f"타당성 {basis['validity_score'] if basis['validity_score'] is not None else '-'}점 · "
                            f"승인 {_voc_status_label(candidate.get('workflow_state', 'DRAFT'))}"
                        ),
                        "tone": "green" if str(candidate.get("validity_status") or "") == "AI_PASS" else "orange",
                        "badge": "평가",
                    },
                    {
                        "icon": "psychology",
                        "label": "독립 LLM",
                        "value": basis["judge_label"],
                        "detail": f"{basis['judge_score'] if basis['judge_score'] is not None else '-'}점",
                        "tone": "green" if str(candidate.get("judge_status") or "") == "PASS" else "orange",
                        "badge": "선행",
                    },
                ],
                columns=4,
                height=112,
            )
            if basis["parent_run_id"]:
                st.caption(f"재시험 회차 · 원본 Run {basis['parent_run_id']}")
            else:
                st.caption(f"Run ID: {basis['run_id']}")
            return
        _render_voc_summary_cards(
            [
                {
                    "icon": "fact_check",
                    "label": "Case",
                    "value": basis["case_id"],
                    "detail": basis["question"],
                    "tone": "blue",
                    "badge": basis["run_type_label"],
                },
                {
                    "icon": "conversion_path",
                    "label": "다음 액션",
                    "value": basis["next_action_label"],
                    "detail": basis["next_action_detail"],
                    "tone": action_tone,
                    "badge": "지금 할 일",
                },
                {
                    "icon": "approval",
                    "label": "개선안 타당성",
                    "value": basis["validity_label"],
                    "detail": f"{basis['validity_score'] if basis['validity_score'] is not None else '-'}점",
                    "tone": "green" if str(candidate.get("validity_status") or "") == "AI_PASS" else "orange",
                    "badge": "평가",
                },
                {
                    "icon": "psychology",
                    "label": "독립 LLM 평가",
                    "value": basis["judge_label"],
                    "detail": f"{basis['judge_score'] if basis['judge_score'] is not None else '-'}점",
                    "tone": "green" if str(candidate.get("judge_status") or "") == "PASS" else "orange",
                    "badge": "선행",
                },
                {
                    "icon": "account_tree",
                    "label": "파이프라인",
                    "value": basis["pipeline_label"],
                    "detail": basis["pipeline_detail"],
                    "tone": "green" if basis["pipeline_success"] else "orange",
                    "badge": "증적",
                },
            ],
            columns=5,
            height=132,
        )
        if basis["parent_run_id"]:
            st.caption(f"재시험 회차입니다. 원본 Run: {basis['parent_run_id']}")
        else:
            st.caption(f"Run ID: {basis['run_id']}")


def _render_validity_history_focus_action(candidate: dict) -> None:
    action_code = str(st.session_state.get("voc_validity_focus_action_code") or "")
    target_key = str(st.session_state.get("voc_validity_focus_target_key") or "")
    if not action_code or target_key != _validity_candidate_key(candidate):
        return

    spec = VALIDITY_FOCUS_ACTION_LABELS.get(
        action_code,
        {
            "label": "다음 액션",
            "detail": "수행 이력에서 선택한 대상의 다음 처리 단계를 이어갑니다.",
            "icon": "conversion_path",
            "tone": "blue",
        },
    )
    tone = {
        "green": "green",
        "blue": "blue",
        "orange": "orange",
        "red": "red",
        "gray": "gray",
    }.get(str(spec.get("tone") or "blue"), "blue")
    with st.container(border=True):
        icon_col, text_col, action_col = st.columns([0.18, 2.45, 0.95], gap="small", vertical_alignment="center")
        with icon_col:
            st.markdown(f"### :material/{spec['icon']}:")
        with text_col:
            st.markdown(f"#### 수행 이력 연결 액션 · {spec['label']}")
            st.caption(spec["detail"])
        with action_col:
            st.markdown(f":{tone}-badge[{spec['label']}]", text_alignment="right")
            if action_code == "RUN_VALIDITY":
                if st.button(
                    "평가 설정으로 이동",
                    icon=":material/fact_check:",
                    width="stretch",
                    key=f"voc_validity_focus_to_evaluation_{_validity_candidate_key(candidate)}",
                ):
                    st.session_state.voc_validity_evaluation_focus_once = True
                    st.rerun()
            if st.button(
                "연결 안내 닫기",
                icon=":material/close:",
                width="stretch",
                key=f"voc_validity_focus_clear_{_validity_candidate_key(candidate)}",
            ):
                st.session_state.pop("voc_validity_focus_action_code", None)
                st.session_state.pop("voc_validity_focus_target_key", None)
                st.rerun()


VALIDITY_SUPPLEMENT_FIELDS = (
    ("owner", "담당/오너", "예: 모바일앱개발팀 리드, CS 운영 담당자"),
    ("schedule", "일정/마일스톤", "예: 2026-08-01 착수, 08-15 QA, 08-22 배포"),
    ("kpi", "정량 KPI", "예: 갱신 오류율 2.1%→0.5%, 주간 재문의율 15%↓"),
    ("priority", "우선순위", "예: P1 · 결제/구독 갱신 VOC 재발 방지"),
    ("execution_plan", "실행계획/적용범위", "예: 앱 청구 상태 조회 API·화면 상태 표시 로직 수정 → QA → 배포"),
    ("evidence", "VOC·실행 Trace 근거", "예: VOC-001, 실행 Trace step 3/4, 고객 문의 유형"),
    ("risk", "리스크/우회방안", "예: 정책 변경 전 안내 배너와 상담 스크립트 병행"),
    ("note", "검토 메모", "예: AI_PASS 근거 확인 완료 · QA 검토 시 확인할 사항"),
)
VALIDITY_SUPPLEMENT_FIELD_ORDER = tuple(field for field, _label, _placeholder in VALIDITY_SUPPLEMENT_FIELDS)
VALIDITY_EVALUATION_SUPPLEMENT_FIELD_KEYS = tuple(
    field for field in VALIDITY_SUPPLEMENT_FIELD_ORDER
    if field != "note"
)
VALIDITY_DIMENSION_SUPPLEMENT_FIELD_KEYS = {
    "cause_linkage": ("evidence", "priority", "execution_plan"),
    "evidence_traceability": ("evidence",),
    "feasibility": ("execution_plan",),
    "ownership_schedule_kpi": ("owner", "schedule", "kpi", "priority"),
    "risk_security_compliance": ("risk",),
}
VALIDITY_HOLD_SUPPLEMENT_FIELD_KEYS = {
    "missing_voc_or_trace_evidence": ("evidence",),
    "unsafe_or_noncompliant_action": ("risk",),
    "unresolved_high_or_critical_defect": ("risk", "execution_plan"),
    "safety_regression_against_baseline": ("risk", "evidence"),
}


def _validity_supplement_field_spec(field: str) -> tuple[str, str, str]:
    return next(
        (
            item for item in VALIDITY_SUPPLEMENT_FIELDS
            if item[0] == field
        ),
        (field, field, ""),
    )


def _validity_ordered_supplement_fields(fields: list[str] | tuple[str, ...] | set[str]) -> list[str]:
    requested = {str(field) for field in fields if str(field or "").strip()}
    return [field for field in VALIDITY_SUPPLEMENT_FIELD_ORDER if field in requested]


def _validity_filled_supplement_field_keys(
    supplement: dict | None,
    *,
    include_note: bool = True,
) -> list[str]:
    supplement = supplement if isinstance(supplement, dict) else {}
    return [
        field for field in VALIDITY_SUPPLEMENT_FIELD_ORDER
        if (include_note or field != "note")
        and str(supplement.get(field) or "").strip()
    ]


def _validity_key_base(run_id: str, case_id: str, key_scope: str = "") -> str:
    scope = f"{key_scope}_" if key_scope else ""
    return f"{scope}{run_id}_{case_id}"


def _validity_supplement_key(run_id: str, case_id: str, field: str, key_scope: str = "") -> str:
    return f"validity_supplement_{_validity_key_base(run_id, case_id, key_scope)}_{field}"


def _validity_supplement_fields(supplement: dict | None) -> list[dict]:
    supplement = supplement or {}
    rows = []
    for key, label, _placeholder in VALIDITY_SUPPLEMENT_FIELDS:
        value = str(supplement.get(key) or "").strip()
        if not value:
            continue
        rows.append({"보완 항목": label, "입력 내용": value})
    return rows


def _validity_iso_datetime(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(str(value).replace("Z", "+00:00"))
    except ValueError:
        return None


def _validity_supplement_applied_to_result(artifacts: dict, result: dict | None) -> bool:
    result = result or {}
    supplement = artifacts.get("validity_supplement", {})
    if not isinstance(supplement, dict):
        return True
    evaluation_fields = [
        field for field in VALIDITY_EVALUATION_SUPPLEMENT_FIELD_KEYS
        if str(supplement.get(field) or "").strip()
    ]
    if not evaluation_fields:
        return True
    if not result or not result.get("supplemental_evidence_applied"):
        return False
    applied = result.get("supplemental_evidence", {})
    if not isinstance(applied, dict):
        return False
    for field in VALIDITY_EVALUATION_SUPPLEMENT_FIELD_KEYS:
        if str(supplement.get(field) or "").strip() != str(applied.get(field) or "").strip():
            return False
    supplement_time = _validity_iso_datetime(supplement.get("updated_at"))
    evaluated_time = _validity_iso_datetime(result.get("evaluated_at"))
    if supplement_time and evaluated_time and supplement_time > evaluated_time:
        return False
    return True


def _validity_score_gap_text(result: dict | None, rubric: dict) -> str:
    result = result or {}
    rule = _validity_ai_pass_rule(rubric)
    threshold = float(rule.get("min_score", 80) or 80)
    score = _validity_score_value(result.get("total_score"))
    if score is None:
        return f"AI_PASS 기준 {threshold:g}점"
    gap = round(max(threshold - score, 0), 2)
    if gap <= 0:
        return "점수 기준 충족"
    return f"AI_PASS까지 {gap:g}점 부족"


def _sync_validity_candidate_from_artifacts(candidate: dict, artifacts: dict | None) -> dict:
    """화면 후보 요약값을 실제 Run/Case 증적 파일 기준으로 보정합니다."""
    synced = dict(candidate or {})
    artifacts = artifacts or {}
    judge_result = artifacts.get("judge_result", {})
    if isinstance(judge_result, dict) and judge_result:
        judge_decision = str(judge_result.get("decision") or "").strip().upper()
        if judge_decision:
            synced["judge_status"] = judge_decision
        if "total_score" in judge_result:
            synced["judge_score"] = judge_result.get("total_score")
        synced["judge_provider"] = judge_result.get("provider") or synced.get("judge_provider")
        synced["judge_model"] = judge_result.get("model") or synced.get("judge_model")

    validity_result = artifacts.get("validity_result", {})
    if isinstance(validity_result, dict) and validity_result:
        validity_decision = str(validity_result.get("decision") or "").strip().upper()
        if validity_decision:
            synced["validity_status"] = validity_decision
        if "total_score" in validity_result:
            synced["validity_score"] = validity_result.get("total_score")
        synced["workflow_state"] = validity_result.get(
            "workflow_state",
            synced.get("workflow_state", "DRAFT"),
        )
        synced["formal_approval"] = bool(validity_result.get("formal_approval"))
        immediate_holds = validity_result.get("immediate_hold_rules_triggered") or []
        if isinstance(immediate_holds, str):
            immediate_hold_count = 1 if immediate_holds.strip() else 0
        else:
            try:
                immediate_hold_count = len(immediate_holds)
            except TypeError:
                immediate_hold_count = int(bool(immediate_holds))
        synced["immediate_hold_count"] = immediate_hold_count
        readiness = validity_human_review_readiness(
            validity_status=synced.get("validity_status", "NOT_RUN"),
            workflow_state=synced.get("workflow_state", "DRAFT"),
            immediate_hold_count=immediate_hold_count,
            formal_approval=bool(synced.get("formal_approval")),
        )
        synced["qa_review_ready"] = readiness["can_qa_review"]
        synced["business_review_ready"] = readiness["can_business_approve"]
        synced["review_action"] = readiness["action"]
        synced["review_action_label"] = readiness["action_label"]
        synced["deployment_decision"] = readiness["deployment_decision"]
        synced["next_action"] = {
            "code": readiness["action"],
            "label": readiness["action_label"],
            "menu": "개선안 타당성 검증",
            "detail": "선택한 Run·Case의 타당성/QA 승인 단계에서 이어서 처리합니다.",
        }
    return synced


def _validity_judge_gate_model(candidate: dict, artifacts: dict | None = None) -> dict:
    artifacts = artifacts or {}
    judge_result = artifacts.get("judge_result", {})
    if not isinstance(judge_result, dict):
        judge_result = {}

    raw_decision = (
        judge_result.get("decision")
        or candidate.get("judge_status")
        or "NOT_RUN"
    )
    decision = str(raw_decision or "NOT_RUN").strip().upper()
    if not decision or decision == "-":
        decision = "NOT_RUN"

    score = judge_result.get("total_score", candidate.get("judge_score"))
    provider = judge_result.get("provider") or candidate.get("judge_provider") or "-"
    model = judge_result.get("model") or candidate.get("judge_model") or "-"
    label = _voc_status_label(decision)
    passed = decision == "PASS"

    if passed:
        tone = "green"
        current = "독립 LLM 평가 PASS"
        next_title = "타당성 평가 진행 가능"
        next_detail = "동일 개선안에 대한 독립 LLM 평가가 PASS이므로 보완 입력과 타당성 평가를 진행할 수 있습니다."
    elif decision == "NOT_RUN":
        tone = "orange"
        current = "독립 LLM 평가 미수행"
        next_title = "독립 LLM 평가 먼저 실행"
        next_detail = "타당성 평가 전에 수동 TC 수행 또는 수행 이력에서 독립 LLM 평가를 먼저 완료하세요."
    elif decision == "ERROR":
        tone = "red"
        current = "독립 LLM 평가 오류"
        next_title = "독립 LLM 평가 오류 조치"
        next_detail = "Provider API 키와 모델 설정을 확인한 뒤 독립 LLM 평가를 다시 실행하세요."
    else:
        tone = "orange"
        current = f"독립 LLM 평가 {label}"
        next_title = "독립 LLM 평가 PASS 필요"
        next_detail = (
            "현재 개선안은 독립 LLM 평가를 통과하지 못했습니다. "
            "타당성 보완 입력이나 재평가를 반복해도 QA 검토로 넘어가지 않으므로 "
            "Agent 개선안 보완, 독립 LLM 재평가 또는 RETEST를 먼저 진행하세요."
        )

    detail_parts = [f"판정 {label}"]
    if score is not None:
        detail_parts.append(f"점수 {score}점")
    if provider != "-":
        detail_parts.append(str(provider))
    if model != "-":
        detail_parts.append(str(model))

    return {
        "passed": passed,
        "blocked": not passed,
        "decision": decision,
        "label": label,
        "score": score,
        "provider": provider,
        "model": model,
        "tone": tone,
        "current": current,
        "next_title": next_title,
        "next_detail": next_detail,
        "summary": " · ".join(detail_parts),
    }


def _validity_workflow_status_model(
    candidate: dict,
    artifacts: dict,
    result: dict | None,
    rubric: dict,
) -> dict:
    result = result or {}
    judge_gate = _validity_judge_gate_model(candidate, artifacts)
    supplement_rows = _validity_supplement_fields(artifacts.get("validity_supplement", {}))
    supplement_applied = _validity_supplement_applied_to_result(artifacts, result)
    reeval_needed = bool(supplement_rows) and bool(result) and not supplement_applied
    gate = _validity_qa_gate_model(candidate, result)
    readiness = validity_human_review_readiness(
        validity_status=result.get("decision", candidate.get("validity_status", "NOT_RUN")),
        workflow_state=result.get("workflow_state", candidate.get("workflow_state", "DRAFT")),
        immediate_hold_count=len(_validity_immediate_holds(result)),
        formal_approval=bool(result.get("formal_approval") or candidate.get("formal_approval")),
    )
    action = readiness["action"]
    review_already_started = (
        readiness["workflow_state"] in {"QA_REVIEWED", "BUSINESS_APPROVED"}
        or readiness["formal_approval"]
    )
    if judge_gate["blocked"] and not review_already_started:
        current = "독립 LLM 평가 보완 필요"
        tone = judge_gate["tone"]
        next_title = judge_gate["next_title"]
        next_detail = judge_gate["next_detail"]
        active_step = "judge"
    elif not result:
        current = "개선안 타당성 평가 전"
        tone = "gray"
        next_title = "개선안 타당성 평가 실행"
        next_detail = "아래 평가 실행 영역에서 Provider와 모델을 확인한 뒤 `타당성 평가 실행`을 누르세요."
        active_step = "evaluate"
    elif reeval_needed:
        current = "재평가 필요"
        tone = "orange"
        next_title = "보완 반영 재평가"
        next_detail = "보완 입력이 저장됐지만 현재 평가 결과에는 아직 반영되지 않았습니다. 아래 개선안 타당성 평가 버튼으로 재평가하세요."
        active_step = "evaluate"
    elif action == "REWORK_REQUIRED":
        current = "보완 필요"
        tone = "red"
        next_title = "보완 입력 후 재평가"
        next_detail = "부족 항목을 보완 입력에 작성하고 저장한 다음 개선안 타당성 평가를 다시 실행하세요."
        active_step = "supplement"
    elif action == "QA_REVIEW":
        current = "QA 검토 가능"
        tone = "green"
        next_title = "QA 검토 저장"
        next_detail = "아래 QA 검토 영역에서 승인, 보완 요청, 반려 중 하나를 선택해 검토 결과를 저장하세요."
        active_step = "qa"
    elif action == "BUSINESS_APPROVAL":
        current = "업무 승인 가능"
        tone = "green"
        next_title = "업무 승인 저장"
        next_detail = "QA 검토가 완료된 건입니다. 아래 업무 승인 영역에서 최종 승인 여부를 저장하세요."
        active_step = "business"
    elif action == "FORMAL_APPROVED":
        current = "정식 승인 완료"
        tone = "green"
        next_title = "보고서/최종 시연 연결"
        next_detail = "품질 보고서와 최종 인수·시연 화면에서 승인 증적을 확인할 수 있습니다."
        active_step = "approved"
    else:
        current = readiness["action_label"]
        tone = "gray"
        next_title = readiness["action_label"]
        next_detail = "현재 선택 대상의 평가 결과와 승인 상태를 확인하세요."
        active_step = "evaluate"

    stage_specs = [
        ("judge", "1. 독립 LLM 평가", "PASS일 때만 타당성 평가 가능"),
        ("supplement", "2. 보완 입력", "부족 근거·담당·일정·KPI를 입력"),
        ("evaluate", "3. 타당성 평가", "보완 내용을 반영해 개선안 타당성 평가"),
        ("qa", "4. QA 검토", "AI_PASS 및 보류 0건이면 QA 검토"),
        ("business", "5. 업무 승인", "QA 승인 후 최종 업무 승인"),
    ]
    supplement_not_required = (
        str(result.get("decision") or "").upper() == "AI_PASS"
        and not _validity_immediate_holds(result)
    )
    completed = {
        "judge": judge_gate["passed"],
        "supplement": bool(supplement_rows) or supplement_not_required,
        "evaluate": bool(result) and not reeval_needed,
        "qa": readiness["workflow_state"] in {"QA_REVIEWED", "BUSINESS_APPROVED"} or readiness["formal_approval"],
        "business": readiness["workflow_state"] == "BUSINESS_APPROVED" or readiness["formal_approval"],
    }
    stages = []
    for key, label, detail in stage_specs:
        if judge_gate["blocked"] and key != "judge":
            status = "잠금"
            stage_tone = "gray"
            icon = "lock"
        elif completed[key]:
            status = "완료"
            stage_tone = "green"
            icon = "check_circle"
        elif key == active_step:
            status = "현재 단계"
            stage_tone = "blue" if tone == "green" else tone
            icon = "play_circle"
        else:
            status = "대기"
            stage_tone = "gray"
            icon = "pending"
        stages.append(
            {
                "key": key,
                "label": label,
                "detail": detail,
                "status": status,
                "tone": stage_tone,
                "icon": icon,
            }
        )
    return {
        "current": current,
        "tone": tone,
        "next_title": next_title,
        "next_detail": next_detail,
        "active_step": active_step,
        "action_code": action,
        "readiness": readiness,
        "stages": stages,
        "score_text": _validity_score_gap_text(result, rubric),
        "gate_summary": gate["summary"],
        "judge_summary": judge_gate["summary"],
        "judge_gate": judge_gate,
        "supplement_count": len(supplement_rows),
        "reeval_needed": reeval_needed,
    }


def _render_validity_workflow_status(
    candidate: dict,
    artifacts: dict,
    result: dict | None,
    rubric: dict,
    *,
    key_scope: str = "main",
    compact: bool = False,
) -> None:
    model = _validity_workflow_status_model(candidate, artifacts, result, rubric)
    badge_color = {
        "green": "green",
        "blue": "blue",
        "orange": "orange",
        "red": "red",
        "gray": "gray",
    }.get(model["tone"], "gray")
    with st.container(border=True):
        heading, state = st.columns([2.6, 1], vertical_alignment="center")
        with heading:
            st.markdown("#### 선택 대상 작업 흐름" if compact else "#### 독립 LLM 평가 → 타당성 평가 → QA 검토 흐름")
            if not compact:
                st.caption("선택한 Run·Case가 지금 어느 단계인지와 다음에 눌러야 할 버튼을 고정해서 보여줍니다.")
        with state:
            st.markdown(f":{badge_color}-badge[{model['current']}]", text_alignment="right")

        columns = st.columns(len(model["stages"]), gap="small")
        for column, stage in zip(columns, model["stages"], strict=False):
            stage_color = {
                "green": "green",
                "blue": "blue",
                "orange": "orange",
                "red": "red",
                "gray": "gray",
            }.get(stage["tone"], "gray")
            with column.container(border=True, height=78 if compact else "stretch"):
                if compact:
                    st.caption(f":material/{stage['icon']}: {stage['label']}")
                else:
                    st.markdown(f":material/{stage['icon']}: **{stage['label']}**")
                st.markdown(f":{stage_color}-badge[{stage['status']}]")
                if not compact:
                    st.caption(stage["detail"])

        if compact:
            st.caption(f"다음 액션 · {model['next_title']} — {model['next_detail']}")
            st.markdown(
                " ".join(
                    (
                        _voc_ui_badge(f"점수 기준 {model['score_text']}", model["tone"]),
                        _voc_ui_badge(f"QA Gate {model['gate_summary']}", model["tone"]),
                    )
                )
            )
            return

        next_col, score_col, gate_col = st.columns([2.5, 0.9, 0.9], gap="small", vertical_alignment="center")
        with next_col:
            st.markdown(f"**다음 액션 · {model['next_title']}**")
            st.caption(model["next_detail"])
            action_model = _validity_post_evaluation_action_model(candidate, result)
            if action_model.get("visible"):
                key_base = _validity_key_base(
                    str(candidate.get("run_id") or ""),
                    str(candidate.get("case_id") or ""),
                    f"{key_scope}_workflow_{action_model['action_code']}",
                )
                if action_model.get("target") == "approval":
                    if st.button(
                        action_model["button_label"],
                        icon=f":material/{action_model['icon']}:",
                        type="primary",
                        width="stretch",
                        key=f"validity_next_action_{key_base}",
                    ):
                        _queue_validity_candidate_focus(
                            candidate,
                            str(action_model["action_code"]),
                            notice=(
                                f"{candidate.get('case_id', '-')} · {action_model['title']} 단계로 이동합니다. "
                                "아래 QA 검토·업무 승인 영역에서 저장하세요."
                            ),
                            scroll_to_approval=True,
                        )
                        st.rerun()
                elif action_model.get("target") == "report":
                    report_col, acceptance_col = st.columns(2, gap="small")
                    with report_col:
                        if st.button(
                            action_model["button_label"],
                            icon=":material/summarize:",
                            width="stretch",
                            key=f"validity_next_report_{key_base}",
                        ):
                            _go_to_voc_report(
                                str(candidate.get("run_id") or ""),
                                str(candidate.get("case_id") or ""),
                            )
                            st.rerun()
                    with acceptance_col:
                        if st.button(
                            action_model.get("secondary_button_label", "최종 인수·시연으로 이동"),
                            icon=":material/approval:",
                            width="stretch",
                            key=f"validity_next_acceptance_{key_base}",
                        ):
                            _go_to_voc_acceptance(
                                str(candidate.get("run_id") or ""),
                                str(candidate.get("case_id") or ""),
                            )
                            st.rerun()
            elif (
                model.get("active_step") == "evaluate"
                and not model.get("judge_gate", {}).get("blocked")
            ):
                key_base = _validity_key_base(
                    str(candidate.get("run_id") or ""),
                    str(candidate.get("case_id") or ""),
                    f"{key_scope}_workflow_evaluate",
                )
                if st.button(
                    "타당성 평가 설정으로 이동",
                    icon=":material/fact_check:",
                    type="primary",
                    width="stretch",
                    key=f"validity_next_evaluate_{key_base}",
                ):
                    st.session_state.voc_validity_evaluation_focus_once = True
                    st.rerun()
        with score_col:
            st.metric("점수 기준", model["score_text"], border=True)
        with gate_col:
            st.metric("QA Gate", model["gate_summary"], border=True)


def _render_validity_judge_prerequisite_notice(candidate: dict, artifacts: dict) -> None:
    judge_gate = _validity_judge_gate_model(candidate, artifacts)
    if not judge_gate["blocked"]:
        return
    supplement_rows = _validity_supplement_fields(artifacts.get("validity_supplement", {}))
    with st.container(border=True):
        title_col, status_col = st.columns([2.5, 1], gap="small", vertical_alignment="center")
        with title_col:
            st.markdown("#### 독립 LLM 평가 선행 필요")
            st.caption(judge_gate["next_detail"])
        with status_col:
            st.markdown(f":{judge_gate['tone']}-badge[{judge_gate['current']}]", text_alignment="right")
            st.caption(judge_gate["summary"])
        if supplement_rows:
            st.caption(
                f"이미 입력한 타당성 보완 정보 {len(supplement_rows)}건은 보존되어 있으며, "
                "독립 LLM 평가가 PASS된 뒤 타당성 평가 입력으로 반영됩니다."
            )


def _validity_rework_records(rubric: dict, result: dict | None) -> list[dict]:
    rows = _validity_rework_items(rubric, result)
    return rows.to_dict("records") if not rows.empty else []


def _validity_required_supplement_fields(
    rubric: dict,
    result: dict | None,
    rework_records: list[dict],
) -> list[str]:
    result = result or {}
    dimensions = rubric.get("dimensions", {}) if isinstance(rubric.get("dimensions"), dict) else {}
    fields: list[str] = []
    mandatory_records = [record for record in rework_records if bool(record.get("_필수보완"))]
    target_records = mandatory_records or rework_records
    for record in target_records:
        dimension_key = str(record.get("_항목키") or "").strip()
        if not dimension_key:
            label = str(record.get("평가 항목") or "")
            dimension_key = next(
                (
                    key for key, spec in dimensions.items()
                    if str(spec.get("label") or key) == label
                ),
                "",
            )
        fields.extend(VALIDITY_DIMENSION_SUPPLEMENT_FIELD_KEYS.get(dimension_key, ()))
    for rule in _validity_immediate_holds(result):
        fields.extend(VALIDITY_HOLD_SUPPLEMENT_FIELD_KEYS.get(str(rule), ()))
    return _validity_ordered_supplement_fields(fields)


def _validity_pass_basis_items(result: dict | None, rubric: dict) -> list[dict]:
    result = result or {}
    rule = _validity_ai_pass_rule(rubric)
    pass_floor = float(rule.get("min_score", 80) or 80)
    total_score = _validity_score_value(result.get("total_score"))
    holds = _validity_immediate_holds(result)
    items = [
        {
            "icon": "score",
            "label": "총점 기준",
            "value": "-" if total_score is None else f"{total_score:g} / 100점",
            "detail": f"AI_PASS 기준 {pass_floor:g}점",
            "tone": "green" if total_score is not None and total_score >= pass_floor else "orange",
        },
        {
            "icon": "rule",
            "label": "항목별 기준",
            "value": "충족" if result.get("all_pass_floors_met") else "확인 필요",
            "detail": "각 평가 항목의 통과 하한 확인",
            "tone": "green" if result.get("all_pass_floors_met") else "orange",
        },
        {
            "icon": "block",
            "label": "즉시 보류",
            "value": "0건" if not holds else f"{len(holds)}건",
            "detail": "보류 규칙이 없어야 QA 검토 가능",
            "tone": "green" if not holds else "red",
        },
    ]
    scores = result.get("dimension_scores", {}) if isinstance(result.get("dimension_scores"), dict) else {}
    passed_dimensions = []
    for key, spec in rubric.get("dimensions", {}).items():
        score_detail = scores.get(key, {}) if isinstance(scores.get(key), dict) else {}
        score = _validity_score_value(score_detail.get("score"))
        pass_floor_value = _validity_score_value(spec.get("pass_floor"))
        if score is not None and pass_floor_value is not None and score >= pass_floor_value:
            passed_dimensions.append(str(spec.get("label") or key))
    if passed_dimensions:
        items.append(
            {
                "icon": "fact_check",
                "label": "통과 항목",
                "value": f"{len(passed_dimensions)}개",
                "detail": ", ".join(passed_dimensions[:2]) + (" 외" if len(passed_dimensions) > 2 else ""),
                "tone": "green",
            }
        )
    return items


def _validity_supplement_context_model(
    candidate: dict,
    artifacts: dict,
    result: dict | None,
    rubric: dict | None,
) -> dict:
    result = result or {}
    rubric = rubric or {}
    supplement = artifacts.get("validity_supplement", {})
    if not isinstance(supplement, dict):
        supplement = {}
    filled_fields = _validity_filled_supplement_field_keys(supplement)
    filled_evaluation_fields = _validity_filled_supplement_field_keys(supplement, include_note=False)
    decision = str(result.get("decision") or candidate.get("validity_status") or "NOT_RUN").strip().upper()
    rework_records = _validity_rework_records(rubric, result) if result else []
    required_fields = _validity_required_supplement_fields(rubric, result, rework_records) if result else []
    required_fields = _validity_ordered_supplement_fields([*required_fields, *filled_evaluation_fields])
    supplement_applied = _validity_supplement_applied_to_result(artifacts, result)

    if not result:
        return {
            "phase": "pre_evaluation",
            "title": "사전 보완 입력",
            "caption": "최초 평가는 보완 없이 먼저 실행해도 됩니다. 이미 알고 있는 담당·일정·근거가 있을 때만 입력하세요.",
            "badge": "선택 사항",
            "tone": "gray",
            "field_keys": list(VALIDITY_EVALUATION_SUPPLEMENT_FIELD_KEYS),
            "expanded": bool(filled_evaluation_fields),
            "rework_records": [],
            "pass_basis": [],
            "save_label": "사전 보완 입력 저장",
            "draft_label": "비어 있는 항목 초안 채우기",
            "empty_message": "사전 보완 입력은 선택 사항입니다. 먼저 평가를 실행하면 부족 항목이 자동으로 표시됩니다.",
        }

    if decision == "AI_PASS" and not _validity_immediate_holds(result):
        return {
            "phase": "pass_review",
            "title": "AI_PASS 근거 확인",
            "caption": "점수를 올리기 위한 보완이 아니라, QA 검토자가 왜 통과 가능한지 확인하는 메모를 남기는 단계입니다.",
            "badge": "확인 메모",
            "tone": "green",
            "field_keys": ["note"],
            "expanded": True,
            "rework_records": [],
            "pass_basis": _validity_pass_basis_items(result, rubric),
            "save_label": "확인 메모 저장",
            "draft_label": "확인 메모 초안 채우기",
            "empty_message": "AI_PASS 상태입니다. QA 검토 전 통과 근거 확인 메모만 남기면 됩니다.",
        }

    if filled_evaluation_fields and not supplement_applied:
        title = "보완 입력 반영 대기"
        caption = "저장된 보완 입력이 현재 타당성 평가 결과에 아직 반영되지 않았습니다. 저장 내용을 확인한 뒤 재평가하세요."
        badge = "재평가 필요"
        tone = "orange"
        phase = "pending_reevaluation"
    else:
        title = "부족 항목 보완 입력"
        caption = "평가 결과에서 부족한 항목에 필요한 입력란만 표시합니다. 비어 있는 값은 초안 채우기로 시작할 수 있습니다."
        badge = "보완 필요" if rework_records else "확인 필요"
        tone = "red" if rework_records else "orange"
        phase = "rework"

    field_keys = required_fields or list(VALIDITY_EVALUATION_SUPPLEMENT_FIELD_KEYS)
    return {
        "phase": phase,
        "title": title,
        "caption": caption,
        "badge": badge,
        "tone": tone,
        "field_keys": field_keys,
        "expanded": True,
        "rework_records": rework_records,
        "pass_basis": [],
        "save_label": "보완 입력 저장",
        "draft_label": "필요 항목 초안 채우기",
        "empty_message": "부족 항목을 기준으로 필요한 보완 정보를 입력하세요.",
    }


def _validity_trace_id(artifacts: dict) -> str:
    trace = artifacts.get("trace", {}) if isinstance(artifacts.get("trace"), dict) else {}
    return str(trace.get("trace_id") or "").strip()


def _validity_draft_dates() -> dict[str, str]:
    base = datetime.now().date()
    return {
        "start": (base + timedelta(days=1)).isoformat(),
        "cause_done": (base + timedelta(days=8)).isoformat(),
        "fix_done": (base + timedelta(days=15)).isoformat(),
        "qa_done": (base + timedelta(days=22)).isoformat(),
        "deploy": (base + timedelta(days=30)).isoformat(),
    }


def _validity_case_subject(candidate: dict) -> str:
    question = str(candidate.get("question") or "").strip()
    if not question:
        return "VOC 대상 업무"
    if "앱" in question or "모바일" in question:
        return "모바일 앱 VOC 처리"
    if "보험" in question or "갱신" in question:
        return "보험 갱신 VOC 처리"
    return "선택 VOC 처리"


def _validity_autofill_value(
    field: str,
    candidate: dict,
    artifacts: dict,
    result: dict | None,
    rubric: dict,
    context: dict,
) -> str:
    result = result or {}
    trace_id = _validity_trace_id(artifacts)
    run_id = candidate.get("run_id", "-")
    case_id = candidate.get("case_id", "-")
    question = candidate.get("question") or "-"
    score = result.get("total_score", "-")
    decision = _voc_status_label(result.get("decision", "평가 전"))
    subject = _validity_case_subject(candidate)
    dates = _validity_draft_dates()
    if field == "owner":
        return "\n".join(
            [
                "담당/오너 확인",
                f"  - 주관 조직 : {subject} 담당 조직",
                "  - 개발 담당 : IT개발팀 담당자 확인 필요",
                "  - 업무 담당 : 상품/서비스 운영 담당자 확인 필요",
                "  - QA 검토 주체 : QA 담당자 확인 필요",
                "  - 업무 승인 주체 : 상품/서비스 오너 확인 필요",
            ]
        )
    if field == "schedule":
        return "\n".join(
            [
                "일정 확인",
                f"  - 착수일 : {dates['start']}",
                f"  - 원인 분석 완료일 : {dates['cause_done']}",
                f"  - 수정 완료일 : {dates['fix_done']}",
                f"  - QA 검증일 : {dates['qa_done']}",
                f"  - 배포 목표일 : {dates['deploy']}",
            ]
        )
    if field == "kpi":
        return "\n".join(
            [
                "정량 KPI 확인",
                "  - 동일 VOC 재문의율 : 현재 확인 필요 → 목표 20% 감소",
                "  - 오류/실패율 : 현재 확인 필요 → 목표 0.5% 이하",
                "  - 평균 처리시간 : 현재 확인 필요 → 목표 30% 단축",
                "  - 고객 재문의 건수 : 현재 확인 필요 → 목표 주간 3건 이하",
                "  - 측정 주기 : 배포 후 4주간 주 1회 확인",
                "  - 완료 기준 : 목표 KPI 2개 이상 충족 및 중대 결함 0건",
            ]
        )
    if field == "priority":
        return "\n".join(
            [
                "우선순위 확인",
                "  - 우선순위 : P1",
                "  - 판단 기준 : 고객 영향도 높음 / 재발 가능성 있음 / 배포 영향 중간",
                "  - 우선 보완 대상 : 통과 기준 미달 항목 먼저 조치",
                "  - 후속 보완 대상 : 보완 권장 항목은 QA 검토 전 확인",
            ]
        )
    if field == "execution_plan":
        return "\n".join(
            [
                "실행계획/적용범위 확인",
                f"  - 적용 범위 : {subject} 관련 화면/API/업무 절차",
                "  - 변경 대상 : 상태 표시 로직, 오류 메시지, 재시도/자동저장 처리, 운영 모니터링 기준",
                "  - 수행 단계 :",
                "    1. VOC·Trace 근거와 재현 조건 확인",
                "    2. 원인 분석 및 변경 대상 확정",
                "    3. 화면/API/업무 절차 수정",
                "    4. QA 검증 및 회귀 테스트 수행",
                "    5. 배포 후 KPI 모니터링",
                "  - 완료 기준 : QA 검증 PASS, 중대 결함 0건, 목표 KPI 충족",
            ]
        )
    if field == "evidence":
        evidence_parts = [f"Run {run_id}", f"Case {case_id}"]
        if trace_id:
            evidence_parts.append(f"Trace {trace_id}")
        evidence_parts.append(f"질문: {question}")
        return "\n".join(
            [
                "VOC·실행 Trace 근거 확인",
                f"  - Run ID : {run_id}",
                f"  - Case ID : {case_id}",
                f"  - Trace ID : {trace_id or '확인 필요'}",
                f"  - 사용자 질문 : {question}",
                "  - 추가 VOC ID : 확인 필요",
                f"  - 기준 근거 : {' · '.join(evidence_parts)}",
            ]
        )
    if field == "risk":
        return "\n".join(
            [
                "리스크/우회방안 확인",
                "  - 고객 영향 리스크 : 오류 재발 시 고객 안내 지연 가능성 확인",
                "  - 운영 리스크 : 장애 재발, 상담 전환 증가, 처리 지연 가능성 확인",
                "  - 보안·규정 리스크 : 개인정보 저장/마스킹/접근통제 영향 확인",
                "  - 임시 우회 방안 : 상담 스크립트, 안내 배너, 수동 처리 절차 준비",
                "  - 롤백 기준 : 배포 후 중대 결함 발생 시 즉시 이전 정책으로 복구",
            ]
        )
    if field == "note":
        if context.get("phase") == "pass_review":
            holds = _validity_immediate_holds(result)
            return "\n".join(
                [
                    "AI_PASS 근거 확인",
                    f"  - 개선안 타당성 점수 : {score}/100점",
                    f"  - 판정 : {decision}",
                    f"  - 즉시 보류 규칙 : {len(holds)}건",
                    "  - QA 검토 전 확인 : VOC·Trace 근거, 담당, 일정, KPI 확인 완료",
                ]
            )
        return "\n".join(
            [
                "검토 메모",
                "  - 보완 입력 출처 : 담당자 확인 필요",
                "  - 추가 확인 필요 사항 : 확인 필요",
                "  - 재평가 전 확인 : 필수 보완 항목 입력 완료 여부 확인",
            ]
        )
    return ""


def _render_validity_rework_need_cards(context: dict) -> None:
    records = context.get("rework_records") or []
    if not records:
        return
    st.caption("아래 부족 항목을 기준으로 입력란을 자동 구성했습니다.")
    columns = st.columns(min(3, len(records)), gap="small")
    for index, record in enumerate(records):
        column = columns[index % len(columns)]
        score = record.get("현재 점수")
        max_points = record.get("최고점")
        with column.container(border=True, height=140):
            st.caption(f":material/report: {record.get('평가 항목', '-')}")
            st.markdown(
                f"**{'-' if score is None else f'{score:g}'} / "
                f"{'-' if max_points is None else f'{max_points:g}'}점**"
            )
            st.caption(f"{record.get('상태', '-')} · {record.get('보완 지시', '-')}")


def _render_validity_pass_basis_cards(context: dict) -> None:
    basis = context.get("pass_basis") or []
    if not basis:
        return
    columns = st.columns(min(4, len(basis)), gap="small")
    for column, item in zip(columns, basis, strict=False):
        tone = {
            "green": "green",
            "orange": "orange",
            "red": "red",
            "gray": "gray",
        }.get(item.get("tone", "gray"), "gray")
        with column.container(border=True, height=122):
            st.caption(f":material/{item.get('icon', 'fact_check')}: {item.get('label', '-')}")
            st.markdown(f":{tone}-badge[{item.get('value', '-')}]")
            st.caption(item.get("detail", "-"))


def _render_validity_supplement_form_body(
    candidate: dict,
    artifacts: dict,
    current: dict,
    context: dict,
    rubric: dict,
    *,
    key_scope: str = "",
    compact: bool = False,
) -> dict | None:
    run_id = candidate["run_id"]
    case_id = candidate["case_id"]
    field_keys = context.get("field_keys") or []
    for field, _label, _placeholder in VALIDITY_SUPPLEMENT_FIELDS:
        key = _validity_supplement_key(run_id, case_id, field, key_scope)
        if key not in st.session_state:
            st.session_state[key] = str(current.get(field) or "")

    if field_keys:
        with st.container(horizontal=True, horizontal_alignment="right"):
            if st.button(
                context.get("draft_label", "초안 채우기"),
                icon=":material/auto_fix_high:",
                width="content",
                key=f"validity_supplement_draft_{_validity_key_base(run_id, case_id, key_scope)}",
            ):
                for field in field_keys:
                    key = _validity_supplement_key(run_id, case_id, field, key_scope)
                    if not str(st.session_state.get(key) or "").strip():
                        st.session_state[key] = _validity_autofill_value(
                            field,
                            candidate,
                            artifacts,
                            artifacts.get("validity_result", {}),
                            rubric,
                            context,
                        )
                st.toast("비어 있는 입력란에 초안을 채웠습니다. 확인 후 저장하세요.", icon=":material/edit:")

    with st.form(f"validity_supplement_form_{_validity_key_base(run_id, case_id, key_scope)}", border=False):
        field_specs = [_validity_supplement_field_spec(field) for field in field_keys]
        if not field_specs:
            st.caption(context.get("empty_message", "입력할 보완 항목이 없습니다."))
        for start in range(0, len(field_specs), 3):
            row_specs = field_specs[start:start + 3]
            columns = st.columns(len(row_specs), gap="small") if len(row_specs) > 1 else [st.container()]
            for column, (field, label, placeholder) in zip(columns, row_specs, strict=False):
                with column:
                    field_label = "QA 확인 메모" if field == "note" and context.get("phase") == "pass_review" else label
                    st.text_area(
                        field_label,
                        placeholder=placeholder,
                        key=_validity_supplement_key(run_id, case_id, field, key_scope),
                        height=(92 if field in {"note", "execution_plan"} else 76) if compact else (104 if field in {"note", "execution_plan"} else 88),
                    )
        submitted = st.form_submit_button(
            context.get("save_label", "보완 입력 저장"),
            type="primary",
            icon=":material/save:",
            width="stretch",
        )
    if not submitted:
        return None
    return {
        field: st.session_state.get(
            _validity_supplement_key(run_id, case_id, field, key_scope),
            "",
        )
        for field, _label, _placeholder in VALIDITY_SUPPLEMENT_FIELDS
    }


def _render_validity_supplement_editor(
    candidate: dict,
    artifacts: dict,
    result: dict | None = None,
    rubric: dict | None = None,
    *,
    key_scope: str = "",
    read_only: bool = False,
    compact: bool = False,
) -> dict:
    current = artifacts.get("validity_supplement", {})
    if not isinstance(current, dict):
        current = {}
    run_id = candidate["run_id"]
    case_id = candidate["case_id"]
    filled_rows = _validity_supplement_fields(current)
    rubric = rubric or load_improvement_validity_rubric()
    result = result or artifacts.get("validity_result", {})
    context = _validity_supplement_context_model(candidate, artifacts, result, rubric)
    badge_color = {
        "green": "green",
        "blue": "blue",
        "orange": "orange",
        "red": "red",
        "gray": "gray",
    }.get(context.get("tone", "gray"), "gray")

    with st.container(border=True):
        heading, state = st.columns([2.6, 1], vertical_alignment="center")
        with heading:
            st.markdown(f"#### {context['title']}")
            if compact:
                field_labels = [
                    _validity_supplement_field_spec(field)[1]
                    for field in context.get("field_keys", [])
                ]
                if field_labels:
                    st.caption("입력 대상: " + " · ".join(field_labels[:4]) + (" 외" if len(field_labels) > 4 else ""))
                else:
                    st.caption(context["caption"])
            else:
                st.caption(context["caption"])
        with state:
            filled_count = len(filled_rows)
            count_text = f" · 저장 {filled_count}건" if filled_count else ""
            st.markdown(
                f":{badge_color}-badge[{context['badge']}{count_text}]",
                text_alignment="right",
            )

        if not compact:
            _render_validity_rework_need_cards(context)
            _render_validity_pass_basis_cards(context)

        if read_only:
            payload = None
            if filled_rows:
                st.caption("저장된 보완 입력과 확인 메모를 조회합니다.")
        elif context["phase"] == "pre_evaluation":
            with st.expander(
                "입력 열기",
                expanded=context["expanded"],
                icon=":material/edit_note:",
            ):
                payload = _render_validity_supplement_form_body(
                    candidate,
                    artifacts,
                    current,
                    context,
                    rubric,
                    key_scope=key_scope,
                    compact=compact,
                )
        else:
            payload = _render_validity_supplement_form_body(
                candidate,
                artifacts,
                current,
                context,
                rubric,
                key_scope=key_scope,
                compact=compact,
            )

        if payload is not None:
            if not any(str(value or "").strip() for value in payload.values()):
                st.toast("저장할 입력을 먼저 작성하세요.", icon=":material/info:")
            else:
                try:
                    saved = save_voc_validity_supplement(run_id, case_id, payload)
                    artifacts["validity_supplement"] = saved.get("validity_supplement", {})
                    _load_validity_candidates.clear()
                    if context["phase"] == "pass_review":
                        st.session_state.voc_validity_focus_notice = (
                            "AI_PASS 근거 확인 메모를 저장했습니다. QA 검토 단계에서 참고할 수 있습니다."
                        )
                        st.session_state.voc_validity_notice = "AI_PASS 근거 확인 메모를 저장했습니다."
                    elif context["phase"] == "pre_evaluation":
                        st.session_state.voc_validity_focus_notice = (
                            "사전 보완 입력을 저장했습니다. 최초 개선안 타당성 평가에 함께 반영됩니다."
                        )
                        st.session_state.voc_validity_notice = "사전 보완 입력을 저장했습니다."
                    else:
                        st.session_state.voc_validity_focus_notice = (
                            "보완 입력을 저장했습니다. 아래 평가 설정에서 재평가를 실행하면 QA 검토 가능 여부가 다시 계산됩니다."
                        )
                        st.session_state.voc_validity_notice = "개선안 타당성 평가 보완 입력을 저장했습니다. 다음 개선안 타당성 평가에 반영됩니다."
                    st.rerun()
                except Exception as exc:
                    st.error(f"보완 입력 저장 실패: {type(exc).__name__}: {exc}")

        if filled_rows:
            if compact:
                with st.expander(
                    f"저장된 입력 {len(filled_rows)}건 보기",
                    expanded=False,
                    icon=":material/list_alt:",
                ):
                    st.dataframe(
                        pd.DataFrame(filled_rows),
                        hide_index=True,
                        width="stretch",
                        height=min(178, 72 + len(filled_rows) * 34),
                        column_config={
                            "보완 항목": st.column_config.TextColumn(width="medium", pinned=True),
                            "입력 내용": st.column_config.TextColumn(width="large"),
                        },
                    )
                return artifacts
            st.dataframe(
                pd.DataFrame(filled_rows),
                hide_index=True,
                width="stretch",
                height=min(178, 72 + len(filled_rows) * 34),
                column_config={
                    "보완 항목": st.column_config.TextColumn(width="medium", pinned=True),
                    "입력 내용": st.column_config.TextColumn(width="large"),
                },
            )
        else:
            st.caption(context["empty_message"])

    return artifacts


def _render_validity_qa_gate_cards(candidate: dict, result: dict | None):
    gate = _validity_qa_gate_model(candidate, result)
    with st.container(border=True):
        heading, summary = st.columns([2.4, 1], vertical_alignment="center")
        with heading:
            st.markdown("#### QA 검토 가능 조건")
            st.caption("QA 검토는 개선안 타당성 평가 이후 아래 조건이 모두 충족될 때만 진행합니다.")
        with summary:
            if gate["ready"] or gate["completed"]:
                st.markdown(f":green-badge[{gate['summary']}]", text_alignment="right")
            elif gate["workflow_state"] == "QA_REVIEWED":
                st.markdown(f":blue-badge[{gate['summary']}]", text_alignment="right")
            else:
                st.markdown(f":gray-badge[{gate['summary']}]", text_alignment="right")
        columns = st.columns(4, gap="small")
        for column, check in zip(columns, gate["checks"], strict=False):
            icon = ":material/check_circle:" if check["ok"] else ":material/pending:"
            with column.container(border=True, height="stretch"):
                st.markdown(f"{icon} **{check['label']}**")
                st.markdown(f"##### {check['value']}")
                st.caption(check["detail"])
        if gate["holds"]:
            st.caption("보류 규칙: " + ", ".join(gate["holds"]))


def _validity_hold_rule_label(rule: str) -> str:
    return VALIDITY_HOLD_RULE_LABELS.get(str(rule), str(rule))


def _validity_criteria_text(criteria: dict | None) -> str:
    if not isinstance(criteria, dict) or not criteria:
        return "-"
    return " · ".join(
        f"{VALIDITY_CRITERIA_LABELS.get(str(key), str(key))} {value}점"
        for key, value in criteria.items()
    )


def _validity_ai_pass_floor(rubric: dict) -> str:
    for item in rubric.get("automatic_decisions", []):
        if item.get("decision") == "AI_PASS":
            score = item.get("min_score", "-")
            floor = "항목별 통과 기준 충족" if item.get("requires_all_pass_floors") else "항목별 통과 기준 미적용"
            return f"{score}점 이상 · {floor}"
    return "80점 이상 · 항목별 통과 기준 충족"


def _validity_ai_pass_rule(rubric: dict) -> dict:
    return next(
        (
            item
            for item in rubric.get("automatic_decisions", [])
            if item.get("decision") == "AI_PASS"
        ),
        {"min_score": 80, "requires_all_pass_floors": True},
    )


def _validity_ai_pass_failure_model(
    result: dict | None,
    rubric: dict,
    artifacts: dict | None = None,
) -> dict:
    result = result or {}
    artifacts = artifacts or {}
    ai_pass_rule = _validity_ai_pass_rule(rubric)
    threshold = float(ai_pass_rule.get("min_score", 80) or 80)
    total_score = _validity_score_value(result.get("total_score"))
    score_gap = (
        round(max(threshold - total_score, 0), 2)
        if total_score is not None
        else None
    )

    dimension_scores = (
        result.get("dimension_scores", {})
        if isinstance(result.get("dimension_scores"), dict)
        else {}
    )
    floor_misses = []
    for key, spec in rubric.get("dimensions", {}).items():
        score_detail = (
            dimension_scores.get(key, {})
            if isinstance(dimension_scores.get(key), dict)
            else {}
        )
        score = _validity_score_value(score_detail.get("score"))
        floor = float(spec.get("pass_floor", 0) or 0)
        if score is not None and score >= floor:
            continue
        floor_misses.append(
            {
                "key": key,
                "label": _voc_display_term(spec.get("label", key)),
                "score": score,
                "floor": floor,
                "gap": None if score is None else round(max(floor - score, 0), 2),
                "reason": _voc_display_term(score_detail.get("reason", "-") or "-"),
            }
        )

    holds = _validity_immediate_holds(result)
    evidence_hold = "missing_voc_or_trace_evidence" in holds
    other_holds = [rule for rule in holds if rule != "missing_voc_or_trace_evidence"]

    trace = artifacts.get("trace", {}) if isinstance(artifacts.get("trace"), dict) else {}
    trace_events = trace.get("events", []) if isinstance(trace.get("events"), list) else []
    evidence_score = dimension_scores.get("evidence_traceability", {})
    if not isinstance(evidence_score, dict):
        evidence_score = {}
    evidence_spec = rubric.get("dimensions", {}).get("evidence_traceability", {})
    evidence_value = _validity_score_value(evidence_score.get("score"))
    evidence_floor = float(evidence_spec.get("pass_floor", 0) or 0)
    evidence_floor_miss = (
        evidence_value is not None
        and evidence_value < evidence_floor
    )
    trace_loaded = "trace" in artifacts and isinstance(artifacts.get("trace"), dict)
    evidence_details = []
    if evidence_hold:
        evidence_details.append("서버 검증에서 VOC 또는 실행 Trace 근거 누락이 감지되었습니다.")
    if trace_loaded and not trace.get("trace_id"):
        evidence_details.append("실행 Trace ID가 없습니다.")
    if trace_loaded and not trace_events:
        evidence_details.append("실행 Trace 이벤트가 없습니다.")
    if evidence_floor_miss:
        evidence_details.append(
            f"VOC·실행 Trace 근거 추적성 {evidence_value:g}점 / 하한 {evidence_floor:g}점"
        )

    evaluated = bool(result)
    score_failed = evaluated and (total_score is None or total_score < threshold)
    floors_failed = evaluated and bool(floor_misses)
    holds_failed = evaluated and bool(other_holds)
    evidence_failed = evaluated and bool(evidence_details)
    categories = [
        {
            "key": "score",
            "label": "점수 부족",
            "icon": ":material/score:",
            "failed": score_failed,
            "value": (
                "평가 전"
                if total_score is None
                else f"{total_score:g} / {threshold:g}점"
            ),
            "summary": (
                "개선안 타당성 평가 점수가 없습니다."
                if total_score is None
                else f"통과까지 {score_gap:g}점 부족합니다."
                if score_failed
                else "총점 기준을 충족했습니다."
            ),
            "details": (
                []
                if not score_failed
                else [f"현재 {total_score:g}점 · AI 평가 통과 기준 {threshold:g}점"]
                if total_score is not None
                else ["평가 점수가 산출되지 않았습니다."]
            ),
            "action": "낮은 평가 항목의 개선안과 실행 근거를 보완하세요.",
        },
        {
            "key": "floors",
            "label": "항목별 하한 미달",
            "icon": ":material/low_priority:",
            "failed": floors_failed,
            "value": "평가 전" if not evaluated else f"{len(floor_misses)}개",
            "summary": (
                "항목별 점수가 없습니다."
                if not evaluated
                else f"{len(floor_misses)}개 항목이 통과 하한에 미달했습니다."
                if floors_failed
                else "모든 항목별 하한을 충족했습니다."
            ),
            "details": [
                (
                    f"{item['label']}: 미평가 / 하한 {item['floor']:g}점"
                    if item["score"] is None
                    else f"{item['label']}: {item['score']:g}점 / 하한 {item['floor']:g}점"
                )
                for item in floor_misses
            ],
            "action": "미달 항목별 판정 근거와 보완 지시를 확인하세요.",
        },
        {
            "key": "holds",
            "label": "즉시 보류 규칙",
            "icon": ":material/block:",
            "failed": holds_failed,
            "value": "평가 전" if not evaluated else f"{len(other_holds)}건",
            "summary": (
                "보류 규칙을 아직 확인하지 않았습니다."
                if not evaluated
                else f"{len(other_holds)}건의 즉시 보류 규칙이 적용됐습니다."
                if holds_failed
                else "근거 부족 외 즉시 보류 규칙은 없습니다."
            ),
            "details": [_validity_hold_rule_label(rule) for rule in other_holds],
            "action": "보류 사유를 해소한 증적을 확보한 뒤 다시 평가하세요.",
        },
        {
            "key": "evidence",
            "label": "VOC·실행 Trace 근거 부족",
            "icon": ":material/account_tree:",
            "failed": evidence_failed,
            "value": "평가 전" if not evaluated else f"{len(evidence_details)}건",
            "summary": (
                "근거 상태를 아직 평가하지 않았습니다."
                if not evaluated
                else "VOC·실행 Trace 근거를 보완해야 합니다."
                if evidence_failed
                else "VOC·실행 Trace 근거 부족이 감지되지 않았습니다."
            ),
            "details": evidence_details,
            "action": "VOC ID, 실행 Trace ID, Agent별 전달 근거와 판단 근거를 연결하세요.",
        },
    ]
    failed_categories = [item for item in categories if item["failed"]]
    return {
        "evaluated": evaluated,
        "passed": result.get("decision") == "AI_PASS" and not failed_categories,
        "threshold": threshold,
        "failed_count": len(failed_categories),
        "categories": categories,
    }


def _render_validity_ai_pass_failures(
    result: dict | None,
    rubric: dict,
    artifacts: dict | None = None,
):
    model = _validity_ai_pass_failure_model(result, rubric, artifacts)
    if not model["evaluated"]:
        return

    with st.container(border=True):
        heading, state = st.columns([2.5, 1], vertical_alignment="center")
        with heading:
            st.markdown("#### AI 평가 통과 진단")
            st.caption("통과 실패 원인을 총점·항목 하한·즉시 보류·VOC·실행 Trace 근거로 나누어 확인합니다.")
        with state:
            if model["passed"]:
                st.markdown(":green-badge[AI 평가 통과 조건 충족]", text_alignment="right")
            else:
                st.markdown(
                    f":red-badge[실패 원인 {model['failed_count']}개 유형]",
                    text_alignment="right",
                )

        columns = st.columns(4, gap="small")
        for column, category in zip(columns, model["categories"], strict=False):
            with column.container(border=True, height="stretch"):
                st.markdown(f"{category['icon']} **{category['label']}**")
                st.markdown(f"##### {category['value']}")
                if category["failed"]:
                    st.markdown(":red-badge[확인 필요]")
                else:
                    st.markdown(":green-badge[충족]")
                st.caption(category["summary"])

        failure_rows = []
        for category in model["categories"]:
            if not category["failed"]:
                continue
            details = category["details"] or [category["summary"]]
            for detail in details:
                failure_rows.append(
                    {
                        "실패 유형": category["label"],
                        "확인된 원인": detail,
                        "필요 조치": category["action"],
                    }
                )
        if failure_rows:
            st.dataframe(
                pd.DataFrame(failure_rows),
                hide_index=True,
                width="stretch",
                height=min(270, 76 + len(failure_rows) * 42),
                column_config={
                    "실패 유형": st.column_config.TextColumn(width="medium", pinned=True),
                    "확인된 원인": st.column_config.TextColumn(width="large"),
                    "필요 조치": st.column_config.TextColumn(width="large"),
                },
            )
        else:
            st.success("총점, 항목별 하한, 즉시 보류, VOC·실행 Trace 근거 조건을 모두 충족했습니다.")


def _validity_dimension_rows(rubric: dict, result: dict | None = None) -> pd.DataFrame:
    result = result or {}
    scores = result.get("dimension_scores", {}) if isinstance(result.get("dimension_scores"), dict) else {}
    rows = []
    for key, spec in rubric.get("dimensions", {}).items():
        score_detail = scores.get(key, {}) if isinstance(scores.get(key), dict) else {}
        max_points = float(spec.get("max_points", 0) or 0)
        pass_floor = float(spec.get("pass_floor", 0) or 0)
        score = score_detail.get("score")
        score_value = None
        if score is not None:
            try:
                score_value = float(score)
            except (TypeError, ValueError):
                score_value = None
        if score_value is None:
            status = "평가 전"
            ratio = None
        elif score_value >= pass_floor:
            status = "기준 충족"
            ratio = round(score_value / max_points * 100, 1) if max_points else None
        else:
            status = "기준 미달"
            ratio = round(score_value / max_points * 100, 1) if max_points else None
        rows.append(
            {
                "평가 항목": _voc_display_term(spec.get("label", key)),
                "배점": max_points,
                "통과 기준": pass_floor,
                "결과 점수": score_value,
                "달성률": ratio,
                "판정": status,
                "세부 지표": _voc_display_term(_validity_criteria_text(spec.get("criteria", {}))),
                "판정 근거": _voc_display_term(score_detail.get("reason", "-") or "-"),
            }
        )
    return pd.DataFrame(rows)


def _render_validity_dimension_scorecard(rubric: dict, result: dict | None = None):
    result = result or {}
    rows = _validity_dimension_rows(rubric, result)
    with st.container(border=True):
        heading, summary = st.columns([2.4, 1.15], vertical_alignment="center")
        with heading:
            st.markdown("#### 평가 항목과 점수 지표")
            st.caption("개선안 타당성 평가는 아래 항목별 배점과 통과 기준을 적용해 최종 개선안을 평가합니다.")
        with summary:
            total_score = result.get("total_score")
            value = "-" if total_score is None else f"{total_score:g} / 100점"
            decision = _voc_status_label(result.get("decision", "평가 전"))
            st.metric("현재 평가 결과", value, delta=decision, border=True)
        with st.container(horizontal=True):
            st.metric("Rubric 버전", rubric.get("version", "-"), border=True)
            st.metric("평가 항목", f"{len(rows)}개", border=True)
            st.metric("총 배점", f"{int(rows['배점'].sum()) if not rows.empty else 0}점", border=True)
            st.metric("AI 평가 통과 기준", _validity_ai_pass_floor(rubric), border=True)
        st.dataframe(
            rows,
            hide_index=True,
            width="stretch",
            height=min(310, 78 + len(rows) * 44),
            column_config={
                "평가 항목": st.column_config.TextColumn(width="medium", pinned=True),
                "배점": st.column_config.NumberColumn(format="%g점", width="small"),
                "통과 기준": st.column_config.NumberColumn(format="%g점", width="small"),
                "결과 점수": st.column_config.NumberColumn(format="%g점", width="small"),
                "달성률": st.column_config.ProgressColumn(
                    min_value=0,
                    max_value=100,
                    format="%.0f%%",
                    width="small",
                ),
                "판정": st.column_config.TextColumn(width="small"),
                "세부 지표": st.column_config.TextColumn(width="large"),
                "판정 근거": st.column_config.TextColumn(width="large"),
            },
        )


def _validity_execution_step_rows(
    candidate: dict,
    artifacts: dict,
    result: dict | None,
    rubric: dict,
) -> pd.DataFrame:
    result = result or {}
    basis = _validity_selection_basis(candidate, artifacts)
    trace = artifacts.get("trace", {}) if isinstance(artifacts.get("trace"), dict) else {}
    trace_events = trace.get("events", []) if isinstance(trace.get("events"), list) else []
    judge = artifacts.get("judge_result", {}) if isinstance(artifacts.get("judge_result"), dict) else {}
    defects = artifacts.get("defects", {}) if isinstance(artifacts.get("defects"), dict) else {}
    defect_items = defects.get("defects", []) if isinstance(defects.get("defects"), list) else []
    supplement_rows = _validity_supplement_fields(artifacts.get("validity_supplement", {}))
    supplement_labels = ", ".join(row["보완 항목"] for row in supplement_rows)
    dimensions = rubric.get("dimensions", {}) if isinstance(rubric.get("dimensions"), dict) else {}
    has_result = bool(result)
    holds = _validity_immediate_holds(result)
    gate = _validity_qa_gate_model(candidate, result)
    attempts = result.get("attempts", []) if isinstance(result.get("attempts"), list) else []
    successful_attempts = sum(1 for item in attempts if item.get("status") == "SUCCESS")
    return pd.DataFrame(
        [
            {
                "순서": 1,
                "수행 절차": "대상 증적 수집",
                "상태": "완료" if basis["pipeline_success"] else "확인 필요",
                "절차별 결과": basis["pipeline_label"],
                "확인 내용": (
                    f"실행 Trace {len(trace_events)}건 · 독립 LLM 평가 {_voc_status_label(judge.get('decision', 'NOT_RUN'))} · "
                    f"미종결 결함 후보 {len(defect_items)}건"
                ),
            },
            {
                "순서": 2,
                "수행 절차": "보완 입력 반영",
                "상태": "완료" if supplement_rows else "미입력",
                "절차별 결과": f"{len(supplement_rows)}개 항목" if supplement_rows else "원본 개선안 기준",
                "확인 내용": (
                    f"다음 개선안 타당성 평가 입력에 포함: {supplement_labels}"
                    if supplement_rows
                    else "보완 입력이 없으면 Agent 파이프라인 최종 개선안만 기준으로 평가합니다."
                ),
            },
            {
                "순서": 3,
                "수행 절차": "평가 기준 구성",
                "상태": "완료" if dimensions else "확인 필요",
                "절차별 결과": f"{len(dimensions)}개 항목 · 총 {sum(float(item.get('max_points', 0) or 0) for item in dimensions.values()):g}점",
                "확인 내용": f"Rubric {rubric.get('version', '-')} · AI 통과 기준 {_validity_ai_pass_floor(rubric)}",
            },
            {
                "순서": 4,
                "수행 절차": "독립 LLM 평가",
                "상태": "완료" if has_result and result.get("decision") != "ERROR" else ("오류" if result.get("decision") == "ERROR" else "대기"),
                "절차별 결과": (
                    f"{result.get('provider', '-')} / {result.get('model', '-')} · "
                    f"{result.get('duration_seconds', '-')}초"
                    if has_result
                    else "실행 전"
                ),
                "확인 내용": (
                    f"시도 {len(attempts)}회 · 성공 {successful_attempts}회"
                    if has_result
                    else "버튼 실행 시 Provider에 평가 요청을 보냅니다."
                ),
            },
            {
                "순서": 5,
                "수행 절차": "점수·판정 산출",
                "상태": "완료" if has_result and result.get("total_score") is not None else "대기",
                "절차별 결과": (
                    f"{result.get('total_score', '-')}점 · {_voc_status_label(result.get('decision', '-'))}"
                    if has_result
                    else "실행 전"
                ),
                "확인 내용": (
                    "항목별 통과 기준 충족"
                    if result.get("all_pass_floors_met")
                    else "항목별 통과 기준 미충족 또는 평가 전"
                ),
            },
            {
                "순서": 6,
                "수행 절차": "즉시 보류 규칙 확인",
                "상태": "완료" if has_result and not holds else ("보류" if holds else "대기"),
                "절차별 결과": "보류 없음" if has_result and not holds else (f"{len(holds)}건 보류" if holds else "실행 전"),
                "확인 내용": ", ".join(_validity_hold_rule_label(rule) for rule in holds) if holds else "보류 규칙이 없으면 QA 검토 조건으로 이동합니다.",
            },
            {
                "순서": 7,
                "수행 절차": "QA Gate 판정",
                "상태": "가능" if gate["ready"] else ("완료" if gate["completed"] else "대기"),
                "절차별 결과": gate["summary"],
                "확인 내용": (
                    "QA 검토를 시작할 수 있습니다."
                    if gate["ready"]
                    else "대기 조건: " + ", ".join(gate["blocked_reasons"])
                    if gate["blocked_reasons"]
                    else "승인 흐름이 완료되었습니다."
                ),
            },
        ]
    )


def _render_validity_execution_steps(
    candidate: dict,
    artifacts: dict,
    result: dict | None,
    rubric: dict,
):
    rows = _validity_execution_step_rows(candidate, artifacts, result, rubric)
    with st.container(border=True):
        heading, status = st.columns([2.4, 1], vertical_alignment="center")
        with heading:
            st.markdown("#### 개선안 타당성 평가 수행 절차")
            st.caption("선택 대상 개선안 타당성 평가 실행 시 진행되는 절차와 각 절차의 결과입니다.")
        with status:
            finished = bool(result)
            st.markdown(
                ":green-badge[평가 완료]" if finished else ":gray-badge[실행 전]",
                text_alignment="right",
            )
        st.dataframe(
            rows,
            hide_index=True,
            width="stretch",
            height=min(330, 76 + len(rows) * 38),
            column_config={
                "순서": st.column_config.NumberColumn(format="%d", width="small"),
                "수행 절차": st.column_config.TextColumn(width="medium", pinned=True),
                "상태": st.column_config.TextColumn(width="small"),
                "절차별 결과": st.column_config.TextColumn(width="large"),
                "확인 내용": st.column_config.TextColumn(width="large"),
            },
        )


VALIDITY_REWORK_ACTIONS = {
    "cause_linkage": "VOC 불만 원인, 근본 원인, 개선안을 1:1로 연결하고 누락된 우선순위와 고객 영향 범위를 보완하세요.",
    "evidence_traceability": "VOC ID, 실행 Trace ID, 검색 근거, 판단 근거를 항목별로 명시하고 근거가 없는 내용은 확인 필요로 분리하세요.",
    "feasibility": "개선안이 실제 실행계획이 되도록 적용 범위, 변경 대상, 단계별 수행 방법을 구체화하세요.",
    "ownership_schedule_kpi": "담당 조직, 일정, 마일스톤, 완료 기준, 정량 KPI를 숫자와 날짜 중심으로 보완하세요.",
    "risk_security_compliance": "장애·보안·법규·고객 안내 리스크와 우회 방안을 함께 제시하세요.",
}

VALIDITY_REWORK_INPUT_GUIDES = {
    "cause_linkage": (
        ("VOC 원인", "고객 불만 문장과 근본 원인을 1:1로 연결"),
        ("개선 조치", "각 원인별로 적용할 정책·화면·프로세스 변경 대상 작성"),
        ("우선순위", "고객 영향도와 재발 가능성 기준으로 P1/P2/P3 지정"),
    ),
    "evidence_traceability": (
        ("VOC·Trace 근거", "VOC ID, Run ID, Trace 단계, 검색 근거를 항목별로 작성"),
        ("근거 부족", "확인되지 않은 내용은 추정하지 말고 확인 필요로 분리"),
    ),
    "feasibility": (
        ("적용 범위", "화면·API·업무 절차·데이터 중 어디를 바꿀지 명시"),
        ("변경 대상", "수정 대상 시스템, 정책, 조직, 산출물을 구체화"),
        ("수행 방법", "분석 → 설계 → 구현 → QA → 배포 순서로 작성"),
        ("완료 기준", "어떤 증적과 테스트로 완료를 판단할지 작성"),
    ),
    "ownership_schedule_kpi": (
        ("담당", "주관 조직, 협업 조직, QA 검토 주체 지정"),
        ("일정", "착수일, 중간 마일스톤, QA일, 배포 목표일을 날짜로 작성"),
        ("KPI", "오류율·재문의율·처리시간 등 개선 전/후 목표값 작성"),
        ("검증 기준", "측정 방법, 확인 주기, 완료 판정 기준 작성"),
    ),
    "risk_security_compliance": (
        ("운영 리스크", "장애 재발, 고객 안내, 우회 처리 방안 작성"),
        ("보안·규정", "개인정보·접근통제·보관기간·규정 영향 작성"),
        ("롤백", "문제 발생 시 되돌릴 방법과 임시 대응 기준 작성"),
    ),
}


def _clip_text(value, limit: int = 180) -> str:
    text = " ".join(str(value or "").split())
    if not text:
        return "-"
    return text if len(text) <= limit else f"{text[: max(0, limit - 1)].rstrip()}…"


def _validity_score_value(value) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _validity_rework_items(rubric: dict, result: dict | None) -> pd.DataFrame:
    result = result or {}
    scores = result.get("dimension_scores", {}) if isinstance(result.get("dimension_scores"), dict) else {}
    rows = []
    for order, (key, spec) in enumerate(rubric.get("dimensions", {}).items()):
        score_detail = scores.get(key, {}) if isinstance(scores.get(key), dict) else {}
        max_points = float(spec.get("max_points", 0) or 0)
        pass_floor = float(spec.get("pass_floor", 0) or 0)
        score = _validity_score_value(score_detail.get("score"))
        ratio = round(score / max_points * 100, 1) if score is not None and max_points else None
        gap = round(max(pass_floor - (score or 0), 0), 1)
        is_floor_miss = score is None or score < pass_floor
        is_weak = bool(ratio is not None and ratio < 85)
        if not is_floor_miss and not is_weak:
            continue
        if score is None:
            status = "미평가"
        elif is_floor_miss:
            status = "통과 기준 미달"
        else:
            status = "보완 권장"
        rows.append(
            {
                "_항목키": str(key),
                "_필수보완": is_floor_miss,
                "_점수비율": ratio if ratio is not None else -1,
                "_정렬순서": order,
                "평가 항목": spec.get("label", key),
                "구분": "필수 보완" if is_floor_miss else "보완 권장",
                "현재 점수": score,
                "통과 기준": pass_floor,
                "최고점": max_points,
                "부족 점수": gap,
                "상태": status,
                "현재 미흡 근거": score_detail.get("reason", "-") or "-",
                "보완 지시": VALIDITY_REWORK_ACTIONS.get(
                    str(key),
                    "평가 근거에서 부족하다고 지적된 내용을 실행 가능한 보완 항목으로 구체화하세요.",
                ),
            }
        )
    rows.sort(
        key=lambda item: (
            0 if item["_필수보완"] else 1,
            -float(item.get("부족 점수") or 0),
            float(item.get("_점수비율") if item.get("_점수비율") is not None else 999),
            int(item.get("_정렬순서") or 0),
        )
    )
    for index, row in enumerate(rows, start=1):
        row["우선순위"] = index
    columns = [
        "_항목키",
        "_필수보완",
        "_점수비율",
        "_정렬순서",
        "우선순위",
        "구분",
        "평가 항목",
        "현재 점수",
        "통과 기준",
        "최고점",
        "부족 점수",
        "상태",
        "현재 미흡 근거",
        "보완 지시",
    ]
    return pd.DataFrame(rows, columns=columns)


def _validity_rework_field_guidance(dimension_key: str) -> str:
    guides = VALIDITY_REWORK_INPUT_GUIDES.get(str(dimension_key), ())
    if not guides:
        return "평가 근거에서 부족하다고 지적된 내용을 실행 가능한 입력값으로 구체화"
    return " / ".join(f"{label}: {guide}" for label, guide in guides)


def _validity_rework_input_guide_rows(rows: pd.DataFrame) -> pd.DataFrame:
    if rows.empty:
        return pd.DataFrame(columns=["우선순위", "구분", "평가 항목", "작성할 보완 정보"])
    guide_rows = []
    for item in rows.to_dict("records"):
        guide_rows.append(
            {
                "우선순위": item.get("우선순위"),
                "구분": item.get("구분"),
                "평가 항목": item.get("평가 항목"),
                "작성할 보완 정보": _validity_rework_field_guidance(str(item.get("_항목키") or "")),
            }
        )
    return pd.DataFrame(guide_rows)


def _validity_rework_instruction(candidate: dict, artifacts: dict, result: dict, rubric: dict) -> str:
    rows = _validity_rework_items(rubric, result)
    execution = artifacts.get("pipeline_result", {}).get("execution", {})
    execution_result = execution.get("result", {}) if isinstance(execution, dict) else {}
    holds = _validity_immediate_holds(result)
    recommendations = result.get("recommendations", []) if isinstance(result.get("recommendations"), list) else []

    lines = [
        f"원본 Run: {candidate.get('run_id', '-')}",
        f"Case ID: {candidate.get('case_id', '-')}",
        f"원 질문: {candidate.get('question') or execution.get('question') or '-'}",
        "",
        "목표: 아래 미흡 항목을 보완해 개선안 타당성 평가에서 AI 통과, 즉시 보류 규칙 0건, QA 검토 가능 상태가 되도록 최종 개선안을 실행계획 형태로 다시 작성하세요.",
        "",
        "현재 평가 요약",
        f"- 개선안 타당성 판정: {_voc_status_label(result.get('decision', 'NOT_RUN'))}",
        f"- 개선안 타당성 점수: {result.get('total_score', '-')} / 100점",
        f"- 승인 단계: {_voc_status_label(result.get('workflow_state', 'DRAFT'))}",
    ]
    if execution_result.get("summary"):
        lines.extend(["", "기존 파이프라인 요약", f"- {execution_result.get('summary')}"])
    if execution_result.get("policy"):
        lines.extend(["", "기존 최종 개선안", f"- {execution_result.get('policy')}"])

    lines.extend(["", "보완 우선순위"])
    if rows.empty:
        lines.append("- 통과 기준 미달 항목은 없지만, QA 검토자가 바로 확인할 수 있도록 근거·담당·일정·KPI를 더 명확히 정리하세요.")
    else:
        lines.append("- 필수 보완은 QA 검토 단계로 진행하기 위해 먼저 해소해야 하는 항목입니다.")
        lines.append("- 보완 권장은 통과 가능성을 높이는 참고 항목이며, 재시험 전달 시에는 필수 보완 중심으로 압축됩니다.")
        for index, item in enumerate(rows.to_dict("records"), start=1):
            score = "-" if item["현재 점수"] is None else f"{item['현재 점수']:g}"
            lines.append(
                f"{index}. [{item['구분']}] {item['평가 항목']}: 현재 {score}/{item['최고점']:g}점, "
                f"통과 기준 {item['통과 기준']:g}점, 부족 {item['부족 점수']:g}점"
            )
            lines.append(f"   - 보완 지시: {item['보완 지시']}")
            lines.append(f"   - 작성 가이드: {_validity_rework_field_guidance(str(item.get('_항목키') or ''))}")
            if item["현재 미흡 근거"] != "-":
                lines.append(f"   - 현재 미흡 근거: {item['현재 미흡 근거']}")

    if recommendations:
        lines.extend(["", "개선안 타당성 평가 권고"])
        for recommendation in recommendations:
            lines.append(f"- {recommendation}")
    if holds:
        lines.extend(["", "즉시 보류 규칙 해소"])
        for rule in holds:
            lines.append(f"- {_validity_hold_rule_label(rule)} 항목을 해소할 수 있는 근거와 조치안을 명시하세요.")

    lines.extend(
        [
            "",
            "반드시 포함할 출력 구조",
            "1. VOC 원인-근거-개선안 매트릭스: VOC ID 또는 실행 Trace 근거, 원인, 개선 조치를 연결",
            "2. 실행계획: 우선순위, 담당 조직, 일정, 마일스톤, 완료 기준",
            "3. 정량 KPI: 개선 전/후 목표값, 측정 방법, 확인 주기",
            "4. 리스크와 우회 방안: 장애·보안·법규·고객 안내 관점",
            "5. 확인 필요 항목: 근거가 부족한 내용은 추정하지 말고 확인 필요로 분리",
        ]
    )
    return "\n".join(lines)


def _validity_retest_instruction_payload(candidate: dict, artifacts: dict, result: dict, rubric: dict) -> str:
    rows = _validity_rework_items(rubric, result)
    holds = _validity_immediate_holds(result)
    required_rows = rows[rows["_필수보완"]] if not rows.empty and "_필수보완" in rows else pd.DataFrame()
    target_rows = required_rows if not required_rows.empty else rows
    execution = artifacts.get("pipeline_result", {}).get("execution", {})
    question = candidate.get("question") or execution.get("question") or "-"
    target_label = "통과 기준 미달 항목" if not required_rows.empty else "보완 권장 항목"

    lines = [
        f"Run/Case: {candidate.get('run_id', '-')} / {candidate.get('case_id', '-')}",
        f"원 질문: {_clip_text(question, 180)}",
        "목표: 개선안 타당성 평가에서 AI_PASS, 즉시 보류 0건, QA 검토 가능 상태가 되도록 최종 개선안을 보완하세요.",
        f"재시험 대상: {target_label} {len(target_rows)}건 중심. 이미 충족한 항목 설명은 반복하지 말고 부족 항목만 보강하세요.",
    ]

    if not target_rows.empty:
        lines.append("필수 보완 지시:")
        for item in target_rows.to_dict("records"):
            score = "-" if item["현재 점수"] is None else f"{item['현재 점수']:g}"
            lines.append(
                f"{item.get('우선순위')}. {item['평가 항목']} "
                f"(현재 {score}/{item['최고점']:g}, 기준 {item['통과 기준']:g}, 부족 {item['부족 점수']:g})"
            )
            lines.append(f"- 보완: {item['보완 지시']}")
            lines.append(f"- 작성: {_validity_rework_field_guidance(str(item.get('_항목키') or ''))}")
            reason = _clip_text(item.get("현재 미흡 근거"), 220)
            if reason != "-":
                lines.append(f"- 미흡 근거: {reason}")
    else:
        lines.append("필수 보완 지시: 점수 미달 항목은 없으므로 담당·일정·KPI와 근거 확인성을 더 명확히 하세요.")

    if holds:
        lines.append("즉시 보류 해소:")
        for rule in holds:
            lines.append(f"- {_validity_hold_rule_label(rule)} 해소 근거와 조치안을 반드시 포함")

    lines.extend(
        [
            "출력 형식:",
            "1) 적용 범위와 변경 대상",
            "2) 단계별 실행계획",
            "3) 담당·일정·마일스톤",
            "4) 정량 KPI와 완료 기준",
            "5) VOC·Trace 근거와 리스크 대응",
        ]
    )
    return "\n".join(lines)


def _render_validity_rework_guide(
    candidate: dict,
    artifacts: dict,
    result: dict | None,
    rubric: dict,
    *,
    key_scope: str = "",
    show_actions: bool = True,
    compact: bool = False,
):
    if not result:
        return
    gate = _validity_qa_gate_model(candidate, result)
    rows = _validity_rework_items(rubric, result)
    holds = _validity_immediate_holds(result)
    default_instruction = _validity_rework_instruction(candidate, artifacts, result, rubric)
    default_retest_instruction = _validity_retest_instruction_payload(candidate, artifacts, result, rubric)
    key_base = _validity_key_base(candidate["run_id"], candidate["case_id"], key_scope)
    default_key = f"validity_rework_default_{key_base}"
    instruction_key = f"validity_rework_instruction_{key_base}"
    retest_default_key = f"validity_retest_default_{key_base}"
    retest_instruction_key = f"validity_retest_instruction_{key_base}"
    if st.session_state.get(default_key) != default_instruction:
        st.session_state[default_key] = default_instruction
        st.session_state[instruction_key] = default_instruction
    if st.session_state.get(retest_default_key) != default_retest_instruction:
        st.session_state[retest_default_key] = default_retest_instruction
        st.session_state[retest_instruction_key] = default_retest_instruction

    required_count = int(rows["_필수보완"].sum()) if not rows.empty and "_필수보완" in rows else 0
    recommended_count = max(len(rows) - required_count, 0)

    if compact:
        if gate["ready"] and required_count == 0 and recommended_count == 0 and not holds:
            return

        with st.container(border=True):
            heading, state = st.columns([2.6, 1], vertical_alignment="center")
            with heading:
                st.markdown("#### 보완 재시험 실행")
                st.caption("상세 점수와 평가 근거는 선택 대상 상세 팝업에서 확인하고, 여기서는 재시험 실행 지시문만 관리합니다.")
            with state:
                badge = ":green-badge[QA 검토 가능]" if gate["ready"] else f":red-badge[필수 보완 {required_count}건]"
                st.markdown(badge, text_alignment="right")

            if rows.empty:
                st.caption("통과 기준 미달 항목은 없습니다. 필요한 경우에만 재시험 지시문을 보완해 실행하세요.")
            else:
                required_rows = rows[rows["_필수보완"]] if "_필수보완" in rows else rows
                focus_labels = [
                    str(label)
                    for label in required_rows.get("평가 항목", pd.Series(dtype="string")).tolist()
                    if str(label).strip()
                ]
                if not focus_labels:
                    focus_labels = [
                        str(label)
                        for label in rows.get("평가 항목", pd.Series(dtype="string")).head(3).tolist()
                        if str(label).strip()
                    ]
                st.caption(
                    "우선 보완 항목: "
                    + (", ".join(focus_labels[:3]) if focus_labels else "상세 팝업에서 확인")
                )

            if show_actions:
                with st.expander(
                    "재시험 지시문 편집·실행",
                    expanded=False,
                    icon=":material/replay:",
                ):
                    retest_instruction = st.text_area(
                        "재시험 전달 지시문",
                        key=retest_instruction_key,
                        height=132,
                        disabled=False,
                        help="재시험 Agent 파이프라인에 실제 전달되는 압축 지시문입니다. 상세 평가 결과는 중복 전달하지 않습니다.",
                    )
                    left, right = st.columns([2, 1], vertical_alignment="bottom")
                    active_state = _active_batch_run_state()
                    active_run_id = active_state["run_id"]
                    active = active_state["active"]
                    judge = artifacts.get("judge_result", {}) if isinstance(artifacts.get("judge_result"), dict) else {}
                    judge_config = {
                        "enabled": True,
                        "provider": judge.get("provider", "anthropic"),
                        "model": judge.get("model", "claude-haiku-4-5"),
                    }
                    with left:
                        st.caption("재시험은 원본 Run과 연결되어 수행 이력과 A/B 비교에서 전후 관계로 표시됩니다.")
                    with right:
                        if st.button(
                            "재시험 실행",
                            type="primary",
                            icon=":material/replay:",
                            disabled=active or not retest_instruction.strip(),
                            width="stretch",
                            key=f"validity_rework_retest_{key_base}",
                        ):
                            _launch_batch(
                                [candidate["case_id"]],
                                parent_run_id=candidate["run_id"],
                                judge_config=judge_config,
                                rework_instruction=retest_instruction,
                            )
                            st.toast("보완 지시 기반 재시험을 시작했습니다.", icon=":material/replay:")
                            st.rerun()
                    if active_run_id:
                        _live_batch_progress()
            else:
                st.caption("팝업에서는 보완 가이드와 재시험 지시문을 조회만 합니다.")
        return

    with st.container(border=True):
        heading, state = st.columns([2.6, 1], vertical_alignment="center")
        with heading:
            st.markdown("#### QA 검토 전 보완 가이드")
            st.caption("필수 보완 항목을 먼저 정렬하고, 재시험에는 압축 지시문만 전달합니다.")
        with state:
            badge = ":green-badge[QA 검토 가능]" if gate["ready"] else ":red-badge[보완 후 재시험]"
            st.markdown(badge, text_alignment="right")

        card_payloads = [
            (":material/rule:", "현재 판정", _voc_status_label(result.get("decision", "NOT_RUN"))),
            (":material/score:", "개선안 타당성 점수", f"{result.get('total_score', '-')} / 100점"),
            (":material/priority_high:", "필수 보완", f"{required_count}건"),
            (":material/edit_note:", "보완 권장", f"{recommended_count}건"),
            (":material/block:", "즉시 보류", f"{len(holds)}건"),
        ]
        for column, (icon, label, value) in zip(st.columns(5, gap="small"), card_payloads, strict=False):
            with column.container(border=True, height="stretch"):
                st.markdown(f"{icon} **{label}**")
                st.markdown(f"##### {value}")

        if rows.empty:
            st.success("통과 기준 미달 항목은 없습니다. QA 검토 전 근거·담당·일정·KPI만 한 번 더 확인하세요.")
        else:
            display_rows = rows.drop(
                columns=["_항목키", "_필수보완", "_점수비율", "_정렬순서"],
                errors="ignore",
            )
            st.dataframe(
                display_rows,
                hide_index=True,
                width="stretch",
                height=min(260, 74 + len(rows) * 42),
                column_config={
                    "우선순위": st.column_config.NumberColumn(format="%d", width="small"),
                    "구분": st.column_config.TextColumn(width="small"),
                    "평가 항목": st.column_config.TextColumn(width="medium", pinned=True),
                    "현재 점수": st.column_config.NumberColumn(format="%g점", width="small"),
                    "통과 기준": st.column_config.NumberColumn(format="%g점", width="small"),
                    "최고점": st.column_config.NumberColumn(format="%g점", width="small"),
                    "부족 점수": st.column_config.NumberColumn(format="%g점", width="small"),
                    "상태": st.column_config.TextColumn(width="small"),
                    "현재 미흡 근거": st.column_config.TextColumn(width="large"),
                    "보완 지시": st.column_config.TextColumn(width="large"),
                },
            )

            input_guide = _validity_rework_input_guide_rows(rows)
            st.markdown("##### 실제 보완 작성 가이드")
            st.dataframe(
                input_guide,
                hide_index=True,
                width="stretch",
                height=min(220, 72 + len(input_guide) * 44),
                column_config={
                    "우선순위": st.column_config.NumberColumn(format="%d", width="small"),
                    "구분": st.column_config.TextColumn(width="small"),
                    "평가 항목": st.column_config.TextColumn(width="medium", pinned=True),
                    "작성할 보완 정보": st.column_config.TextColumn(width="large"),
                },
            )

        with st.expander("상세 보완 가이드 보기", expanded=False, icon=":material/article:"):
            st.text_area(
                "사용자 확인용 상세 가이드",
                key=instruction_key,
                height=220,
                disabled=not show_actions,
                help="사람이 확인하기 위한 상세 가이드입니다. 기존 요약과 평가 근거를 함께 보여주며 재시험에는 그대로 전달하지 않습니다.",
            )

        retest_instruction = st.text_area(
            "재시험 전달 지시문",
            key=retest_instruction_key,
            height=168,
            disabled=not show_actions,
            help="재시험 Agent 파이프라인에 실제 전달되는 압축 지시문입니다. 통과 기준 미달 항목 중심으로 자동 구성됩니다.",
        )
        if show_actions:
            st.caption(
                "재시험에는 위 압축 지시문만 전달합니다. 상세 평가 결과·기존 개선안 전문은 화면 확인용으로 분리해 중복 전달을 줄입니다."
            )
        else:
            st.caption("팝업에서는 보완 가이드와 재시험 지시문을 조회만 합니다.")
        instruction = retest_instruction
        active_state = _active_batch_run_state()
        active_run_id = active_state["run_id"]
        active = active_state["active"]
        judge = artifacts.get("judge_result", {}) if isinstance(artifacts.get("judge_result"), dict) else {}
        judge_config = {
            "enabled": True,
            "provider": judge.get("provider", "anthropic"),
            "model": judge.get("model", "claude-haiku-4-5"),
        }
        if show_actions:
            left, right = st.columns([2, 1], vertical_alignment="bottom")
            with left:
                st.caption("재시험은 원본 Run과 연결되어 수행 이력과 A/B 비교에서 전후 관계로 표시됩니다.")
            with right:
                if st.button(
                    "보완 지시 기반 재시험 실행",
                    type="primary",
                    icon=":material/replay:",
                    disabled=active or not instruction.strip(),
                    width="stretch",
                    key=f"validity_rework_retest_{key_base}",
                ):
                    _launch_batch(
                        [candidate["case_id"]],
                        parent_run_id=candidate["run_id"],
                        judge_config=judge_config,
                        rework_instruction=instruction,
                    )
                    st.toast("보완 지시 기반 재시험을 시작했습니다.", icon=":material/replay:")
                    st.rerun()
        if show_actions and active_run_id:
            _live_batch_progress()


def _render_validity_auto_evaluation_controls(
    candidate: dict,
    artifacts: dict,
    validity: dict,
    validity_rubric: dict,
    *,
    key_scope: str = "",
    compact: bool = False,
) -> dict:
    run_id = candidate["run_id"]
    case_id = candidate["case_id"]
    key_base = _validity_key_base(run_id, case_id, key_scope)
    _render_validity_evaluation_focus_anchor_once()
    with st.container(border=True):
        heading, supplement_state = st.columns([2.4, 1], vertical_alignment="center")
        with heading:
            st.markdown("#### 평가 실행")
            if not compact:
                st.caption("선택 대상의 Agent 파이프라인 결과와 보완 입력을 기준으로 개선안 타당성을 평가합니다.")
        with supplement_state:
            supplement_rows = _validity_supplement_fields(artifacts.get("validity_supplement", {}))
            badge = (
                f":blue-badge[보완 입력 적용 {len(supplement_rows)}건]"
                if supplement_rows
                else ":gray-badge[보완 입력 미입력]"
            )
            st.markdown(badge, text_alignment="right")
        judge_gate = _validity_judge_gate_model(candidate, artifacts)
        config = _validity_config_controls(f"{key_base}_validity_auto", compact=compact)
        if judge_gate["blocked"]:
            st.caption(
                f"독립 LLM 평가 상태: {judge_gate['summary']} · "
                "PASS 후 개선안 타당성 평가를 실행할 수 있습니다."
            )
        if st.button(
            "타당성 평가 실행" if not judge_gate["blocked"] else "독립 LLM PASS 후 실행 가능",
            type="primary",
            icon=":material/fact_check:",
            disabled=not config["credential_configured"] or judge_gate["blocked"],
            width="stretch",
            key=f"validity_auto_evaluate_{key_base}",
        ):
            with st.status(
                "개선안 타당성 평가 수행 중",
                expanded=True,
                state="running",
            ) as status:
                st.write("1. 대상 증적 수집: Agent 파이프라인 결과, 실행 Trace, 독립 LLM 평가 결과, 결함 후보를 확인합니다.")
                st.write(
                    "2. 보완 입력 반영: "
                    + (
                        f"사용자 보완 입력 {len(supplement_rows)}건을 평가 입력에 포함합니다."
                        if supplement_rows
                        else "보완 입력 없이 원본 Agent 파이프라인 최종 개선안을 기준으로 평가합니다."
                    )
                )
                st.write(
                    f"3. 평가 기준 구성: Rubric {validity_rubric.get('version', '-')} · "
                    f"{len(validity_rubric.get('dimensions', {}))}개 항목 · 100점 기준을 적용합니다."
                )
                st.write(f"4. 타당성 평가 LLM 요청: {config['provider']} / {config['model']}")
                evaluated = evaluate_voc_improvement_validity(run_id, case_id, config)
                validity = evaluated.get("validity_result", {})
                artifacts["validity_result"] = validity
                candidate["validity_status"] = validity.get(
                    "decision",
                    candidate.get("validity_status", "NOT_RUN"),
                )
                candidate["validity_score"] = validity.get("total_score")
                candidate["workflow_state"] = validity.get(
                    "workflow_state",
                    candidate.get("workflow_state", "DRAFT"),
                )
                candidate["formal_approval"] = bool(validity.get("formal_approval"))
                candidate["immediate_hold_count"] = len(_validity_immediate_holds(validity))
                st.write(
                    f"5. 점수·판정 산출: {validity.get('total_score', '-')}점 · "
                    f"{_voc_status_label(validity.get('decision', '-'))}"
                )
                holds = _validity_immediate_holds(validity)
                st.write(
                    "6. 즉시 보류 규칙 확인: "
                    + ("보류 없음" if not holds else ", ".join(_validity_hold_rule_label(rule) for rule in holds))
                )
                gate = _validity_qa_gate_model(candidate, validity)
                st.write(f"7. QA Gate 판정: {gate['summary']}")
                status.update(
                    label="개선안 타당성 평가 완료" if validity.get("decision") != "ERROR" else "개선안 타당성 평가 오류",
                    state="complete" if validity.get("decision") != "ERROR" else "error",
                    expanded=True,
                )
            _load_validity_candidates.clear()
            _load_voc_history_rows.clear()
            if validity.get("decision") == "ERROR":
                st.session_state.voc_validity_error_notice = (
                    validity.get("error") or "개선안 타당성 평가 중 오류가 발생했습니다."
                )
            else:
                gate = _validity_qa_gate_model(candidate, validity)
                readiness = validity_human_review_readiness(
                    validity_status=validity.get("decision", candidate.get("validity_status", "NOT_RUN")),
                    workflow_state=validity.get("workflow_state", candidate.get("workflow_state", "DRAFT")),
                    immediate_hold_count=len(_validity_immediate_holds(validity)),
                    formal_approval=bool(validity.get("formal_approval") or candidate.get("formal_approval")),
                )
                action_code = str(readiness.get("action") or "")
                if action_code in {"QA_REVIEW", "BUSINESS_APPROVAL", "FORMAL_APPROVED"}:
                    _queue_validity_candidate_focus(
                        candidate,
                        action_code,
                        notice=(
                            f"개선안 타당성 평가가 {gate['summary']} 상태로 저장됐습니다. "
                            f"다음 액션은 {readiness['action_label']}입니다."
                        ),
                        scroll_to_approval=action_code in {"QA_REVIEW", "BUSINESS_APPROVAL"},
                    )
                    st.session_state.voc_validity_notice = "개선안 타당성 평가를 저장했습니다."
                else:
                    st.session_state.voc_validity_notice = (
                        f"개선안 타당성 평가를 저장했습니다. 현재 상태: {gate['summary']}"
                    )
            st.rerun()
    return validity


def _render_validity_trace_evidence(artifacts: dict):
    trace = artifacts.get("trace", {}) if isinstance(artifacts.get("trace"), dict) else {}
    trace_events = trace.get("events", []) if isinstance(trace.get("events"), list) else []
    with st.container(border=True):
        st.markdown("#### :material/account_tree: 실행 Trace")
        st.caption(f"실행 Trace ID: {trace.get('trace_id', '-')} · 이벤트 {len(trace_events)}건")
        if not trace_events:
            st.info("표시할 Agent 실행 Trace 이벤트가 없습니다.")
            return
        trace_rows = pd.DataFrame(trace_events)
        visible = [
            column for column in ("source", "target", "status", "duration_ms", "message")
            if column in trace_rows.columns
        ]
        display_rows = trace_rows[visible].rename(columns={
            "source": "출발 Agent",
            "target": "도착 Agent",
            "status": "상태",
            "duration_ms": "소요 시간(ms)",
            "message": "메시지",
        })
        if "상태" in display_rows.columns:
            display_rows["상태"] = display_rows["상태"].map(_voc_status_label)
        st.dataframe(
            display_rows,
            hide_index=True,
            width="stretch",
            height=min(310, 76 + len(display_rows) * 38),
            column_config={
                "출발 Agent": st.column_config.TextColumn(width="small"),
                "도착 Agent": st.column_config.TextColumn(width="small"),
                "상태": st.column_config.TextColumn(width="small"),
                "소요 시간(ms)": st.column_config.NumberColumn(width="small", format="%d"),
                "메시지": st.column_config.TextColumn(width="large"),
            },
        )


def _render_validity_judge_panel(candidate: dict, judge: dict):
    judge = judge if isinstance(judge, dict) else {}
    with st.container(horizontal=True):
        st.metric(
            "독립 LLM 평가",
            _voc_status_label(judge.get("decision", candidate.get("judge_status", "NOT_RUN"))),
            border=True,
        )
        st.metric("독립 LLM 평가 점수", f"{judge.get('total_score', '-')}점", border=True)
        st.metric("독립성", _judge_independence_grade_label(judge.get("independence_grade", "-")), border=True)
        st.metric("Provider", judge.get("provider", "-"), border=True)
    with st.container(border=True):
        st.markdown("#### :material/fact_check: 독립 LLM 평가 근거")
        st.caption(f"모델 {judge.get('model', '-')} · 실행 시각 {judge.get('evaluated_at', '-')}")
        summary = (
            judge.get("summary")
            or judge.get("reason")
            or judge.get("rationale")
            or judge.get("comment")
            or "-"
        )
        st.write(summary)
        score_rows = []
        raw_scores = judge.get("dimension_scores") or judge.get("scores") or {}
        if isinstance(raw_scores, dict):
            for label, detail in raw_scores.items():
                if isinstance(detail, dict):
                    score_rows.append({
                        "평가 항목": label,
                        "점수": detail.get("score", detail.get("value")),
                        "판정 근거": detail.get("reason", detail.get("comment", "-")),
                    })
                else:
                    score_rows.append({"평가 항목": label, "점수": detail, "판정 근거": "-"})
        if score_rows:
            st.dataframe(pd.DataFrame(score_rows), hide_index=True, width="stretch")


@st.dialog(
    "개선안 타당성 검증 대상 상세",
    width="large",
    icon=":material/fact_check:",
    on_dismiss=_dismiss_validity_candidate_dialog,
)
def _render_validity_candidate_dialog(candidate: dict):
    artifacts = load_voc_case_history_detail(candidate["run_id"], candidate["case_id"])
    execution = artifacts.get("pipeline_result", {}).get("execution", {})
    result = execution.get("result", {}) if isinstance(execution, dict) else {}
    judge = artifacts.get("judge_result", {}) if isinstance(artifacts.get("judge_result"), dict) else {}
    validity = artifacts.get("validity_result", {}) if isinstance(artifacts.get("validity_result"), dict) else {}
    validity_rubric = load_improvement_validity_rubric()

    _render_contextual_jira_action_menu(
        area_label="개선안 타당성 검증 상세",
        target_label="선택 Case",
        run_id=candidate.get("run_id", ""),
        case_id=candidate.get("case_id", ""),
        status_label=_voc_status_label(candidate.get("validity_status", "NOT_RUN")),
        question=candidate.get("question", ""),
        extra_detail=(
            f"다음 조치: {((candidate.get('next_action') or {}).get('label') or candidate.get('review_action_label') or _candidate_review_readiness(candidate)['action_label'])}\n"
            f"독립 LLM 평가: {_voc_status_label(candidate.get('judge_status', 'NOT_RUN'))}\n"
            f"독립 LLM 점수: {judge.get('total_score', candidate.get('llm_score', '-'))}\n"
            f"개선안 타당성 평가: {_voc_status_label(candidate.get('validity_status', 'NOT_RUN'))}\n"
            f"개선안 타당성 점수: {validity.get('total_score', candidate.get('validity_score', '-'))}\n"
            f"승인 단계: {_voc_status_label(candidate.get('workflow_state', 'DRAFT'))}"
        ),
        key=f"validity_dialog_{_validity_candidate_key(candidate)}",
    )

    st.markdown(
        f"""
        <div class="detail-status">
            검증 대상: {_voc_status_label(candidate.get('run_type', '-'))} ·
            {candidate['run_id']} · {candidate['case_id']} ·
            {_dashboard_timestamp(candidate.get('started_at', ''))}
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.container(horizontal=True):
        st.metric("Case", candidate["case_id"], border=True)
        st.metric(
            "독립 LLM 평가",
            _voc_status_label(judge.get("decision", candidate.get("judge_status", "NOT_RUN"))),
            border=True,
        )
        st.metric(
            "개선안 타당성 평가",
            _voc_status_label(validity.get("decision", candidate.get("validity_status", "NOT_RUN"))),
            border=True,
        )
        st.metric(
            "승인 단계",
            _voc_status_label(validity.get("workflow_state", candidate.get("workflow_state", "DRAFT"))),
            border=True,
        )

    tab_target, tab_a2a, tab_judge, tab_validity, tab_approval = st.tabs(
        ["대상 요약", "Agent 파이프라인 결과", "독립 LLM 평가", "개선안 타당성 평가", "QA 검토·승인"]
    )

    with tab_target:
        _render_validity_selection_basis(candidate, artifacts)
        _render_validity_qa_gate_cards(candidate, validity)

    with tab_a2a:
        st.markdown("### Agent 파이프라인 수행 결과")
        st.markdown("#### :material/help: 검증 질문")
        st.write(candidate.get("question") or execution.get("question") or "-")
        answer_columns = st.columns(2, gap="medium")
        with answer_columns[0].container(border=True, height="stretch"):
            st.markdown("#### :material/summarize: Agent 파이프라인 요약")
            st.write(result.get("summary", "-") or "-")
        with answer_columns[1].container(border=True, height="stretch"):
            st.markdown("#### :material/lightbulb: 최종 개선안")
            st.write(result.get("policy", "-") or "-")
        _render_validity_trace_evidence(artifacts)

    with tab_judge:
        st.markdown("### 독립 LLM 평가")
        _render_validity_judge_panel(candidate, judge)

    with tab_validity:
        st.markdown("### 개선안 타당성 평가 및 보완")
        _render_validity_workflow_status(
            candidate,
            artifacts,
            validity,
            validity_rubric,
            key_scope="dialog",
        )
        artifacts = _render_validity_supplement_editor(
            candidate,
            artifacts,
            result=validity,
            rubric=validity_rubric,
            key_scope="dialog",
            read_only=True,
        )
        _render_validity_result(validity, validity_rubric, artifacts)
        _render_validity_execution_steps(candidate, artifacts, validity, validity_rubric)
        _render_validity_rework_guide(
            candidate,
            artifacts,
            validity,
            validity_rubric,
            key_scope="dialog",
            show_actions=False,
        )

    with tab_approval:
        st.markdown("### QA 검토·업무 승인")
        _render_validity_qa_gate_cards(candidate, validity)
        if validity:
            _render_human_validity_review(
                candidate["run_id"],
                candidate["case_id"],
                validity,
                key_scope="dialog",
            )
        else:
            st.info("개선안 타당성 평가 결과가 있어야 QA 검토와 업무 승인을 진행할 수 있습니다.")


def _render_validity_result(
    result: dict,
    rubric: dict | None = None,
    artifacts: dict | None = None,
):
    if not result:
        st.info("아직 개선안 타당성 평가가 없습니다.")
        return
    with st.container(horizontal=True):
        st.metric("자동 판정", _voc_status_label(result.get("decision", "NOT_RUN")), border=True)
        st.metric("개선안 타당성 점수", result.get("total_score", "-"), border=True)
        st.metric("승인 상태", _voc_status_label(result.get("workflow_state", "DRAFT")), border=True)
        st.metric("정식 승인", "승인" if result.get("formal_approval") else "미승인", border=True)
    if result.get("supplemental_evidence_applied"):
        rows = _validity_supplement_fields(result.get("supplemental_evidence", {}))
        st.caption(
            f"이 평가 결과는 사용자 보완 입력 {len(rows)}건을 함께 반영했습니다."
            if rows
            else "이 평가 결과는 사용자 보완 입력을 함께 반영했습니다."
        )
    if result.get("approval_history_preserved") and result.get("approval_reset_reason"):
        st.caption(result["approval_reset_reason"])
    if result.get("error"):
        st.error(result["error"])
    holds = result.get("immediate_hold_rules_triggered", [])
    if holds:
        st.error("즉시 승인 보류: " + ", ".join(_validity_hold_rule_label(rule) for rule in holds))
    if rubric:
        _render_validity_ai_pass_failures(result, rubric, artifacts)
    if rubric and result.get("dimension_scores"):
        _render_validity_dimension_scorecard(rubric, result)
    if result.get("recommendations"):
        st.markdown("#### 보완 권고")
        for recommendation in result["recommendations"]:
            st.write(f"- {recommendation}")
    reviews = result.get("human_reviews", [])
    if reviews:
        st.markdown("#### 사람 검토 감사 이력")
        review_rows = _history_validity_review_rows(reviews)
        if not review_rows.empty:
            st.dataframe(
                review_rows,
                hide_index=True,
                width="stretch",
                column_config={
                    "단계": st.column_config.TextColumn(width=92),
                    "결정": st.column_config.TextColumn(width=86),
                    "검토자": st.column_config.TextColumn(width=120),
                    "상태 변화": st.column_config.TextColumn(width=190),
                    "검토 의견": st.column_config.TextColumn(width="large"),
                    "검토 시각": st.column_config.TextColumn(width=150),
                },
            )


DEPLOYMENT_DECISION_LABELS = {
    "NOT_EVALUATED": "미평가",
    "HUMAN_REVIEW_REQUIRED": "QA 검토 필요",
    "BUSINESS_REVIEW_REQUIRED": "업무 승인 필요",
    "FORMAL_QUALITY_APPROVED": "정식 품질 승인",
    "REMAINING_CASE_REVIEW_REQUIRED": "잔여 Case 검토 필요",
    "REVISION_REQUIRED": "보완 필요",
    "REJECTED": "반려",
    "미판정": "미판정",
}

VALIDITY_FOCUS_ACTION_LABELS = {
    "RUN_VALIDITY": {
        "label": "개선안 타당성 평가",
        "detail": "평가 설정에서 Provider와 모델을 확인한 뒤 평가를 실행합니다.",
        "icon": "fact_check",
        "tone": "blue",
    },
    "REWORK_AND_RETEST": {
        "label": "보완 입력·재시험",
        "detail": "부족한 담당·일정·KPI·근거를 보완하고 필요하면 연결 재시험을 실행합니다.",
        "icon": "edit_note",
        "tone": "red",
    },
    "QA_REVIEW": {
        "label": "QA 검토 저장",
        "detail": "승인/보완 요청/반려 중 하나를 선택해 QA 검토 결과를 저장합니다.",
        "icon": "rate_review",
        "tone": "green",
    },
    "BUSINESS_APPROVAL": {
        "label": "업무 승인 저장",
        "detail": "QA 검토 완료 건을 업무 관점에서 최종 승인하거나 보완/반려 처리합니다.",
        "icon": "verified",
        "tone": "green",
    },
    "CHECK_REMAINING_CASES": {
        "label": "잔여 Case 검토",
        "detail": "아직 승인되지 않은 Case를 선택해 평가·보완·승인을 이어갑니다.",
        "icon": "pending_actions",
        "tone": "orange",
    },
    "REPORT_READY": {
        "label": "보고서/최종 시연",
        "detail": "업무 승인 완료 증적을 품질 보고서와 최종 인수·시연 화면으로 연결합니다.",
        "icon": "summarize",
        "tone": "green",
    },
}

VALIDITY_REVIEW_DECISION_CARDS = {
    "QA": (
        {
            "decision": "APPROVE",
            "title": "QA 승인",
            "detail": "평가 근거가 충분하면 업무 승인 단계로 넘깁니다.",
            "button": "QA 승인 저장",
            "icon": "rate_review",
            "type": "primary",
        },
        {
            "decision": "REVISION_REQUIRED",
            "title": "보완 요청",
            "detail": "담당·일정·KPI·근거가 부족하면 보완 후 재평가로 돌립니다.",
            "button": "보완 요청",
            "icon": "edit_note",
            "type": "secondary",
        },
        {
            "decision": "REJECTED",
            "title": "반려",
            "detail": "현 개선안으로 진행할 수 없으면 반려 상태로 종료합니다.",
            "button": "반려",
            "icon": "block",
            "type": "secondary",
        },
    ),
    "BUSINESS": (
        {
            "decision": "APPROVE",
            "title": "업무 승인",
            "detail": "업무 적용 가능성이 충분하면 정식 승인으로 확정합니다.",
            "button": "업무 승인 저장",
            "icon": "verified",
            "type": "primary",
        },
        {
            "decision": "REVISION_REQUIRED",
            "title": "보완 요청",
            "detail": "배포 전 실행계획 보완이 필요하면 보완 상태로 되돌립니다.",
            "button": "보완 요청",
            "icon": "edit_note",
            "type": "secondary",
        },
        {
            "decision": "REJECTED",
            "title": "반려",
            "detail": "업무 적용이 부적합하면 반려 상태로 확정합니다.",
            "button": "반려",
            "icon": "block",
            "type": "secondary",
        },
    ),
}


def _validity_approval_workflow_model(result: dict | None) -> dict:
    result = result or {}
    holds = _validity_immediate_holds(result)
    readiness = validity_human_review_readiness(
        validity_status=result.get("decision", "NOT_RUN"),
        workflow_state=result.get("workflow_state", "DRAFT"),
        immediate_hold_count=len(holds),
        formal_approval=bool(result.get("formal_approval")),
    )
    state = readiness["workflow_state"]
    action = readiness["action"]
    has_result = bool(result)
    needs_rework = action == "REWORK_REQUIRED"
    qa_done = state in {"QA_REVIEWED", "BUSINESS_APPROVED"} or readiness["formal_approval"]
    business_done = state == "BUSINESS_APPROVED" or readiness["formal_approval"]
    evaluation_history = result.get("evaluation_history", []) or []
    human_reviews = result.get("human_reviews", []) or []
    if not isinstance(evaluation_history, list):
        evaluation_history = []
    if not isinstance(human_reviews, list):
        human_reviews = []
    audit_cards = [
        {
            "icon": "history",
            "label": "개선안 타당성 재평가",
            "value": f"{len(evaluation_history)}회",
            "detail": "이전 개선안 타당성 평가 결과는 이력으로 보존됩니다.",
        },
        {
            "icon": "assignment_ind",
            "label": "검토 감사 이력",
            "value": f"{len(human_reviews)}건",
            "detail": "QA·업무 결정은 append-only로 누적됩니다.",
        },
        {
            "icon": "schedule",
            "label": "최신 평가 시각",
            "value": _dashboard_timestamp(result.get("evaluated_at", "")) if has_result else "-",
            "detail": f"Rubric {result.get('rubric_version', '-')}" if has_result else "평가 후 표시됩니다.",
        },
        {
            "icon": "published_with_changes",
            "label": "재승인 기준",
            "value": "재승인 필요" if result.get("approval_history_preserved") and not business_done else "최신 상태 기준",
            "detail": result.get("approval_reset_reason") or "재평가 후에는 최신 판정 상태를 기준으로 승인합니다.",
        },
    ]
    stages = [
        {
            "label": "개선안 타당성 평가",
            "status": "완료" if has_result else "대기",
            "tone": "green" if has_result and not needs_rework else ("red" if needs_rework else "gray"),
            "detail": (
                f"{_voc_status_label(result.get('decision', 'NOT_RUN'))} · "
                f"{result.get('total_score', '-')}점"
                if has_result
                else "개선안 타당성 평가를 먼저 수행합니다."
            ),
        },
        {
            "label": "QA 검토",
            "status": "현재 단계" if readiness["can_qa_review"] else ("완료" if qa_done else "대기"),
            "tone": "blue" if readiness["can_qa_review"] else ("green" if qa_done else ("red" if needs_rework else "gray")),
            "detail": "AI 평가 통과와 보류 규칙 0건을 확인합니다." if readiness["can_qa_review"] else readiness["action_label"],
        },
        {
            "label": "업무 승인",
            "status": "현재 단계" if readiness["can_business_approve"] else ("완료" if business_done else "대기"),
            "tone": "blue" if readiness["can_business_approve"] else ("green" if business_done else "gray"),
            "detail": "QA 검토 완료 건을 최종 업무 승인합니다." if readiness["can_business_approve"] else readiness["action_label"],
        },
        {
            "label": "최종 배포 판정",
            "status": "완료" if business_done else "대기",
            "tone": "green" if business_done else ("red" if needs_rework else "gray"),
            "detail": DEPLOYMENT_DECISION_LABELS.get(
                readiness["deployment_decision"],
                readiness["deployment_decision"],
            ),
        },
    ]
    return {
        "readiness": readiness,
        "stages": stages,
        "holds": holds,
        "audit_cards": audit_cards,
    }


def _render_validity_approval_workflow(result: dict | None, *, compact: bool = False):
    model = _validity_approval_workflow_model(result)
    with st.container(border=True):
        heading, state = st.columns([2.3, 1], vertical_alignment="center")
        with heading:
            st.markdown("#### QA 검토·업무 승인 흐름")
            if not compact:
                st.caption("권한 구분은 없으며 현재 사용자가 QA 검토와 업무 승인을 순차로 수행합니다.")
        with state:
            readiness = model["readiness"]
            badge = {
                "QA_REVIEW": ":blue-badge[QA 검토 가능]",
                "BUSINESS_APPROVAL": ":blue-badge[업무 승인 가능]",
                "FORMAL_APPROVED": ":green-badge[정식 승인 완료]",
                "REWORK_REQUIRED": ":red-badge[보완 필요]",
                "VALIDITY_EVALUATION_REQUIRED": ":gray-badge[평가 필요]",
            }.get(readiness["action"], f":gray-badge[{readiness['action_label']}]")
            st.markdown(badge, text_alignment="right")
        columns = st.columns(4, gap="small")
        for column, stage in zip(columns, model["stages"], strict=False):
            tone = stage["tone"]
            badge = {
                "green": f":green-badge[{stage['status']}]",
                "blue": f":blue-badge[{stage['status']}]",
                "red": f":red-badge[{stage['status']}]",
                "gray": f":gray-badge[{stage['status']}]",
            }[tone]
            with column.container(border=True, height="stretch"):
                st.markdown(f"**{stage['label']}**")
                st.markdown(badge)
                if not compact:
                    st.caption(stage["detail"])
        if not compact:
            audit_columns = st.columns(4, gap="small")
            for column, audit in zip(audit_columns, model["audit_cards"], strict=False):
                with column.container(border=True, height=112):
                    st.markdown(f":material/{audit['icon']}: **{audit['label']}**")
                    st.markdown(f"##### {audit['value']}")
                    st.caption(audit["detail"])
        if model["holds"] and not compact:
            st.caption("즉시 보류 규칙: " + ", ".join(_validity_hold_rule_label(rule) for rule in model["holds"]))


def _render_human_validity_review(
    run_id: str,
    case_id: str,
    result: dict,
    *,
    key_scope: str = "",
    compact: bool = False,
):
    _render_validity_approval_focus_anchor_once()
    _render_validity_approval_workflow(result, compact=compact)
    state = result.get("workflow_state", "DRAFT")
    if result.get("decision") != "AI_PASS" or result.get("immediate_hold_rules_triggered"):
        st.warning("AI 평가 통과 상태이고 즉시 보류 규칙이 없어야 QA 검토를 시작할 수 있습니다.")
        return
    if state == "AI_REVIEWED":
        role, heading = "QA", "QA 검토"
    elif state == "QA_REVIEWED":
        role, heading = "BUSINESS", "업무 담당자 승인"
    elif state == "BUSINESS_APPROVED":
        st.success("QA와 업무 담당자 승인이 모두 완료되어 정식 운영 승인 상태입니다.")
        return
    else:
        st.info(f"현재 상태({state})에서는 추가 승인을 진행할 수 없습니다.")
        return

    st.markdown(f"### {heading}")
    if not compact:
        st.caption("현재 사용자가 이 단계의 검토자입니다. 저장 시 역할·시각·의견이 감사 이력에 추가됩니다.")
    with st.form(f"validity_review_{_validity_key_base(run_id, case_id, key_scope)}_{role}"):
        reviewer_col, comment_col = st.columns([0.8, 2.2], gap="small", vertical_alignment="top")
        with reviewer_col:
            reviewer = st.text_input("검토자", value="현재 사용자", max_chars=100)
        with comment_col:
            comment = st.text_area(
                "검토 의견",
                value="현재 화면에서 개선안 타당성 평가 결과, 즉시 보류 규칙, 보완 가이드와 증적을 확인했습니다.",
                max_chars=1000,
                height=96,
            )

        submitted_decision = None
        action_columns = st.columns(3, gap="small")
        for column, card in zip(action_columns, VALIDITY_REVIEW_DECISION_CARDS[role], strict=False):
            with column.container(border=True, height=128 if compact else 154):
                st.markdown(f":material/{card['icon']}: **{card['title']}**")
                st.caption(card["detail"])
                if st.form_submit_button(
                    card["button"],
                    key=(
                        f"validity_review_decision_"
                        f"{_validity_key_base(run_id, case_id, key_scope)}_{role}_{card['decision']}"
                    ),
                    type=card["type"],
                    icon=f":material/{card['icon']}:",
                    width="stretch",
                ):
                    submitted_decision = card["decision"]
    if submitted_decision:
        try:
            review_voc_improvement_validity(
                run_id,
                case_id,
                reviewer_role=role,
                reviewer_name_or_id=reviewer,
                decision=submitted_decision,
                comment=comment,
            )
        except Exception as exc:
            st.error(str(exc))
            return
        _load_validity_candidates.clear()
        _load_voc_history_rows.clear()
        st.session_state.pop("voc_validity_focus_action_code", None)
        st.session_state.pop("voc_validity_focus_target_key", None)
        candidate_ref = {"run_id": run_id, "case_id": case_id}
        if submitted_decision == "APPROVE" and role == "QA":
            _queue_validity_candidate_focus(
                candidate_ref,
                "BUSINESS_APPROVAL",
                notice=(
                    f"{case_id} QA 검토가 완료됐습니다. "
                    "다음 액션은 업무 승인 저장입니다."
                ),
                scroll_to_approval=True,
            )
            st.session_state.voc_validity_notice = "QA 검토 결과를 저장했습니다."
        elif submitted_decision == "APPROVE" and role == "BUSINESS":
            _queue_validity_candidate_focus(
                candidate_ref,
                "REPORT_READY",
                notice=(
                    f"{case_id} 업무 승인이 완료됐습니다. "
                    "품질 보고서와 최종 인수·시연 화면에서 증적을 확인할 수 있습니다."
                ),
            )
            st.session_state.voc_validity_notice = "업무 승인 결과를 저장했습니다."
        elif submitted_decision == "REVISION_REQUIRED":
            _queue_validity_candidate_focus(
                candidate_ref,
                "REWORK_AND_RETEST",
                notice=(
                    f"{case_id} 보완 요청이 저장됐습니다. "
                    "보완 입력 후 개선안 타당성 평가를 다시 실행하세요."
                ),
            )
            st.session_state.voc_validity_notice = f"{heading} 결과를 저장했습니다."
        else:
            st.session_state.voc_validity_notice = f"{heading} 결과를 저장했습니다. 최신 승인 단계가 갱신되었습니다."
        st.rerun()


def _render_improvement_ab(candidates: list[dict], baseline: dict, artifacts: dict):
    st.markdown("## 기존·개선 답변 A/B 비교")
    st.caption(
        "현재 선택한 결과를 A로 고정합니다. A에서 실행한 연결 재시험만 B 후보로 표시하므로 "
        "두 Run을 직접 맞춰 고를 필요가 없습니다."
    )
    with st.container(border=True):
        st.markdown("#### 사용 순서")
        st.write("1. 위에서 기준 Run·Case를 선택합니다.")
        st.write("2. 아래 버튼으로 같은 TC의 연결 재시험을 실행합니다.")
        st.write("3. 재시험 완료 후 B 결과에 독립 LLM 평가와 개선안 타당성 평가를 수행합니다.")
        st.write("4. B 후보를 선택하면 원본 A를 자동으로 찾아 점수와 답변을 비교합니다.")
        st.caption("재시험은 실제 Agent·LLM을 호출하므로 수행 시간과 API 비용이 발생합니다.")

    active_state = _active_batch_run_state()
    active_run_id = active_state["run_id"]
    active = active_state["active"]
    judge = artifacts.get("judge_result", {})
    judge_config = {
        "enabled": True,
        "provider": judge.get("provider", "anthropic"),
        "model": judge.get("model", "claude-haiku-4-5"),
    }
    with st.container(horizontal=True):
        if st.button(
            "현재 Case 연결 재시험 실행",
            icon=":material/replay:",
            disabled=active,
            key=f"validity_linked_retest_{baseline['run_id']}_{baseline['case_id']}",
        ):
            _launch_batch(
                [baseline["case_id"]],
                parent_run_id=baseline["run_id"],
                judge_config=judge_config,
            )
            st.rerun()
        if st.button("B 후보 새로고침", icon=":material/refresh:", key="validity_ab_refresh"):
            _load_validity_candidates.clear()
            _load_voc_history_rows.clear()
            st.rerun()
    if active_run_id:
        _live_batch_progress()

    retests = [
        item for item in candidates
        if item.get("run_type") == "RETEST"
        and item.get("parent_run_id") == baseline["run_id"]
        and item.get("case_id") == baseline["case_id"]
    ]
    if not retests:
        st.info("현재 선택한 A와 연결된 완료 재시험이 없습니다. 위 버튼으로 B 후보를 먼저 생성하세요.")
        return
    labels = [
        f"{item['run_id']} · 독립 LLM 평가 {_voc_status_label(item['judge_status'])} · 개선안 타당성 {_voc_status_label(item['validity_status'])}"
        for item in retests
    ]
    candidate_label = st.selectbox("B · 연결 재시험", labels, key="validity_ab_candidate")
    candidate = retests[labels.index(candidate_label)]
    comparison = compare_voc_improvement_answers(
        baseline["run_id"], candidate["run_id"], baseline["case_id"]
    )
    if not comparison["compatible"]:
        st.error("동일 조건 A/B 비교 불가: " + ", ".join(comparison["compatibility_differences"]))
        return
    st.success("동일 질문·TC·Catalog·Rubric과 재시험 부모 연결이 확인됐습니다.")
    before, after = st.columns(2)
    with before:
        st.markdown("### A · 기존 답변")
        st.write(comparison["baseline"]["policy"] or "-")
        st.caption(
            f"독립 LLM 평가 {comparison['baseline']['judge_score']} · "
            f"개선안 타당성 {comparison['baseline']['validity_score']} · "
            f"{_voc_status_label(comparison['baseline']['workflow_state'])}"
        )
    with after:
        st.markdown("### B · 개선 답변")
        st.write(comparison["candidate"]["policy"] or "-")
        st.caption(
            f"독립 LLM 평가 {comparison['candidate']['judge_score']} · "
            f"개선안 타당성 {comparison['candidate']['validity_score']} · "
            f"{_voc_status_label(comparison['candidate']['workflow_state'])}"
        )
    with st.container(horizontal=True):
        judge_delta = comparison["score_deltas"]["judge_score"]
        validity_delta = comparison["score_deltas"]["validity_score"]
        st.metric("독립 LLM 평가 점수 변화", "-" if judge_delta is None else judge_delta)
        st.metric("개선안 타당성 점수 변화", "-" if validity_delta is None else validity_delta)


def render_improvement_validity():
    _apply_pending_validity_candidate_filters()
    notice = st.session_state.pop("voc_validity_notice", None)
    if notice:
        st.success(notice)
    error_notice = st.session_state.pop("voc_validity_error_notice", None)
    if error_notice:
        st.error(error_notice)
    focus_notice = st.session_state.pop("voc_validity_focus_notice", None)
    if focus_notice:
        st.info(focus_notice, icon=":material/conversion_path:")
    st.markdown("## 검증 대상 선택")
    st.caption("성공한 VOC 파이프라인 결과를 자동 채점한 뒤 QA와 업무 담당자가 순서대로 검토합니다.")
    candidates = _load_validity_candidates()
    if not candidates:
        st.info("타당성을 검증할 수 있는 완료 VOC Case가 없습니다.")
        return

    _render_voc_summary_cards(
        _validity_focus_cards(candidates),
        columns=5,
        height=118,
    )

    st.session_state.setdefault("voc_validity_run_type", "전체")
    st.session_state.setdefault("voc_validity_candidate_status", "전체")

    with st.container(border=True):
        filter_columns = st.columns([0.75, 1.15, 1.75], gap="small", vertical_alignment="bottom")
        with filter_columns[0]:
            query = st.text_input(
                "대상 검색",
                placeholder="Run · Case · 질문",
                key="voc_validity_candidate_query",
                icon=":material/search:",
            )
        with filter_columns[1]:
            run_type_filter = st.segmented_control(
                "회차 유형",
                VALIDITY_RUN_TYPE_FILTERS,
                required=True,
                key="voc_validity_run_type",
                width="stretch",
                persist_state="session",
            )
        with filter_columns[2]:
            status_filter = st.segmented_control(
                "평가 상태",
                ("전체", "평가 전", "평가 완료", "QA 검토 가능", "업무 승인 가능", "정식 승인"),
                required=True,
                key="voc_validity_candidate_status",
                width="stretch",
                persist_state="session",
            )

    filtered = _filter_validity_candidates(
        candidates,
        query=query,
        status_filter=status_filter or "전체",
        run_type_filter=run_type_filter or "전체",
    )
    if not filtered:
        st.info("현재 검색·상태 조건에 맞는 검증 대상이 없습니다.")
        return

    selected_key = st.session_state.get("voc_validity_selected_key", "")
    if not any(_validity_candidate_key(item) == selected_key for item in filtered):
        selected_key = _validity_candidate_key(filtered[0])
        st.session_state.voc_validity_selected_key = selected_key

    with st.container(border=True):
        _render_voc_section_heading(
            "검증 대상 목록",
            "",
            icon="fact_check",
            badges=_validity_status_count_badges(filtered),
        )
        candidate_frame = _validity_candidate_rows(filtered, selected_key)
        candidate_keys = tuple(_validity_candidate_key(item) for item in filtered)
        table_key = _validity_candidate_table_widget_key()
        st.dataframe(
            candidate_frame,
            hide_index=True,
            width="stretch",
            height=min(465, 76 + len(candidate_frame) * 38),
            on_select=partial(_remember_validity_candidate_selection, table_key, candidate_keys),
            selection_mode=["single-row", "single-cell"],
            key=table_key,
            column_order=[
                "Case ID",
                "질문",
                "다음 조치",
                "개선안 타당성",
                "타당성 점수",
                "승인 단계",
                "독립 LLM 평가",
                "독립 LLM 점수",
                "수행 유형",
                "수행 일시",
                "정식 승인",
                "Run ID",
            ],
            column_config={
                "Case ID": st.column_config.TextColumn(width=78, pinned=True),
                "질문": st.column_config.TextColumn(width=320),
                "다음 조치": st.column_config.TextColumn(width=138),
                "개선안 타당성": st.column_config.TextColumn(width=106),
                "타당성 점수": st.column_config.ProgressColumn(
                    width=84, min_value=0, max_value=100, format="%g점"
                ),
                "승인 단계": st.column_config.TextColumn(width=112),
                "독립 LLM 평가": st.column_config.TextColumn(width=96),
                "독립 LLM 점수": st.column_config.ProgressColumn(
                    width=84, min_value=0, max_value=100, format="%g점"
                ),
                "수행 유형": st.column_config.TextColumn(width=86),
                "수행 일시": st.column_config.TextColumn(width=118),
                "정식 승인": st.column_config.TextColumn(width=78),
                "Run ID": st.column_config.TextColumn(width=196),
            },
        )
    _render_validity_review_queue(filtered)

    selected = next(
        (
            item for item in candidates
            if _validity_candidate_key(item) == st.session_state.voc_validity_selected_key
        ),
        filtered[0],
    )
    st.session_state.voc_validity_selected_key = _validity_candidate_key(selected)
    artifacts = load_voc_case_history_detail(selected["run_id"], selected["case_id"])
    selected = _sync_validity_candidate_from_artifacts(selected, artifacts)
    validity = artifacts.get("validity_result", {})
    validity_rubric = load_improvement_validity_rubric()
    _render_validity_history_focus_action(selected)
    _render_validity_selection_basis(selected, artifacts, compact=True)
    with st.container(horizontal=True, horizontal_alignment="right"):
        _render_contextual_jira_action_menu(
            area_label="개선안 타당성 검증",
            target_label="선택 Case",
            run_id=selected.get("run_id", ""),
            case_id=selected.get("case_id", ""),
            status_label=_voc_status_label(selected.get("validity_status", "NOT_RUN")),
            question=selected.get("question", ""),
            extra_detail=(
                f"다음 조치: {((selected.get('next_action') or {}).get('label') or selected.get('review_action_label') or _candidate_review_readiness(selected)['action_label'])}\n"
                f"독립 LLM 평가: {_voc_status_label(selected.get('judge_status', 'NOT_RUN'))}\n"
                f"독립 LLM 점수: {selected.get('llm_score', '-')}\n"
                f"개선안 타당성 평가: {_voc_status_label(selected.get('validity_status', 'NOT_RUN'))}\n"
                f"개선안 타당성 점수: {selected.get('validity_score', '-')}\n"
                f"승인 단계: {_voc_status_label(selected.get('workflow_state', 'DRAFT'))}"
            ),
            key=f"validity_{_validity_candidate_key(selected)}",
        )
        if st.button(
            "선택 대상 상세 보기",
            icon=":material/open_in_new:",
            key=f"validity_selected_detail_{_validity_candidate_key(selected)}",
        ):
            _open_validity_candidate_dialog(selected)
            st.rerun()
    _render_validity_workflow_status(selected, artifacts, validity, validity_rubric, compact=True)
    workflow_model = _validity_workflow_status_model(selected, artifacts, validity, validity_rubric)
    action_tone = {
        "green": "green",
        "blue": "blue",
        "orange": "orange",
        "red": "red",
        "gray": "gray",
    }.get(workflow_model.get("tone", "gray"), "gray")
    with st.container(border=True):
        _render_voc_section_heading(
            "선택 대상 실행/입력",
            "필요한 입력과 실행 버튼만 표시합니다. 평가 결과·증적·상세 로그는 상세 팝업에서 확인하세요.",
            icon="edit_note",
            badges=((workflow_model["next_title"], action_tone),),
        )

    judge_gate = workflow_model["judge_gate"]
    if judge_gate["blocked"]:
        _render_validity_judge_prerequisite_notice(selected, artifacts)
    else:
        active_step = workflow_model.get("active_step")
        if active_step in {"supplement", "evaluate"} or workflow_model.get("reeval_needed"):
            artifacts = _render_validity_supplement_editor(
                selected,
                artifacts,
                result=validity,
                rubric=validity_rubric,
                compact=True,
            )

            validity = _render_validity_auto_evaluation_controls(
                selected,
                artifacts,
                validity,
                validity_rubric,
                compact=True,
            )

    latest_workflow = _validity_workflow_status_model(selected, artifacts, validity, validity_rubric)
    if latest_workflow.get("action_code") == "REWORK_REQUIRED":
        _render_validity_rework_guide(selected, artifacts, validity, validity_rubric, compact=True)
    if latest_workflow.get("action_code") in {"QA_REVIEW", "BUSINESS_APPROVAL", "FORMAL_APPROVED"} and validity:
        _render_human_validity_review(selected["run_id"], selected["case_id"], validity, compact=True)
    with st.expander("추가 분석 · 동일 조건 A/B 비교", expanded=False, icon=":material/compare_arrows:"):
        st.caption("재시험 전후 비교가 필요할 때만 펼쳐서 확인합니다.")
        if st.toggle("동일 조건 A/B 비교 보기", key="validity_show_ab"):
            _render_improvement_ab(candidates, selected, artifacts)

    dialog_candidate_key = st.session_state.get(VALIDITY_DETAIL_DIALOG_CANDIDATE_KEY)
    if dialog_candidate_key:
        dialog_candidate = next(
            (item for item in candidates if _validity_candidate_key(item) == dialog_candidate_key),
            None,
        )
        if dialog_candidate:
            _render_validity_candidate_dialog(dialog_candidate)
        else:
            _dismiss_validity_candidate_dialog()

def _build_testcase_group_chart(group_rows: pd.DataFrame) -> alt.Chart:
    return (
        alt.Chart(group_rows)
        .mark_bar(color="#2F6FB0", cornerRadiusEnd=5, size=18)
        .encode(
            y=alt.Y(
                "검증 영역:N",
                title=None,
                sort="-x",
                axis=alt.Axis(
                    domain=False,
                    ticks=False,
                    labelColor="#425b76",
                    labelLimit=150,
                    labelPadding=8,
                ),
            ),
            x=alt.X(
                "Case 수:Q",
                title=None,
                axis=alt.Axis(
                    domain=False,
                    gridColor="#e6eef6",
                    format="d",
                    tickCount=5,
                ),
            ),
            tooltip=[
                alt.Tooltip("검증 영역:N", title="검증 영역"),
                alt.Tooltip("Case 수:Q", title="Case 수", format="d"),
            ],
        )
        .properties(height=130)
        .configure_view(strokeWidth=0)
    )


def _catalog_upload_missing_fields(detail: str) -> list[str]:
    marker = "필수 필드 누락"
    if marker not in detail:
        return []
    raw_fields = detail.split(marker, 1)[1].strip()
    try:
        parsed = ast.literal_eval(raw_fields)
    except (SyntaxError, ValueError):
        parsed = []
    if isinstance(parsed, (list, tuple, set)):
        return [str(item) for item in parsed if str(item).strip()]
    return []


def _catalog_upload_error_field(detail: str) -> str:
    detail = str(detail or "").strip()
    missing_fields = _catalog_upload_missing_fields(detail)
    if missing_fields:
        return ", ".join(missing_fields)
    if detail.startswith("execution."):
        return detail.split(" ", 1)[0].strip()
    if "등록되지 않은 검증 영역" in detail:
        return "group"
    if "객체 형식" in detail:
        return "Case 형식"
    if "API 키" in detail or "인증정보" in detail:
        return "보안 정보"
    if "JSON 파일을 해석" in detail:
        return "JSON 형식"
    if detail.endswith("가 필요합니다."):
        return detail.removesuffix("가 필요합니다.").strip()
    for field in (
        "version",
        "suite_id",
        "groups",
        "cases",
        "case_id",
        "implementation_status",
        "execution_type",
        "total_cases",
    ):
        if detail.startswith(field) or field in detail:
            return field
    return "검증 오류"


def _catalog_upload_error_cards(errors: list[str]) -> list[dict]:
    cards = []
    case_prefixes = ("TC-", "FT-", "AG-", "QG-", "cases[")
    for error in errors:
        text = str(error or "").strip()
        case_ref = "카탈로그"
        detail = text
        if ": " in text:
            prefix, message = text.split(": ", 1)
            prefix = prefix.strip()
            if prefix.startswith(case_prefixes):
                case_ref = prefix
                detail = message.strip()
        cards.append({
            "case_ref": case_ref,
            "field": _catalog_upload_error_field(detail),
            "message": detail,
        })
    return cards


def _render_testcase_catalog_upload_errors(errors: list[str]) -> None:
    cards = _catalog_upload_error_cards(errors)
    st.error(f"카탈로그 검증 실패 · {len(cards)}건", icon=":material/error:")
    st.caption("아래 Case와 필드를 수정한 뒤 다시 업로드하세요.")
    card_area_height = min(360, max(150, len(cards) * 86))
    with st.container(height=card_area_height, border=False):
        for card in cards:
            with st.container(border=True):
                case_col, field_col = st.columns([0.85, 1.15], vertical_alignment="center")
                with case_col:
                    st.markdown(f"**{card['case_ref']}**")
                with field_col:
                    st.caption(f"필드 · {card['field']}")
                st.write(card["message"])


def render_testcases():
    payload = load_unified_quality_cases()
    catalog = load_quality_test_catalog()
    cases = catalog.get("cases", [])
    groups = catalog.get("groups", {})
    testcase_details = {
        item.get("case_id"): item for item in payload.get("cases", [])
    }

    if not cases:
        st.warning("통합 테스트 카탈로그에 등록된 Case가 없습니다.")
        return

    implemented_count = sum(
        item.get("implementation_status") == "IMPLEMENTED" for item in cases
    )
    defined_count = len(cases) - implemented_count
    voc_count = sum(item.get("group") == "voc_functional" for item in cases)
    additional_count = len(cases) - voc_count

    group_rows = pd.DataFrame(
        [
            {
                "검증 영역": group.get("label", group_key),
                "Case 수": sum(item.get("group") == group_key for item in cases),
            }
            for group_key, group in groups.items()
        ]
    )
    st.html(
        """
        <style>
        .st-key-voc_testcase_metrics [data-testid="stMetric"]{
            height:130px!important;min-height:130px!important;padding:11px 10px!important;
        }
        .st-key-voc_testcase_metrics [data-testid="stMetricLabel"] p{
            font-size:.72rem!important;line-height:1.2!important;
        }
        .st-key-voc_testcase_metrics [data-testid="stMetricValue"]{
            font-size:1.25rem!important;line-height:1.25!important;
        }
        .st-key-voc_testcase_metrics [data-testid="stMetricDelta"]{
            font-size:.66rem!important;
        }
        .st-key-voc_testcase_metrics>div[data-testid="stVerticalBlock"]{
            gap:.45rem!important;
        }
        .st-key-voc_testcase_metrics_header h4{
            margin:0!important;padding:0!important;line-height:1.25!important;
        }
        .st-key-voc_testcase_metrics_header button{
            min-height:32px!important;margin-top:0!important;
        }
        .st-key-voc_testcase_search>div[data-testid="stVerticalBlock"],
        .st-key-voc_testcase_browser>div[data-testid="stVerticalBlock"],
        .st-key-voc_testcase_detail>div[data-testid="stVerticalBlock"]{
            gap:.55rem!important;
        }
        .st-key-voc_testcase_browser [data-testid="stDataFrame"]{
            margin-top:.1rem;
        }
        .st-key-voc_testcase_detail [data-testid="stTabs"] button{
            min-height:34px!important;
        }
        </style>
        """
    )
    overview_columns = st.columns(
        [1.6, 1],
        gap="small",
        vertical_alignment="top",
    )
    with overview_columns[0].container(
        border=True,
            height=210,
        key="voc_testcase_metrics",
    ):
        with st.container(
            horizontal=True,
            horizontal_alignment="distribute",
            vertical_alignment="center",
            gap="small",
            key="voc_testcase_metrics_header",
        ):
            st.markdown("#### :material/target: 실행 대상 요약")
            with st.container(horizontal=True, horizontal_alignment="right", vertical_alignment="center"):
                st.download_button(
                    "TC Download",
                    data=json.dumps(catalog, ensure_ascii=False, indent=2),
                    file_name="quality_test_catalog.json",
                    mime="application/json",
                    icon=":material/download:",
                    type="primary",
                    width="content",
                    key="voc_testcase_catalog_download",
                )
                with st.popover(
                    "TC Upload",
                    icon=":material/upload_file:",
                    width="content",
                ):
                    uploaded_catalog = st.file_uploader(
                        "통합 테스트케이스 JSON",
                        type=["json"],
                        key="voc_testcase_catalog_upload",
                    )
                    if uploaded_catalog is not None:
                        try:
                            uploaded_payload = json.loads(
                                uploaded_catalog.getvalue().decode("utf-8-sig")
                            )
                            upload_errors = validate_quality_test_catalog(uploaded_payload)
                        except Exception as exc:
                            uploaded_payload = None
                            upload_errors = [f"JSON 파일을 해석할 수 없습니다: {exc}"]
                        if upload_errors:
                            _render_testcase_catalog_upload_errors(upload_errors)
                        elif st.button(
                            "검증 완료 · 적용",
                            type="primary",
                            key="voc_testcase_catalog_apply_upload",
                        ):
                            result = save_quality_test_catalog(uploaded_payload, source="json_upload")
                            st.success(
                                f"통합 테스트케이스 {result.get('total_cases', 0)}건을 저장했습니다."
                            )
                            st.rerun()
        metric_row = st.columns(4, gap="small")
        metric_row[0].metric("전체 실행 대상", f"{len(cases)}건", border=True)
        metric_row[1].metric("VOC 질문형", f"{voc_count}건", border=True)
        metric_row[2].metric("추가 검증 Case", f"{additional_count}건", border=True)
        metric_row[3].metric(
            "구현 상태",
            f"{implemented_count}건 완료",
            delta=f"{defined_count}건 후속 구현",
            delta_color="off",
            border=True,
    )
    with overview_columns[1].container(
        border=True,
            height=210,
        key="voc_testcase_group_chart",
    ):
        st.markdown("#### :material/bar_chart: 검증 영역별 Case 구성")
        st.altair_chart(_build_testcase_group_chart(group_rows))

    browser_columns = st.columns([0.5, 1.1, 1], gap="small", vertical_alignment="top")
    with browser_columns[0].container(
        border=True,
        height=380,
        key="voc_testcase_search",
    ):
        st.markdown("#### :material/search: Case 탐색")
        search_text = st.text_input(
            "검색",
            placeholder="ID·이름·판정 기준",
            key="voc_testcase_catalog_search",
        ).strip().lower()
        group_filter = st.selectbox(
            "검증 영역",
            options=["전체", *groups.keys()],
            format_func=lambda key: (
                "전체" if key == "전체" else groups[key].get("label", key)
            ),
            key="voc_testcase_catalog_group",
        )
        status_filter = st.selectbox(
            "구현 상태",
            options=["전체", "IMPLEMENTED", "DEFINED"],
            format_func=lambda value: {
                "전체": "전체",
                "IMPLEMENTED": "실행 구현 완료",
                "DEFINED": "정의됨 · 후속 구현",
            }[value],
            key="voc_testcase_catalog_status",
        )

    status_labels = {
        "IMPLEMENTED": "실행 구현 완료",
        "DEFINED": "정의됨 · 후속 구현",
    }
    rows = []
    visible_cases = []
    for case in cases:
        group_key = case.get("group", "")
        status = case.get("implementation_status", "")
        searchable = " ".join(
            str(case.get(key, ""))
            for key in ("case_id", "name", "acceptance", "source_ref")
        ).lower()
        if group_filter != "전체" and group_key != group_filter:
            continue
        if status_filter != "전체" and status != status_filter:
            continue
        if search_text and search_text not in searchable:
            continue
        visible_cases.append(case)
        rows.append(
            {
                "Case ID": case.get("case_id"),
                "검증 영역": groups.get(group_key, {}).get("label", group_key),
                "이름": case.get("name"),
                "구현 상태": status_labels.get(status, status or "-"),
            }
        )

    with browser_columns[1].container(
        border=True,
        height=380,
        key="voc_testcase_browser",
    ):
        st.markdown("#### :material/list_alt: Case 목록")
        st.caption(f"검색 결과 {len(rows)}건 · 행을 선택하면 우측에서 상세 확인")
        visible_case_ids = [case.get("case_id", "") for case in visible_cases]
        table_key = "voc_testcase_catalog_table"
        remembered_case_id = st.session_state.get("voc_testcase_selected_case_id")
        default_index = (
            visible_case_ids.index(remembered_case_id)
            if remembered_case_id in visible_case_ids
            else 0
        )
        event = st.dataframe(
            pd.DataFrame(rows),
            hide_index=True,
            width="stretch",
            height=270,
            on_select=(
                partial(
                    _remember_catalog_case_selection,
                    table_key,
                    visible_case_ids,
                )
                if visible_case_ids
                else "rerun"
            ),
            selection_mode=(
                ["single-row-required", "single-cell"]
                if visible_case_ids
                else "single-row"
            ),
            selection_default=(
                {"selection": {"rows": [default_index]}}
                if visible_case_ids
                else None
            ),
            key=table_key,
            column_config={
                "Case ID": st.column_config.TextColumn(width=82, pinned=True),
                "검증 영역": st.column_config.TextColumn(width=110),
                "이름": st.column_config.TextColumn(width=150),
                "구현 상태": st.column_config.TextColumn(width=130),
            },
        )

    selected_row = _table_selected_row_index(
        event,
        len(visible_cases),
    )
    if selected_row is None and visible_cases:
        selected_row = default_index
    with browser_columns[2].container(
        border=True,
        height=380,
        key="voc_testcase_detail",
    ):
        st.markdown("#### :material/description: Case 상세")
        if selected_row is None:
            st.caption("목록에서 Case ID를 선택하세요.")
            st.markdown(
                """
                :blue-badge[선택한 Case의 검증 기준]

                선택 후 이 영역에서 등록 원본, 판정 기준과 VOC 입출력 조건을 간단히 확인할 수 있습니다.
                """
            )
        else:
            selected = visible_cases[selected_row]
            detail = testcase_details.get(selected.get("case_id"), {})
            group_label = groups.get(selected.get("group"), {}).get(
                "label", selected.get("group", "-")
            )
            status_label = status_labels.get(
                selected.get("implementation_status"), "-"
            )
            status_badge = (
                "green"
                if selected.get("implementation_status") == "IMPLEMENTED"
                else "orange"
            )
            st.markdown(
                f"**{selected.get('case_id')}** · {selected.get('name', '-')}"
            )
            st.markdown(
                f":blue-badge[{group_label}] "
                f":{status_badge}-badge[{status_label}]"
            )
            criteria_tab, conditions_tab = st.tabs(["판정 기준", "VOC 조건"])
            with criteria_tab:
                st.caption(f"등록 원본 · {selected.get('source_ref', '-')}")
                st.write(selected.get("acceptance", "-") or "-")
            with conditions_tab:
                if not detail:
                    st.caption("이 Case에는 별도의 VOC 질문 조건이 없습니다.")
                else:
                    st.markdown("**질문**")
                    st.write(detail.get("question", "-") or "-")
                    st.caption(f"예상 의도 · {detail.get('expected_intent', '-')}")
                    with st.expander(
                        "필수·금지 출력 조건",
                        icon=":material/rule:",
                    ):
                        output_columns = st.columns(2, gap="small")
                        with output_columns[0]:
                            st.markdown("**필수 출력**")
                            st.write(
                                "\n".join(
                                    f"- {item}"
                                    for item in detail.get("required_output", [])
                                )
                                or "-"
                            )
                        with output_columns[1]:
                            st.markdown("**금지 출력**")
                            st.write(
                                "\n".join(
                                    f"- {item}"
                                    for item in detail.get("prohibited_output", [])
                                )
                                or "-"
                            )


def render_analysis():
    with st.container(border=True):
        st.info("통합 런타임 Agent(6101~6106)가 실행 중이어야 합니다.", icon=":material/hub:")
        question = st.text_area(
            "VOC 질문",
            placeholder="예: 모바일 앱에서 보험 갱신이 되지 않는 불만을 요약하고 개선안을 제안해 주세요.",
            height=120,
        )
        save_report = st.checkbox(
            "질문과 결과를 Reports/VOC에 저장합니다.",
            value=False,
            help="질문에 개인정보나 민감정보가 있으면 선택하지 마세요.",
        )
        if st.button("VOC 분석 실행", type="primary", icon=":material/analytics:", width="stretch"):
            with st.spinner("6개 Agent가 VOC를 처리하고 있습니다. 최대 3분 정도 걸릴 수 있습니다..."):
                st.session_state.voc_analysis_result = run_voc_analysis(question, save_report)

    payload = st.session_state.get("voc_analysis_result")
    if not payload:
        return
    result = payload.get("result", {})
    if result.get("ok"):
        st.success(result.get("message", "VOC 분석 완료"))
    else:
        st.warning(result.get("message") or result.get("error") or "VOC 분석 결과가 없습니다.")
    result_columns = st.columns(2, gap="medium")
    with result_columns[0].container(border=True, height="stretch"):
        st.markdown("### :material/summarize: 요약")
        st.write(result.get("summary", "-") or "-")
    with result_columns[1].container(border=True, height="stretch"):
        st.markdown("### :material/lightbulb: 정책 개선안")
        st.write(result.get("policy", "-") or "-")
    with st.expander("의도·평가·비평·실행 Trace", icon=":material/account_tree:"):
        st.json({
            "intent_json": result.get("intent_json", "{}"),
            "eval_json": result.get("eval_json", "{}"),
            "summary_critic_json": result.get("summary_critic_json", "{}"),
            "trace": result.get("trace", ""),
            "error": result.get("error", ""),
        })
    if payload.get("reports"):
        st.caption(f"저장된 Report: {payload['reports']}")


def _rubric_rows(items: dict) -> list[dict]:
    rows = []
    for key, value in items.items():
        row = {
            "ID": key,
            "평가 항목": _voc_display_term(value.get("label")),
            "배점": value.get("max_points"),
            "세부 기준": ", ".join(
                f"{name} {score}" for name, score in value.get("criteria", {}).items()
            ),
        }
        if "pass_floor" in value:
            row["통과 하한"] = value.get("pass_floor")
        rows.append(row)
    return rows


def _render_rules(title: str, rules: list[str]):
    st.markdown(f"### {title}")
    for rule in rules:
        st.markdown(f"- `{rule}`")


def _render_internal_pipeline_rubric():
    rubric = load_system_rubric()
    st.caption(
        "6개 Agent 80점 + Agent 연계 10점 + 장애·로그 5점 + 성능 5점으로 "
        "파이프라인 내부 실행 품질을 평가합니다."
    )
    st.dataframe(pd.DataFrame(_rubric_rows(rubric.get("categories", {}))), hide_index=True)
    st.markdown("### 배포 판정")
    st.dataframe(pd.DataFrame(rubric.get("deployment_decisions", [])), hide_index=True)
    _render_rules("점수와 무관한 즉시 배포 보류", rubric.get("immediate_deployment_hold", []))
    st.info(
        "현재는 평가 기준 정의 단계입니다. case별 점수·Run ID·실행 Trace·Rubric 버전 저장은 "
        "Step 2~4 실행 이력에서 연결합니다."
    )


def _render_independent_judge_rubric():
    rubric = load_independent_judge_rubric()
    st.caption(
        f"{rubric.get('title')} · 기본 Provider: {rubric.get('default_provider')} "
        f"(실행 시 변경 가능) · Rubric {rubric.get('version')} · "
        "Evaluator·Critic과 분리된 외부 판정입니다."
    )
    st.dataframe(pd.DataFrame(_rubric_rows(rubric.get("dimensions", {}))), hide_index=True)
    st.markdown("### 독립 LLM 평가 판정")
    st.dataframe(pd.DataFrame(rubric.get("decisions", [])), hide_index=True)
    _render_rules("점수와 무관한 즉시 실패", rubric.get("immediate_fail_rules", []))
    st.markdown("### 품질 점수와 분리하는 실행 상태")
    st.dataframe(
        pd.DataFrame(
            [
                {"상태": status, "의미": description}
                for status, description in rubric.get("non_quality_statuses", {}).items()
            ]
        ),
        hide_index=True,
    )
    st.info(
        "기준은 Step 1에서 등록했습니다. 실제 독립 LLM 평가 호출, 모델 독립성 등급, case별 판정과 "
        "비용·시간 증적은 Step 5에서 구현합니다."
    )


def _render_improvement_validity_rubric():
    rubric = load_improvement_validity_rubric()
    st.caption(
        f"{rubric.get('title')} · Rubric {rubric.get('version')} · "
        "최종 개선안의 근거·실행 가능성과 사람 승인을 검증합니다."
    )
    st.dataframe(pd.DataFrame(_rubric_rows(rubric.get("dimensions", {}))), hide_index=True)
    st.markdown("### AI 자동 판정")
    st.dataframe(pd.DataFrame(rubric.get("automatic_decisions", [])), hide_index=True)
    st.markdown("### 승인 흐름")
    st.write(" → ".join(rubric.get("workflow_states", [])))
    st.warning(rubric.get("formal_approval_rule", ""))
    _render_rules("즉시 승인 보류", rubric.get("immediate_hold_rules", []))
    st.info(
        "기준은 Step 1에서 등록했습니다. VOC·실행 Trace 연결, QA 검토, 업무 담당자 승인과 변경 이력은 "
        "Step 6에서 구현합니다."
    )


def _plain_editor_value(value):
    if value is None or (not isinstance(value, (list, dict)) and pd.isna(value)):
        return None
    if hasattr(value, "item"):
        value = value.item()
    return value


def _score_value(value):
    number = float(_plain_editor_value(value))
    return int(number) if number.is_integer() else number


def _rubric_editor_rows(payload: dict, items_key: str) -> tuple[list[dict], list[dict]]:
    items = payload.get(items_key, {})
    include_pass_floor = any("pass_floor" in item for item in items.values())
    item_rows = []
    criterion_rows = []
    for item_id, item in items.items():
        row = {
            "ID": item_id,
            "평가 항목": item.get("label", ""),
            "배점": item.get("max_points", 0),
        }
        if include_pass_floor:
            row["통과 하한"] = item.get("pass_floor")
        item_rows.append(row)
        for criterion_id, points in item.get("criteria", {}).items():
            criterion_rows.append(
                {
                    "평가 항목 ID": item_id,
                    "세부 기준 ID": criterion_id,
                    "점수": points,
                }
            )
    return item_rows, criterion_rows


def _build_edited_rubric(
    payload: dict,
    spec: dict,
    *,
    version: str,
    title: str | None,
    item_rows: list[dict],
    criterion_rows: list[dict],
    decision_rows: list[dict],
    hold_rules_text: str,
    default_provider: str | None = None,
) -> dict:
    edited = deepcopy(payload)
    edited["version"] = version.strip()
    if title is not None:
        edited["title"] = title.strip()
    if default_provider is not None:
        edited["default_provider"] = default_provider.strip()

    items = edited[spec["items_key"]]
    for row in item_rows:
        item = items[str(row["ID"])]
        item["label"] = str(row["평가 항목"]).strip()
        item["max_points"] = _score_value(row["배점"])
        if "통과 하한" in row:
            item["pass_floor"] = _score_value(row["통과 하한"])
    for row in criterion_rows:
        item_id = str(row["평가 항목 ID"])
        criterion_id = str(row["세부 기준 ID"])
        items[item_id]["criteria"][criterion_id] = _score_value(row["점수"])

    edited_decisions = []
    for row in decision_rows:
        edited_decisions.append(
            {
                key: _plain_editor_value(value)
                for key, value in row.items()
                if _plain_editor_value(value) is not None
            }
        )
    edited[spec["decisions_key"]] = edited_decisions
    edited[spec["hold_rules_key"]] = [
        line.strip() for line in hold_rules_text.splitlines() if line.strip()
    ]
    return edited


def _show_rubric_save_result(rubric_type: str, result: dict, saved_payload: dict):
    if not result.get("ok"):
        for error in result.get("errors", ["품질 평가 기준을 저장하지 못했습니다."]):
            st.error(error)
        return False
    if result.get("changed", True):
        saved_signature = _rubric_signature(saved_payload)
        st.session_state[f"voc_rubric_last_save_message_{rubric_type}"] = "변경완료"
        st.session_state[f"voc_rubric_last_saved_signature_{rubric_type}"] = saved_signature
        st.session_state[f"rubric_edit_{rubric_type}_draft"] = deepcopy(saved_payload)
        st.session_state[f"rubric_edit_{rubric_type}_source"] = saved_signature
    else:
        st.session_state[f"voc_rubric_last_save_message_{rubric_type}"] = "변경없음"
    return True


def _rubric_signature(payload: dict) -> str:
    return json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)


def _rubric_save_state_pill(label: str, *, tone: str) -> str:
    colors = {
        "red": ("#b42318", "#fff1f0", "#f2b8b5"),
        "gray": ("#617083", "#f3f6fa", "#d8e1eb"),
    }
    color, background, border = colors.get(tone, colors["gray"])
    return (
        "<div style=\""
        "display:inline-flex;align-items:center;justify-content:center;"
        "height:38px;min-width:86px;padding:0 14px;border-radius:999px;"
        f"border:1px solid {border};background:{background};color:{color};"
        "font-size:12px;font-weight:800;line-height:1;white-space:nowrap;"
        "box-sizing:border-box;max-width:100%;overflow:hidden;text-overflow:ellipsis;"
        f"\">{escape(label)}</div>"
    )


def _rubric_save_control_state(
    *,
    has_changes: bool,
    needs_version_change: bool,
    validation_errors: list[str],
    last_save_message: str | None,
    saved_signature: str | None,
    draft_signature: str,
) -> dict:
    if not has_changes:
        saved = last_save_message == "변경완료" and saved_signature == draft_signature
        return {
            "label": "변경완료" if saved else "변경없음",
            "tone": "gray",
            "disabled": True,
            "help": "저장할 변경사항이 없습니다.",
        }
    if needs_version_change:
        return {
            "label": "변경발생",
            "tone": "red",
            "disabled": True,
            "help": "Rubric 버전을 변경해야 저장할 수 있습니다.",
            "focus_version": True,
        }
    if validation_errors:
        return {
            "label": "변경발생",
            "tone": "red",
            "disabled": True,
            "help": f"저장 전 확인 필요: {validation_errors[0]}",
        }
    return {
        "label": "변경발생",
        "tone": "red",
        "disabled": False,
        "help": "변경된 평가 기준을 저장합니다.",
    }


def _stabilize_rubric_header_layout(rubric_type: str) -> None:
    widget_keys = (
        f"rubric_edit_{rubric_type}_widget_version",
        f"rubric_edit_{rubric_type}_widget_title",
        f"rubric_edit_{rubric_type}_widget_provider",
    )
    selectors = ",\n".join(
        f".st-key-{widget_key} [data-testid='stTextInput'], "
        f".st-key-{widget_key} [data-testid='stSelectbox']"
        for widget_key in widget_keys
    )
    label_selectors = ",\n".join(
        f".st-key-{widget_key} [data-testid='stWidgetLabel']"
        for widget_key in widget_keys
    )
    st.markdown(
        f"""
        <style>
        {selectors} {{
            min-height:78px!important;
        }}
        {label_selectors} {{
            min-height:32px!important;
            display:flex!important;
            align-items:flex-start!important;
            margin-bottom:4px!important;
        }}
        .st-key-rubric_download_{rubric_type},
        .st-key-rubric_upload_popover_{rubric_type},
        .st-key-rubric_edit_{rubric_type}_save,
        .st-key-rubric_edit_{rubric_type}_save_state {{
            min-height:78px!important;
            display:flex!important;
            flex-direction:column!important;
            justify-content:flex-end!important;
        }}
        .st-key-rubric_edit_{rubric_type}_save_state {{
            align-items:center!important;
            justify-content:flex-end!important;
            overflow:hidden!important;
        }}
        .st-key-rubric_edit_{rubric_type}_save_state [data-testid="stMarkdownContainer"] {{
            width:100%!important;
            display:flex!important;
            justify-content:flex-end!important;
        }}
        .st-key-rubric_download_{rubric_type} button,
        .st-key-rubric_upload_popover_{rubric_type} button,
        .st-key-rubric_edit_{rubric_type}_save button {{
            min-height:38px!important;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _render_rubric_stage_tab_style() -> None:
    st.html(
        """
        <style>
        .st-key-voc_quality_rubric_stage div[data-testid="stRadio"] > label {
            display:none;
        }
        .st-key-voc_quality_rubric_stage div[role="radiogroup"] {
            gap:4px;
            border-bottom:1px solid #c7d8ef;
            margin-bottom:14px;
            flex-wrap:wrap;
        }
        .st-key-voc_quality_rubric_stage div[role="radiogroup"] label {
            padding:10px 14px 9px;
            border-bottom:3px solid transparent;
            border-radius:6px 6px 0 0;
            margin-bottom:-1px;
            color:#5b6b82;
            font-weight:700;
            transition:color .15s ease, background-color .15s ease;
        }
        .st-key-voc_quality_rubric_stage div[role="radiogroup"] label:hover {
            color:#245f99;
            background:#f4f8fd;
        }
        .st-key-voc_quality_rubric_stage div[role="radiogroup"] label:has(input:checked) {
            color:#174b7a;
            background:#eaf3fc;
            border-bottom-color:#2f6fb0;
        }
        </style>
        """
    )


def _render_report_mode_tab_style() -> None:
    st.html(
        """
        <style>
        .st-key-voc_report_mode div[data-testid="stRadio"] > label {
            display:none;
        }
        .st-key-voc_report_mode div[role="radiogroup"] {
            gap:4px;
            border-bottom:1px solid #c7d8ef;
            margin-bottom:14px;
            flex-wrap:wrap;
        }
        .st-key-voc_report_mode div[role="radiogroup"] label {
            padding:10px 14px 9px;
            border-bottom:3px solid transparent;
            border-radius:6px 6px 0 0;
            margin-bottom:-1px;
            color:#5b6b82;
            font-weight:700;
            transition:color .15s ease, background-color .15s ease;
        }
        .st-key-voc_report_mode div[role="radiogroup"] label:hover {
            color:#245f99;
            background:#f4f8fd;
        }
        .st-key-voc_report_mode div[role="radiogroup"] label:has(input:checked) {
            color:#174b7a;
            background:#eaf3fc;
            border-bottom-color:#2f6fb0;
        }
        </style>
        """
    )


def _highlight_rubric_version_input(rubric_type: str) -> None:
    widget_class = f"st-key-rubric_edit_{rubric_type}_widget_version"
    st.markdown(
        f"""
        <style>
        .st-key-rubric_header_{rubric_type}>div[data-testid="stVerticalBlock"] {{
            gap:.35rem!important;
        }}
        .st-key-rubric_header_{rubric_type} [data-testid="stTextInput"],
        .st-key-rubric_header_{rubric_type} [data-testid="stSelectbox"] {{
            min-height:72px!important;
        }}
        .st-key-rubric_header_{rubric_type} [data-testid="stElementContainer"] {{
            margin-bottom:0!important;
        }}
        .{widget_class} input {{
            border-color:#d83f36!important;
            box-shadow:0 0 0 2px rgba(216,63,54,.14)!important;
            background:#fffafa!important;
        }}
        .{widget_class} label p {{
            color:#b42318!important;
            font-weight:800!important;
            line-height:1.2!important;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def _ordered_decision_rows(decisions: list[dict], spec: dict) -> list[dict]:
    """Return decision bands from the highest score to the lowest score."""
    min_key = spec["decision_min_key"]
    return sorted(
        (deepcopy(row) for row in decisions),
        key=lambda row: float(row.get(min_key, 0)),
        reverse=True,
    )


def _link_decision_ranges(
    decisions: list[dict],
    spec: dict,
    boundary_index: int,
    boundary_score: float,
) -> list[dict]:
    """Move one boundary and keep adjoining decision bands contiguous."""
    min_key = spec["decision_min_key"]
    max_key = spec["decision_max_key"]
    ordered = _ordered_decision_rows(decisions, spec)
    boundaries = [float(row[min_key]) for row in ordered[:-1]]
    boundaries[boundary_index] = round(float(boundary_score), 2)

    for index, row in enumerate(ordered):
        minimum = boundaries[index] if index < len(boundaries) else 0.0
        maximum = 100.0 if index == 0 else round(boundaries[index - 1] - 0.01, 2)
        row[min_key] = _score_value(minimum)
        row[max_key] = _score_value(maximum)
    return ordered


def _decision_display_frame(decisions: list[dict], spec: dict) -> pd.DataFrame:
    """Build the preview table with the user-facing decision first."""
    rows = []
    for row in _ordered_decision_rows(decisions, spec):
        rows.append({"decision": row.get("decision", ""), **{k: v for k, v in row.items() if k != "decision"}})
    return pd.DataFrame(rows)


def _rubric_draft(payload: dict, rubric_type: str) -> dict:
    draft_key = f"rubric_edit_{rubric_type}_draft"
    source_key = f"rubric_edit_{rubric_type}_source"
    source = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    if st.session_state.get(source_key) != source:
        widget_prefix = f"rubric_edit_{rubric_type}_widget_"
        for key in list(st.session_state):
            if str(key).startswith(widget_prefix):
                del st.session_state[key]
        st.session_state[draft_key] = deepcopy(payload)
        st.session_state[source_key] = source
    return st.session_state[draft_key]


def _rubric_item_total(item: dict) -> float:
    return sum(float(value) for value in item.get("criteria", {}).values())


def _rubric_total(items: dict) -> float:
    return sum(_rubric_item_total(item) for item in items.values())


RUBRIC_CRITERION_TEMPORARY_DELTA = 2


def _rubric_criterion_range(
    items: dict,
    item_id: str,
    criterion_id: str,
    *,
    total_budget: int = 100,
    temporary_delta: int = RUBRIC_CRITERION_TEMPORARY_DELTA,
    current_score: int | float | None = None,
) -> tuple[int, int]:
    """Return a narrow edit range that permits temporary over/under allocation.

    The screen should not trap users while they redistribute points, but it
    should also prevent accidental large swings. The final 100-point contract
    is still enforced by validate_quality_rubric/save_quality_rubric when
    saving.
    """
    try:
        source_score = (
            current_score
            if current_score is not None
            else items[item_id]["criteria"][criterion_id]
        )
        current = int(round(float(source_score)))
    except (KeyError, TypeError, ValueError):
        current = 0
    delta = max(0, int(temporary_delta))
    return max(0, current - delta), min(int(total_budget), current + delta)


def _rubric_detail_range_anchors(rubric_type: str, item_id: str, item: dict) -> dict:
    selected_key = f"rubric_edit_{rubric_type}_selected_item"
    dialog_key = f"{selected_key}_detail_dialog_item"
    anchor_key = f"{dialog_key}_range_anchors"
    anchors = st.session_state.setdefault(anchor_key, {})
    item_anchor = anchors.setdefault(item_id, {})
    for criterion_id, points in item.get("criteria", {}).items():
        if criterion_id not in item_anchor:
            try:
                item_anchor[criterion_id] = int(round(float(points)))
            except (TypeError, ValueError):
                item_anchor[criterion_id] = 0
    return item_anchor


def _sync_rubric_score_widgets_to_draft(draft: dict, rubric_type: str, spec: dict) -> None:
    items = draft.get(spec["items_key"], {})
    if not isinstance(items, dict):
        return
    for item_id, item in items.items():
        if not isinstance(item, dict):
            continue
        criteria = item.get("criteria")
        if not isinstance(criteria, dict):
            continue
        prefix = f"rubric_edit_{rubric_type}_widget_{item_id}"
        changed = False
        for criterion_id in list(criteria):
            widget_key = f"{prefix}_criterion_{criterion_id}"
            if widget_key not in st.session_state:
                continue
            score = st.session_state.get(widget_key)
            try:
                score_value = _score_value(float(score))
            except (TypeError, ValueError):
                continue
            if criteria.get(criterion_id) != score_value:
                criteria[criterion_id] = score_value
                changed = True
        if changed:
            item["max_points"] = _score_value(_rubric_item_total(item))
        pass_floor_key = f"{prefix}_pass_floor"
        if "pass_floor" in item and pass_floor_key in st.session_state:
            try:
                item_total = max(0, int(round(_rubric_item_total(item))))
                if item_total > 0:
                    item["pass_floor"] = max(
                        1,
                        min(int(st.session_state[pass_floor_key]), item_total),
                    )
            except (TypeError, ValueError):
                pass


def _render_rubric_transfer_tools(
    draft: dict,
    rubric_type: str,
    spec: dict,
    download_container=None,
    upload_container=None,
):
    if download_container is None or upload_container is None:
        download_container, upload_container = st.columns(2, gap="small")
    with download_container:
        st.download_button(
            "JSON D/L",
            data=(json.dumps(draft, ensure_ascii=False, indent=2) + "\n").encode("utf-8"),
            file_name=Path(spec["relative_path"]).name,
            mime="application/json",
            icon=":material/download:",
            key=f"rubric_download_{rubric_type}",
            width="stretch",
        )
    with upload_container, st.popover(
            "JSON Up",
            icon=":material/upload_file:",
            key=f"rubric_upload_popover_{rubric_type}",
            width="stretch",
        ):
        uploaded = st.file_uploader(
            "Rubric JSON 파일",
            type=["json"],
            max_upload_size=1,
            key=f"rubric_upload_{rubric_type}",
            help="현재 선택한 평가 단계와 같은 구조의 JSON만 적용할 수 있습니다.",
        )
        if uploaded is None:
            st.caption("업로드한 파일은 검증을 통과한 뒤에만 저장됩니다.")
            return
        try:
            uploaded_payload = json.loads(uploaded.getvalue().decode("utf-8-sig"))
            upload_errors = validate_quality_rubric(rubric_type, uploaded_payload)
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            uploaded_payload = None
            upload_errors = [f"JSON 파일을 해석할 수 없습니다: {exc}"]
        if upload_errors:
            for error in upload_errors:
                st.error(error)
            return
        st.success("업로드 파일 검증을 통과했습니다.")
        if st.button(
            "업로드 기준 적용",
            type="primary",
            icon=":material/check:",
            key=f"rubric_apply_upload_{rubric_type}",
            width="stretch",
        ):
            result = save_quality_rubric(rubric_type, uploaded_payload, source="json_upload")
            saved = _show_rubric_save_result(rubric_type, result, uploaded_payload)
            if saved:
                draft.clear()
                draft.update(deepcopy(uploaded_payload))
                st.rerun()


RUBRIC_CRITERION_KO_LABELS = {
    "intent": "의도 파악",
    "keywords": "핵심어 추출",
    "search_conditions": "검색 조건",
    "defaults_and_format": "기본값·형식",
    "recall": "검색 재현율",
    "precision": "검색 정밀도",
    "source_preservation": "출처 보존",
    "limit_and_error_handling": "제한·오류 처리",
    "factual_consistency": "사실 일관성",
    "coverage": "핵심 내용 포함",
    "deduplication_and_conciseness": "중복 제거·간결성",
    "format_and_readability": "형식·가독성",
    "criteria_consistency": "평가 기준 일관성",
    "winner_correctness": "우수안 판정 정확성",
    "evidence_validity": "근거 타당성",
    "format_and_repeatability": "형식·반복 가능성",
    "defect_detection": "결함 탐지",
    "risk_detection": "위험 탐지",
    "edit_actionability": "수정 실행 가능성",
    "false_positive_control": "오탐 통제",
    "voc_grounding": "VOC 근거 반영",
    "specificity": "구체성",
    "feasibility": "실행 가능성",
    "measurability": "측정 가능성",
    "priority": "우선순위",
    "trace_completeness": "실행 Trace 완전성",
    "data_transfer_integrity": "데이터 전달 무결성",
    "upstream_result_usage": "이전 단계 결과 활용",
    "no_duplicate_calls": "중복 호출 방지",
    "explicit_failure_response": "명확한 실패 응답",
    "traceable_error_log": "추적 가능한 오류 로그",
    "recovery_and_cleanup": "복구·정리",
    "end_to_end_response_time": "전체 응답시간",
    "timeout_compliance": "타임아웃 준수",
    "per_agent_timing_visibility": "Agent별 소요시간 표시",
    "complaint_type_and_cause": "불만 유형·원인 정확성",
    "impact_consistency": "영향 일관성",
    "policy_statement_correctness": "정책 설명 정확성",
    "question_relevance": "질문 관련성",
    "voc_source_traceability": "VOC 출처 추적성",
    "no_unsupported_claim": "근거 없는 주장 방지",
    "source_meaning_preservation": "원문 의미 보존",
    "uncertainty_disclosure": "불확실성 명시",
    "required_issue_coverage": "필수 이슈 포함",
    "cause_impact_action_coverage": "원인·영향·조치 포함",
    "compound_complaint_coverage": "복합 불만 포함",
    "no_critical_omission": "핵심 누락 방지",
    "action_detail": "조치 상세성",
    "owner_and_priority": "담당·우선순위",
    "measurable_kpi": "측정 가능한 KPI",
    "clear_language": "명확한 표현",
    "no_sensitive_data_exposure": "민감정보 노출 방지",
    "no_fabricated_guarantee": "허위 보장 방지",
    "failure_transparency": "실패 투명성",
    "safe_escalation": "안전한 상향 보고",
    "complaint_to_root_cause": "불만-근본원인 연결",
    "root_cause_to_action": "근본원인-조치 연결",
    "expected_customer_impact": "예상 고객 영향",
    "voc_id_reference": "VOC ID 참조",
    "trace_and_agent_reference": "실행 Trace·Agent 참조",
    "no_unsupported_evidence": "근거 없는 증적 방지",
    "process_feasibility": "업무 절차 실행 가능성",
    "technical_feasibility": "기술 실행 가능성",
    "resource_and_dependency_awareness": "자원·의존성 고려",
    "responsible_owner": "담당자 명확성",
    "target_schedule": "목표 일정",
    "customer_and_operational_risk": "고객·운영 위험",
    "privacy_and_security": "개인정보·보안",
    "compliance_and_escalation": "규제 준수·상향 보고",
}


RUBRIC_ITEM_PANEL_HEIGHT = 460
RUBRIC_CRITERIA_PANEL_MIN_HEIGHT = 430
RUBRIC_WEIGHT_CHART_HEIGHT = 112
RUBRIC_DETAIL_DIALOG_WIDTH = "medium"
RUBRIC_DETAIL_NAV_STYLE = """
<style>
[class*="st-key-rubric_detail_previous_"] button,
[class*="st-key-rubric_detail_next_"] button {
    background: #F2F6FB;
    border: 1px solid #B9CBE0;
    color: #345F8A;
    box-shadow: none;
    font-weight: 600;
}
[class*="st-key-rubric_detail_previous_"] button:hover,
[class*="st-key-rubric_detail_next_"] button:hover {
    background: #E8F0F8;
    border-color: #8EADCC;
    color: #244F78;
}
[class*="st-key-rubric_detail_previous_"] button:focus,
[class*="st-key-rubric_detail_next_"] button:focus {
    border-color: #789DC2;
    box-shadow: 0 0 0 0.12rem rgba(72, 116, 158, 0.16);
}
[class*="st-key-rubric_criteria_panel_"] {
    min-height: %dpx;
    overflow: visible;
}
</style>
""" % RUBRIC_CRITERIA_PANEL_MIN_HEIGHT


def _rubric_criterion_label(criterion_id: str) -> str:
    korean = RUBRIC_CRITERION_KO_LABELS.get(
        criterion_id,
        criterion_id.replace("_", " ").title(),
    )
    return f"{korean} ({criterion_id})"


def _build_rubric_weight_chart(item_label: str, item_total: float) -> alt.LayerChart:
    selected_score = max(0.0, min(float(item_total), 100.0))
    chart_rows = pd.DataFrame(
        [
            {
                "구분": item_label,
                "배점": selected_score,
                "색상": "선택 평가 항목",
                "순서": 1,
            },
            {
                "구분": "나머지 평가 항목",
                "배점": 100.0 - selected_score,
                "색상": "나머지 평가 항목",
                "순서": 2,
            },
        ]
    )
    arc = (
        alt.Chart(chart_rows)
        .mark_arc(innerRadius=34, outerRadius=50, cornerRadius=4)
        .encode(
            theta=alt.Theta("배점:Q", stack=True),
            color=alt.Color(
                "색상:N",
                scale=alt.Scale(
                    domain=["선택 평가 항목", "나머지 평가 항목"],
                    range=["#1769AA", "#E3EAF2"],
                ),
                legend=None,
            ),
            order=alt.Order("순서:O"),
            tooltip=[
                alt.Tooltip("구분:N"),
                alt.Tooltip("배점:Q", format=".0f", title="배점"),
            ],
        )
    )
    score_text = (
        alt.Chart(pd.DataFrame([{"text": f"{selected_score:g}점"}]))
        .mark_text(fontSize=18, fontWeight=700, color="#0B4F91", dy=-6)
        .encode(text="text:N")
    )
    total_text = (
        alt.Chart(pd.DataFrame([{"text": "전체 100점 중"}]))
        .mark_text(fontSize=10, color="#66788A", dy=13)
        .encode(text="text:N")
    )
    return (arc + score_text + total_text).properties(height=RUBRIC_WEIGHT_CHART_HEIGHT)


def _selected_rubric_item_id(
    item_ids: list[str],
    selection_state: dict | None,
    fallback: str,
) -> str:
    selected_row = _table_selected_row_index(selection_state, len(item_ids))
    if selected_row is not None:
        return item_ids[selected_row]
    return fallback


def _sync_rubric_item_selection(
    table_key: str,
    selected_key: str,
    item_ids: list[str],
):
    suppress_key = f"{selected_key}_suppress_detail_dialog_once"
    if st.session_state.pop(suppress_key, False):
        return
    selected_row = _promote_table_cell_to_row_selection(
        table_key,
        len(item_ids),
    )
    if selected_row is not None:
        selected_id = item_ids[selected_row]
        st.session_state[selected_key] = selected_id
        st.session_state[f"{selected_key}_detail_dialog_request"] = selected_id


def _navigate_rubric_detail_dialog(
    rubric_type: str,
    item_ids: list[str],
    item_id: str,
):
    """팝업 항목만 이동하고 목록은 선택 항목 상태로 다음 전체 렌더에서 맞춥니다."""
    if item_id not in item_ids:
        return
    selected_key = f"rubric_edit_{rubric_type}_selected_item"
    dialog_key = f"{selected_key}_detail_dialog_item"
    opened_key = f"{dialog_key}_opened"
    st.session_state[selected_key] = item_id
    st.session_state[dialog_key] = item_id
    st.session_state[opened_key] = True


def _clear_rubric_detail_dialog_state(rubric_type: str):
    selected_key = f"rubric_edit_{rubric_type}_selected_item"
    dialog_key = f"{selected_key}_detail_dialog_item"
    st.session_state.pop(f"{selected_key}_detail_dialog_request", None)
    st.session_state.pop(dialog_key, None)
    st.session_state.pop(f"{dialog_key}_opened", None)
    st.session_state.pop(f"{dialog_key}_range_anchors", None)


def _dismiss_rubric_detail_dialog():
    for rubric_type in QUALITY_RUBRIC_SPECS:
        _clear_rubric_detail_dialog_state(rubric_type)


def _complete_rubric_detail_dialog(rubric_type: str):
    selected_key = f"rubric_edit_{rubric_type}_selected_item"
    dialog_key = f"{selected_key}_detail_dialog_item"
    st.session_state.pop(f"{selected_key}_detail_dialog_request", None)
    st.session_state.pop(dialog_key, None)
    st.session_state.pop(f"{dialog_key}_range_anchors", None)
    # Dialog 조각 rerun이 먼저 발생하므로 전체 앱 rerun을 유도할 표식을 남깁니다.
    st.session_state[f"{dialog_key}_opened"] = True


@st.dialog(
    "세부 배점 설정",
    width=RUBRIC_DETAIL_DIALOG_WIDTH,
    icon=":material/tune:",
    on_dismiss=_dismiss_rubric_detail_dialog,
)
def _rubric_item_detail_dialog(
    draft: dict,
    rubric_type: str,
    spec: dict,
    selected_id: str,
):
    items = draft.get(spec["items_key"], {})
    item_ids = list(items)
    selected_key = f"rubric_edit_{rubric_type}_selected_item"
    dialog_key = f"{selected_key}_detail_dialog_item"
    opened_key = f"{dialog_key}_opened"
    if st.session_state.get(dialog_key) not in items:
        if st.session_state.get(opened_key):
            st.session_state.pop(opened_key, None)
            st.rerun(scope="app")
        st.session_state[dialog_key] = selected_id
        st.session_state[opened_key] = True
    selected_id = st.session_state[dialog_key]
    if selected_id not in items or not item_ids:
        st.warning("선택한 평가 항목을 찾을 수 없습니다.")
        return
    selected_index = item_ids.index(selected_id)
    previous_id = item_ids[(selected_index - 1) % len(item_ids)]
    next_id = item_ids[(selected_index + 1) % len(item_ids)]
    item = items[selected_id]
    range_anchors = _rubric_detail_range_anchors(rubric_type, selected_id, item)
    prefix = f"rubric_edit_{rubric_type}_widget_{selected_id}"
    item_total = _rubric_item_total(item)
    st.html(RUBRIC_DETAIL_NAV_STYLE)
    previous_col, title_col, chart_col, next_col = st.columns(
        [1.35, 4.6, 2.4, 1.35],
        gap="small",
        vertical_alignment="center",
    )
    with previous_col:
        st.button(
            "< 이전",
            type="secondary",
            key=f"rubric_detail_previous_{rubric_type}_{selected_id}",
            help=f"이전 · {items[previous_id].get('label', previous_id)}",
            on_click=_navigate_rubric_detail_dialog,
            args=(rubric_type, item_ids, previous_id),
            width="stretch",
        )
    with title_col:
        st.markdown(f"#### {item.get('label', selected_id)}")
        st.caption(
            f"{selected_index + 1} / {len(item_ids)} · "
            f"현재 점수 기준 ±{RUBRIC_CRITERION_TEMPORARY_DELTA}점 임시 조정 · 저장 시 총점 100점 검증"
        )
    with chart_col:
        chart_slot = st.empty()
    with next_col:
        st.button(
            "다음 >",
            type="secondary",
            key=f"rubric_detail_next_{rubric_type}_{selected_id}",
            help=f"다음 · {items[next_id].get('label', next_id)}",
            on_click=_navigate_rubric_detail_dialog,
            args=(rubric_type, item_ids, next_id),
            width="stretch",
        )

    with st.container(
        border=False,
        key=f"rubric_criteria_panel_{rubric_type}_{selected_id}",
    ):
        for criterion_id, points in list(item.get("criteria", {}).items()):
            minimum, maximum = _rubric_criterion_range(
                items,
                selected_id,
                criterion_id,
                current_score=range_anchors.get(criterion_id, points),
            )
            score = st.slider(
                _rubric_criterion_label(criterion_id),
                min_value=minimum,
                max_value=maximum,
                value=max(minimum, min(int(round(float(points))), maximum)),
                step=1,
                format="%d점",
                key=f"{prefix}_criterion_{criterion_id}",
                help=(
                    f"현재 설정 점수에서 ±{RUBRIC_CRITERION_TEMPORARY_DELTA}점 범위로 임시 초과·부족을 허용합니다. "
                    "최종 저장은 전체 합계가 100점일 때만 가능합니다."
                ),
            )
            if score != points:
                item["criteria"][criterion_id] = score
                item["max_points"] = _score_value(_rubric_item_total(item))
                if "pass_floor" in item and float(item["pass_floor"]) > float(item["max_points"]):
                    item["pass_floor"] = item["max_points"]

        item_total = _rubric_item_total(item)
        item["max_points"] = _score_value(item_total)
        if "pass_floor" in item:
            pass_floor = st.slider(
                "통과 하한",
                min_value=1,
                max_value=max(1, int(round(item_total))),
                value=max(1, min(int(round(float(item.get("pass_floor", 1)))), int(round(item_total)))),
                step=1,
                format="%d점",
                key=f"{prefix}_pass_floor",
                help="평가 항목 배점을 넘지 않는 범위에서 설정합니다.",
            )
            item["pass_floor"] = pass_floor

    chart_slot.altair_chart(
        _build_rubric_weight_chart(
            str(item.get("label", selected_id)),
            item_total,
        )
    )

    with st.container(horizontal=True, horizontal_alignment="right"):
        st.button(
            "설정 완료",
            type="primary",
            icon=":material/check:",
            key=f"rubric_detail_done_{rubric_type}_{selected_id}",
            on_click=_complete_rubric_detail_dialog,
            args=(rubric_type,),
        )


def _render_rubric_items(draft: dict, rubric_type: str, spec: dict):
    items = draft.get(spec["items_key"], {})
    item_ids = list(items)
    selected_key = f"rubric_edit_{rubric_type}_selected_item"
    if st.session_state.get(selected_key) not in items:
        st.session_state[selected_key] = item_ids[0]
    selected_id = st.session_state[selected_key]

    with st.container(
        border=True,
        height=RUBRIC_ITEM_PANEL_HEIGHT,
        key=f"rubric_item_list_{rubric_type}",
    ):
        with st.container(
            horizontal=True,
            horizontal_alignment="distribute",
            vertical_alignment="center",
        ):
            st.markdown("### 항목별 배점 설정")
            _render_rubric_total_summary(draft, spec)
        item_frame = pd.DataFrame(
            [
                {
                    "ID": item_id,
                    "평가 항목": item.get("label", ""),
                    "배점": _score_value(_rubric_item_total(item)),
                }
                for item_id, item in items.items()
            ]
        )
        score_column = item_frame.columns[-1]
        item_scores = item_frame[score_column].astype(float).tolist()
        item_score_max = max(item_scores, default=1.0)
        item_score_max = max(item_score_max, 1.0)
        default_row = item_ids.index(selected_id)
        table_key = f"rubric_edit_{rubric_type}_widget_item_table"
        st.dataframe(
            item_frame,
            hide_index=True,
            on_select=partial(
                _sync_rubric_item_selection,
                table_key,
                selected_key,
                item_ids,
            ),
            selection_mode=["single-row-required", "single-cell"],
            selection_default={"selection": {"rows": [default_row]}},
            key=table_key,
            column_config={
                "ID": None,
                "배점": st.column_config.ProgressColumn(
                    "배점",
                    min_value=0,
                    max_value=item_score_max,
                    format="%g점",
                ),
            },
            height=min(430, 38 + len(item_frame) * 35),
        )
    request_key = f"{selected_key}_detail_dialog_request"
    dialog_key = f"{selected_key}_detail_dialog_item"
    opened_key = f"{dialog_key}_opened"
    if st.session_state.pop(f"{selected_key}_suppress_detail_dialog_once", False):
        st.session_state.pop(request_key, None)
    requested_item = st.session_state.pop(request_key, None)
    if requested_item in items:
        st.session_state[dialog_key] = requested_item
        st.session_state[opened_key] = True
    elif dialog_key not in st.session_state:
        st.session_state.pop(opened_key, None)
    dialog_item = st.session_state.get(dialog_key)
    if dialog_item in items:
        _rubric_item_detail_dialog(
            draft,
            rubric_type,
            spec,
            dialog_item,
        )


def _render_rubric_total_summary(draft: dict, spec: dict):
    total = _rubric_total(draft.get(spec["items_key"], {}))
    complete = abs(total - 100.0) < 0.001
    if complete:
        first_line = f"{total:g} / 100점 배점 구성 완료"
        second_line = "저장 가능한 배점 구성입니다."
        accent = "#155a96"
        background = "#f3f8fd"
        border = "#b9cee2"
    else:
        first_line = f"{total:g} / 100점 배점 조정 필요"
        second_line = f"100점까지 {100 - total:+g}점 조정이 필요합니다."
        accent = "#b42318"
        background = "#fff7f6"
        border = "#f2b8b5"
    st.markdown(
        f"""
        <div style="
            display:flex;flex-direction:column;align-items:flex-end;justify-content:center;gap:3px;
            min-width:190px;max-width:260px;padding:5px 9px;border:1px solid {border};
            border-radius:10px;background:{background};box-sizing:border-box;line-height:1.2;
        ">
            <div style="display:block;font-size:12px;font-weight:850;color:{accent};white-space:nowrap;">
                {escape(first_line)}
            </div>
            <div style="display:block;font-size:10px;font-weight:700;color:{accent};opacity:.82;white-space:nowrap;">
                {escape(second_line)}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


@st.dialog(
    "판정 구간 미리보기",
    width="large",
    icon=":material/table_view:",
)
def _rubric_decision_preview_dialog(draft: dict, spec: dict):
    decision_frame = _decision_display_frame(
        draft.get(spec["decisions_key"], []),
        spec,
    )
    ordered_columns = [
        "decision",
        spec["decision_min_key"],
        spec["decision_max_key"],
        *[
            column
            for column in decision_frame.columns
            if column
            not in {
                "decision",
                spec["decision_min_key"],
                spec["decision_max_key"],
            }
        ],
    ]
    st.dataframe(
        decision_frame,
        hide_index=True,
        column_order=ordered_columns,
        column_config={
            "decision": st.column_config.TextColumn("decision", pinned=True),
        },
    )


def _render_decision_gauges(draft: dict, rubric_type: str, spec: dict):
    decisions_key = spec["decisions_key"]
    min_key = spec["decision_min_key"]
    ordered = _ordered_decision_rows(draft.get(decisions_key, []), spec)

    with st.container(
        border=True,
        height=RUBRIC_ITEM_PANEL_HEIGHT,
        key=f"rubric_decision_section_{rubric_type}",
    ):
        st.markdown("### 판정 구간")
        st.caption("시작 점수를 조정하면 맞닿은 상·하위 판정의 최소·최대 점수가 함께 변경됩니다.")
        for index, row in enumerate(ordered[:-1]):
            lower_bound = round(float(ordered[index + 1][min_key]) + 0.01, 2)
            upper_bound = 100.0 if index == 0 else round(float(ordered[index - 1][min_key]) - 0.01, 2)
            boundary = st.slider(
                f"{row.get('decision', '-')} 시작 점수",
                min_value=lower_bound,
                max_value=upper_bound,
                value=float(row[min_key]),
                step=0.01,
                format="%.2f점",
                key=f"rubric_edit_{rubric_type}_widget_decision_{index}",
            )
            if abs(boundary - float(row[min_key])) > 0.001:
                draft[decisions_key] = _link_decision_ranges(
                    ordered,
                    spec,
                    index,
                    boundary,
                )
                st.rerun()

        if st.button(
            "판정 구간 미리보기",
            icon=":material/visibility:",
            key=f"rubric_decision_preview_{rubric_type}",
            width="stretch",
        ):
            _rubric_decision_preview_dialog(draft, spec)

        with st.expander("즉시 실패·보류 규칙", icon=":material/warning:"):
            hold_rules_text = st.text_area(
                "한 줄에 한 규칙",
                value="\n".join(draft.get(spec["hold_rules_key"], [])),
                height=140,
                key=f"rubric_edit_{rubric_type}_widget_hold_rules",
            )
            draft[spec["hold_rules_key"]] = [
                line.strip()
                for line in hold_rules_text.splitlines()
                if line.strip()
            ]


def _render_rubric_management(stage: str):
    rubric_type = RUBRIC_STAGE_TYPES[stage]
    spec = QUALITY_RUBRIC_SPECS[rubric_type]
    payload = load_quality_rubric(rubric_type)
    draft = _rubric_draft(payload, rubric_type)
    _stabilize_rubric_header_layout(rubric_type)
    _sync_rubric_score_widgets_to_draft(draft, rubric_type, spec)

    (
        version_col,
        title_col,
        provider_col,
        download_col,
        upload_col,
        save_state_col,
        save_col,
    ) = st.columns(
        [1.0, 1.8, 1.4, 0.8, 0.8, 0.78, 1.32],
        gap="small",
        vertical_alignment="bottom",
    )
    with version_col:
        version_widget_key = f"rubric_edit_{rubric_type}_widget_version"
        previous_version_key = f"{version_widget_key}_previous"
        version = st.text_input(
            "Rubric 버전",
            value=str(draft.get("version", "")),
            key=version_widget_key,
            help="기준 내용을 변경해 저장할 때는 이전과 다른 버전을 입력해야 합니다.",
        )
        previous_version = st.session_state.get(previous_version_key)
        if previous_version is not None and str(previous_version) != str(version):
            selected_key = f"rubric_edit_{rubric_type}_selected_item"
            st.session_state[f"{selected_key}_suppress_detail_dialog_once"] = True
            st.session_state.pop(f"{selected_key}_detail_dialog_request", None)
        st.session_state[previous_version_key] = str(version)
        draft["version"] = version.strip()
        original_version = str(payload.get("version", "")).strip()
    with title_col:
        default_titles = {
            "internal_pipeline": "내부 파이프라인 품질 평가 기준",
            "independent_judge": "독립 LLM 평가 100점 기준",
            "improvement_validity": "개선안 타당성 평가 100점 기준",
        }
        title = st.text_input(
            "기준명",
            value=str(draft.get("title") or default_titles[rubric_type]),
            key=f"rubric_edit_{rubric_type}_widget_title",
        )
        draft["title"] = title.strip()
    with provider_col:
        if rubric_type == "independent_judge":
            current_provider = str(draft.get("default_provider") or "anthropic")
            provider_options = list(
                dict.fromkeys(
                    [current_provider, "anthropic", "openai", "google", "azure_openai"]
                )
            )
            draft["default_provider"] = st.selectbox(
                "기본 평가 Provider",
                provider_options,
                index=provider_options.index(current_provider),
                accept_new_options=True,
                key=f"rubric_edit_{rubric_type}_widget_provider",
            )
        else:
            st.selectbox(
                "기본 평가 Provider",
                ["해당 없음"],
                disabled=True,
                key=f"rubric_edit_{rubric_type}_widget_provider",
            )

    header_validation_errors = validate_quality_rubric(
        rubric_type,
        draft,
    )
    draft_signature = _rubric_signature(draft)
    payload_signature = _rubric_signature(payload)
    saved_signature = st.session_state.get(
        f"voc_rubric_last_saved_signature_{rubric_type}"
    )
    has_rubric_changes = draft_signature != payload_signature
    needs_version_change = has_rubric_changes and draft.get("version", "") == original_version
    save_state = _rubric_save_control_state(
        has_changes=has_rubric_changes,
        needs_version_change=needs_version_change,
        validation_errors=header_validation_errors,
        last_save_message=st.session_state.get(
            f"voc_rubric_last_save_message_{rubric_type}"
        ),
        saved_signature=saved_signature,
        draft_signature=draft_signature,
    )
    if save_state.get("focus_version"):
        _highlight_rubric_version_input(rubric_type)

    _render_rubric_transfer_tools(
        draft,
        rubric_type,
        spec,
        download_container=download_col,
        upload_container=upload_col,
    )
    with save_state_col:
        with st.container(
            horizontal=True,
            horizontal_alignment="right",
            key=f"rubric_edit_{rubric_type}_save_state",
        ):
            st.markdown(
                _rubric_save_state_pill(
                    save_state["label"],
                    tone=save_state["tone"],
                ),
                unsafe_allow_html=True,
            )
    with save_col:
        if st.button(
            "평가 기준 저장",
            type="primary",
            icon=":material/save:",
            key=f"rubric_edit_{rubric_type}_save",
            width="stretch",
            help=save_state["help"],
            disabled=bool(save_state["disabled"]),
        ):
            saved_payload = deepcopy(draft)
            saved = _show_rubric_save_result(
                rubric_type,
                save_quality_rubric(
                    rubric_type,
                    saved_payload,
                    source="screen_editor",
                ),
                saved_payload,
            )
            if saved:
                st.rerun()

    score_col, decision_col = st.columns(
        2,
        gap="small",
        vertical_alignment="top",
    )
    with score_col:
        _render_rubric_items(draft, rubric_type, spec)
    with decision_col:
        _render_decision_gauges(draft, rubric_type, spec)

    validation_errors = validate_quality_rubric(rubric_type, draft)
    if validation_errors:
        with st.expander(f"저장 전 확인 필요 · {len(validation_errors)}건", icon=":material/error:"):
            for error in validation_errors:
                st.markdown(f"- {error}")


def render_rubric():
    _render_rubric_stage_tab_style()
    current_stage = st.session_state.get("voc_quality_rubric_stage")
    if current_stage in RUBRIC_STAGE_DISPLAY_ALIASES:
        st.session_state.voc_quality_rubric_stage = RUBRIC_STAGE_DISPLAY_ALIASES[current_stage]
    stage = st.radio(
        "수정할 평가 단계",
        RUBRIC_STAGE_OPTIONS,
        index=0,
        horizontal=True,
        key="voc_quality_rubric_stage",
        label_visibility="collapsed",
        on_change=_dismiss_rubric_detail_dialog,
    )
    selected_stage = stage or RUBRIC_STAGE_OPTIONS[0]
    _render_rubric_management(selected_stage)


@st.cache_data(ttl=3, max_entries=1, show_spinner=False)
def _load_voc_defect_rows():
    return list_voc_defects()


def _clear_voc_defect_caches():
    _load_voc_defect_rows.clear()
    _load_voc_history_rows.clear()


def _defect_status_label(status: str) -> str:
    return {
        "OPEN": "접수",
        "ANALYZED": "분석 완료",
        "FIXED": "조치 완료",
        "RETESTED": "재시험 완료",
        "CLOSED": "종료",
    }.get(status, status or "-")


def _render_defect_create():
    st.markdown("### 신규 결함 등록")
    st.caption("미확인 이슈는 PENDING으로 등록하고 원본 Run·Case·실행 Trace 확인 후 CONFIRMED로 전환합니다.")
    history = [row for row in _load_voc_history_rows() if row.get("status") != "RUNNING"]
    run_options = [""] + [row["run_id"] for row in history]
    selected_run_id = st.selectbox(
        "원본 Run ID (선택)",
        run_options,
        format_func=lambda value: value or "연결하지 않음",
        key="defect_create_run",
    )
    selected_cases = []
    if selected_run_id:
        selected_row = next(row for row in history if row["run_id"] == selected_run_id)
        selected_cases = st.multiselect(
            "관련 Case ID", selected_row.get("selected_case_ids", []), key="defect_create_cases"
        )

    with st.form("voc_defect_create_form", border=True):
        title = st.text_input("결함 제목")
        columns = st.columns(3)
        severity = columns[0].selectbox("심각도", ["CRITICAL", "HIGH", "MEDIUM", "LOW"], index=1)
        category = columns[1].selectbox(
            "결함 분류",
            ["INTERFACE_BRANCH", "API_RATE_LIMIT", "AGENT_FAILURE", "DATA", "PERFORMANCE", "OTHER"],
        )
        evidence_status = columns[2].selectbox("증적 상태", ["PENDING", "CONFIRMED"])
        description = st.text_area("현상 및 재현 정보", height=120)
        trace_text = st.text_input("관련 실행 Trace ID", help="여러 건은 쉼표로 구분합니다.")
        metadata_columns = st.columns(3)
        actor = metadata_columns[0].text_input("등록자", value="QA")
        candidate_key = metadata_columns[1].text_input("후보 결함 키 (선택)")
        jira_key = metadata_columns[2].text_input("Jira Key (선택)")
        submitted = st.form_submit_button("결함 등록", type="primary", icon=":material/add_circle:")

    if submitted:
        try:
            defect = create_voc_defect(
                title=title,
                severity=severity,
                category=category,
                description=description,
                actor=actor,
                evidence_status=evidence_status,
                related_run_ids=[selected_run_id] if selected_run_id else [],
                related_case_ids=selected_cases,
                related_trace_ids=[value.strip() for value in trace_text.split(",") if value.strip()],
                candidate_key=candidate_key,
                jira_key=jira_key,
            )
        except Exception as exc:
            st.error(f"결함을 등록하지 못했습니다: {exc}")
        else:
            _clear_voc_defect_caches()
            st.session_state.voc_selected_defect_id = defect["defect_id"]
            st.success(f"결함을 등록했습니다: {defect['defect_id']}")


def _eligible_retest_runs(defect: dict) -> list[str]:
    originals = set(defect.get("related_run_ids", []))
    eligible = []
    for row in _load_voc_history_rows():
        if row.get("run_type") != "RETEST" or row.get("status") == "RUNNING":
            continue
        try:
            detail = load_voc_run_history_detail(row["run_id"])
        except Exception:
            continue
        parent = detail.get("manifest", {}).get("run_metadata", {}).get("parent_run_id")
        if parent in originals:
            eligible.append(row["run_id"])
    return eligible


def _change_defect_status(defect_id: str, target: str, actor: str, comment: str, fields: dict):
    try:
        transition_voc_defect(
            defect_id, target_status=target, actor=actor, comment=comment, fields=fields
        )
    except Exception as exc:
        st.error(f"상태를 변경하지 못했습니다: {exc}")
        return
    _clear_voc_defect_caches()
    st.success(f"결함 상태를 {_defect_status_label(target)}로 변경했습니다.")
    st.rerun()


def _render_isolated_fault_tests():
    st.markdown("### 격리 장애시험")
    st.caption("운영 Agent와 실제 키를 변경하지 않는 격리 모드로 6개 장애를 재현합니다.")
    scenarios = [
        ("FT-01", "Retriever 종료"), ("FT-02", "포트 충돌"), ("FT-03", "CSV 파일 누락"),
        ("FT-04", "API 키 오류"), ("FT-05", "응답 지연"), ("FT-06", "빈 검색 결과"),
    ]
    st.dataframe(pd.DataFrame(scenarios, columns=["ID", "장애 상황"]), hide_index=True, width="stretch")
    if st.button("장애 진단 6종 실행", type="primary"):
        _run_and_store(run_diagnostics, "fault")
    _show_command_result()
    reports = list_reports("장애 진단 Fault")
    latest = next((item for item in reports if item["name"] == "latest.md"), None)
    if latest:
        st.markdown(read_report(latest["path"]))


def _render_defect_transition(defect: dict):
    defect_id = defect["defect_id"]
    status = defect.get("status")
    st.markdown("#### 다음 상태 처리")

    if status == "OPEN":
        with st.form(f"analyze_{defect_id}", border=True):
            root_cause = st.text_area("원인 분석")
            impact = st.text_area("영향 범위")
            evidence_status = st.selectbox(
                "증적 상태", ["PENDING", "CONFIRMED"],
                index=0 if defect.get("evidence_status") == "PENDING" else 1,
            )
            actor = st.text_input("처리자", value="QA")
            comment = st.text_input("처리 의견", value="원인 및 영향 분석 완료")
            submitted = st.form_submit_button("분석 완료 처리", type="primary")
        if submitted:
            _change_defect_status(
                defect_id, "ANALYZED", actor, comment,
                {"root_cause": root_cause, "impact": impact, "evidence_status": evidence_status},
            )
        return

    if status == "ANALYZED":
        with st.form(f"fix_{defect_id}", border=True):
            corrective_action = st.text_area("조치 내용")
            columns = st.columns(2)
            owner = columns[0].text_input("담당자", value="QA")
            due_date = columns[1].date_input("조치 기한")
            actor = st.text_input("처리자", value="QA")
            comment = st.text_input("처리 의견", value="조치 반영 완료")
            submitted = st.form_submit_button("조치 완료 처리", type="primary")
        if submitted:
            _change_defect_status(
                defect_id, "FIXED", actor, comment,
                {"corrective_action": corrective_action, "owner": owner, "due_date": due_date.isoformat()},
            )
        return

    if status == "FIXED":
        related_cases = defect.get("related_case_ids", [])
        original_runs = []
        for run_id in defect.get("related_run_ids", []):
            try:
                detail = load_voc_run_history_detail(run_id)
            except Exception:
                continue
            if detail.get("manifest", {}).get("run_type") != "RETEST":
                original_runs.append(run_id)

        if original_runs and related_cases:
            st.info("원본 Run과 동일한 Case로 연결된 재시험을 실행한 뒤 통과 결과를 선택하세요.")
            parent_run_id = st.selectbox(
                "재시험 기준 원본 Run", original_runs, key=f"retest_parent_{defect_id}"
            )
            active_state = _active_batch_run_state()
            active_run_id = active_state["run_id"]
            is_running = active_state["active"]
            if st.button(
                f"관련 Case {len(related_cases)}건 재시험 시작",
                type="primary", icon=":material/replay:", disabled=is_running,
                key=f"start_retest_{defect_id}",
            ):
                _launch_batch(
                    related_cases, parent_run_id=parent_run_id, judge_config={"enabled": False}
                )
                st.rerun()
            if active_run_id:
                _live_batch_progress()
        else:
            st.warning("재시험을 만들려면 결함에 원본 Run과 관련 Case가 모두 연결되어 있어야 합니다.")

        retest_runs = _eligible_retest_runs(defect)
        if not retest_runs:
            st.caption("연결 가능한 완료 재시험이 없습니다. 재시험을 먼저 실행하세요.")
            return
        with st.form(f"retested_{defect_id}", border=True):
            retest_run_id = st.selectbox("통과 재시험 Run", retest_runs)
            actor = st.text_input("처리자", value="QA")
            comment = st.text_input("처리 의견", value="연결 재시험 결과 확인")
            submitted = st.form_submit_button("재시험 완료 처리", type="primary")
        if submitted:
            _change_defect_status(
                defect_id, "RETESTED", actor, comment, {"retest_run_id": retest_run_id}
            )
        return

    if status == "RETESTED":
        with st.form(f"close_{defect_id}", border=True):
            closure_comment = st.text_area("종료 근거")
            actor = st.text_input("처리자", value="QA")
            comment = st.text_input("처리 의견", value="통과 재시험 증적 확인 후 종료")
            submitted = st.form_submit_button("결함 종료", type="primary")
        if submitted:
            _change_defect_status(
                defect_id, "CLOSED", actor, comment, {"closure_comment": closure_comment}
            )
        return

    st.success("통과 재시험 증적을 근거로 종료된 결함입니다.")


def _render_defect_list():
    defects = _load_voc_defect_rows()
    if not defects:
        st.info("등록된 결함이 없습니다. 신규 등록에서 첫 결함을 등록하세요.")
        return

    columns = st.columns(4)
    columns[0].metric("전체", len(defects))
    columns[1].metric("미종료", sum(item.get("status") != "CLOSED" for item in defects))
    columns[2].metric(
        "미종료 중요 결함",
        sum(item.get("status") != "CLOSED" and item.get("severity") in {"CRITICAL", "HIGH"} for item in defects),
    )
    columns[3].metric("종료", sum(item.get("status") == "CLOSED" for item in defects))

    filters = st.columns(2)
    status_filter = filters[0].multiselect(
        "상태", ["OPEN", "ANALYZED", "FIXED", "RETESTED", "CLOSED"]
    )
    severity_filter = filters[1].multiselect(
        "심각도", ["CRITICAL", "HIGH", "MEDIUM", "LOW"]
    )
    filtered = [
        item for item in defects
        if (not status_filter or item.get("status") in status_filter)
        and (not severity_filter or item.get("severity") in severity_filter)
    ]
    rows = pd.DataFrame([
        {
            "결함 ID": item.get("defect_id"), "제목": item.get("title"),
            "심각도": item.get("severity"), "상태": _defect_status_label(item.get("status")),
            "증적": item.get("evidence_status"), "담당자": item.get("owner") or "-",
            "갱신 시각": item.get("updated_at"),
        }
        for item in filtered
    ])
    st.dataframe(rows, hide_index=True, width="stretch")
    if not filtered:
        st.info("선택한 조건에 해당하는 결함이 없습니다.")
        return

    ids = [item["defect_id"] for item in filtered]
    remembered = st.session_state.get("voc_selected_defect_id")
    selected_id = st.selectbox(
        "상세 조회 결함", ids, index=ids.index(remembered) if remembered in ids else 0,
        format_func=lambda value: f"{value} · {next(item['title'] for item in filtered if item['defect_id'] == value)}",
    )
    st.session_state.voc_selected_defect_id = selected_id
    defect = load_voc_defect(selected_id)

    st.markdown(f"### {defect['title']}")
    with st.container(border=True):
        detail_columns = st.columns(4)
        detail_columns[0].metric("상태", _defect_status_label(defect.get("status")))
        detail_columns[1].metric("심각도", defect.get("severity", "-"))
        detail_columns[2].metric("증적", defect.get("evidence_status", "-"))
        detail_columns[3].metric("담당자", defect.get("owner") or "미지정")
        st.write(defect.get("description") or "-")
        st.caption(
            f"Run: {', '.join(defect.get('related_run_ids', [])) or '-'}  |  "
            f"Case: {', '.join(defect.get('related_case_ids', [])) or '-'}  |  "
            f"실행 Trace: {', '.join(defect.get('related_trace_ids', [])) or '-'}"
        )
        if defect.get("jira_key"):
            st.caption(f"Jira: {defect['jira_key']}")

    tabs = st.tabs(["원인·조치", "처리 이력", "재시험 증적"])
    with tabs[0]:
        st.write({
            "원인": defect.get("root_cause") or "미분석",
            "영향": defect.get("impact") or "미분석",
            "조치": defect.get("corrective_action") or "미조치",
            "조치 기한": defect.get("due_date") or "-",
            "종료 근거": defect.get("closure_comment") or "-",
        })
    with tabs[1]:
        st.dataframe(pd.DataFrame(defect.get("history", [])), hide_index=True, width="stretch")
    with tabs[2]:
        evidence = defect.get("retest_evidence", [])
        st.json(evidence) if evidence else st.info("등록된 재시험 증적이 없습니다.")
    _render_defect_transition(defect)


def render_fault():
    mode = st.segmented_control(
        "관리 구분", ["결함 목록", "신규 등록", "격리 장애시험"],
        default="결함 목록", key="voc_defect_mode",
    )
    if mode == "신규 등록":
        _render_defect_create()
    elif mode == "격리 장애시험":
        _render_isolated_fault_tests()
    else:
        _render_defect_list()


def render_a2a():
    summary = audit_summary()
    cols = st.columns(4)
    cols[0].metric("실행 Trace", summary["traces"], border=True)
    cols[1].metric("이벤트", summary["events"], border=True)
    cols[2].metric("성공", summary["success"], border=True)
    cols[3].metric("실패", summary["failure"], border=True)
    with st.container(horizontal=True, horizontal_alignment="distribute", vertical_alignment="center"):
        st.caption(f"원시 감사 로그: {summary['path']}")
        create_report = st.button(
            "Agent 파이프라인 Report 생성",
            type="primary",
            icon=":material/description:",
        )
    if create_report:
        _run_and_store(run_diagnostics, "a2a")
    _show_command_result()
    reports = list_reports("Agent 연결 A2A")
    latest = next((item for item in reports if item["name"] == "latest.md"), None)
    if latest:
        st.markdown(read_report(latest["path"]))
    elif not summary["exists"]:
        st.info("아직 실행 Trace가 없습니다. 실제 VOC 요청을 처리한 뒤 Report를 생성하세요.")


LEGACY_REPORT_GUIDE = {
    "종합 Summary": {
        "purpose": "실행 환경과 진단 결과를 한 번에 요약",
        "when": "전체 점검 결과를 빠르게 확인할 때",
        "icon": "summarize",
    },
    "정의 검증 Validation": {
        "purpose": "데이터 계약·정의·필수 파일 검증",
        "when": "스키마나 기준 파일 누락을 확인할 때",
        "icon": "fact_check",
    },
    "장애 진단 Fault": {
        "purpose": "장애 주입·결함 후보 진단 결과",
        "when": "실패 원인과 결함 후보를 추적할 때",
        "icon": "bug_report",
    },
    "Agent 연결 A2A": {
        "purpose": "Agent 간 호출 실행 Trace와 연결 상태",
        "when": "Agent 파이프라인 흐름이 정상 연결됐는지 볼 때",
        "icon": "hub",
    },
    "VOC 분석 결과": {
        "purpose": "VOC 분석 실행 산출물 원본",
        "when": "VOC 분석 결과 원문을 확인할 때",
        "icon": "psychology",
    },
}


def _format_legacy_report_size(size: int | float | None) -> str:
    value = float(size or 0)
    if value >= 1024 * 1024:
        return f"{value / (1024 * 1024):,.1f} MB"
    if value >= 1024:
        return f"{value / 1024:,.1f} KB"
    return f"{int(value):,} B"


def _format_legacy_report_time(timestamp: int | float | None) -> str:
    if not timestamp:
        return "-"
    try:
        return datetime.fromtimestamp(float(timestamp)).strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return "-"


def _legacy_report_category_rows() -> list[dict]:
    rows = []
    for category in REPORT_CATEGORIES:
        reports = list_reports(category)
        guide = LEGACY_REPORT_GUIDE.get(category, {})
        latest = reports[0] if reports else {}
        rows.append(
            {
                "분류": category,
                "파일 수": f"{len(reports)}건",
                "최근 파일": latest.get("name") or "-",
                "최근 생성": _format_legacy_report_time(latest.get("modified")),
                "용도": guide.get("purpose", "진단 산출물 원본 확인"),
            }
        )
    return rows


def _render_legacy_json_summary(content: str) -> None:
    try:
        payload = json.loads(content)
    except Exception:
        st.code(content, language="json")
        return

    if isinstance(payload, dict):
        summary_rows = [
            {
                "항목": str(key),
                "값 유형": type(value).__name__,
                "요약": (
                    f"{len(value)}개 항목"
                    if isinstance(value, (dict, list))
                    else str(value)[:120] or "-"
                ),
            }
            for key, value in payload.items()
        ]
        st.dataframe(
            pd.DataFrame(summary_rows),
            hide_index=True,
            width="stretch",
            height=min(300, 42 + max(len(summary_rows), 1) * 36),
        )
    elif isinstance(payload, list):
        st.caption(f"목록형 JSON · {len(payload)}건")
        if payload and isinstance(payload[0], dict):
            st.dataframe(pd.DataFrame(payload[:50]), hide_index=True, width="stretch", height=300)
        else:
            st.json(payload[:50], expanded=False)
    else:
        st.json(payload, expanded=False)

    with st.expander("원본 JSON 보기", icon=":material/data_object:"):
        st.json(payload, expanded=False)


@st.dialog("운영 진단 산출물", width="large", icon=":material/folder_open:")
def _render_legacy_report_dialog(category: str, report: dict):
    content = read_report(report["path"])
    guide = LEGACY_REPORT_GUIDE.get(category, {})
    header_col, download_col = st.columns([4, 1], gap="small", vertical_alignment="center")
    with header_col:
        st.caption(f"{category} · {report.get('name', '-')}")
        st.markdown(f"#### :material/{guide.get('icon', 'description')}: {guide.get('purpose', '진단 산출물 원본')}")
        st.caption(f"최근 생성 {_format_legacy_report_time(report.get('modified'))} · {_format_legacy_report_size(report.get('size'))}")
    with download_col:
        st.download_button(
            "원본 D/L",
            data=content.encode("utf-8"),
            file_name=report["name"],
            mime="application/json" if str(report["name"]).endswith(".json") else "text/markdown",
            icon=":material/download:",
            width="stretch",
            on_click="ignore",
            key=f"download_legacy_report_{category}_{report.get('name', '-')}",
        )

    st.info(
        "이 산출물은 최종 제출용 품질 보고서가 아니라, 예전 진단 스크립트가 남긴 운영·개발 점검용 원본입니다.",
        icon=":material/info:",
    )
    if str(report["name"]).endswith(".json"):
        _render_legacy_json_summary(content)
    else:
        with st.container(border=True, height=460):
            st.markdown(content)


def _render_legacy_reports():
    with st.expander("운영 진단 산출물 보기", icon=":material/folder_open:", expanded=False):
        st.caption(
            "기존 진단 보고서는 최종 품질 보고서가 아니라 Summary·Validation·Fault·Agent 파이프라인·VOC 분석 원본을 확인하는 보조 영역입니다."
        )
        category_rows = _legacy_report_category_rows()
        st.dataframe(
            pd.DataFrame(category_rows),
            hide_index=True,
            width="stretch",
            height=220,
            column_config={
                "분류": st.column_config.TextColumn("분류", width="small"),
                "파일 수": st.column_config.TextColumn("파일 수", width="small"),
                "최근 파일": st.column_config.TextColumn("최근 파일", width="medium"),
                "최근 생성": st.column_config.TextColumn("최근 생성", width="medium"),
                "용도": st.column_config.TextColumn("용도", width="large"),
            },
        )

        category_options = list(REPORT_CATEGORIES)
        if st.session_state.get("legacy_report_category") not in category_options:
            st.session_state.legacy_report_category = category_options[0]
        selected_category = st.segmented_control(
            "진단 산출물 분류",
            category_options,
            required=True,
            width="stretch",
            key="legacy_report_category",
        )
        if not selected_category:
            return

        reports = list_reports(str(selected_category))
        if not reports:
            st.info("이 분류의 산출물이 없습니다. 해당 진단을 먼저 실행하면 여기에 표시됩니다.", icon=":material/info:")
            return

        report_labels = [
            f"{item['name']} · {_format_legacy_report_time(item.get('modified'))} · {_format_legacy_report_size(item.get('size'))}"
            for item in reports
        ]
        selected_label = st.selectbox(
            "산출물 파일",
            report_labels,
            index=0,
            key=f"legacy_report_file_{selected_category}",
        )
        selected_report = reports[report_labels.index(selected_label)]
        meta_col, action_col = st.columns([4, 1], gap="small", vertical_alignment="center")
        with meta_col:
            guide = LEGACY_REPORT_GUIDE.get(str(selected_category), {})
            st.markdown(f"**{selected_report['name']}**")
            st.caption(f"{guide.get('when', '원본 진단 산출물 확인')} · {_format_legacy_report_size(selected_report.get('size'))}")
        with action_col:
            if st.button(
                "원본 보기",
                icon=":material/visibility:",
                width="stretch",
                key=f"open_legacy_report_{selected_category}_{selected_report['name']}",
            ):
                _render_legacy_report_dialog(str(selected_category), selected_report)


@st.cache_data(ttl=3, max_entries=20, show_spinner=False)
def _load_voc_quality_report_model(run_id: str, baseline_run_id: str):
    return build_voc_quality_report(run_id, baseline_run_id)


def _report_download_basename(model: dict) -> str:
    raw = str(model.get("report_id") or model.get("run", {}).get("run_id") or "voc_quality_report")
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", raw).strip("._")
    return safe or "voc_quality_report"


def _report_overview_rows(model: dict) -> list[tuple[str, str]]:
    run = model.get("run", {}) if isinstance(model.get("run"), dict) else {}
    counts = run.get("counts", {}) if isinstance(run.get("counts"), dict) else {}
    evaluation = model.get("evaluation", {}) if isinstance(model.get("evaluation"), dict) else {}
    integrity = model.get("integrity", {}) if isinstance(model.get("integrity"), dict) else {}
    release_scope = model.get("release_scope", {}) if isinstance(model.get("release_scope"), dict) else {}
    return [
        ("보고서 ID", str(model.get("report_id", "-"))),
        ("보고서 상태", _voc_status_label(model.get("report_state", "-"))),
        ("최종 판정", _voc_status_label(model.get("release_decision", "-"))),
        ("생성 일시", str(model.get("generated_at", "-"))),
        ("Run ID", str(run.get("run_id", "-"))),
        ("실행 유형", _voc_status_label(run.get("run_type", "-"))),
        ("테스트 묶음", str(run.get("suite_id") or "-")),
        ("카탈로그 버전", str(run.get("catalog_version") or "-")),
        ("대상 Case", f"{int(run.get('selected_count') or 0)}건"),
        ("인수 기준", str(release_scope.get("basis") or "실행 가능 Case PASS + 후속 구현 Case 승인")),
        ("실행 가능/후속", f"{int(release_scope.get('executable_count') or 0)}건 / {int(release_scope.get('pending_count') or 0)}건"),
        ("내부 판정", " · ".join(f"{_voc_status_label(key)} {int(value or 0)}건" for key, value in counts.items()) or "-"),
        ("독립 LLM 평가", " · ".join(f"{_voc_status_label(key)} {int(value or 0)}건" for key, value in (evaluation.get("judge_counts") or {}).items()) or "-"),
        ("개선안 타당성 평가", " · ".join(f"{_voc_status_label(key)} {int(value or 0)}건" for key, value in (evaluation.get("validity_counts") or {}).items()) or "-"),
        ("증적 무결성", "정상" if integrity.get("ok") else "확인 필요"),
    ]


def _report_stage_rows(model: dict) -> list[tuple[str, str, str, str]]:
    run = model.get("run", {}) if isinstance(model.get("run"), dict) else {}
    evaluation = model.get("evaluation", {}) if isinstance(model.get("evaluation"), dict) else {}
    release_scope = model.get("release_scope", {}) if isinstance(model.get("release_scope"), dict) else {}
    has_release_scope = bool(release_scope)
    total = int(run.get("selected_count") or 0)
    judge_required_total = int(release_scope.get("judge_required_count") or release_scope.get("voc_count") or total)
    validity_required_total = int(release_scope.get("validity_required_count") or release_scope.get("voc_count") or total)
    fault_counts = release_scope.get("fault_counts", {}) if isinstance(release_scope.get("fault_counts"), dict) else {}
    fault_total = int(release_scope.get("fault_count") or 0)
    fault_confirmed = int(fault_counts.get("PASS") or 0) + int(fault_counts.get("REVIEW_REQUIRED") or 0)
    rows = [
        ("VOC 분석 및 개선안", str(len(evaluation.get("voc_examples", []) or [])), str(total), "대표 산출물·개선안 연결"),
        ("6개 에이전트 내부 진단", str(int(evaluation.get("trace_cases", 0) or 0)), str(total), f"추적 이벤트 {int(evaluation.get('trace_events', 0) or 0)}건"),
        (
            "독립 LLM 평가",
            str(int((release_scope.get("executable_judge_counts") or {}).get("PASS", 0) or 0)) if has_release_scope else str(int(evaluation.get("judge_evaluated", 0) or 0)),
            str(judge_required_total if has_release_scope else total),
            "VOC 개선 Case 기준 PASS 집계" if has_release_scope else " · ".join(f"{_voc_status_label(key)} {int(value or 0)}건" for key, value in (evaluation.get("judge_counts") or {}).items()) or "-",
        ),
        (
            "개선안 타당성 평가",
            str(int((release_scope.get("executable_validity_counts") or {}).get("BUSINESS_APPROVED", 0) or 0)) if has_release_scope else str(int(evaluation.get("validity_evaluated", 0) or 0)),
            str(validity_required_total if has_release_scope else total),
            "VOC 개선 Case 기준 업무 승인 집계" if has_release_scope else " · ".join(f"{_voc_status_label(key)} {int(value or 0)}건" for key, value in (evaluation.get("validity_counts") or {}).items()) or "-",
        ),
    ]
    if has_release_scope:
        rows.append(("장애 검증 실행", str(fault_confirmed), str(fault_total), "장애 보호 동작 실행 확인"))
    return rows


def _report_claim_rows(model: dict) -> list[tuple[str, str, str]]:
    release_scope = model.get("release_scope", {}) if isinstance(model.get("release_scope"), dict) else {}
    voc_counts = release_scope.get("voc_counts", {}) if isinstance(release_scope.get("voc_counts"), dict) else {}
    fault_counts = release_scope.get("fault_counts", {}) if isinstance(release_scope.get("fault_counts"), dict) else {}
    pending_counts = release_scope.get("pending_counts", {}) if isinstance(release_scope.get("pending_counts"), dict) else {}
    voc_total = int(release_scope.get("voc_count") or 0)
    fault_total = int(release_scope.get("fault_count") or 0)
    fault_confirmed = int(fault_counts.get("PASS") or 0) + int(fault_counts.get("REVIEW_REQUIRED") or 0)
    pending_total = int(release_scope.get("pending_count") or 0)
    rows = [
        ("VOC 개선 Case", "충족" if release_scope.get("voc_pass_ready", release_scope.get("executable_pass_ready")) else "확인 필요", f"PASS {int(voc_counts.get('PASS') or 0)}/{voc_total}건"),
        ("장애 검증 Case", "충족" if release_scope.get("fault_execution_ready", fault_total == 0) else "확인 필요", f"실행 확인 {fault_confirmed}/{fault_total}건"),
        ("후속 구현 Case", "승인" if release_scope.get("pending_plan_approved") else "확인 필요", f"NOT_RUN {int(pending_counts.get('NOT_RUN') or 0)}/{pending_total}건 · 후속 구현 계획 기준"),
    ]
    linked_retest_count = int(release_scope.get("linked_retest_count") or 0)
    if linked_retest_count:
        linked_labels = [
            f"{item.get('case_id', '-')} → {item.get('retest_run_id', '-')}"
            for item in (release_scope.get("linked_retest_evidence") or [])[:2]
            if isinstance(item, dict)
        ]
        rows.append(("연결 RETEST", "반영", f"{linked_retest_count}건 · {', '.join(linked_labels) or '보완 재시험 반영'}"))
    rows.append(
        ("최종 인수 범위", "검증 완료" if release_scope.get("release_scope_ready") else "미검증", str(release_scope.get("basis") or "VOC 개선 Case PASS·승인 + 장애 검증 실행 확인 + 후속 구현 Case 승인"))
    )
    return rows


def _report_table_rows(items: list[dict], columns: list[tuple[str, str]], *, empty_text: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for item in items:
        rows.append([str(item.get(key) if item.get(key) not in (None, "") else "-") for key, _ in columns])
    return rows or [[empty_text] + [""] * (len(columns) - 1)]


def _pdf_font_name() -> str:
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        font_name = "MalgunGothic"
        try:
            pdfmetrics.getFont(font_name)
            return font_name
        except Exception:
            pass
        for font_path in (
            Path("C:/Windows/Fonts/malgun.ttf"),
            Path("C:/Windows/Fonts/malgunbd.ttf"),
        ):
            if font_path.exists():
                pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
                return font_name
    except Exception:
        pass
    return "Helvetica"


def _build_voc_quality_report_pdf_bytes(model: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=34, rightMargin=34, topMargin=34, bottomMargin=34)
    font_name = _pdf_font_name()
    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = font_name
        if hasattr(style, "leading"):
            style.leading = max(style.leading, style.fontSize + 4)
    styles["Title"].fontSize = 18
    styles["Heading2"].textColor = colors.HexColor("#0C4B7E")

    def paragraph(value: object, style_name: str = "BodyText") -> Paragraph:
        return Paragraph(escape(str(value if value not in (None, "") else "-")).replace("\n", "<br/>"), styles[style_name])

    def table(headers: list[str], rows: list[list[object]], widths: list[int] | None = None) -> Table:
        data = [[paragraph(header, "BodyText") for header in headers]]
        data.extend([[paragraph(cell, "BodyText") for cell in row] for row in rows])
        result = Table(data, colWidths=widths, repeatRows=1)
        result.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF3FC")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0C4B7E")),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8D9EE")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return result

    defects = model.get("defects", []) if isinstance(model.get("defects"), list) else []
    risks = model.get("risks", []) if isinstance(model.get("risks"), list) else []
    story = [
        paragraph("VOC 품질진단 결과 보고서", "Title"),
        paragraph("VOC 개선안 생성부터 에이전트 내부 진단, 독립 LLM 평가, 업무 승인까지 연결된 품질 증적입니다."),
        Spacer(1, 10),
        paragraph("1. 보고서 요약", "Heading2"),
        table(["항목", "내용"], [[label, value] for label, value in _report_overview_rows(model)], [120, 380]),
        Spacer(1, 10),
        paragraph("2. 품질 평가 단계", "Heading2"),
        table(["단계", "증적", "대상", "근거"], _report_stage_rows(model), [155, 55, 55, 235]),
        Spacer(1, 10),
        paragraph("3. 개선 추이", "Heading2"),
        table(["구분", "상태", "판정 근거"], _report_claim_rows(model), [95, 80, 325]),
        Spacer(1, 10),
        paragraph("4. 결함 요약", "Heading2"),
        table(
            ["결함 ID", "제목", "심각도", "상태", "증적"],
            _report_table_rows(
                defects,
                [("defect_id", "결함 ID"), ("title", "제목"), ("severity", "심각도"), ("status", "상태"), ("evidence_status", "증적")],
                empty_text="등록된 결함이 없습니다.",
            ),
            [85, 190, 65, 65, 95],
        ),
        Spacer(1, 10),
        paragraph("5. 잔여 위험", "Heading2"),
        table(
            ["등급", "위험", "권고 조치"],
            _report_table_rows(
                risks,
                [("level", "등급"), ("risk", "위험"), ("action", "권고 조치")],
                empty_text="현재 집계된 잔여 위험이 없습니다.",
            ),
            [70, 215, 215],
        ),
    ]
    doc.build(story)
    return buffer.getvalue()


def _set_docx_font(document) -> None:
    try:
        from docx.oxml.ns import qn

        for style in document.styles:
            if getattr(style, "font", None):
                style.font.name = "맑은 고딕"
                if getattr(style, "_element", None) is not None and getattr(style._element, "rPr", None) is not None:
                    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    except Exception:
        pass


def _build_voc_quality_report_docx_bytes(model: dict) -> bytes:
    from docx import Document

    document = Document()
    _set_docx_font(document)
    document.add_heading("VOC 품질진단 결과 보고서", level=0)
    document.add_paragraph("VOC 개선안 생성부터 에이전트 내부 진단, 독립 LLM 평가, 업무 승인까지 연결된 품질 증적입니다.")

    def add_table(title: str, headers: list[str], rows: list[list[object]]) -> None:
        document.add_heading(title, level=1)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = str(header)
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = str(value if value not in (None, "") else "-")

    defects = model.get("defects", []) if isinstance(model.get("defects"), list) else []
    risks = model.get("risks", []) if isinstance(model.get("risks"), list) else []
    add_table("1. 보고서 요약", ["항목", "내용"], [[label, value] for label, value in _report_overview_rows(model)])
    add_table("2. 품질 평가 단계", ["단계", "증적", "대상", "근거"], _report_stage_rows(model))
    add_table("3. 개선 추이", ["구분", "상태", "판정 근거"], _report_claim_rows(model))
    add_table(
        "4. 결함 요약",
        ["결함 ID", "제목", "심각도", "상태", "증적"],
        _report_table_rows(
            defects,
            [("defect_id", "결함 ID"), ("title", "제목"), ("severity", "심각도"), ("status", "상태"), ("evidence_status", "증적")],
            empty_text="등록된 결함이 없습니다.",
        ),
    )
    add_table(
        "5. 잔여 위험",
        ["등급", "위험", "권고 조치"],
        _report_table_rows(
            risks,
            [("level", "등급"), ("risk", "위험"), ("action", "권고 조치")],
            empty_text="현재 집계된 잔여 위험이 없습니다.",
        ),
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@st.dialog("VOC 품질진단 결과 보고서", width="large", icon=":material/article:")
def _render_quality_report_dialog(model: dict):
    run = model.get("run", {}) if isinstance(model.get("run"), dict) else {}
    report_html = build_voc_quality_report_html(model)
    basename = _report_download_basename(model)
    meta_col, html_col, pdf_col, word_col = st.columns(
        [4.2, 0.78, 0.78, 0.78],
        gap="small",
        vertical_alignment="center",
    )
    with meta_col:
        st.caption(
            f"{model.get('report_id', '-')} · "
            f"Run {run.get('run_id', '-')} · "
            f"{_voc_status_label(model.get('report_state', '-'))}"
        )
    with html_col:
        st.download_button(
            "HTML",
            data=report_html.encode("utf-8"),
            file_name=f"{basename}.html",
            mime="text/html",
            icon=":material/download:",
            width="stretch",
            on_click="ignore",
            key=f"download_report_html_{basename}",
        )
    with pdf_col:
        st.download_button(
            "PDF",
            data=_build_voc_quality_report_pdf_bytes(model),
            file_name=f"{basename}.pdf",
            mime="application/pdf",
            icon=":material/download:",
            width="stretch",
            on_click="ignore",
            key=f"download_report_pdf_{basename}",
        )
    with word_col:
        st.download_button(
            "WORD",
            data=_build_voc_quality_report_docx_bytes(model),
            file_name=f"{basename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            icon=":material/download:",
            width="stretch",
            on_click="ignore",
            key=f"download_report_word_{basename}",
        )
    st.markdown(report_html, unsafe_allow_html=True)


def _render_quality_report_preview(model: dict):
    report_state = _voc_status_label(model.get("report_state", "-"))
    release_decision = _voc_status_label(model.get("release_decision", "-"))
    run = model.get("run", {}) if isinstance(model.get("run"), dict) else {}
    release_scope = model.get("release_scope", {}) if isinstance(model.get("release_scope"), dict) else {}
    if model["release_decision"] == "FORMAL_APPROVED":
        st.success("모든 정식 품질 승인 조건을 충족했습니다.", icon=":material/verified:")
    else:
        st.warning(
            "현재 보고서는 증적 초안입니다. 검증되지 않은 수치와 충족하지 못한 승인 조건을 성공으로 표시하지 않습니다.",
            icon=":material/release_alert:",
        )

    with st.container(border=True):
        heading, status_col, action_col = st.columns(
            [2.2, 1.1, 1],
            gap="small",
            vertical_alignment="center",
        )
        with heading:
            st.markdown("#### VOC 품질진단 결과 보고서")
            st.caption("보고서 본문은 가독성을 위해 팝업에서 확인합니다.")
        with status_col:
            st.markdown(
                f":blue-badge[{report_state}] :{'green' if model.get('release_decision') == 'FORMAL_APPROVED' else 'orange'}-badge[{release_decision}]",
                text_alignment="right",
            )
            st.caption(
                f"Run {run.get('run_id', '-')}",
                text_alignment="right",
            )
        with action_col:
            if st.button(
                "보고서 팝업 보기",
                icon=":material/open_in_new:",
                type="primary",
                width="stretch",
                key=f"open_voc_quality_report_dialog_{model.get('report_id', '-')}_{run.get('run_id', '-')}",
            ):
                _render_quality_report_dialog(model)

    _render_release_scope_basis(
        release_scope,
        release_decision=str(model.get("release_decision") or ""),
        title="품질 보고서 판정 기준",
        caption="보고서는 35건 전체를 단순 PASS 수가 아니라 인수 범위별 승인 기준으로 판단합니다.",
    )

    with st.expander("판정 산식 원문", icon=":material/calculate:"):
        st.json(model["formula"], expanded=False)


REPORT_MODE_FINAL = "최종 품질 보고서"
REPORT_MODE_DRAFT = "증적 초안"


def _report_candidate_summary_by_run() -> dict[str, dict]:
    summary: dict[str, dict] = {}
    for candidate in _load_validity_candidates():
        run_id = str(candidate.get("run_id") or "")
        if not run_id:
            continue
        bucket = summary.setdefault(
            run_id,
            {
                "candidate_count": 0,
                "approved_count": 0,
                "judge_missing_count": 0,
                "validity_missing_count": 0,
                "qa_waiting_count": 0,
                "business_waiting_count": 0,
            },
        )
        bucket["candidate_count"] += 1
        workflow_state = str(candidate.get("workflow_state") or "DRAFT")
        if candidate.get("formal_approval") or workflow_state == "BUSINESS_APPROVED":
            bucket["approved_count"] += 1
        judge_status = str(candidate.get("judge_status") or "NOT_RUN")
        validity_status = str(candidate.get("validity_status") or "NOT_RUN")
        if judge_status in {"", "NOT_RUN", "ERROR"}:
            bucket["judge_missing_count"] += 1
        if validity_status in {"", "NOT_RUN", "ERROR"}:
            bucket["validity_missing_count"] += 1
        if candidate.get("qa_review_ready"):
            bucket["qa_waiting_count"] += 1
        if candidate.get("business_review_ready"):
            bucket["business_waiting_count"] += 1
    return summary


def _report_run_is_final(row: dict, candidate_summary: dict[str, dict]) -> bool:
    run_id = str(row.get("run_id") or "")
    selected_count = int(row.get("selected_count") or 0)
    scope = row.get("verification_scope", {}) if isinstance(row.get("verification_scope"), dict) else {}
    catalog_total = int(scope.get("catalog_total_cases") or 35)
    executable_count = int(scope.get("executable_count") or selected_count)
    approved_count = int(candidate_summary.get(run_id, {}).get("approved_count", 0) or 0)
    if str(row.get("status") or "") != "COMPLETED" or selected_count != catalog_total:
        return False
    if str(row.get("deployment_decision") or "") == "FORMAL_QUALITY_APPROVED":
        return True
    try:
        model = _load_voc_quality_report_model(run_id, "")
        if model.get("release_decision") == "FORMAL_APPROVED":
            return True
    except Exception:
        pass
    return approved_count >= executable_count


def _report_run_label(row: dict, candidate_summary: dict[str, dict], *, final_scope: bool) -> str:
    run_id = str(row.get("run_id") or "-")
    selected_count = int(row.get("selected_count") or 0)
    bucket = candidate_summary.get(run_id, {})
    approved_count = int(bucket.get("approved_count", 0) or 0)
    scope = row.get("verification_scope", {}) if isinstance(row.get("verification_scope"), dict) else {}
    executable_count = int(scope.get("executable_count") or selected_count)
    status = _voc_status_label(row.get("deployment_decision") or row.get("validity_state") or "미판정")
    if final_scope:
        return f"{run_id} · 정식 승인 완료 · 실행 가능 {executable_count}/{selected_count}건"
    return (
        f"{run_id} · 실행 가능 승인 {approved_count}/{executable_count or '-'}건 · "
        f"{status}"
    )


def _render_report_scope_cards(
    *,
    final_rows: list[dict],
    draft_rows: list[dict],
    candidate_summary: dict[str, dict],
) -> None:
    approved_case_count = sum(int(item.get("approved_count", 0) or 0) for item in candidate_summary.values())
    judge_missing_count = sum(int(item.get("judge_missing_count", 0) or 0) for item in candidate_summary.values())
    validity_missing_count = sum(int(item.get("validity_missing_count", 0) or 0) for item in candidate_summary.values())
    cards = [
        {
            "icon": "verified",
            "label": "최종 보고서 대상",
            "value": f"{len(final_rows)}건",
            "detail": "실행 가능 Case 승인 + 후속 구현 승인",
            "tone": "green" if final_rows else "gray",
            "badge": "최종",
        },
        {
            "icon": "draft",
            "label": "증적 초안",
            "value": f"{len(draft_rows)}건",
            "detail": "평가·승인 진행 중 Run",
            "tone": "blue" if draft_rows else "gray",
            "badge": "진행",
        },
        {
            "icon": "approval",
            "label": "업무 승인 Case",
            "value": f"{approved_case_count}건",
            "detail": "보고서/시연 연결 가능",
            "tone": "green" if approved_case_count else "gray",
            "badge": "승인",
        },
        {
            "icon": "pending_actions",
            "label": "남은 평가",
            "value": f"{judge_missing_count + validity_missing_count}건",
            "detail": f"독립 LLM {judge_missing_count} · 개선안 타당성 {validity_missing_count}",
            "tone": "orange" if judge_missing_count + validity_missing_count else "green",
            "badge": "확인 필요" if judge_missing_count + validity_missing_count else "없음",
        },
    ]
    _render_voc_summary_cards(cards, columns=4, height=118)


def _render_voc_quality_report(report_mode: str):
    focus_notice = st.session_state.pop("voc_report_focus_notice", None)
    if focus_notice:
        st.info(focus_notice, icon=":material/summarize:")
    history = [row for row in _load_voc_history_rows() if row.get("status") != "RUNNING"]
    if not history:
        st.info("보고서를 생성할 완료 Run이 없습니다.")
        return
    candidate_summary = _report_candidate_summary_by_run()
    final_rows = [row for row in history if _report_run_is_final(row, candidate_summary)]
    draft_rows = [row for row in history if not _report_run_is_final(row, candidate_summary)]
    _render_report_scope_cards(
        final_rows=final_rows,
        draft_rows=draft_rows,
        candidate_summary=candidate_summary,
    )

    final_scope = report_mode == REPORT_MODE_FINAL
    report_rows = final_rows if final_scope else draft_rows
    if final_scope:
        st.caption("실행 가능 Case가 통과·독립 LLM 평가 통과·업무 승인 완료되고, 후속 구현 Case가 승인된 Run만 최종 보고서 대상으로 표시합니다.")
    else:
        st.caption("승인 전 Run은 최종 보고서가 아니라 누락 평가와 보완 위치를 확인하는 증적 초안으로만 표시합니다.")

    if not report_rows:
        with st.container(border=True):
            if final_scope:
                st.markdown("#### 최종 품질 보고서 대상 Run이 아직 없습니다.")
                st.caption("35건 검증 회차 실행 → 실행 가능 Case 독립 LLM 평가 → 개선안 타당성 평가 → QA 검토 → 업무 승인 완료 후 이 목록에 표시됩니다.")
                st.markdown(":gray-badge[정식 승인 완료 Run 없음]")
                _render_approved_demo_flow_panel("report")
            else:
                st.markdown("#### 증적 초안 대상 Run이 없습니다.")
                st.caption("승인 전 완료 Run이 없거나 모든 Run이 정식 승인 완료 상태입니다.")
        return

    run_ids = [row["run_id"] for row in report_rows]
    full_suite = [row["run_id"] for row in history if row.get("selected_count") == 35]
    default_id = next((run_id for run_id in full_suite if run_id in run_ids), run_ids[0])
    if st.session_state.get("voc_report_run_id") not in run_ids:
        st.session_state.voc_report_run_id = default_id
    row_map = {row["run_id"]: row for row in report_rows}
    selected_run_id = st.selectbox(
        "보고 대상 Run",
        run_ids,
        index=None,
        format_func=lambda value: _report_run_label(
            row_map.get(value, {}),
            candidate_summary,
            final_scope=final_scope,
        ),
        key="voc_report_run_id",
    )
    if not selected_run_id:
        return
    if final_scope:
        _render_approved_demo_flow_panel(
            "report",
            focus_run_id=str(selected_run_id),
            focus_case_id=str(st.session_state.get("voc_report_focus_case_id") or ""),
        )
    else:
        bucket = candidate_summary.get(str(selected_run_id), {})
        with st.container(border=True):
            st.markdown("#### 이 Run은 아직 최종 보고서 대상이 아닙니다.")
            st.caption("아래 미완료 항목을 해소하면 최종 품질 보고서 대상으로 승격됩니다.")
            status_columns = st.columns(4, gap="small")
            status_columns[0].metric("업무 승인", f"{int(bucket.get('approved_count', 0) or 0)}건", border=True)
            status_columns[1].metric("독립 LLM 미완료", f"{int(bucket.get('judge_missing_count', 0) or 0)}건", border=True)
            status_columns[2].metric("개선안 타당성 평가 미완료", f"{int(bucket.get('validity_missing_count', 0) or 0)}건", border=True)
            status_columns[3].metric("현재 판정", _voc_status_label(row_map[selected_run_id].get("deployment_decision", "미판정")), border=True)
    baseline_options = [""] + [value for value in full_suite if value != selected_run_id]
    if st.session_state.get("voc_report_baseline_run_id") not in baseline_options:
        st.session_state.voc_report_baseline_run_id = ""
    baseline_run_id = st.selectbox(
        "33 통과 / 2 실패 기준선 실행 (선택)", baseline_options,
        index=None,
        format_func=lambda value: value or "연결하지 않음 · 현재 기준선 증적 없음",
        key="voc_report_baseline_run_id",
        help="동일한 35개 케이스·목록·테스트케이스 해시·평가 기준과 결함 연결이 확인되는 실행만 유효합니다.",
    )
    with st.expander("보고서 템플릿 안내", expanded=False, icon=":material/article:"):
        st.caption(
            "사용자 제공 최종 보고서 양식은 아직 전달되지 않아 기본 증적 템플릿을 사용합니다. "
            "양식을 받으면 같은 보고서 데이터 모델에 적용할 수 있습니다."
        )
    model = _load_voc_quality_report_model(selected_run_id, baseline_run_id)
    _render_quality_report_preview(model)
    _render_retest_formal_approval_flow(
        run_id=str(selected_run_id),
        release_scope=model.get("release_scope", {}) if isinstance(model.get("release_scope"), dict) else {},
        location="report",
    )

    if st.button(
        "TXT·XML·HTML 증적 생성", type="primary", icon=":material/description:",
        key=f"generate_voc_report_{selected_run_id}_{baseline_run_id}",
    ):
        with st.spinner("수행 이력과 증적 수치를 다시 대조하고 있습니다..."):
            st.session_state.voc_generated_quality_report = generate_voc_quality_report(
                selected_run_id, baseline_run_id
            )
        st.success("세 형식의 증적을 같은 보고서 데이터 모델에서 생성했습니다.")

    generated = st.session_state.get("voc_generated_quality_report")
    if (
        not generated
        or generated.get("model", {}).get("run", {}).get("run_id") != selected_run_id
        or generated.get("manifest", {}).get("baseline_run_id", "") != baseline_run_id
    ):
        return
    st.caption(f"저장 위치: {Path(generated['paths']['txt']).parent}")
    with st.container(horizontal=True):
        st.download_button(
            "TXT 다운로드", generated["contents"]["txt"], file_name="result.txt",
            mime="text/plain", icon=":material/download:",
        )
        st.download_button(
            "JUnit XML 다운로드", generated["contents"]["xml"], file_name="junit.xml",
            mime="application/xml", icon=":material/download:",
        )
        st.download_button(
            "HTML 다운로드", generated["contents"]["html"], file_name="report.html",
            mime="text/html", icon=":material/download:",
        )


def render_reports():
    report_modes = (REPORT_MODE_FINAL, REPORT_MODE_DRAFT)
    current_mode = st.session_state.get("voc_report_mode")
    if current_mode not in report_modes:
        st.session_state.pop("voc_report_mode", None)
        current_mode = REPORT_MODE_FINAL
    _render_report_mode_tab_style()
    mode = st.radio(
        "보고서 구분",
        list(report_modes),
        index=list(report_modes).index(current_mode),
        horizontal=True,
        key="voc_report_mode",
        label_visibility="collapsed",
    )
    _render_voc_quality_report(str(mode or REPORT_MODE_FINAL))
    _render_legacy_reports()


def render_guide():
    with st.container(border=True):
        guide_name = st.segmented_control(
            "가이드 구분",
            ["사용자 가이드", "품질진단 실행", "이식 가이드", "이식 체크리스트"],
            default="사용자 가이드",
            key="voc_user_guide_type",
            width="stretch",
        )
    with st.container(border=True):
        st.markdown(load_guide(guide_name))


def render_acceptance():
    focus_notice = st.session_state.pop("voc_acceptance_focus_notice", None)
    if focus_notice:
        st.info(focus_notice, icon=":material/approval:")
    _render_approved_demo_flow_panel(
        "acceptance",
        focus_run_id=str(st.session_state.get("voc_acceptance_focus_run_id") or ""),
        focus_case_id=str(st.session_state.get("voc_acceptance_focus_case_id") or ""),
    )

    history = [
        row for row in _load_voc_history_rows()
        if row.get("status") == "COMPLETED" and row.get("selected_count") == 35
    ]
    if not history:
        st.warning("최종 인수 판정에 사용할 완료된 35건 Run이 없습니다.")
        return

    run_ids = [row["run_id"] for row in history]
    default_id = latest_voc_full_run_id()
    if default_id not in run_ids:
        default_id = run_ids[0]
    if st.session_state.get("voc_acceptance_run_id") not in run_ids:
        st.session_state.voc_acceptance_run_id = default_id
    run_id = st.selectbox(
        "최종 인수 대상 Run",
        run_ids,
        index=None,
        key="voc_acceptance_run_id",
        help="완료된 35건 Run의 저장 증적만 최종 품질 게이트에 사용합니다.",
    )
    if not run_id:
        return
    baseline_ids = [value for value in run_ids if value != run_id]
    baseline_options = [""] + baseline_ids
    if st.session_state.get("voc_acceptance_baseline_run_id") not in baseline_options:
        st.session_state.voc_acceptance_baseline_run_id = ""
    baseline_run_id = st.selectbox(
        "33 통과 / 2 실패 기준선 Run (선택)",
        baseline_options,
        index=None,
        format_func=lambda value: value or "연결하지 않음",
        key="voc_acceptance_baseline_run_id",
    )
    with st.spinner("Run·Case·독립 LLM 평가·개선안 타당성 평가·결함·회귀 증적을 대조하고 있습니다..."):
        snapshot = build_voc_acceptance_snapshot(run_id, baseline_run_id)

    quantitative = snapshot["quantitative"]
    release_scope = quantitative.get("release_scope", {}) if isinstance(quantitative.get("release_scope"), dict) else {}
    _render_acceptance_formal_connection(snapshot)
    _render_retest_formal_approval_flow(
        run_id=str(run_id),
        release_scope=release_scope,
        location="acceptance",
    )
    if snapshot.get("release_report_decision") == "FORMAL_APPROVED" and snapshot["decision"] == "READY_FOR_UAT":
        st.success("정식 승인 완료 결과가 최종 인수·시연 게이트까지 연결되었습니다. 사용자 UAT와 최종 서명이 남았습니다.")
    elif snapshot.get("release_report_decision") == "FORMAL_APPROVED":
        st.warning("품질 보고서는 정식 승인 완료 상태입니다. 다만 최종 인수 보조 점검 중 HOLD 항목이 남아 있습니다.")
    elif snapshot["decision"] == "READY_FOR_UAT":
        st.success("모든 자동 품질 게이트를 통과했습니다. 사용자 UAT와 최종 서명이 남았습니다.")
    else:
        st.error("현재 최종 판정은 HOLD입니다. 미충족 게이트를 보완하기 전 정식 배포할 수 없습니다.")

    _render_release_scope_basis(
        release_scope,
        release_decision=str(snapshot.get("release_report_decision") or ""),
        title="35건 최종 인수 기준",
        caption="시연 중 가장 먼저 설명할 기준입니다. PASS, 실행 확인, 후속 승인, RETEST 연결을 분리해서 봅니다.",
    )

    scope_summary = snapshot.get("release_scope_summary", {}) if isinstance(snapshot.get("release_scope_summary"), dict) else {}
    with st.container(border=True):
        _render_voc_section_heading(
            "최종 판정 요약",
            "최종 인수·시연에서 바로 확인해야 하는 자동 게이트 결과입니다.",
            icon="analytics",
            badges=((_voc_status_label(snapshot["decision"]), "green" if snapshot["decision"] == "READY_FOR_UAT" else "orange"),),
        )
        with st.container(horizontal=True):
            st.metric("인수 판정", _voc_status_label(snapshot["decision"]), border=True)
            st.metric("정식 승인", _voc_status_label(snapshot.get("release_report_decision")), border=True)
            st.metric("품질 게이트", f"{snapshot['gate_summary']['pass']}/{snapshot['gate_summary']['total']}", border=True)
            st.metric("HOLD", snapshot["gate_summary"]["hold"], border=True)
            st.metric("연결 RETEST", f"{int(scope_summary.get('linked_retest_count') or 0)}건", border=True)

    with st.container(border=True):
        _render_voc_section_heading(
            "최종 품질 게이트",
            "자동 판정에 사용된 게이트입니다. HOLD가 있으면 최종 인수 설명 전에 원인을 먼저 해소합니다.",
            icon="fact_check",
            badges=((f"PASS {snapshot['gate_summary']['pass']}/{snapshot['gate_summary']['total']}", "green"),),
        )
        st.dataframe(
            pd.DataFrame(snapshot["gates"]).rename(columns={
                "label": "완료 조건", "status": "상태", "evidence": "증적", "gate_id": "ID",
            })[["상태", "완료 조건", "증적"]],
            hide_index=True,
            width="stretch",
        )

    with st.container(border=True):
        _render_voc_section_heading(
            "핵심 업무 흐름 인수 범위",
            "시연에서 설명할 업무 흐름별 증적 준비 상태입니다.",
            icon="account_tree",
        )
        st.dataframe(
            pd.DataFrame(snapshot["workflow_coverage"]).rename(columns={
                "workflow": "업무 흐름", "status": "상태", "evidence": "증적",
            }),
            hide_index=True,
            width="stretch",
        )

    voc_counts = release_scope.get("voc_counts", {}) if isinstance(release_scope.get("voc_counts"), dict) else {}
    fault_counts = release_scope.get("fault_counts", {}) if isinstance(release_scope.get("fault_counts"), dict) else {}
    judge_counts = release_scope.get("executable_judge_counts", {}) if isinstance(release_scope.get("executable_judge_counts"), dict) else {}
    validity_counts = release_scope.get("executable_validity_counts", {}) if isinstance(release_scope.get("executable_validity_counts"), dict) else {}
    voc_total = int(release_scope.get("voc_count") or 0)
    fault_total = int(release_scope.get("fault_count") or 0)
    judge_total = int(release_scope.get("judge_required_count") or voc_total)
    validity_total = int(release_scope.get("validity_required_count") or voc_total)
    fault_confirmed = int(fault_counts.get("PASS") or 0) + int(fault_counts.get("REVIEW_REQUIRED") or 0)
    with st.container(horizontal=True):
        st.metric("VOC 개선 PASS", f"{int(voc_counts.get('PASS') or 0)}/{voc_total}", border=True)
        st.metric("장애 검증 실행", f"{fault_confirmed}/{fault_total}", border=True)
        st.metric("독립 LLM PASS", f"{int(judge_counts.get('PASS') or 0)}/{judge_total}", border=True)
        st.metric("업무 승인 완료", f"{int(validity_counts.get('BUSINESS_APPROVED') or 0)}/{validity_total}", border=True)
    st.caption(
        "비용은 현재 저장 증적에 공통 필드가 없어 확인 불가로 표시합니다. "
        "응답시간은 수행 이력의 Run·Case 시작/종료 시각으로 확인합니다."
    )

    st.markdown("### 잔여 위험과 운영 권고")
    if snapshot["remaining_risks"]:
        st.dataframe(pd.DataFrame(snapshot["remaining_risks"]), hide_index=True, width="stretch")
    else:
        st.success("저장 증적 기준으로 식별된 잔여 위험이 없습니다.")

    if st.button(
        "최종 판정 증적 저장",
        type="primary",
        icon=":material/fact_check:",
        key=f"generate_voc_acceptance_{run_id}_{baseline_run_id}",
    ):
        st.session_state.voc_acceptance_evidence = generate_voc_acceptance_evidence(snapshot)
        st.success("최종 판정 JSON·Markdown 증적을 Run evidence 경로에 저장했습니다.")
    generated = st.session_state.get("voc_acceptance_evidence")
    if generated and generated.get("snapshot", {}).get("run_id") == run_id:
        st.caption(f"저장 위치: {Path(generated['paths']['json']).parent}")
        with st.container(horizontal=True):
            st.download_button(
                "판정 JSON 다운로드", generated["contents"]["json"],
                file_name="step10_acceptance.json", mime="application/json",
                icon=":material/download:",
            )
            st.download_button(
                "판정 Markdown 다운로드", generated["contents"]["markdown"],
                file_name="step10_acceptance.md", mime="text/markdown",
                icon=":material/download:",
            )


ROUTES = {
    "Dashboard": render_dashboard,
    "수동 TC 수행": render_goal_monitor,
    "일괄 TC 수행": render_batch_execution,
    "수행 이력": render_voc_history,
    "개선안 타당성 검증": render_improvement_validity,
    "Agent 관리": render_agents,
    "VOC 분석": render_analysis,
    "테스트케이스": render_testcases,
    "품질 평가 기준": render_rubric,
    "장애·결함 관리": render_fault,
    "A2A Trace": render_a2a,
    "품질 보고서": render_reports,
    "사용자 가이드": render_guide,
    "최종 인수·시연": render_acceptance,
}


def render_voc_quality_view(sub_menu):
    renderer = ROUTES.get(sub_menu)
    if not renderer:
        return False
    try:
        _render_voc_design_system()
        _render_voc_page_header(sub_menu)
        with st.container(key="voc_page_content"):
            renderer()
    except Exception as exc:
        st.error(f"VOC 품질진단 화면을 불러오지 못했습니다: {type(exc).__name__}: {exc}")
    return True
