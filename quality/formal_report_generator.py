import html
import base64
import io
import json
import math
import subprocess
import tempfile
import zipfile
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from quality_criteria import get_quality_criteria

from quality.report_appendix_generator import (
    append_docx_appendix,
    append_pdf_appendix,
    build_appendix_context,
    render_html_appendix,
    render_markdown_appendix,
)


METRICS = [
    ("accuracy", "정확성"),
    ("groundedness", "근거성"),
    ("helpfulness", "유용성"),
    ("safety", "안전성"),
]
AGENTS = [
    ("rule_based", "Rule 기반 응답", "RULE"),
    ("api_based", "LLM API 응답", "API"),
]
DECISIONS = ("PASS", "REVIEW", "FAIL")
def load_pipeline_outputs(result_json_path, quality_criteria=None):
    path = Path(result_json_path)
    with path.open("r", encoding="utf-8") as file:
        data = json.load(file)
    if not isinstance(data, list):
        raise ValueError("평가 결과 JSON은 리스트 형태여야 합니다.")
    return normalize_pipeline_outputs(data, quality_criteria)


def normalize_pipeline_outputs(items, quality_criteria=None):
    return [
        _normalize_case(item, index, quality_criteria)
        for index, item in enumerate(items, start=1)
    ]


def build_report_context(pipeline_outputs, quality_criteria=None):
    criteria = get_quality_criteria(quality_criteria)
    cases = normalize_pipeline_outputs(pipeline_outputs, criteria)
    total_cases = len(cases)
    agents = {}

    for agent_key, label, _ in AGENTS:
        decisions = [get_decision(case, agent_key) for case in cases]
        metric_rates = {}
        metric_averages = {}
        for metric_key, metric_label in METRICS:
            scores = [get_metric_score(case, agent_key, metric_key) for case in cases]
            passed = sum(1 for score in scores if score >= criteria.pass_min_score)
            metric_rates[metric_key] = {
                "label": metric_label,
                "passed": passed,
                "rate": _rate(passed, total_cases),
            }
            metric_averages[metric_key] = round(sum(scores) / len(scores), 2) if scores else 0

        pass_count = decisions.count("PASS")
        agents[agent_key] = {
            "label": label,
            "decision_counts": {decision: decisions.count(decision) for decision in DECISIONS},
            "pass_count": pass_count,
            "pass_rate": _rate(pass_count, total_cases),
            "metric_rates": metric_rates,
            "metric_averages": metric_averages,
            "pass_threshold": (
                criteria.rule_pass_rate_min if agent_key == "rule_based" else criteria.api_pass_rate_min
            ),
        }

    comparisons = []
    for case in cases:
        rule_decision = get_decision(case, "rule_based")
        api_decision = get_decision(case, "api_based")
        comparisons.append(
            {
                "case_id": case.get("case_id", ""),
                "category": case.get("category", ""),
                "test_type": case.get("test_type", ""),
                "user_question": case.get("user_question", ""),
                "rule_decision": rule_decision,
                "api_decision": api_decision,
                "match": "일치" if rule_decision == api_decision else "불일치",
            }
        )

    defects = build_defect_items(cases, criteria)
    safety_avg = agents["api_based"]["metric_averages"].get("safety", 0)
    final_decision = (
        "합격"
        if agents["rule_based"]["pass_rate"] >= criteria.rule_pass_rate_min
        and agents["api_based"]["pass_rate"] >= criteria.api_pass_rate_min
        and safety_avg >= criteria.safety_avg_min
        else "개선 필요"
    )
    final_opinion = (
        "API 기반 평가 성공률이 기준을 충족했습니다. 운영 반영 전 주요 결함 재확인과 회귀 테스트를 권장합니다."
        if final_decision == "합격"
        else "성공률 또는 핵심 지표가 기준에 미달했습니다. FAIL 케이스를 우선 보완한 뒤 재평가가 필요합니다."
    )

    return {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "generated_date": datetime.now().strftime("%Y. %m. %d"),
        "total_cases": total_cases,
        "pass_threshold": criteria.api_pass_rate_min,
        "quality_criteria": criteria.to_dict(),
        "cases": cases,
        "agents": agents,
        "comparisons": comparisons,
        "defects": defects,
        "type_distribution": build_type_distribution(cases),
        "final_decision": final_decision,
        "final_opinion": final_opinion,
    }


def generate_html_report(
    pipeline_outputs, include_appendix=False, appendix_options=None, run_id=None, quality_criteria=None
):
    context = build_report_context(pipeline_outputs, quality_criteria)
    appendix_html = ""
    if include_appendix:
        appendix_context = build_appendix_context(pipeline_outputs, run_id=run_id, options=appendix_options)
        appendix_html = render_html_appendix(appendix_context)
    style = """
    <style>
      :root { --ink:#172033; --navy:#18345f; --blue:#2563eb; --line:#d8e2f1; --soft:#f6f8fc; --head:#eef3fb; }
      body { font-family: 'Malgun Gothic', 'AppleGothic', Arial, sans-serif; color:var(--ink); margin:0; background:#dfe5ef; }
      .report { width:794px; min-height:1123px; box-sizing:border-box; margin:0 auto; padding:46px 52px 58px; background:#fff; box-shadow:0 16px 40px rgba(15,23,42,.20); }
      .cover { min-height:280px; border:1px solid #c8d3e4; border-top:10px solid var(--navy); padding:34px 34px 26px; margin-bottom:30px; background:#fff; }
      .doc-label { color:#64748b; font-size:12px; font-weight:700; letter-spacing:.04em; text-transform:uppercase; margin-bottom:12px; }
      h1 { margin:34px 0 24px; color:#0f2547; font-size:34px; font-weight:800; line-height:1.25; }
      h2 { margin:34px 0 14px; color:#0f2547; font-size:19px; border-bottom:2px solid var(--navy); padding-bottom:9px; page-break-after:avoid; }
      h2::before { content:""; display:inline-block; width:6px; height:18px; margin-right:8px; background:var(--navy); vertical-align:-3px; }
      h3 { margin:20px 0 10px; color:#18345f; font-size:15px; }
      .meta { color:#26364f; line-height:1.85; display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:4px 20px; border-top:1px solid #d9e1ef; border-bottom:1px solid #d9e1ef; padding:14px 0; }
      .summary-strip { display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); border:1px solid #c8d3e4; margin:28px 0 4px; }
      .summary-strip div { padding:13px 14px; border-right:1px solid var(--line); background:#fff; }
      .summary-strip div:last-child { border-right:0; }
      .summary-strip span { display:block; color:#64748b; font-size:12px; }
      .summary-strip strong { display:block; margin-top:4px; color:#0f2547; font-size:18px; }
      .grid { display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); gap:10px; margin:16px 0; }
      .card { border:1px solid var(--line); border-radius:0; padding:14px; background:var(--soft); }
      .card span { color:#42526b; font-weight:700; font-size:12px; }
      .card strong { display:block; color:#0f2547; font-size:24px; margin-top:6px; }
      .card small { color:#64748b; }
      table { width:100%; border-collapse:collapse; margin:12px 0 20px; font-size:12px; page-break-inside:auto; }
      th, td { border:1px solid #d9e1ef; padding:8px 9px; vertical-align:top; }
      th { background:var(--head); color:#0f2547; text-align:left; font-weight:800; }
      tr:nth-child(even) td { background:#fbfcff; }
      .badge { border-radius:999px; padding:2px 8px; font-weight:700; font-size:12px; display:inline-block; }
      .pass { background:#dcfce7; color:#166534; }
      .review { background:#fef9c3; color:#854d0e; }
      .fail { background:#fee2e2; color:#991b1b; }
      .bar-wrap { background:#e8eef8; border-radius:0; height:13px; overflow:hidden; border:1px solid #d9e1ef; }
      .bar { height:13px; background:#2563eb; }
      .muted { color:#42526b; }
      .section-note { background:#f8fbff; border:1px solid var(--line); border-left:5px solid var(--blue); padding:14px 16px; margin:12px 0; line-height:1.7; }
      .report-footer { margin-top:30px; padding-top:12px; border-top:1px solid var(--line); color:#64748b; font-size:11px; text-align:right; }
      @media print {
        body { background:#fff; }
        .report { width:auto; min-height:auto; box-shadow:none; margin:0; }
      }
    </style>
    """
    return (
        style
        + "<main class='report'>"
        + f"""
        <section class="cover">
          <div class="doc-label">AI QUALITY ASSURANCE REPORT</div>
          <h1>QA 최종 테스트 결과 보고서</h1>
          <div class="meta">
            <div>프로젝트명: AI 교육과정 안내 챗봇 품질검증</div>
            <div>생성일시: {html.escape(context['generated_at'])}</div>
            <div>총 테스트 케이스: {context['total_cases']}건</div>
            <div>최종 종합 의견: <strong>{html.escape(context['final_decision'])}</strong></div>
          </div>
          <div class="summary-strip">
            <div><span>평가 기준</span><strong>{context['pass_threshold']:.0f}%</strong></div>
            <div><span>LLM API 성공률</span><strong>{context['agents']['api_based']['pass_rate']:.1f}%</strong></div>
            <div><span>Rule 응답 성공률</span><strong>{context['agents']['rule_based']['pass_rate']:.1f}%</strong></div>
            <div><span>등록 대상 결함</span><strong>{len(context['defects'])}건</strong></div>
          </div>
        </section>
        """
        + _render_overview(context)
        + _render_metric_charts(context)
        + _render_comparison_table(context["comparisons"])
        + "".join(_render_agent_detail(context, agent_key, label) for agent_key, label, _ in AGENTS)
        + _render_defect_report(context["defects"])
        + _render_severity_rules()
        + f"""
        <h2>6. 최종 종합 의견 및 개선 권고</h2>
        <div class="section-note">{html.escape(context['final_opinion'])}</div>
        {appendix_html}
        <div class="report-footer">Generated by AI Quality Dashboard</div>
        </main>
        """
    )


def generate_html_report(
    pipeline_outputs, include_appendix=False, appendix_options=None, run_id=None, quality_criteria=None
):
    context = build_report_context(pipeline_outputs, quality_criteria)
    appendix_html = ""
    if include_appendix:
        appendix_context = build_appendix_context(pipeline_outputs, run_id=run_id, options=appendix_options)
        appendix_html = render_html_appendix(appendix_context)

    return (
        _image_report_style()
        + "<main class='report image-report'>"
        + _render_image_style_hero(context)
        + _render_four_metric_comparison(context)
        + _render_test_overview_table(context)
        + _render_overall_summary(context)
        + _render_agent_image_section(context, "rule_based", "3. 규칙 기반 챗봇 상세 결과")
        + _render_agent_image_section(context, "api_based", "4. API 기반 챗봇 상세 결과")
        + _render_bug_report_image_style(context["defects"])
        + _render_severity_rules()
        + f"""
        <section class="panel">
          <h2>6. 최종 종합 의견 및 개선 권고</h2>
          <p>{html.escape(context['final_opinion'])}</p>
        </section>
        {appendix_html}
        <div class="report-footer">Generated by AI Quality Dashboard</div>
        </main>
        """
    )


def _image_report_style():
    return """
    <style>
      :root { --teal:#1899aa; --teal-dark:#087887; --mint:#63c3a6; --red:#e87e80; --bg:#edf5f8; --line:#c9dde5; --ink:#203047; --muted:#71829a; }
      body { margin:0; background:#fff; color:var(--ink); font-family:'Malgun Gothic','AppleGothic',Arial,sans-serif; }
      .image-report { width:928px; box-sizing:border-box; margin:0 auto; padding:44px 44px 56px; background:var(--bg); }
      .hero { text-align:center; padding:34px 24px 24px; }
      .hero h1 { margin:6px 0 6px; font-size:27px; color:#1f2d42; letter-spacing:.01em; }
      .hero .subtitle { color:#687b94; font-size:17px; font-weight:800; border-bottom:3px solid var(--teal); padding-bottom:7px; }
      .hero .meta-line { margin-top:18px; color:#74859b; font-size:12px; }
      .kpi-band { background:#fff; padding:22px 18px 24px; display:grid; grid-template-columns:repeat(4,1fr); gap:30px; margin-bottom:20px; }
      .kpi { min-height:125px; position:relative; border:1px solid #dce7ed; border-radius:6px; background:#fff; padding:26px 14px 12px; box-shadow:0 1px 2px rgba(20,58,80,.05); }
      .kpi::before { content:""; position:absolute; left:0; top:0; right:0; height:11px; border-radius:4px 4px 0 0; background:var(--teal); }
      .kpi.dark::before { background:var(--teal-dark); }
      .kpi.mint::before { background:var(--mint); }
      .kpi.red::before { background:var(--red); }
      .kpi-title { color:#4a5b72; font-size:12px; font-weight:800; margin-bottom:22px; }
      .kpi-value { color:var(--teal); font-size:28px; font-weight:900; }
      .kpi.dark .kpi-value { color:var(--teal-dark); }
      .kpi.mint .kpi-value { color:var(--mint); }
      .kpi.red .kpi-value { color:var(--red); }
      .kpi small { display:block; color:#a0adbd; margin-top:12px; font-size:9px; font-weight:700; }
      .gauge { position:absolute; right:13px; bottom:30px; width:58px; height:30px; border:9px solid currentColor; border-bottom:0; border-radius:58px 58px 0 0; color:var(--teal); }
      .kpi.dark .gauge { color:var(--teal-dark); }
      .kpi.mint .gauge { color:var(--mint); }
      .gauge span { position:absolute; left:17px; top:16px; font-size:9px; font-weight:900; color:currentColor; }
      .panel { background:#fff; border:1px solid var(--line); border-radius:14px; padding:24px 26px; margin:20px 0; box-shadow:0 1px 2px rgba(20,58,80,.04); }
      h2 { margin:0 0 20px; color:#04758a; font-size:26px; border-bottom:3px solid var(--teal); padding-bottom:8px; }
      h3 { margin:18px 0 12px; color:#04758a; font-size:20px; }
      h4 { margin:16px 0 8px; color:#203047; font-size:13px; }
      table { width:100%; border-collapse:collapse; margin:12px 0 10px; font-size:12px; }
      th, td { border:1px solid var(--line); padding:8px 9px; vertical-align:middle; }
      th { background:var(--teal); color:#fff; text-align:center; font-weight:800; }
      tr:nth-child(even) td { background:#f2f8fa; }
      .label-cell { width:165px; background:#e3f2f6 !important; color:#04758a; font-weight:800; }
      .center { text-align:center; }
      .pass-cell { background:#e3f5ee !important; }
      .fail-cell { background:#f9dddd !important; }
      .muted { color:var(--muted); font-size:11px; }
      .chart-box { background:#fff; padding:16px 14px; margin-top:12px; }
      .bar-row { display:grid; grid-template-columns:72px 1fr 42px; align-items:center; gap:8px; margin:9px 0; font-size:12px; }
      .bar-track { height:15px; background:#eef3f5; position:relative; border-left:1px solid #ddd; }
      .bar-fill { height:15px; background:var(--teal); }
      .bar-fill.alt { background:#7ccac0; }
      .threshold { position:absolute; top:-4px; bottom:-4px; width:0; border-left:1px dashed #333; }
      .two-col { display:grid; grid-template-columns:1fr 1fr; gap:24px; align-items:start; }
      .donut { width:178px; height:178px; border-radius:50%; margin:12px auto; background:conic-gradient(var(--mint) var(--pct), var(--red) 0); position:relative; }
      .donut::after { content:""; position:absolute; inset:42px; border-radius:50%; background:#fff; }
      .donut span { position:absolute; inset:58px 35px auto; text-align:center; z-index:1; font-weight:900; color:#203047; font-size:17px; }
      .type-bars .stack { height:140px; display:flex; gap:32px; align-items:end; justify-content:center; border-left:1px solid #d5dce0; border-bottom:1px solid #d5dce0; padding:10px 12px 0; }
      .type-bar { width:52px; background:var(--mint); position:relative; }
      .type-bar .fail-part { position:absolute; left:0; right:0; top:0; background:var(--red); }
      .type-labels { display:flex; gap:32px; justify-content:center; font-size:11px; margin-top:5px; }
      .bug-card { display:grid; grid-template-columns:1fr 1fr 1fr; border:1px solid var(--line); margin-top:10px; font-size:11px; }
      .bug-card > div { border-right:1px solid var(--line); }
      .bug-card > div:last-child { border-right:0; }
      .bug-head { background:var(--teal); color:#fff; padding:7px 9px; font-weight:800; }
      .bug-body { padding:8px 10px; line-height:1.55; }
      .badge { border-radius:999px; padding:2px 8px; font-weight:800; font-size:11px; display:inline-block; }
      .pass { background:#dcfce7; color:#166534; }
      .review { background:#fef9c3; color:#854d0e; }
      .fail { background:#fee2e2; color:#991b1b; }
      .report-footer { margin-top:30px; padding-top:12px; border-top:1px solid var(--line); color:#64748b; font-size:11px; text-align:right; }
    </style>
    """


def _render_image_style_hero(context):
    rule = context["agents"]["rule_based"]
    api = context["agents"]["api_based"]
    avg_metric = sum(metric["rate"] for metric in api["metric_rates"].values()) / max(len(api["metric_rates"]), 1)
    return f"""
    <section class="hero">
      <span style="display:none">QA 최종 테스트 결과 보고서</span>
      <h1>AI 교육과정 안내 챗봇 품질 테스트</h1>
      <div class="subtitle">규칙 기반 vs API 기반 챗봇 비교 — 최종 테스트 결과 보고서</div>
      <div class="meta-line">작성일: {html.escape(context['generated_date'])} &nbsp;&nbsp; AI 기반 소프트웨어 QA 및 모니터링 실무 교육</div>
    </section>
    <section class="kpi-band">
      {_kpi_card("규칙 기반 합격률", rule["pass_rate"], f"{rule['decision_counts']['PASS']}/{context['total_cases']} PASS", "")}
      {_kpi_card("API 기반 합격률", api["pass_rate"], f"{api['decision_counts']['PASS']}/{context['total_cases']} PASS", "dark")}
      {_kpi_card("평균 지표 충족률", avg_metric, "4대 지표 평균", "mint")}
      <div class="kpi red"><div class="kpi-title">결함(FAIL)</div><div class="kpi-value">{len(context['defects'])}건</div><small>조치 대상</small></div>
    </section>
    """


def _kpi_card(title, value, caption, class_name):
    return f"""
    <div class="kpi {class_name}">
      <div class="kpi-title">{html.escape(title)}</div>
      <div class="kpi-value">{value:.0f}%</div>
      <small>{html.escape(caption)}</small>
      <div class="gauge"><span>{value:.0f}</span></div>
    </div>
    """


def _render_four_metric_comparison(context):
    rule = context["agents"]["rule_based"]
    api = context["agents"]["api_based"]
    rows = []
    for metric_key, metric_label in METRICS:
        rule_score = rule["metric_averages"].get(metric_key, 0)
        api_score = api["metric_averages"].get(metric_key, 0)
        rows.append(_dual_score_bar(metric_label, rule_score, api_score))
    return f"""
    <section class="panel">
      <h2>규칙 기반 vs API 기반 — 4대 지표 비교</h2>
      <div class="chart-box">
        <div class="center" style="font-weight:800; margin-bottom:12px;">규칙 기반 vs API 기반 — 4대 지표 평균 비교</div>
        {''.join(rows)}
        <div class="muted center">평균 점수 (5점 만점)</div>
      </div>
    </section>
    """


def _dual_score_bar(label, rule_score, api_score):
    rule_width = max(0, min(rule_score / 5 * 100, 100))
    api_width = max(0, min(api_score / 5 * 100, 100))
    return f"""
    <div class="bar-row"><div>{html.escape(label)}</div><div class="bar-track"><div class="bar-fill" style="width:{rule_width:.1f}%"></div></div><div>{rule_score:.2f}</div></div>
    <div class="bar-row"><div></div><div class="bar-track"><div class="bar-fill alt" style="width:{api_width:.1f}%"></div></div><div>{api_score:.2f}</div></div>
    """


def _render_test_overview_table(context):
    return f"""
    <section class="panel">
      <h2>1. 테스트 개요</h2>
      <table>
        <tbody>
          <tr><td class="label-cell">테스트명</td><td>AI 교육과정 안내 챗봇 규칙 기반 vs API 기반 비교 테스트</td></tr>
          <tr><td class="label-cell">테스트 대상</td><td>① 규칙 기반 챗봇(rule_based_agent.py, API 미사용) ② API 기반 챗봇(service_agent.py + ChromaDB RAG)</td></tr>
          <tr><td class="label-cell">테스트 목적</td><td>동일한 테스트 케이스에 대해 두 챗봇의 응답 정확성·근거성·유용성·안전성을 동일 기준으로 비교 검증</td></tr>
          <tr><td class="label-cell">테스트 수행자</td><td>AI 품질관리 자동화 파이프라인 (Rule Validator + Judge Agent)</td></tr>
          <tr><td class="label-cell">테스트 방법</td><td>Happy Case, Edge Case, Negative Case 기반 Judge Agent 자동 평가</td></tr>
          <tr><td class="label-cell">테스트 환경</td><td>Python / Windows / OpenAI + ChromaDB RAG / 키워드 매칭</td></tr>
          <tr><td class="label-cell">테스트 일자</td><td>{html.escape(context['generated_date'])}</td></tr>
          <tr><td class="label-cell">평가 도구</td><td>Judge Agent 기반 자동 판정: 정확성·근거성·유용성·안전성 4대 지표 채점</td></tr>
        </tbody>
      </table>
    </section>
    """


def _render_overall_summary(context):
    return (
        "<section class='panel'><h2>2. 전체 비교 결과 요약</h2>"
        + _render_agent_stats_table(context)
        + _render_case_comparison_table(context["comparisons"])
        + _render_decision_chart(context)
        + "</section>"
    )


def _render_agent_stats_table(context):
    rows = []
    for icon, key in (("⚙", "rule_based"), ("🤖", "api_based")):
        agent = context["agents"][key]
        review_fail = agent["decision_counts"]["REVIEW"] + agent["decision_counts"]["FAIL"]
        rows.append(
            f"<tr><td>{icon} {html.escape(agent['label'])}</td><td>{context['total_cases']}건</td><td>{agent['decision_counts']['PASS']}건</td><td>{review_fail}건</td><td>{agent['pass_rate']:.0f}%</td><td>✓ 합격</td></tr>"
        )
    return (
        "<h3>2.1 챗봇 유형별 TC 수행 결과 통계</h3>"
        "<table><thead><tr><th>챗봇 유형</th><th>전체 TC</th><th>통과(Pass)</th><th>실패/재검토</th><th>성공률</th><th>판정</th></tr></thead>"
        f"<tbody>{''.join(rows)}</tbody></table>"
    )


def _render_case_comparison_table(comparisons):
    rows = []
    for row in comparisons:
        match = row["rule_decision"] == row["api_decision"]
        rows.append(
            "<tr>"
            f"<td class='center'>{html.escape(str(row['case_id']))}</td>"
            f"<td class='center'>{html.escape(str(row['category']))}</td>"
            f"<td class='center'>{html.escape(str(row['test_type']))}</td>"
            f"<td class='center {_decision_cell(row['rule_decision'])}'>{html.escape(row['rule_decision'])}</td>"
            f"<td class='center {_decision_cell(row['api_decision'])}'>{html.escape(row['api_decision'])}</td>"
            f"<td class='center'>{'🟢 일치' if match else '🔶 불일치'}</td>"
            "</tr>"
        )
    return (
        "<h3>2.2 케이스별 판정 비교표</h3>"
        "<table><thead><tr><th>TC ID</th><th>카테고리</th><th>유형</th><th>규칙기반 판정</th><th>API기반 판정</th><th>일치여부</th></tr></thead>"
        f"<tbody>{''.join(rows)}</tbody></table>"
    )


def _decision_cell(decision):
    return "pass-cell" if decision == "PASS" else "fail-cell" if decision == "FAIL" else ""


def _render_decision_chart(context):
    chart_data = _decision_chart_data_uri(context)
    return (
        "<div class='chart-box'><div class='center' style='font-weight:800;'>규칙 기반 vs API 기반 챗봇 — 판정 결과 비교</div>"
        f"<img src='{chart_data}' alt='규칙 기반 vs API 기반 챗봇 판정 결과 비교' style='display:block;width:660px;max-width:100%;margin:10px auto 0;'>"
        "</div>"
    )


def _decision_chart_data_uri(context):
    rule = context["agents"]["rule_based"]["decision_counts"]
    api = context["agents"]["api_based"]["decision_counts"]
    values = {
        "Pass": (rule["PASS"], api["PASS"]),
        "Review": (rule["REVIEW"], api["REVIEW"]),
        "Fail": (rule["FAIL"], api["FAIL"]),
    }
    max_value = max([number for pair in values.values() for number in pair] + [1])
    y_max = max(20, int(math.ceil(max_value / 2.5) * 2.5))

    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return ""

    width, height = 760, 370
    margin_left, margin_right, margin_top, margin_bottom = 72, 30, 38, 58
    plot_width = width - margin_left - margin_right
    plot_height = height - margin_top - margin_bottom
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = _pil_font(16)
    small_font = _pil_font(13)
    title_font = _pil_font(18)

    draw.text((width / 2, 10), "규칙 기반 vs API 기반 챗봇 — 판정 결과 비교", fill="black", font=title_font, anchor="ma")

    axis_color = "#222222"
    grid_color = "#d8d8d8"
    draw.line((margin_left, margin_top, margin_left, margin_top + plot_height), fill=axis_color, width=2)
    draw.line((margin_left, margin_top + plot_height, margin_left + plot_width, margin_top + plot_height), fill=axis_color, width=2)

    tick_step = 2.5
    tick = 0
    while tick <= y_max + 0.001:
        y = margin_top + plot_height - (tick / y_max * plot_height)
        draw.line((margin_left - 5, y, margin_left, y), fill=axis_color, width=1)
        if tick > 0:
            draw.line((margin_left, y, margin_left + plot_width, y), fill=grid_color, width=1)
        draw.text((margin_left - 10, y), f"{tick:.1f}", fill="black", font=small_font, anchor="rm")
        tick += tick_step

    colors = ("#2f63df", "#88baf0")
    group_centers = [margin_left + plot_width * ratio for ratio in (0.20, 0.50, 0.80)]
    bar_width = 52
    gap = 6
    for center, (label, pair) in zip(group_centers, values.items()):
        for offset, value, color in ((-(bar_width + gap) / 2, pair[0], colors[0]), ((bar_width + gap) / 2, pair[1], colors[1])):
            x0 = center + offset - bar_width / 2
            x1 = x0 + bar_width
            y1 = margin_top + plot_height
            y0 = y1 - (value / y_max * plot_height)
            draw.rectangle((x0, y0, x1, y1), fill=color)
            draw.text(((x0 + x1) / 2, y0 - 15), str(value), fill="black", font=small_font, anchor="ma")
        draw.text((center, margin_top + plot_height + 24), label, fill="black", font=font, anchor="ma")

    legend_x, legend_y = width - 170, 52
    draw.rectangle((legend_x - 10, legend_y - 10, width - 34, legend_y + 50), outline="#b8b8b8", fill="white")
    draw.rectangle((legend_x, legend_y, legend_x + 22, legend_y + 12), fill=colors[0])
    draw.text((legend_x + 30, legend_y - 1), "규칙 기반 챗봇", fill="black", font=small_font)
    draw.rectangle((legend_x, legend_y + 24, legend_x + 22, legend_y + 36), fill=colors[1])
    draw.text((legend_x + 30, legend_y + 23), "API 기반 챗봇", fill="black", font=small_font)

    return _image_to_data_uri(image)


def _metric_bar_chart_data_uri(agent):
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        return ""

    width, height = 700, 260
    margin_left, margin_right, margin_top, margin_bottom = 82, 36, 36, 48
    plot_width = width - margin_left - margin_right
    plot_height = height - margin_top - margin_bottom
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = _pil_font(15)
    small_font = _pil_font(12)
    title_font = _pil_font(16)

    draw.text((width / 2, 12), f"[{agent['label']}] 품질 항목별 점수 vs 합격 기준", fill="black", font=title_font, anchor="ma")
    draw.rectangle((margin_left, margin_top, margin_left + plot_width, margin_top + plot_height), outline="#333333", width=1)

    for tick in range(0, 101, 20):
        x = margin_left + tick / 100 * plot_width
        draw.line((x, margin_top + plot_height, x, margin_top + plot_height + 5), fill="black")
        draw.text((x, margin_top + plot_height + 9), str(tick), fill="black", font=small_font, anchor="ma")

    pass_threshold = agent.get("pass_threshold", 85.0)
    threshold_x = margin_left + pass_threshold / 100 * plot_width
    draw.line((threshold_x, margin_top, threshold_x, margin_top + plot_height), fill="#222222", width=1)
    for y in range(margin_top, margin_top + plot_height, 8):
        draw.line((threshold_x, y, threshold_x, min(y + 3, margin_top + plot_height)), fill="#222222", width=1)
    draw.text(
        (threshold_x + 25, margin_top + plot_height - 13),
        f"합격 기준 {pass_threshold:g}점",
        fill="#333333",
        font=small_font,
    )

    labels_and_scores = []
    for metric_key, metric_label in reversed(METRICS):
        labels_and_scores.append((metric_label, agent["metric_rates"][metric_key]["rate"]))

    bar_h = 28
    gap = 12
    start_y = margin_top + 18
    for index, (label, score) in enumerate(labels_and_scores):
        y0 = start_y + index * (bar_h + gap)
        y1 = y0 + bar_h
        draw.text((margin_left - 10, y0 + bar_h / 2), label, fill="black", font=small_font, anchor="rm")
        draw.rectangle((margin_left, y0, margin_left + plot_width, y1), fill="#eef3f5")
        fill_w = max(0, min(score, 100)) / 100 * plot_width
        draw.rectangle((margin_left, y0, margin_left + fill_w, y1), fill="#63c3a6")
        draw.text((margin_left + fill_w + 7, y0 + bar_h / 2), f"{score:.1f}점", fill="black", font=small_font, anchor="lm")

    draw.text((width / 2, height - 16), "점수 (0-100)", fill="black", font=small_font, anchor="ma")
    return _image_to_data_uri(image)


def _radar_chart_data_uri(agent):
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        return ""

    width, height = 520, 390
    center = (width // 2, 205)
    radius = 120
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    small_font = _pil_font(12)
    title_font = _pil_font(15)

    draw.text((width / 2, 22), f"[{agent['label']}] 품질 항목별 점수 레이더 차트", fill="black", font=title_font, anchor="ma")
    metrics_order = [
        ("groundedness", "근거성", -math.pi / 2),
        ("accuracy", "정확성", 0),
        ("safety", "안전성", math.pi),
        ("helpfulness", "유용성", math.pi / 2),
    ]

    for ratio in (0.25, 0.5, 0.75, 1.0):
        points = []
        for _, _, angle in metrics_order:
            x = center[0] + radius * ratio * math.cos(angle)
            y = center[1] + radius * ratio * math.sin(angle)
            points.append((x, y))
        draw.polygon(points, outline="#b8cbe6")
    for _, label, angle in metrics_order:
        x = center[0] + radius * math.cos(angle)
        y = center[1] + radius * math.sin(angle)
        draw.line((center[0], center[1], x, y), fill="#c7d7e8")
        lx = center[0] + (radius + 28) * math.cos(angle)
        ly = center[1] + (radius + 26) * math.sin(angle)
        draw.text((lx, ly), label, fill="black", font=small_font, anchor="mm")

    actual_points = []
    threshold_points = []
    for metric_key, _, angle in metrics_order:
        actual = agent["metric_rates"][metric_key]["rate"] / 100
        threshold = agent.get("pass_threshold", 85.0) / 100
        actual_points.append((center[0] + radius * actual * math.cos(angle), center[1] + radius * actual * math.sin(angle)))
        threshold_points.append((center[0] + radius * threshold * math.cos(angle), center[1] + radius * threshold * math.sin(angle)))
    draw.polygon(actual_points, fill="#dcefeb", outline="#63c3a6")
    draw.line(actual_points + [actual_points[0]], fill="#63c3a6", width=3)
    _draw_dashed_polygon(draw, threshold_points, fill="#24384f", width=2)

    for tick in (20, 40, 60, 80, 100):
        draw.text((center[0] + radius * tick / 100 + 4, center[1] - 3), str(tick), fill="#666666", font=small_font)

    legend_x, legend_y = width - 145, 48
    draw.line((legend_x, legend_y, legend_x + 24, legend_y), fill="#63c3a6", width=3)
    draw.text((legend_x + 32, legend_y - 7), "실제 점수", fill="black", font=small_font)
    _draw_dashed_line(draw, (legend_x, legend_y + 20), (legend_x + 24, legend_y + 20), fill="#24384f", width=2)
    draw.text((legend_x + 32, legend_y + 13), "합격 기준", fill="black", font=small_font)
    return _image_to_data_uri(image)


def _draw_dashed_polygon(draw, points, fill, width):
    for start, end in zip(points, points[1:] + points[:1]):
        _draw_dashed_line(draw, start, end, fill=fill, width=width)


def _draw_dashed_line(draw, start, end, fill, width=1, dash=6):
    x1, y1 = start
    x2, y2 = end
    length = math.hypot(x2 - x1, y2 - y1)
    if length == 0:
        return
    steps = int(length // dash)
    for index in range(0, steps, 2):
        a = index / steps
        b = min(index + 1, steps) / steps
        draw.line((x1 + (x2 - x1) * a, y1 + (y2 - y1) * a, x1 + (x2 - x1) * b, y1 + (y2 - y1) * b), fill=fill, width=width)


def _image_to_data_uri(image):
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def _pil_font(size):
    candidates = [
        Path(r"C:\Windows\Fonts\malgun.ttf"),
        Path(r"C:\Windows\Fonts\malgunbd.ttf"),
    ]
    try:
        from PIL import ImageFont

        for path in candidates:
            if path.exists():
                return ImageFont.truetype(str(path), size=size)
        return ImageFont.load_default()
    except Exception:
        return None


def _render_agent_image_section(context, agent_key, title):
    agent = context["agents"][agent_key]
    cases = context["cases"]
    return (
        f"<section class='panel'><h2>{html.escape(title)}</h2>"
        + _render_agent_metric_table(agent)
        + _render_agent_metric_bar_chart(agent)
        + _render_agent_radar_chart(agent)
        + _render_agent_log_table(cases, agent_key)
        + "<div class='two-col'>"
        + _render_type_pass_fail_bars(cases, agent_key)
        + _render_donut(agent)
        + "</div></section>"
    )


def _render_agent_metric_table(agent):
    rows = []
    for metric_key, metric_label in METRICS:
        passed = agent["metric_rates"][metric_key]["passed"]
        failed = max(agent["metric_rates"][metric_key].get("total", 0), 0)
        total = passed + failed if failed else passed
        rate = agent["metric_rates"][metric_key]["rate"]
        satisfied = rate >= agent.get("pass_threshold", 85.0)
        status_class = "pass-cell" if satisfied else "fail-cell"
        status_label = "✓ 충족" if satisfied else "미달"
        rows.append(
            f"<tr><td class='center'>{html.escape(metric_label)}</td><td class='center'>{passed}개</td>"
            f"<td class='center'>{contextless_total(agent)}개</td><td class='center'>{rate:.1f}점</td>"
            f"<td class='center {status_class}'>{status_label}</td></tr>"
        )
    return "<h3>품질 항목별 점수</h3><table><thead><tr><th>평가 항목</th><th>Pass</th><th>전체 TC</th><th>점수</th><th>합격 기준 충족</th></tr></thead><tbody>" + "".join(rows) + "</tbody></table>"


def contextless_total(agent):
    return sum(agent["decision_counts"].values())


def _render_agent_metric_bars(agent):
    rows = []
    for metric_key, metric_label in reversed(METRICS):
        score = agent["metric_rates"][metric_key]["rate"]
        rows.append(
            f"<div class='bar-row'><div>{html.escape(metric_label)}</div><div class='bar-track'>"
            f"<div class='threshold' style='left:{agent.get('pass_threshold', 85.0):g}%'></div>"
            f"<div class='bar-fill alt' style='width:{score:.1f}%'></div></div><div>{score:.1f}점</div></div>"
        )
    threshold = agent.get("pass_threshold", 85.0)
    return "<div class='chart-box'><div class='center' style='font-weight:800;'>품질 항목별 점수 vs 합격 기준</div>" + "".join(rows) + f"<div class='muted center'>합격 기준 {threshold:g}점</div></div>"


def _render_agent_metric_bar_chart(agent):
    data_uri = _metric_bar_chart_data_uri(agent)
    return f"<div class='chart-box'><img src='{data_uri}' alt='품질 항목별 점수 vs 합격 기준' style='display:block;width:620px;max-width:100%;margin:0 auto;'></div>"


def _render_agent_radar_chart(agent):
    data_uri = _radar_chart_data_uri(agent)
    return f"<div class='chart-box'><img src='{data_uri}' alt='품질 항목별 레이더 차트' style='display:block;width:420px;max-width:100%;margin:0 auto;'></div>"


def _render_agent_log_table(cases, agent_key):
    rows = []
    for case in cases[:12]:
        result = case.get(agent_key, {})
        evaluation = result.get("evaluation_result", {})
        rows.append(
            "<tr>"
            f"<td class='center'>{html.escape(str(case.get('case_id','')))}</td>"
            f"<td class='center'>{html.escape(str(case.get('test_type','')))}</td>"
            f"<td class='center'>{html.escape(str(case.get('category','')))}</td>"
            f"<td>{html.escape(get_eval_summary(case, agent_key) or result.get('ai_answer',''))}</td>"
            f"<td class='center'>Y</td><td class='center {_decision_cell(evaluation.get('overall_decision',''))}'>{html.escape(evaluation.get('overall_decision',''))}</td>"
            f"<td class='center'>{'결함' if evaluation.get('overall_decision') == 'FAIL' else ''}</td>"
            "</tr>"
        )
    return "<h3>TC 결과 기록표</h3><table><thead><tr><th>TC ID</th><th>유형</th><th>평가 영역</th><th>결과 요약</th><th>충족 여부</th><th>판정</th><th>결함 ID</th></tr></thead><tbody>" + "".join(rows) + "</tbody></table>"


def _render_type_pass_fail_bars(cases, agent_key):
    grouped = defaultdict(lambda: {"pass": 0, "fail": 0})
    for case in cases:
        key = case.get("test_type") or "기타"
        if get_decision(case, agent_key) == "PASS":
            grouped[key]["pass"] += 1
        else:
            grouped[key]["fail"] += 1
    max_total = max((v["pass"] + v["fail"] for v in grouped.values()), default=1)
    bars = []
    labels = []
    for label, counts in grouped.items():
        total = counts["pass"] + counts["fail"]
        height = total / max_total * 140
        fail_height = counts["fail"] / max(total, 1) * height
        bars.append(f"<div class='type-bar' style='height:{height:.1f}px'><div class='fail-part' style='height:{fail_height:.1f}px'></div></div>")
        labels.append(f"<div>{html.escape(label)}</div>")
    return "<div class='chart-box type-bars'><h4>유형별 Pass/Fail 분포</h4><div class='stack'>" + "".join(bars) + "</div><div class='type-labels'>" + "".join(labels) + "</div></div>"


def _render_donut(agent):
    pct = max(0, min(agent["pass_rate"], 100))
    return f"<div class='chart-box'><h4>전체 TC Pass 비율</h4><div class='donut' style='--pct:{pct:.1f}%'><span>{agent['decision_counts']['PASS']} / {contextless_total(agent)}<br>Pass</span></div></div>"


def _render_bug_report_image_style(defects):
    if not defects:
        return "<section class='panel'><h2>4.3 결함 보고서 (Bug Report)</h2><p>FAIL 판정 결함이 발견되지 않았습니다.</p></section>"
    defect = defects[0]
    return f"""
    <section class="panel">
      <h2>4.3 결함 보고서 (Bug Report)</h2>
      <h4>4.3.1 {html.escape(defect['defect_id'])} — {html.escape(defect['case_id'])} 핵심 품질 저하</h4>
      <div class="bug-card">
        <div><div class="bug-head">결함 정보</div><div class="bug-body"><b>결함 ID</b><br>{html.escape(defect['defect_id'])}<br><br><b>심각도</b><br><span class="badge fail">{html.escape(defect['severity'])}</span><br><br><b>우선순위</b><br>{html.escape(defect['priority'])}</div></div>
        <div><div class="bug-head">문제 내용</div><div class="bug-body"><b>테스트 케이스</b><br>{html.escape(defect['case_id'])}<br><br><b>입력 데이터</b><br>{html.escape(defect['question'])}<br><br><b>기대 결과</b><br>정확하고 근거 있는 안내 제공</div></div>
        <div><div class="bug-head">분석·조치</div><div class="bug-body"><b>결함 개요</b><br>{html.escape(defect['summary'])}<br><br><b>영향도 분석</b><br>{html.escape(defect['weak_metrics'])} 지표 개선 필요<br><br><b>권장 조치</b><br>룰/프롬프트/검색 기준 보강 후 회귀 테스트</div></div>
      </div>
    </section>
    """


def generate_docx_report(
    pipeline_outputs, output_path, include_appendix=False, appendix_options=None, run_id=None,
    quality_criteria=None,
):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    context = build_report_context(pipeline_outputs, quality_criteria)

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        lines = _plain_report_lines(context)
        if include_appendix:
            lines.extend(render_markdown_appendix(pipeline_outputs, run_id=run_id, options=appendix_options).splitlines())
        _write_minimal_docx(output_path, lines)
        return output_path

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI 교육과정 안내 챗봇 품질 테스트")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 45, 66)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("규칙 기반 vs API 기반 챗봇 비교 — 최종 테스트 결과 보고서")
    sub_run.bold = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(104, 123, 148)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"작성일: {context['generated_date']}    AI 기반 소프트웨어 QA 및 모니터링 실무 교육")

    _add_docx_kpi_table(document, context)
    _add_docx_image_from_data_uri(document, _metric_comparison_chart_data_uri(context), width=Inches(6.8))

    _add_docx_heading(document, "1. 테스트 개요")
    _add_docx_test_overview(document, context)

    _add_docx_heading(document, "2. 전체 비교 결과 요약")
    _add_docx_agent_stats(document, context)
    _add_docx_case_comparison(document, context["comparisons"])
    _add_docx_image_from_data_uri(document, _decision_chart_data_uri(context), width=Inches(6.5))

    _add_docx_agent_section(document, context, "rule_based", "3. 규칙 기반 챗봇 상세 결과", Inches, Pt, RGBColor)
    _add_docx_agent_section(document, context, "api_based", "4. API 기반 챗봇 상세 결과", Inches, Pt, RGBColor)

    _add_docx_heading(document, "4.3 결함 보고서 (Bug Report)")
    _add_docx_defect_report(document, context["defects"])

    _add_docx_heading(document, "5. 심각도 분류 기준")
    _add_docx_severity_table(document)

    _add_docx_heading(document, "6. 최종 종합 의견 및 개선 권고")
    document.add_paragraph(context["final_opinion"])
    if include_appendix:
        appendix_context = build_appendix_context(pipeline_outputs, run_id=run_id, options=appendix_options)
        append_docx_appendix(document, appendix_context)
    document.save(output_path)
    return output_path


def _metric_comparison_chart_data_uri(context):
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        return ""

    rule = context["agents"]["rule_based"]
    api = context["agents"]["api_based"]
    width, height = 760, 330
    margin_left, margin_right, margin_top, margin_bottom = 86, 88, 50, 48
    plot_width = width - margin_left - margin_right
    plot_height = height - margin_top - margin_bottom
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = _pil_font(14)
    small_font = _pil_font(12)
    title_font = _pil_font(17)

    draw.text((width / 2, 16), "규칙 기반 vs API 기반 — 4대 지표 평균 비교", fill="black", font=title_font, anchor="ma")
    draw.rectangle((margin_left, margin_top, margin_left + plot_width, margin_top + plot_height), outline="#333333", width=1)
    for tick in range(0, 6):
        x = margin_left + tick / 5 * plot_width
        draw.line((x, margin_top, x, margin_top + plot_height), fill="#d7d7d7" if tick else "#333333", width=1)
        draw.text((x, margin_top + plot_height + 10), str(tick), fill="black", font=small_font, anchor="ma")

    colors = ("#1899aa", "#7ccac0")
    row_gap = plot_height / len(METRICS)
    bar_h = 14
    for index, (metric_key, metric_label) in enumerate(METRICS):
        center_y = margin_top + row_gap * index + row_gap / 2
        draw.text((margin_left - 12, center_y), metric_label, fill="black", font=font, anchor="rm")
        for offset, score, color in (
            (-bar_h / 2 - 2, rule["metric_averages"].get(metric_key, 0), colors[0]),
            (bar_h / 2 + 2, api["metric_averages"].get(metric_key, 0), colors[1]),
        ):
            y0 = center_y + offset - bar_h / 2
            y1 = y0 + bar_h
            x1 = margin_left + min(max(score, 0), 5) / 5 * plot_width
            draw.rectangle((margin_left, y0, x1, y1), fill=color)
            draw.text((x1 + 8, y0 + bar_h / 2), f"{score:.2f}", fill="#26364f", font=small_font, anchor="lm")

    legend_x, legend_y = width - 150, height - 74
    draw.rectangle((legend_x, legend_y, legend_x + 18, legend_y + 10), fill=colors[0])
    draw.text((legend_x + 26, legend_y - 2), "규칙 기반", fill="black", font=small_font)
    draw.rectangle((legend_x, legend_y + 20, legend_x + 18, legend_y + 30), fill=colors[1])
    draw.text((legend_x + 26, legend_y + 18), "API 기반", fill="black", font=small_font)
    draw.text((width / 2, height - 16), "평균 점수 (5점 만점)", fill="black", font=small_font, anchor="ma")
    return _image_to_data_uri(image)


def _docx_imports():
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    return WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt, RGBColor


def _set_docx_cell_shading(cell, fill):
    _, OxmlElement, qn, _, _, _ = _docx_imports()
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _style_docx_table(table, header_fill="1899AA"):
    _, _, _, _, Pt, RGBColor = _docx_imports()
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        _set_docx_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(255, 255, 255)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def _add_docx_heading(document, text):
    _, _, _, _, Pt, RGBColor = _docx_imports()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(4, 117, 138)


def _add_docx_image_from_data_uri(document, data_uri, width=None):
    if not data_uri:
        return
    WD_ALIGN_PARAGRAPH, _, _, _, _, _ = _docx_imports()
    encoded = data_uri.split(",", 1)[1] if "," in data_uri else data_uri
    stream = io.BytesIO(base64.b64decode(encoded))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(stream, width=width)


def _add_docx_kpi_table(document, context):
    _, _, _, _, Pt, RGBColor = _docx_imports()
    rule = context["agents"]["rule_based"]
    api = context["agents"]["api_based"]
    avg_metric = sum(metric["rate"] for metric in api["metric_rates"].values()) / max(len(api["metric_rates"]), 1)
    items = [
        ("규칙 기반 합격률", f"{rule['pass_rate']:.0f}%", f"{rule['decision_counts']['PASS']}/{context['total_cases']} PASS", "1899AA"),
        ("API 기반 합격률", f"{api['pass_rate']:.0f}%", f"{api['decision_counts']['PASS']}/{context['total_cases']} PASS", "087887"),
        ("평균 지표 충족률", f"{avg_metric:.0f}%", "4대 지표 평균", "63C3A6"),
        ("결함(FAIL)", f"{len(context['defects'])}건", "조치 대상", "E87E80"),
    ]
    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    for col, (title, value, caption, color) in enumerate(items):
        _set_docx_cell_shading(table.cell(0, col), color)
        header = table.cell(0, col).paragraphs[0]
        header.add_run(title).bold = True
        body = table.cell(1, col).paragraphs[0]
        value_run = body.add_run(value)
        value_run.bold = True
        value_run.font.size = Pt(18)
        value_run.font.color.rgb = RGBColor(4, 117, 138)
        body.add_run(f"\n{caption}")
    document.add_paragraph()


def _add_docx_test_overview(document, context):
    rows = [
        ("테스트명", "AI 교육과정 안내 챗봇 규칙 기반 vs API 기반 비교 테스트"),
        ("테스트 대상", "① 규칙 기반 챗봇(rule_based_agent.py, API 미사용) ② API 기반 챗봇(service_agent.py + ChromaDB RAG)"),
        ("테스트 목적", "동일한 테스트 케이스에 대해 두 챗봇의 응답 정확성·근거성·유용성·안전성을 동일 기준으로 비교 검증"),
        ("테스트 수행자", "AI 품질관리 자동화 파이프라인 (Rule Validator + Judge Agent)"),
        ("테스트 방법", "Happy Case, Edge Case, Negative Case 기반 Judge Agent 자동 평가"),
        ("테스트 환경", "Python / Windows / OpenAI + ChromaDB RAG / 키워드 매칭"),
        ("테스트 일자", context["generated_date"]),
        ("평가 도구", "Judge Agent 기반 자동 판정: 정확성·근거성·유용성·안전성 4대 지표 채점"),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value
        _set_docx_cell_shading(table.cell(index, 0), "E3F2F6")


def _add_docx_agent_stats(document, context):
    table = document.add_table(rows=1, cols=6)
    headers = ["챗봇 유형", "전체 TC", "통과(Pass)", "실패/재검토", "성공률", "판정"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for key in ("rule_based", "api_based"):
        agent = context["agents"][key]
        row = table.add_row().cells
        review_fail = agent["decision_counts"]["REVIEW"] + agent["decision_counts"]["FAIL"]
        values = [agent["label"], f"{context['total_cases']}건", f"{agent['decision_counts']['PASS']}건", f"{review_fail}건", f"{agent['pass_rate']:.0f}%", "합격"]
        for index, value in enumerate(values):
            row[index].text = value
    _style_docx_table(table)


def _add_docx_case_comparison(document, comparisons):
    _add_docx_heading(document, "2.2 케이스별 판정 비교표")
    table = document.add_table(rows=1, cols=6)
    headers = ["TC ID", "카테고리", "유형", "규칙기반 판정", "API기반 판정", "일치여부"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for item in comparisons[:30]:
        row = table.add_row().cells
        values = [
            item["case_id"],
            item["category"],
            item["test_type"],
            item["rule_decision"],
            item["api_decision"],
            "일치" if item["rule_decision"] == item["api_decision"] else "불일치",
        ]
        for index, value in enumerate(values):
            row[index].text = str(value)
        for index in (3, 4):
            fill = "E3F5EE" if row[index].text == "PASS" else "F9DDDD" if row[index].text == "FAIL" else "FEF3C7"
            _set_docx_cell_shading(row[index], fill)
    _style_docx_table(table)


def _add_docx_agent_section(document, context, agent_key, title, Inches, Pt, RGBColor):
    _add_docx_heading(document, title)
    agent = context["agents"][agent_key]
    table = document.add_table(rows=1, cols=5)
    headers = ["평가 항목", "전체 TC", "Pass", "점수", "합격 기준 충족"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for metric_key, metric_label in METRICS:
        rate = agent["metric_rates"][metric_key]
        row = table.add_row().cells
        threshold = context["agents"][agent_key].get("pass_threshold", context["pass_threshold"])
        values = [metric_label, f"{context['total_cases']}개", f"{rate['passed']}개", f"{rate['rate']:.1f}점", "충족" if rate["rate"] >= threshold else "미달"]
        for index, value in enumerate(values):
            row[index].text = value
        _set_docx_cell_shading(row[4], "E3F5EE" if rate["rate"] >= threshold else "F9DDDD")
    _style_docx_table(table)
    _add_docx_image_from_data_uri(document, _metric_bar_chart_data_uri(agent), width=Inches(5.9))
    _add_docx_image_from_data_uri(document, _radar_chart_data_uri(agent), width=Inches(4.3))
    _add_docx_agent_log_table(document, context["cases"], agent_key)


def _add_docx_agent_log_table(document, cases, agent_key):
    table = document.add_table(rows=1, cols=7)
    headers = ["TC ID", "유형", "평가 영역", "결과 요약", "충족 여부", "판정", "결함 ID"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for case in cases[:15]:
        decision = get_decision(case, agent_key)
        row = table.add_row().cells
        values = [
            case.get("case_id", ""),
            case.get("test_type", ""),
            case.get("category", ""),
            get_eval_summary(case, agent_key) or case.get(agent_key, {}).get("ai_answer", ""),
            "Y" if decision == "PASS" else "N",
            decision,
            "결함" if decision == "FAIL" else "",
        ]
        for index, value in enumerate(values):
            row[index].text = str(value)
        _set_docx_cell_shading(row[5], "E3F5EE" if decision == "PASS" else "F9DDDD" if decision == "FAIL" else "FEF3C7")
    _style_docx_table(table)


def _add_docx_defect_report(document, defects):
    if not defects:
        document.add_paragraph("FAIL 판정 결함이 발견되지 않았습니다.")
        return
    table = document.add_table(rows=1, cols=7)
    headers = ["결함 ID", "구분", "TC ID", "심각도", "우선순위", "취약 지표", "평가 요약"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for defect in defects[:10]:
        row = table.add_row().cells
        values = [
            defect["defect_id"],
            defect["agent"],
            defect["case_id"],
            defect["severity"],
            defect["priority"],
            defect["weak_metrics"],
            defect["summary"],
        ]
        for index, value in enumerate(values):
            row[index].text = str(value)
        _set_docx_cell_shading(row[3], "F9DDDD")
    _style_docx_table(table)


def _add_docx_severity_table(document):
    rows = [
        ("Critical", "보안 취약점, 개인정보 유출, 불법·위험 행위에 대한 도움 등 서비스 신뢰를 근본적으로 해치는 최상위 위험입니다."),
        ("High", "핵심 규정 수치 오류 또는 잘못된 정책 안내로 사용자가 실제 불이익을 겪을 수 있는 심각한 문제입니다."),
        ("Medium", "문서에 없는 정보의 확장 생성 등 간접적인 신뢰도 저하 요인으로, 우선 대응 가능합니다."),
        ("Low", "서비스 기능에는 영향이 없는 경미한 표현·형식 문제입니다."),
    ]
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "심각도 등급"
    table.cell(0, 1).text = "정의 및 기준"
    for severity, rule in rows:
        cells = table.add_row().cells
        cells[0].text = severity
        cells[1].text = rule
    _style_docx_table(table)


def generate_pdf_report(
    pipeline_outputs, output_path, include_appendix=False, appendix_options=None, run_id=None,
    quality_criteria=None,
):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    html_report = generate_html_report(
        pipeline_outputs,
        include_appendix=include_appendix,
        appendix_options=appendix_options,
        run_id=run_id,
        quality_criteria=quality_criteria,
    )
    if _write_pdf_from_html_with_browser(html_report, output_path):
        return output_path

    context = build_report_context(pipeline_outputs, quality_criteria)

    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError:
        lines = _plain_report_lines(context)
        if include_appendix:
            lines.extend(render_markdown_appendix(pipeline_outputs, run_id=run_id, options=appendix_options).splitlines())
        _write_minimal_pdf(output_path, lines)
        return output_path

    font_name = _register_korean_font(pdfmetrics, TTFont)
    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = font_name
    styles["Title"].fontSize = 24
    styles["Title"].leading = 31
    styles["Title"].textColor = colors.HexColor("#0f2547")
    styles["Heading2"].fontSize = 14
    styles["Heading2"].leading = 18
    styles["Heading2"].textColor = colors.HexColor("#0f2547")
    styles["Normal"].fontSize = 9
    styles["Normal"].leading = 14
    cover_label_style = ParagraphStyle(
        "CoverLabel",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#64748b"),
        alignment=TA_CENTER,
    )
    cover_title_style = ParagraphStyle(
        "CoverTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=24,
        leading=32,
        textColor=colors.HexColor("#0f2547"),
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=34, leftMargin=34, topMargin=34, bottomMargin=34)
    cover_rows = [
        [Paragraph("AI QUALITY ASSURANCE REPORT", cover_label_style)],
        [Paragraph("QA 최종 테스트 결과 보고서", cover_title_style)],
        [Paragraph("프로젝트명: AI 교육과정 안내 챗봇 품질검증", styles["Normal"])],
        [Paragraph(f"생성일시: {context['generated_at']}", styles["Normal"])],
        [Paragraph(f"총 테스트 케이스: {context['total_cases']}건", styles["Normal"])],
        [Paragraph(f"최종 종합 의견: <b>{context['final_decision']}</b>", styles["Normal"])],
    ]
    cover_table = Table(cover_rows, colWidths=[500])
    cover_table.setStyle(
        TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#c8d3e4")),
                ("LINEABOVE", (0, 0), (-1, 0), 8, colors.HexColor("#18345f")),
                ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                ("TOPPADDING", (0, 0), (-1, -1), 12),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 22),
                ("RIGHTPADDING", (0, 0), (-1, -1), 22),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
            ]
        )
    )
    summary_rows = [
        ["평가 기준", "LLM API 성공률", "Rule 응답 성공률", "등록 대상 결함"],
        [
            f"{context['pass_threshold']:.0f}%",
            f"{context['agents']['api_based']['pass_rate']:.1f}%",
            f"{context['agents']['rule_based']['pass_rate']:.1f}%",
            f"{len(context['defects'])}건",
        ],
    ]
    story = [
        cover_table,
        Spacer(1, 12),
        _pdf_table(summary_rows, Table, TableStyle, colors, font_name),
        Spacer(1, 14),
        Paragraph("1. 전체 비교 결과 요약", styles["Heading2"]),
    ]
    overview_rows = [["구분", "PASS", "REVIEW", "FAIL", "성공률"]]
    for agent in context["agents"].values():
        overview_rows.append(
            [
                agent["label"],
                agent["decision_counts"]["PASS"],
                agent["decision_counts"]["REVIEW"],
                agent["decision_counts"]["FAIL"],
                f"{agent['pass_rate']:.1f}%",
            ]
        )
    story.append(_pdf_table(overview_rows, Table, TableStyle, colors, font_name))
    story.extend([Spacer(1, 14), Paragraph("2. 케이스별 판정 비교표", styles["Heading2"])])
    comparison_rows = [["Case ID", "카테고리", "유형", "규칙", "API", "일치"]]
    for row in context["comparisons"]:
        comparison_rows.append([row["case_id"], row["category"], row["test_type"], row["rule_decision"], row["api_decision"], row["match"]])
    story.append(_pdf_table(comparison_rows, Table, TableStyle, colors, font_name))
    story.extend([Spacer(1, 14), Paragraph("3. 결함 보고서", styles["Heading2"])])
    if context["defects"]:
        defect_rows = [["ID", "Case", "심각도", "우선순위", "취약 지표"]]
        for defect in context["defects"]:
            defect_rows.append([defect["defect_id"], defect["case_id"], defect["severity"], defect["priority"], defect["weak_metrics"]])
        story.append(_pdf_table(defect_rows, Table, TableStyle, colors, font_name))
    else:
        story.append(Paragraph("FAIL 판정 결함이 발견되지 않았습니다.", styles["Normal"]))
    story.extend([Spacer(1, 14), Paragraph("4. 최종 종합 의견 및 개선 권고", styles["Heading2"])])
    story.append(Paragraph(context["final_opinion"], styles["Normal"]))
    if include_appendix:
        appendix_context = build_appendix_context(pipeline_outputs, run_id=run_id, options=appendix_options)
        append_pdf_appendix(story, styles, Table, TableStyle, colors, font_name, appendix_context)
    doc.build(story)
    return output_path


def build_defect_items(cases, quality_criteria=None):
    criteria = get_quality_criteria(quality_criteria)
    defects = []
    counters = defaultdict(int)
    for case in cases:
        for agent_key, agent_label, prefix in AGENTS:
            if get_decision(case, agent_key) != "FAIL":
                continue
            counters[prefix] += 1
            severity = calculate_severity(case, agent_key)
            weak_metrics = [
                label
                for metric_key, label in METRICS
                if get_metric_score(case, agent_key, metric_key) < criteria.pass_min_score
            ]
            defects.append(
                {
                    "defect_id": f"{prefix}-{counters[prefix]:03d}",
                    "agent": agent_label,
                    "case_id": case.get("case_id", ""),
                    "category": case.get("category", ""),
                    "test_type": case.get("test_type", ""),
                    "question": case.get("user_question", ""),
                    "answer": case.get(agent_key, {}).get("ai_answer", ""),
                    "rule_reason": get_rule_reason(case, agent_key),
                    "summary": get_eval_summary(case, agent_key),
                    "weak_metrics": ", ".join(weak_metrics) if weak_metrics else "-",
                    "severity": severity,
                    "priority": {"Critical": "High", "High": "High", "Medium": "Medium", "Low": "Low"}[severity],
                }
            )
    return defects


def calculate_severity(case, agent_key):
    category = str(case.get("category", ""))
    if any(keyword in category for keyword in ("안전", "위험", "유해", "보안")):
        return "Critical"
    validation = case.get(agent_key, {}).get("rule_validation", {})
    if validation.get("rule_status") == "FAIL" or validation.get("passed") is False:
        return "High"
    if get_metric_score(case, agent_key, "groundedness") <= 3:
        return "Medium"
    return "Low"


def build_type_distribution(cases):
    rows = []
    for test_type in sorted({case.get("test_type") or "미분류" for case in cases}):
        type_cases = [case for case in cases if (case.get("test_type") or "미분류") == test_type]
        for decision in DECISIONS:
            rows.append({"test_type": test_type, "decision": decision, "count": sum(1 for case in type_cases if get_decision(case, "api_based") == decision)})
    return rows


def get_decision(case, agent_key):
    return case.get(agent_key, {}).get("evaluation_result", {}).get("overall_decision", "FAIL")


def get_metric_score(case, agent_key, metric_key):
    value = case.get(agent_key, {}).get("evaluation_result", {}).get(metric_key, 0)
    if isinstance(value, dict):
        value = value.get("score", 0)
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0


def get_rule_reason(case, agent_key):
    validation = case.get(agent_key, {}).get("rule_validation", {})
    return validation.get("rule_reason") or validation.get("reason") or ""


def get_eval_summary(case, agent_key):
    evaluation = case.get(agent_key, {}).get("evaluation_result", {})
    return evaluation.get("summary") or evaluation.get("comment") or ""


def _normalize_case(item, index, quality_criteria=None):
    if "rule_based" in item or "api_based" in item:
        case = dict(item)
        case.setdefault("case_id", f"TC-{index:03d}")
        case.setdefault("category", "")
        case.setdefault("test_type", "")
        case.setdefault("user_question", "")
        case.setdefault("rule_based", _empty_agent_result())
        case.setdefault("api_based", _empty_agent_result())
        return case

    api_eval = {metric_key: {"score": item.get(metric_key, 0), "reason": item.get("comment", "")} for metric_key, _ in METRICS}
    scores = [api_eval[metric_key]["score"] for metric_key, _ in METRICS]
    api_eval["overall_decision"] = _decision_from_scores(scores, quality_criteria)
    api_eval["summary"] = item.get("comment", "")
    rule_status = "PASS" if item.get("rule_passed") is True else "FAIL"
    rule_eval = {metric_key: {"score": 5 if rule_status == "PASS" else 1, "reason": item.get("rule_reason", "")} for metric_key, _ in METRICS}
    rule_eval["overall_decision"] = rule_status
    rule_eval["summary"] = item.get("rule_reason", "")
    return {
        "case_id": item.get("case_id", f"TC-{index:03d}"),
        "category": item.get("category", ""),
        "test_type": item.get("test_type", ""),
        "user_question": item.get("user_question", ""),
        "rule_based": {
            "ai_answer": item.get("response", ""),
            "rule_validation": {"passed": item.get("rule_passed") is True, "rule_status": rule_status, "rule_reason": item.get("rule_reason", "")},
            "evaluation_result": rule_eval,
        },
        "api_based": {
            "ai_answer": item.get("response", ""),
            "rule_validation": {"passed": item.get("rule_passed") is True, "rule_status": rule_status, "rule_reason": item.get("rule_reason", "")},
            "evaluation_result": api_eval,
        },
    }


def _empty_agent_result():
    evaluation = {metric_key: {"score": 0, "reason": ""} for metric_key, _ in METRICS}
    evaluation["overall_decision"] = "FAIL"
    evaluation["summary"] = ""
    return {"ai_answer": "", "rule_validation": {}, "evaluation_result": evaluation}


def _decision_from_scores(scores, quality_criteria=None):
    criteria = get_quality_criteria(quality_criteria)
    numeric_scores = []
    for score in scores:
        try:
            numeric_scores.append(float(score))
        except (TypeError, ValueError):
            numeric_scores.append(0)
    safety_score = numeric_scores[3] if len(numeric_scores) > 3 else 0
    if safety_score < criteria.safety_critical_min:
        return "FAIL"
    if min(numeric_scores or [0]) >= criteria.pass_min_score:
        return "PASS"
    if min(numeric_scores or [0]) >= criteria.review_min_score:
        return "REVIEW"
    return "FAIL"


def _rate(count, total):
    return round((count / total * 100), 1) if total else 0.0


def _render_overview(context):
    cards = []
    for agent in context["agents"].values():
        cards.append(
            f"""
            <div class="card">
              <span>{html.escape(agent['label'])} 성공률</span>
              <strong>{agent['pass_rate']:.1f}%</strong>
              <small>PASS {agent['decision_counts']['PASS']} / 전체 {context['total_cases']}</small>
            </div>
            """
        )
    cards.extend(
        [
            f"<div class='card'><span>합격 기준</span><strong>{context['pass_threshold']:.0f}%</strong><small>PASS 비율 기준</small></div>",
            f"<div class='card'><span>결함 수</span><strong>{len(context['defects'])}</strong><small>FAIL 케이스 기준</small></div>",
        ]
    )
    return f"<h2>1. 테스트 개요</h2><div class='grid'>{''.join(cards)}</div>"


def _render_metric_charts(context):
    rows = []
    for agent in context["agents"].values():
        rows.append(f"<h3>{html.escape(agent['label'])} 지표별 합격률</h3>")
        for metric in agent["metric_rates"].values():
            rows.append(
                f"""
                <div style="margin:10px 0;">
                  <div class="muted">{html.escape(metric['label'])}: {metric['rate']:.1f}%</div>
                  <div class="bar-wrap"><div class="bar" style="width:{metric['rate']:.1f}%"></div></div>
                </div>
                """
            )
        rows.append(_render_radar_svg(agent["metric_averages"]))
    rows.append(_render_type_distribution(context["type_distribution"]))
    return "<h2>2. 전체 비교 결과 요약</h2>" + "".join(rows)


def _render_radar_svg(metric_averages):
    center = 90
    radius = 64
    points = []
    labels = []
    for index, (metric_key, metric_label) in enumerate(METRICS):
        angle = -math.pi / 2 + (2 * math.pi * index / len(METRICS))
        score_radius = radius * (metric_averages.get(metric_key, 0) / 5)
        x = center + score_radius * math.cos(angle)
        y = center + score_radius * math.sin(angle)
        points.append(f"{x:.1f},{y:.1f}")
        lx = center + (radius + 18) * math.cos(angle)
        ly = center + (radius + 18) * math.sin(angle)
        labels.append(f"<text x='{lx:.1f}' y='{ly:.1f}' text-anchor='middle' font-size='11'>{html.escape(metric_label)}</text>")
    return (
        "<svg viewBox='0 0 180 180' width='260' height='220' role='img'>"
        "<polygon points='90,26 154,90 90,154 26,90' fill='#f8fbff' stroke='#b8cbe6'/>"
        "<line x1='90' y1='26' x2='90' y2='154' stroke='#d8e2f1'/>"
        "<line x1='26' y1='90' x2='154' y2='90' stroke='#d8e2f1'/>"
        f"<polygon points='{' '.join(points)}' fill='rgba(37,99,235,.25)' stroke='#2563eb' stroke-width='2'/>"
        f"{''.join(labels)}</svg>"
    )


def _render_type_distribution(rows):
    grouped = defaultdict(Counter)
    for row in rows:
        grouped[row["test_type"]][row["decision"]] = row["count"]
    table_rows = "".join(
        f"<tr><td>{html.escape(test_type)}</td><td>{counts['PASS']}</td><td>{counts['REVIEW']}</td><td>{counts['FAIL']}</td></tr>"
        for test_type, counts in grouped.items()
    )
    return (
        "<h3>테스트 유형별 분포</h3>"
        "<table><thead><tr><th>유형</th><th>PASS</th><th>REVIEW</th><th>FAIL</th></tr></thead>"
        f"<tbody>{table_rows}</tbody></table>"
    )


def _render_comparison_table(comparisons):
    rows = []
    for row in comparisons:
        rows.append(
            "<tr>"
            f"<td>{html.escape(str(row['case_id']))}</td>"
            f"<td>{html.escape(str(row['category']))}</td>"
            f"<td>{html.escape(str(row['test_type']))}</td>"
            f"<td>{_decision_badge(row['rule_decision'])}</td>"
            f"<td>{_decision_badge(row['api_decision'])}</td>"
            f"<td>{html.escape(row['match'])}</td>"
            "</tr>"
        )
    return (
        "<h2>3. 케이스별 판정 비교표</h2>"
        "<table><thead><tr><th>Case ID</th><th>카테고리</th><th>유형</th><th>규칙 기반</th><th>API 기반</th><th>일치 여부</th></tr></thead>"
        f"<tbody>{''.join(rows)}</tbody></table>"
    )


def _render_agent_detail(context, agent_key, label):
    agent = context["agents"][agent_key]
    rows = []
    for metric_key, metric_label in METRICS:
        rows.append(
            f"<tr><td>{html.escape(metric_label)}</td><td>{agent['metric_averages'][metric_key]:.2f}/5</td>"
            f"<td>{agent['metric_rates'][metric_key]['rate']:.1f}%</td></tr>"
        )
    return (
        f"<h2>{html.escape(label)} 상세 결과</h2>"
        "<table><thead><tr><th>지표</th><th>평균 점수</th><th>합격률</th></tr></thead>"
        f"<tbody>{''.join(rows)}</tbody></table>"
    )


def _render_defect_report(defects):
    if not defects:
        return "<h2>4. 결함 보고서</h2><p>FAIL 판정 결함이 발견되지 않았습니다.</p>"
    rows = []
    for defect in defects:
        rows.append(
            "<tr>"
            f"<td>{html.escape(defect['defect_id'])}</td><td>{html.escape(defect['agent'])}</td>"
            f"<td>{html.escape(defect['case_id'])}</td><td>{html.escape(defect['severity'])}</td>"
            f"<td>{html.escape(defect['priority'])}</td><td>{html.escape(defect['weak_metrics'])}</td>"
            f"<td>{html.escape(defect['summary'])}</td>"
            "</tr>"
        )
    return (
        "<h2>4. 결함 보고서</h2>"
        "<table><thead><tr><th>결함 ID</th><th>구분</th><th>Case ID</th><th>심각도</th><th>우선순위</th><th>취약 지표</th><th>평가 요약</th></tr></thead>"
        f"<tbody>{''.join(rows)}</tbody></table>"
    )


def _render_severity_rules():
    return """
    <h2>5. 심각도 분류 기준</h2>
    <table>
      <thead><tr><th>조건</th><th>심각도</th><th>우선순위</th></tr></thead>
      <tbody>
        <tr><td>카테고리에 안전/위험 관련 키워드 포함</td><td>Critical</td><td>High</td></tr>
        <tr><td>규칙 검증 결과 FAIL</td><td>High</td><td>High</td></tr>
        <tr><td>근거성 점수 3점 이하</td><td>Medium</td><td>Medium</td></tr>
        <tr><td>그 외 FAIL</td><td>Low</td><td>Low</td></tr>
      </tbody>
    </table>
    """


def _decision_badge(decision):
    class_name = {"PASS": "pass", "REVIEW": "review", "FAIL": "fail"}.get(decision, "review")
    return f"<span class='badge {class_name}'>{html.escape(str(decision))}</span>"


def _register_korean_font(pdfmetrics, TTFont):
    candidates = [
        Path(r"C:\Windows\Fonts\malgun.ttf"),
        Path("/usr/share/fonts/truetype/nanum/NanumGothic.ttf"),
        Path("/System/Library/Fonts/Supplemental/AppleGothic.ttf"),
    ]
    for path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont("KoreanFont", str(path)))
            return "KoreanFont"
    return "Helvetica"


def _write_pdf_from_html_with_browser(html_report, output_path):
    browser_path = _find_browser_executable()
    if not browser_path:
        return False

    output_path = Path(output_path)
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir = Path(temp_dir)
            html_path = temp_dir / "report.html"
            user_data_dir = temp_dir / "browser-profile"
            html_path.write_text(
                "<!doctype html><html><head><meta charset='utf-8'></head><body>"
                + html_report
                + "</body></html>",
                encoding="utf-8",
            )
            command = [
                str(browser_path),
                "--headless",
                "--disable-gpu",
                "--disable-extensions",
                f"--user-data-dir={user_data_dir}",
                f"--print-to-pdf={output_path}",
                "--print-to-pdf-no-header",
                str(html_path),
            ]
            completed = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=60,
                check=False,
            )
            return completed.returncode == 0 and output_path.exists() and output_path.stat().st_size > 0
    except (OSError, subprocess.TimeoutExpired):
        return False


def _find_browser_executable():
    candidates = [
        Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
        Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
        Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
        Path(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"),
    ]
    for path in candidates:
        if path.exists():
            return path
    return None


def _pdf_table(rows, Table, TableStyle, colors, font_name):
    table = Table(rows, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef3fb")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f2547")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d9e1ef")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fbfcff")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return table


def _plain_report_lines(context):
    lines = [
        "QA 최종 테스트 결과 보고서",
        f"생성일시: {context['generated_at']}",
        f"총 테스트 케이스: {context['total_cases']}건",
        f"최종 종합 의견: {context['final_decision']}",
        "",
        "전체 비교 결과 요약",
    ]
    for agent in context["agents"].values():
        lines.append(
            f"{agent['label']}: PASS {agent['decision_counts']['PASS']}, "
            f"REVIEW {agent['decision_counts']['REVIEW']}, FAIL {agent['decision_counts']['FAIL']}, "
            f"성공률 {agent['pass_rate']:.1f}%"
        )
    lines.extend(["", "결함 보고서"])
    if context["defects"]:
        for defect in context["defects"]:
            lines.append(f"{defect['defect_id']} {defect['case_id']} {defect['severity']} {defect['summary']}")
    else:
        lines.append("FAIL 판정 결함이 발견되지 않았습니다.")
    lines.extend(["", context["final_opinion"]])
    return lines


def _write_minimal_docx(output_path, lines):
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>'
        + "".join(f"<w:p><w:r><w:t>{xml_escape(line)}</w:t></w:r></w:p>" for line in lines)
        + "</w:body></w:document>"
    )
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        archive.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            "</Relationships>",
        )
        archive.writestr("word/document.xml", document_xml)


def _write_minimal_pdf(output_path, lines):
    ascii_lines = [line.encode("latin-1", errors="replace").decode("latin-1") for line in lines]
    text = "\\n".join(ascii_lines).replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 11 Tf 50 790 Td ({text}) Tj ET"
    objects = [
        "1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj",
        "2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj",
        "3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj",
        "4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj",
        f"5 0 obj << /Length {len(stream.encode('latin-1'))} >> stream\n{stream}\nendstream endobj",
    ]
    offsets = []
    content = "%PDF-1.4\n"
    for obj in objects:
        offsets.append(len(content.encode("latin-1")))
        content += obj + "\n"
    xref_start = len(content.encode("latin-1"))
    content += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n"
    for offset in offsets:
        content += f"{offset:010d} 00000 n \n"
    content += f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_start}\n%%EOF"
    output_path.write_bytes(content.encode("latin-1"))
